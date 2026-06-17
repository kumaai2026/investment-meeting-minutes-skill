#!/usr/bin/env python3
"""Sanitize Chinese investment meeting minutes into DOCX and JSONL outputs."""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable


ANONYMIZATION_LEVEL = "L2_FACT_PRESERVED"
SUPPORTED_SUFFIXES = {".md", ".txt"}

FILLER_PATTERNS = [
    r"\bOK\b",
    r"\bok\b",
    r"嗯+",
    r"呃+",
    r"啊+",
    r"哈+",
    r"哈哈+",
    r"这个",
    r"那个",
    r"就是说",
    r"就是",
    r"其实",
    r"然后",
    r"对吧",
    r"是不是",
    r"怎么说呢",
    r"坦白讲",
    r"老实说",
]

GENERIC_SPEAKER_WORDS = [
    "某某",
    "某人",
    "发言人",
    "发言人A",
    "发言人B",
    "嘉宾",
    "专家",
    "老师",
    "主持人",
]

IDENTITY_LINE_RE = re.compile(
    r"^\s*(发言人|姓名|身份|职位|职务|机构|公司|部门|所在地|地区|履历|背景|介绍|会议预定人|预定人|主持人)\s*[:：]"
)

TIMESTAMP_RES = [
    re.compile(r"[（(]\s*录音约\s*\d{1,2}:\d{2}(?::\d{2})?\s*[）)]"),
    re.compile(r"\[\s*\d{1,2}:\d{2}(?::\d{2})?\s*\]"),
    re.compile(r"(?<!\d)\d{1,2}:\d{2}:\d{2}(?!\d)"),
    re.compile(r"(?<!\d)\d{1,2}:\d{2}(?!\d)"),
    re.compile(r"\d{4}[-/年]\d{1,2}[-/月]\d{1,2}日?\s+\d{1,2}:\d{2}(?::\d{2})?"),
]

ATTRIBUTION_VERBS = "认为|表示|说|讲|提到|指出|判断|反馈|强调|分享|补充|称|介绍"
GENERIC_ATTRIBUTION_RE = re.compile(
    rf"(?:{'|'.join(map(re.escape, GENERIC_SPEAKER_WORDS))})\s*(?:{ATTRIBUTION_VERBS})[，,：:\s]*"
)


@dataclass
class TopicUnit:
    full_topic: str
    topic: str
    target: str
    text: str
    entities: list[str]


def read_input(path: Path) -> str:
    if path.suffix.lower() not in SUPPORTED_SUFFIXES:
        raise SystemExit("Only .md and .txt input is supported. Convert .docx to Markdown or plain text first.")
    return path.read_text(encoding="utf-8-sig")


def parse_meeting_date(text: str, override: str | None = None) -> str:
    if override:
        return normalize_date(override)

    plain_text = strip_markdown_emphasis(text)
    labeled = re.search(
        r"(?:会议日期|日期|会议时间|时间)\s*[:：]?\s*(\d{4})[-/年](\d{1,2})[-/月](\d{1,2})日?",
        plain_text,
    )
    if labeled:
        return normalize_date("-".join(labeled.groups()))

    first_date = re.search(r"(\d{4})[-/年](\d{1,2})[-/月](\d{1,2})日?", plain_text)
    if first_date:
        return normalize_date("-".join(first_date.groups()))

    return "unknown"


def normalize_date(raw: str) -> str:
    match = re.search(r"(\d{4})[-/年](\d{1,2})[-/月](\d{1,2})日?", raw)
    if not match:
        return raw
    year, month, day = match.groups()
    return f"{int(year):04d}-{int(month):02d}-{int(day):02d}"


def parse_meeting_type(text: str) -> str:
    match = re.search(r"(?:会议类型|类型)\s*[:：]\s*([^\n\r]+)", strip_markdown_emphasis(text))
    if not match:
        return "未识别"
    return neutralize_text(match.group(1), []).rstrip("。") or "未识别"


def strip_markdown_emphasis(text: str) -> str:
    return re.sub(r"[*_`]+", "", text)


def collect_speaker_names(text: str) -> list[str]:
    names: list[str] = []
    for match in re.finditer(r"^\s*###\s+(.+?)\s*$", text, flags=re.MULTILINE):
        raw_name = re.sub(r"[#*`_>\s]+", "", match.group(1))
        raw_name = re.sub(r"[（(].*?[）)]", "", raw_name)
        if is_probable_person_speaker(raw_name):
            names.append(raw_name)
    return dedupe(names)


def is_probable_person_speaker(name: str) -> bool:
    if not name or len(name) > 8:
        return False
    company_markers = "公司|股份|集团|证券|基金|资本|银行|保险|科技|电子|光电|材料|产业链"
    if re.search(company_markers, name):
        return False
    if re.search(r"\d{6}\.(?:SZ|SH|BJ)", name, flags=re.I):
        return False
    return bool(re.search(r"[\u4e00-\u9fffA-Za-z]", name))


def remove_timestamps(text: str) -> str:
    cleaned = text
    for pattern in TIMESTAMP_RES:
        cleaned = pattern.sub("", cleaned)
    return cleaned


def strip_speaker_headings_and_identity(text: str) -> str:
    kept_lines: list[str] = []
    for raw_line in text.splitlines():
        line = remove_timestamps(raw_line).strip()
        if not line:
            kept_lines.append("")
            continue
        if re.match(r"^###\s+.+$", line):
            continue
        if re.match(r"^\*\*[^*\n]{2,80}[？?]\*\*$", line):
            continue
        if IDENTITY_LINE_RE.match(line):
            continue
        line = re.sub(r"^\s*(?:发言人[A-Z]?|嘉宾|专家|老师|主持人|[\u4e00-\u9fff]{2,4}(?:总|老师|博士|经理)?)\s*[:：]\s*", "", line)
        kept_lines.append(line)
    return "\n".join(kept_lines)


def split_pending_section(text: str) -> tuple[str, list[str]]:
    pattern = re.compile(r"^\s{0,3}#{0,6}\s*(?:[一二三四五六七八九十]、)?(?:存疑与待确认|待确认业务事项|业务存疑事项|存疑事项)\s*$", re.MULTILINE)
    match = pattern.search(text)
    if not match:
        return text, []

    main_text = text[: match.start()]
    pending_text = text[match.end() :]
    pending_items: list[str] = []
    for line in pending_text.splitlines():
        item = line.strip()
        if not item or re.match(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?$", item):
            continue
        if item.startswith("#"):
            break
        item = re.sub(r"^\s*[-*]\s*", "", item)
        item = re.sub(r"^\|", "", item).rstrip("|")
        item = "；".join(part.strip() for part in item.split("|") if part.strip())
        compact_item = re.sub(r"[\s；;|]", "", item)
        if "发言人" in compact_item and ("事项" in compact_item or "时间" in compact_item):
            continue
        if "原始表述" in compact_item and "当前判断" in compact_item:
            continue
        pending_items.append(item)
    return main_text, pending_items


def split_topic_units(text: str, speaker_names: list[str]) -> list[TopicUnit]:
    matches = list(re.finditer(r"【([^】]+)】", text))
    if not matches:
        cleaned = neutralize_text(text, speaker_names)
        return [make_topic_unit("未分主题", cleaned)] if cleaned else []

    units: list[TopicUnit] = []
    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        topic_label = match.group(1).strip()
        body = neutralize_text(text[start:end], speaker_names)
        if body:
            units.append(make_topic_unit(topic_label, body))
    return units


def make_topic_unit(topic_label: str, body: str) -> TopicUnit:
    topic_label = strip_inline_markup(remove_timestamps(topic_label)).strip()
    topic, target = split_topic_label(topic_label)
    return TopicUnit(
        full_topic=topic_label,
        topic=topic,
        target=target,
        text=body,
        entities=extract_entities(topic, target, body),
    )


def split_topic_label(label: str) -> tuple[str, str]:
    parts = re.split(r"[｜|]", label, maxsplit=1)
    topic = parts[0].strip() if parts else label.strip()
    target = parts[1].strip() if len(parts) > 1 else ""
    return topic, target


def neutralize_text(text: str, speaker_names: list[str]) -> str:
    text = remove_timestamps(text)
    text = strip_inline_markup(text)
    text = strip_quote_marks(text)
    text = GENERIC_ATTRIBUTION_RE.sub("", text)

    for speaker in speaker_variants(speaker_names):
        escaped = re.escape(speaker)
        text = re.sub(rf"{escaped}\s*(?:{ATTRIBUTION_VERBS})[，,：:\s]*", "", text)
        text = re.sub(rf"据\s*{escaped}\s*(?:{ATTRIBUTION_VERBS})[，,：:\s]*", "", text)
        text = re.sub(rf"{escaped}\s*[:：]", "", text)
        text = re.sub(escaped, "", text)

    text = re.sub(r"据\s*(?:反馈|了解|调研)\s*[，,：:]*", "", text)
    text = re.sub(r"(?:我|个人|我们|咱们)\s*(?:认为|觉得|判断|感觉|看)\s*[，,：:]*", "", text)
    text = re.sub(r"(?:我的|我们的|个人的)\s*(?:观点|判断|感觉)\s*(?:是|为)?\s*[，,：:]*", "", text)
    text = re.sub(r"[^。！？]*不要传出去[^。！？]*[。！？]?", "", text)
    text = re.sub(r"[^。！？]*以我为准[^。！？]*[。！？]?", "", text)
    text = re.sub(r"我周末聊下来[，,]?", "周末交流显示，", text)
    text = re.sub(r"我这边了解的情况", "当前了解的情况", text)
    text = re.sub(r"因为?我今天在[^。！？]*[。！？]?", "", text)
    text = re.sub(r"我不太方便展开", "暂不展开", text)
    text = re.sub(r"我不确定", "不确定", text)
    text = re.sub(r"我了解到", "了解到", text)
    text = re.sub(r"我了解的", "了解的", text)
    text = re.sub(r"我建议咱就不用争了，可以好好看看", "可关注", text)
    text = re.sub(r"我简单再给大家说一说", "", text)
    text = re.sub(r"给大家说一说", "", text)
    text = re.sub(r"等我研究好了再跟大家汇报", "后续进展待跟踪", text)
    text = re.sub(r"我会讲一讲", "后续可跟踪", text)
    text = re.sub(r"我主要看", "主要关注", text)
    text = re.sub(r"我计划", "计划", text)
    text = re.sub(r"我看", "", text)
    text = re.sub(r"我觉得", "", text)
    text = re.sub(r"我感觉", "", text)
    text = re.sub(r"我认为", "", text)
    text = re.sub(r"我认可", "", text)
    text = re.sub(r"我并不认同", "不认同", text)
    text = re.sub(r"我不明白", "尚不明确", text)
    text = re.sub(r"我有点不能理解", "市场理解存在分歧", text)
    text = re.sub(r"我都有点不能理解", "市场理解存在分歧", text)
    text = re.sub(r"不能理解市场", "市场理解存在分歧", text)
    text = re.sub(r"我一直", "", text)
    text = re.sub(r"我现在", "当前", text)
    text = re.sub(r"我本来希望", "原本预期", text)
    text = re.sub(r"我真的觉得", "", text)
    text = re.sub(r"我这边", "", text)
    text = re.sub(r"我这里", "", text)
    text = re.sub(r"我自己", "", text)
    text = re.sub(r"结合我对", "结合对", text)
    text = re.sub(r"咱们|咱", "", text)
    text = re.sub(r"大家", "市场", text)
    text = re.sub(r"市场好[。！？]?", "", text)
    text = re.sub(r"上周四天风国际的郭明錤在推特上有一个观点，他说", "", text)
    text = re.sub(r"郭明錤也讲了一些[，,。]?", "", text)
    text = re.sub(r"他也说了[，,。]?", "", text)
    text = re.sub(r"陈涛他们周五、周六、周日都在搞的", "", text)
    text = re.sub(r"陈涛身边的人", "相关人士", text)
    text = re.sub(r"二爷[^。！？]*[。！？]?", "", text)
    text = re.sub(r"其他老师还有补充吗[^。！？]*[。！？]?", "", text)
    text = re.sub(r"[^。！？]*会议就结束[^。！？]*[。！？]?", "", text)
    text = re.sub(r"拜拜[。！？]?", "", text)
    text = re.sub(r"市场买它不买吗[？?]", "市场关注点在于高端产品。", text)
    text = re.sub(r"当前有点不能理解市场[。！？]?", "市场对该逻辑理解存在分歧。", text)
    text = re.sub(r"人家已经把公告发出来了", "相关公告已发布", text)
    text = re.sub(r"方案都是你们研发的，那器件不采购你们，还采购谁[？?]", "方案由炬光研发，器件采购具备延展空间。", text)
    text = re.sub(r"自己的这么大一个事件", "该事件", text)
    text = re.sub(r"我的|我们的|我们|我[都也]?", "", text)
    text = re.sub(r"发言人\s*未在", "原文未在", text)
    text = re.sub(r"发言人\s*未", "原文未", text)

    for pattern in FILLER_PATTERNS:
        text = re.sub(pattern, "", text)

    lines = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if re.search(r"(会议就结束|拜拜|其他老师还有补充吗)", line):
            continue
        if re.search(r"谢谢.+接下来.+(?:开始分享|分享)", line):
            continue
        if re.match(r"^#{1,6}\s+", line):
            continue
        if IDENTITY_LINE_RE.match(line):
            continue
        line = re.sub(r"^\s*[-*]\s*", "", line)
        line = re.sub(r"\s+", " ", line)
        line = re.sub(r"\s*([，。；：、,.!?！？])\s*", r"\1", line)
        line = re.sub(r"[，,；;：:]+$", "。", line)
        lines.append(line)

    return normalize_sentences(" ".join(lines))


def speaker_variants(speaker_names: list[str]) -> list[str]:
    variants: list[str] = []
    for name in speaker_names:
        add_entity(variants, name)
        for suffix in ("老师", "博士"):
            if name.endswith(suffix) and len(name) > len(suffix) + 1:
                add_entity(variants, name[: -len(suffix)])
        if name.endswith("博士") and len(name) >= 3:
            add_entity(variants, name[:-1])
    return sorted(variants, key=len, reverse=True)


def strip_inline_markup(text: str) -> str:
    text = re.sub(r"</?u>", "", text, flags=re.I)
    text = re.sub(r"</?strong>", "", text, flags=re.I)
    text = re.sub(r"</?b>", "", text, flags=re.I)
    return re.sub(r"[*_`]+", "", text)


def strip_quote_marks(text: str) -> str:
    replacements = {
        "“": "",
        "”": "",
        "‘": "",
        "’": "",
        "「": "",
        "」": "",
        "『": "",
        "』": "",
        '"': "",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def normalize_sentences(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"也[，,]\s*从", "从", text)
    text = re.sub(r"([。！？]){2,}", r"\1", text)
    text = re.sub(r"([，,]){2,}", r"\1", text)
    text = text.strip(" ，,；;：:")
    if text and text[-1] not in "。！？.!?":
        text += "。"
    return text


def extract_entities(topic: str, target: str, body: str) -> list[str]:
    entities: list[str] = []
    target_without_codes = target
    for match in re.finditer(r"([\u4e00-\u9fffA-Za-z0-9·]{2,20})[（(](\d{5,6}\.(?:SZ|SH|BJ|HK|US))[）)]", target, flags=re.I):
        add_entity(entities, match.group(1))
        add_entity(entities, match.group(2).upper())
        target_without_codes = target_without_codes.replace(match.group(0), " ")

    for part in re.split(r"[、,，/；;\s]+", target_without_codes):
        add_entity(entities, part)

    for match in re.finditer(r"([\u4e00-\u9fffA-Za-z0-9·]{2,20})[（(](\d{5,6}\.(?:SZ|SH|BJ|HK|US))[）)]", body, flags=re.I):
        add_entity(entities, match.group(1))
        add_entity(entities, match.group(2).upper())

    for match in re.finditer(r"(?:^|[，,。；;\s])([\u4e00-\u9fffA-Za-z0-9·]{2,8}?)(?:已发出|发出|公告|披露|订单|下修|上修|扩产|量产|验证|交付)", body):
        add_entity(entities, match.group(1))

    if not entities and topic:
        add_entity(entities, topic)
    return entities


def add_entity(entities: list[str], value: str) -> None:
    value = value.strip(" 。，,；;：:（）()[]【】")
    value = re.sub(r"^[和与及对像给把在从看]+", "", value)
    if not value or value in entities:
        return
    stopwords = {"行业", "公司", "客户", "市场", "业务", "价格", "材料", "产品", "订单", "产能", "良率", "上游", "下游"}
    if value in stopwords:
        return
    entities.append(value)


def build_json_rows(units: list[TopicUnit], meeting_date: str) -> list[dict]:
    date_id = meeting_date.replace("-", "") if meeting_date != "unknown" else "unknown"
    rows: list[dict] = []
    for index, unit in enumerate(units, start=1):
        rows.append(
            {
                "chunk_id": f"meeting_{date_id}_unit_{index:03d}",
                "text": f"主题：{unit.full_topic}。{unit.text}",
                "metadata": {
                    "source_type": "meeting_minutes",
                    "meeting_date": meeting_date,
                    "topic": unit.topic,
                    "entities": unit.entities,
                    "speaker_identity": "removed",
                    "speaker_style": "neutralized",
                    "anonymization_level": ANONYMIZATION_LEVEL,
                    "business_facts": "preserved",
                },
            }
        )
    return rows


def quality_check(units: list[TopicUnit], pending_items: list[str], json_rows: list[dict], speaker_names: list[str]) -> None:
    combined = "\n".join([unit.full_topic + "\n" + unit.text for unit in units] + pending_items + [row["text"] for row in json_rows])
    issues: list[str] = []

    if re.search(r"^\s*###\s+", combined, flags=re.MULTILINE):
        issues.append("Markdown speaker heading remains.")
    if re.search(r"(?:某某|某人|发言人[A-Z]?|嘉宾|专家|老师)\s*(?:认为|表示|说|讲|提到|指出|反馈|强调)", combined):
        issues.append("Generic speaker attribution remains.")
    if re.search(r"发言人[A-Z]?", combined):
        issues.append("Speaker marker remains.")
    if re.search(r"[“\"「『][^”\"」』]{20,}[”\"」』]", combined):
        issues.append("Long direct quote remains.")
    for pattern in TIMESTAMP_RES:
        if pattern.search(combined):
            issues.append("Raw timestamp remains.")
            break
    for name in speaker_names:
        if name and re.search(re.escape(name), combined):
            issues.append(f"Collected speaker name remains: {name}")

    for row in json_rows:
        json.loads(json.dumps(row, ensure_ascii=False))

    if len(units) != len(json_rows):
        issues.append("DOCX topic units and JSONL rows differ in count.")

    if issues:
        raise SystemExit("Quality check failed:\n- " + "\n- ".join(dedupe(issues)))


def write_jsonl(rows: Iterable[dict], path: Path) -> None:
    with path.open("w", encoding="utf-8") as handle:
        for row in rows:
            handle.write(json.dumps(row, ensure_ascii=False) + "\n")


def write_docx(units: list[TopicUnit], pending_items: list[str], meeting_date: str, meeting_type: str, path: Path) -> None:
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except ImportError as exc:
        raise SystemExit("python-docx is required to create the sanitized DOCX. Install python-docx and rerun.") from exc

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    doc.add_heading("脱敏会议纪要", level=0)

    doc.add_heading("一、文档信息", level=1)
    for line in [
        f"会议日期：{meeting_date}",
        f"会议类型：{meeting_type}",
        f"脱敏等级：{ANONYMIZATION_LEVEL}",
        "处理说明：已删除发言人身份和发言风格，保留业务事实",
    ]:
        doc.add_paragraph(line)

    doc.add_heading("二、主题纪要", level=1)
    for unit in units:
        doc.add_paragraph(f"主题：{unit.full_topic}")
        doc.add_paragraph(unit.text)

    doc.add_heading("三、待确认业务事项", level=1)
    if pending_items:
        for item in pending_items:
            doc.add_paragraph(item)
    else:
        doc.add_paragraph("无。")

    doc.save(path)


def dedupe(values: Iterable[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for value in values:
        if value not in seen:
            seen.add(value)
            result.append(value)
    return result


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Sanitize Chinese investment meeting minutes into DOCX and JSONL.")
    parser.add_argument("input_file", help="Input .md or .txt meeting-minutes file")
    parser.add_argument("--output-dir", default="outputs", help="Output directory, default: outputs")
    parser.add_argument("--meeting-date", help="Override meeting date, format YYYY-MM-DD")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    input_path = Path(args.input_file).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()

    raw_text = read_input(input_path)
    meeting_date = parse_meeting_date(raw_text, args.meeting_date)
    meeting_type = parse_meeting_type(raw_text)
    speaker_names = collect_speaker_names(raw_text)

    stripped = strip_speaker_headings_and_identity(raw_text)
    topic_text, pending_raw = split_pending_section(stripped)
    units = split_topic_units(topic_text, speaker_names)
    pending_items = [neutralize_text(item, speaker_names) for item in pending_raw]
    pending_items = [item for item in pending_items if item]

    if not units:
        raise SystemExit("No usable topic content found after sanitization.")

    json_rows = build_json_rows(units, meeting_date)
    quality_check(units, pending_items, json_rows, speaker_names)

    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / f"{input_path.stem}_sanitized.docx"
    jsonl_path = output_dir / f"{input_path.stem}_rag.jsonl"
    write_docx(units, pending_items, meeting_date, meeting_type, docx_path)
    write_jsonl(json_rows, jsonl_path)

    print(f"Wrote {docx_path}")
    print(f"Wrote {jsonl_path}")
    print(f"Topic chunks: {len(units)}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
