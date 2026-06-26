#!/usr/bin/env python3
"""
Export a finalized meeting note to the user's Obsidian workflow as Markdown + Word.
"""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

DEFAULT_EXPORT_DIR = Path("/Users/kumaai/Documents/Codex/workspace/投资纪要工作流/01 Projects/会议纪要")
DEFAULT_REMOTE_DIR = "gdrive:投资纪要工作流存档/投资纪要工作流"
INVALID_FILENAME_CHARS = r'[\\/:*?"<>|]+'
CJK_PATTERN = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]")


def validate_utf8_text_file(path: Path, *, require_cjk: bool = False) -> tuple[bool, str]:
    try:
        text = path.read_bytes().decode("utf-8")
    except UnicodeDecodeError as exc:
        return False, f"{path}: 不是有效 UTF-8: {exc}"
    if "\ufffd" in text:
        return False, f"{path}: 检测到 Unicode 替换字符 U+FFFD，疑似编码损坏"
    if require_cjk and not CJK_PATTERN.search(text):
        return False, f"{path}: 未检测到中文字符"
    return True, "ok"


def validate_docx_utf8(path: Path, *, require_cjk: bool = False) -> tuple[bool, str]:
    xml_parts: list[str] = []
    try:
        with zipfile.ZipFile(path) as archive:
            for name in archive.namelist():
                if name.startswith("word/") and name.endswith(".xml"):
                    try:
                        xml_parts.append(archive.read(name).decode("utf-8"))
                    except UnicodeDecodeError as exc:
                        return False, f"{path}:{name}: DOCX XML 不是有效 UTF-8: {exc}"
    except zipfile.BadZipFile as exc:
        return False, f"{path}: 不是有效 DOCX/ZIP 文件: {exc}"
    text = "\n".join(xml_parts)
    if "\ufffd" in text:
        return False, f"{path}: DOCX XML 检测到 Unicode 替换字符 U+FFFD，疑似编码损坏"
    if require_cjk and not CJK_PATTERN.search(text):
        return False, f"{path}: DOCX XML 未检测到中文字符"
    return True, "ok"


@dataclass
class ExportResult:
    md_path: Path
    md_created: bool
    md_message: str
    docx_path: Path
    docx_created: bool
    docx_message: str
    sync_created: bool
    sync_message: str


def sanitize_filename(name: str) -> str:
    cleaned = re.sub(INVALID_FILENAME_CHARS, "-", name).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned or "未命名会议"


def markdown_field(markdown: str, field: str, fallback: str = "") -> str:
    pattern = re.compile(rf"^\*\*{re.escape(field)}\*\*[:：]\s*(.+?)\s*$", re.MULTILINE)
    match = pattern.search(markdown)
    return match.group(1).strip() if match else fallback


def detect_title(content: str, fallback: str) -> str:
    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            return sanitize_filename(stripped.lstrip("#").strip())
    return sanitize_filename(fallback)


def detect_filename_title(content: str, fallback: str) -> str:
    h1_title = detect_title(content, fallback)
    meeting_title_raw = markdown_field(content, "会议标题", "").strip()
    meeting_type_raw = markdown_field(content, "会议类型", "").strip()
    meeting_title = sanitize_filename(meeting_title_raw) if meeting_title_raw else ""
    meeting_type = sanitize_filename(meeting_type_raw) if meeting_type_raw else ""
    generic_titles = {"投资会议纪要", "会议纪要", "未命名会议"}

    if h1_title in generic_titles:
        display_title = meeting_title or sanitize_filename(fallback)
    elif h1_title.startswith("投资会议纪要｜"):
        topic_title = h1_title.removeprefix("投资会议纪要｜").strip()
        if meeting_title and meeting_title not in topic_title:
            display_title = f"{meeting_title}｜{topic_title}"
        else:
            display_title = topic_title or meeting_title or h1_title
    else:
        display_title = h1_title

    if meeting_type and meeting_type not in display_title:
        display_title = f"{display_title} - {meeting_type}"
    return sanitize_filename(display_title)


def normalize_meeting_date(date_override: str | None) -> str:
    raw = (date_override or "").strip()
    if raw:
        try:
            return datetime.strptime(raw, "%Y-%m-%d").strftime("%Y-%m-%d")
        except ValueError:
            pass
    return datetime.now().strftime("%Y-%m-%d")


def next_available_output_pair(export_dir: Path, filename_base: str) -> tuple[Path, Path]:
    """Return an output pair that does not overwrite an existing note."""
    md_path = export_dir / f"{filename_base}.md"
    docx_path = export_dir / f"{filename_base}.docx"
    if not md_path.exists() and not docx_path.exists():
        return md_path, docx_path

    stamp = datetime.now().strftime("%H%M%S")
    for idx in range(1, 1000):
        suffix = f"-{stamp}" if idx == 1 else f"-{stamp}-{idx}"
        candidate_md = export_dir / f"{filename_base}{suffix}.md"
        candidate_docx = export_dir / f"{filename_base}{suffix}.docx"
        if not candidate_md.exists() and not candidate_docx.exists():
            return candidate_md, candidate_docx
    raise FileExistsError(f"无法为 {filename_base} 生成未占用的输出文件名")


def escape_text(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def format_inline_html(text: str) -> str:
    chunks: list[str] = []
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        start, end = match.span()
        if start > cursor:
            chunks.append(escape_text(text[cursor:start]))
        chunks.append(f"<b>{escape_text(match.group(1))}</b>")
        cursor = end
    if cursor < len(text):
        chunks.append(escape_text(text[cursor:]))
    return "".join(chunks)


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator_row(cells: list[str]) -> bool:
    if not cells:
        return False
    return all(cell and set(cell) <= {"-", ":", " "} for cell in cells)


def _add_docx_runs(paragraph, text: str) -> None:
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        start, end = match.span()
        if start > cursor:
            _style_run(paragraph.add_run(text[cursor:start]))
        run = paragraph.add_run(match.group(1))
        run.bold = True
        run.underline = True
        _style_run(run)
        cursor = end
    if cursor < len(text):
        _style_run(paragraph.add_run(text[cursor:]))


def _set_run_font(run, font_name: str = "PingFang SC", size_pt: int | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = font_name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = r_pr._add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font_name)


def _style_run(run, font_name: str = "PingFang SC", size_pt: int | None = None):
    _set_run_font(run, font_name=font_name, size_pt=size_pt)
    return run


def _set_style_font(style, font_name: str, size_pt: int | None = None, bold: bool | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    style.font.name = font_name
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = r_pr._add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font_name)


def _set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_like_paragraph_shading(paragraph, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _format_paragraph(paragraph, *, space_after: int = 6, line_spacing: float = 1.18) -> None:
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt

    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = line_spacing


def _format_table(table) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Pt

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    _set_run_font(run, size_pt=9 if row_idx else 10)
                    if row_idx == 0:
                        run.bold = True
            if row_idx == 0:
                _set_cell_shading(cell, "1F4E79")
            elif row_idx % 2 == 0:
                _set_cell_shading(cell, "F6F8FB")


def convert_markdown_to_docx(source_md: Path, target_docx: Path) -> tuple[bool, str]:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt, RGBColor
    except Exception as exc:
        return False, f"缺少 Word 依赖 python-docx: {exc}"

    lines = source_md.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    _set_style_font(doc.styles["Normal"], "PingFang SC", 10)
    for style_name, size, bold in [
        ("Heading 1", 18, True),
        ("Heading 2", 15, True),
        ("Heading 3", 12, True),
        ("List Bullet", 10, None),
    ]:
        if style_name in doc.styles:
            _set_style_font(doc.styles[style_name], "PingFang SC", size, bold)

    idx = 0
    while idx < len(lines):
        stripped = lines[idx].strip()

        if not stripped or stripped == "---":
            doc.add_paragraph("")
            idx += 1
            continue

        if stripped.startswith("# "):
            paragraph = doc.add_heading(stripped[2:].strip(), level=1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(14)
            for run in paragraph.runs:
                _set_run_font(run, size_pt=18)
                run.bold = True
                run.font.color.rgb = RGBColor(31, 78, 121)
            idx += 1
            continue
        if stripped.startswith("## "):
            paragraph = doc.add_heading(stripped[3:].strip(), level=2)
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(8)
            for run in paragraph.runs:
                _set_run_font(run, size_pt=15)
                run.bold = True
                run.font.color.rgb = RGBColor(31, 78, 121)
            idx += 1
            continue
        if stripped.startswith("### "):
            paragraph = doc.add_heading(stripped[4:].strip(), level=3)
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
            for run in paragraph.runs:
                _set_run_font(run, size_pt=12)
                run.bold = True
                run.font.color.rgb = RGBColor(64, 64, 64)
            idx += 1
            continue

        if re.match(r"^\*\*[^*]+\*\*：", stripped):
            label, value = stripped.split("：", 1)
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"{label.strip('*')}：")
            run.bold = True
            _style_run(run, size_pt=10)
            _add_docx_runs(paragraph, value.strip())
            _format_paragraph(paragraph, space_after=3)
            idx += 1
            continue

        if stripped.startswith("- "):
            while idx < len(lines) and lines[idx].strip().startswith("- "):
                bullet_text = lines[idx].strip()[2:].strip()
                paragraph = doc.add_paragraph(style="List Bullet")
                _add_docx_runs(paragraph, bullet_text)
                _format_paragraph(paragraph, space_after=2, line_spacing=1.1)
                idx += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx].strip())
                idx += 1
            rows = [split_table_row(item) for item in table_lines]
            if len(rows) >= 2 and is_separator_row(rows[1]):
                rows.pop(1)
            if rows:
                col_count = max(len(row) for row in rows)
                normalized = [row + [""] * (col_count - len(row)) for row in rows]
                table = doc.add_table(rows=1, cols=col_count)
                table.style = "Table Grid"
                for col_idx, value in enumerate(normalized[0]):
                    _add_docx_runs(table.rows[0].cells[col_idx].paragraphs[0], value)
                for row in normalized[1:]:
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row):
                        _add_docx_runs(row_cells[col_idx].paragraphs[0], value)
                _format_table(table)
            continue

        paragraph = doc.add_paragraph()
        _add_docx_runs(paragraph, stripped)
        if stripped.startswith("【") and stripped.endswith("】"):
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(3)
            _set_cell_like_paragraph_shading(paragraph, "EAF2F8")
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(31, 78, 121)
        else:
            _format_paragraph(paragraph)
        idx += 1

    doc.save(target_docx)
    return True, "ok"


def sync_vault_to_gdrive(local_dir: Path, remote_dir: str | None = None) -> tuple[bool, str]:
    sync_script = Path(__file__).with_name("sync_obsidian_to_gdrive.py")
    if not sync_script.exists():
        return False, f"同步脚本不存在: {sync_script}"

    log_file = Path("/Users/kumaai/Library/Logs/kumaai-sync/investment-workflow-rclone-launch.log")
    log_file.parent.mkdir(parents=True, exist_ok=True)
    try:
        handle = log_file.open("a", encoding="utf-8")
        command = [sys.executable, str(sync_script), "--local-dir", str(local_dir)]
        if remote_dir:
            command.extend(["--remote-dir", remote_dir])
        subprocess.Popen(
            command,
            stdout=handle,
            stderr=subprocess.STDOUT,
            start_new_session=True,
        )
    except Exception as exc:
        return False, f"Google Drive 后台同步启动失败: {exc}"
    target = remote_dir or DEFAULT_REMOTE_DIR
    return True, f"已触发 Google Drive 后台同步: {local_dir} -> {target}"


def organize_archive_before_sync() -> tuple[bool, str]:
    organizer = Path(__file__).with_name("organize_archive_by_date.py")
    if not organizer.exists():
        return False, f"归档整理脚本不存在: {organizer}"
    result = subprocess.run([sys.executable, str(organizer)], text=True, capture_output=True)
    message = (result.stdout or result.stderr or "").strip()
    if result.returncode != 0:
        return False, message or "归档整理失败"
    return True, message or "归档整理完成"


def export_note(source_file: Path, export_dir: Path, date_override: str | None, *, sync: bool = True) -> ExportResult:
    raw_content = source_file.read_text(encoding="utf-8")
    source_encoding_ok, source_encoding_message = validate_utf8_text_file(source_file, require_cjk=True)
    if not source_encoding_ok:
        raise UnicodeError(source_encoding_message)
    meeting_date = normalize_meeting_date(date_override)
    title = detect_filename_title(raw_content, source_file.stem)
    filename_base = f"{meeting_date} - {title}"
    export_dir = export_dir / meeting_date
    export_dir.mkdir(parents=True, exist_ok=True)

    md_path, docx_path = next_available_output_pair(export_dir, filename_base)

    try:
        shutil.copy2(source_file, md_path)
        md_ok, md_message = validate_utf8_text_file(md_path, require_cjk=True)
    except Exception as exc:
        md_ok = False
        md_message = str(exc)

    docx_source = md_path if md_ok else source_file
    docx_ok, docx_message = convert_markdown_to_docx(docx_source, docx_path)
    if docx_ok:
        docx_ok, docx_message = validate_docx_utf8(docx_path, require_cjk=True)
    sync_ok = False
    sync_message = "跳过同步：Markdown 或 Word 未全部生成"
    if md_ok and docx_ok and sync:
        archive_ok, archive_message = organize_archive_before_sync()
        sync_remote_dir = f"{DEFAULT_REMOTE_DIR}/01 Projects/会议纪要/{meeting_date}"
        sync_ok, sync_message = sync_vault_to_gdrive(export_dir, sync_remote_dir)
        if not archive_ok:
            sync_message = f"{sync_message}; 归档整理未完成: {archive_message}"
    elif md_ok and docx_ok:
        sync_message = "跳过同步：--no-sync"

    return ExportResult(
        md_path=md_path,
        md_created=md_ok,
        md_message=md_message,
        docx_path=docx_path,
        docx_created=docx_ok,
        docx_message=docx_message,
        sync_created=sync_ok,
        sync_message=sync_message,
    )


def main() -> int:
    parser = argparse.ArgumentParser(description="导出投资会议纪要到 Obsidian 目录（仅 Markdown+Word）")
    parser.add_argument("input_file", help="已整理完成的 Markdown 文件")
    parser.add_argument("--export-dir", default=str(DEFAULT_EXPORT_DIR), help=f"导出目录，默认 {DEFAULT_EXPORT_DIR}")
    parser.add_argument("--meeting-date", help="覆盖系统日期，格式 YYYY-MM-DD")
    parser.add_argument("--no-sync", action="store_true", help="只生成本地 Markdown+Word，不触发归档整理或 Google Drive 同步")
    args = parser.parse_args()

    source_file = Path(args.input_file).expanduser().resolve()
    if not source_file.exists():
        print(f"输入文件不存在: {source_file}", file=sys.stderr)
        return 1

    export_dir = Path(args.export_dir).expanduser().resolve()
    result = export_note(source_file, export_dir, args.meeting_date, sync=not args.no_sync)

    if result.md_created:
        print(f"Markdown: {result.md_path}")
    else:
        print(f"Markdown: 未生成 ({result.md_message})")
    if result.docx_created:
        print(f"Word: {result.docx_path}")
    else:
        print(f"Word: 未生成 ({result.docx_message})")
    if result.sync_created:
        print(f"Google Drive: {result.sync_message}")
    else:
        print(f"Google Drive: 未同步 ({result.sync_message})")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
