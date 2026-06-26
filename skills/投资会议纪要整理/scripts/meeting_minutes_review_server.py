#!/usr/bin/env python3
"""Local review/history server for Dify-generated meeting-minute drafts."""

from __future__ import annotations

import html
import hashlib
import hmac
import json
import mimetypes
import os
import re
import secrets
import shutil
import sqlite3
import subprocess
import sys
import tempfile
import threading
import time
import traceback
import urllib.parse
import urllib.request
from collections import Counter, defaultdict
from http import cookies
from datetime import datetime, timedelta
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any
from uuid import uuid4

SCRIPT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPT_DIR))

try:
    import cgi  # type: ignore
except ModuleNotFoundError:
    import compat_cgi as cgi  # type: ignore

from export_to_obsidian import DEFAULT_EXPORT_DIR, export_note, normalize_meeting_date  # noqa: E402
from archive_raw_inputs import DEFAULT_ARCHIVE_ROOT as RAW_ARCHIVE_ROOT, archive_files as archive_raw_files  # noqa: E402
from sync_minutes_to_dify_dataset import (  # noqa: E402
    DEFAULT_CONFIG_PATH,
    build_metadata,
    load_config,
    load_env_file,
    load_mapping,
    mapping_path_from_config,
    sync_markdown_to_dify,
)
from query_symbol_candidates import DEFAULT_ALIAS_PATH, DEFAULT_SYMBOL_ROOT, query_symbols  # noqa: E402
from generate_post_review_artifacts import generate_artifacts  # noqa: E402

DEFAULT_HOST = "0.0.0.0"
DEFAULT_PORT = 8767
KUMA_FRONTEND_PORT = 1000
MAX_POST_BYTES = 500 * 1024 * 1024
MAX_IMPORT_CHUNK_BYTES = max(1, int(os.environ.get("MEETING_IMPORT_CHUNK_MB", "8") or "8")) * 1024 * 1024
DEFAULT_STORAGE_DIR = Path("/Users/kumaai/Library/Application Support/kumaai-sync/meeting-minutes-review/drafts")
DEFAULT_OUTPUT_DIR = Path("/Users/kumaai/Documents/Codex/workspace/投资纪要工作流/01 Projects/会议纪要")
DEFAULT_STRUCTURED_TABLE_ARCHIVE_DIR = Path("/Users/kumaai/Documents/Codex/workspace/投资纪要工作流/01 Projects/会议结构化表格")
STRUCTURED_TABLE_SKILL_DIR = Path("/Users/kumaai/.codex/skills/meeting-minutes-structured-table")
STRUCTURED_TABLE_SCRIPT = STRUCTURED_TABLE_SKILL_DIR / "scripts" / "generate_table.py"
DEFAULT_PLANNER_STATUS_PATH = Path(
    "/Users/kumaai/Library/Application Support/kumaai-sync/planner-agent/status.json"
)
DEFAULT_ACCESS_CONFIG_PATH = Path(
    "/Users/kumaai/Library/Application Support/kumaai-sync/meeting-minutes-review/access-control.json"
)
DEFAULT_ACCESS_DB_PATH = Path(
    "/Users/kumaai/Library/Application Support/kumaai-sync/meeting-minutes-review/access-users.sqlite3"
)
DEFAULT_ADMIN_TOKEN_PATH = Path(
    "/Users/kumaai/Library/Application Support/kumaai-sync/meeting-minutes-review/admin-access-token.txt"
)
DEFAULT_ADMIN_PASSWORD_PATH = Path(
    "/Users/kumaai/Library/Application Support/kumaai-sync/meeting-minutes-review/admin-access-password.txt"
)
DEFAULT_ACCESS_SESSION_DIR = Path(
    "/Users/kumaai/Library/Application Support/kumaai-sync/meeting-minutes-review/sessions"
)
DEFAULT_MCP_CONFIG_PATH = Path(
    "/Users/kumaai/Library/Application Support/kumaai-sync/meeting-minutes-review/mcp-servers.json"
)
DIFY_BASE_URL = "http://localhost:18081"
DIFY_DB_CONTAINER = "docker-db_postgres-1"
DIFY_WORKFLOW_APP_ID = "8b8b90b1-432c-414a-842a-7426c628ae39"
SITE_NAME = "D91 AI投研平台"
PRODUCT_FAMILY = "OpenFinance"
MEETING_TOOL_NAME = "OpenMinute"
MEETING_SERIES_DEFAULT = "研究所周会"
DEFAULT_SENSEVOICE_MODEL_CACHE = os.environ.get(
    "SENSEVOICE_MODEL_CACHE",
    os.environ.get(
        "FUNASR_MODEL_CACHE",
        str(Path.home() / ".cache/modelscope/hub"),
    ),
)
DEFAULT_SENSEVOICE_PYTHON = os.environ.get("SENSEVOICE_PYTHON", "")
LOG_DIR = Path("/Users/kumaai/Library/Logs/kumaai-sync")
RCLONE_LOG_PATH = LOG_DIR / "investment-workflow-rclone.log"
ACCESS_AUDIT_LOG_PATH = LOG_DIR / "meeting-minutes-access-audit.jsonl"
VAULT_DIR = Path("/Users/kumaai/Documents/Codex/workspace/投资纪要工作流")
WECHAT_ARTICLE_DIR = VAULT_DIR / "01 Projects/微信公众号文章归档"
WECHAT_CHAT_DIR = VAULT_DIR / "01 Projects/微信聊天分析"
EXTERNAL_SOURCE_ROOTS = [
    ("wechat_article", "微信公众号文章", WECHAT_ARTICLE_DIR),
    ("wechat_chat", "微信聊天记录", WECHAT_CHAT_DIR),
]
PUBLIC_HOSTS = {"kuma.d91.global", "dify.d91.global"}
PLATFORM_PROXY_BASE_URL = os.environ.get("KUMA_PLATFORM_PROXY_BASE_URL", "http://127.0.0.1:8080").rstrip("/")
PLATFORM_PAGE_PATHS = ("/", "/dashboards", "/minutes", "/mp")
PUBLIC_PLATFORM_PAGE_PATHS = {"/", "/dashboards"}
PLATFORM_API_PREFIXES = ("/api/v1/dashboards", "/api/v1/mp")
PLATFORM_ASSET_PREFIXES = ("/_next/", "/images/")
AUTH_COOKIE_NAME = "workflow_access_token"
PASSWORD_HASH_ITERATIONS = 260000
VALID_ACCESS_ROLES = {"admin", "editor", "viewer", "external_viewer"}
MEETING_TYPES = ["多人复盘会", "上市公司交流", "专家交流"]
MEETING_SERIES = [MEETING_SERIES_DEFAULT]
CUSTOM_CHOICE_VALUE = "__custom__"
TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".csv", ".tsv"}
DOCX_SUFFIXES = {".docx"}
AUDIO_SUFFIXES = {".mp3", ".m4a", ".wav", ".aac", ".flac", ".ogg", ".wma", ".mp4"}
VIDEO_SUFFIXES = {".mp4", ".webm", ".mpeg", ".mpga"}
IMPORT_JOB_FILENAME = "import-job.json"
CONFIRM_INPUT_JOB_FILENAME = "confirm-input-job.json"
LOGIN_RATE_WINDOW_SECONDS = 10 * 60
LOGIN_RATE_LOCK_SECONDS = 15 * 60
LOGIN_RATE_MAX_FAILURES_PER_IP = 20
LOGIN_RATE_MAX_FAILURES_PER_USER = 8
LOGIN_ATTEMPTS: dict[str, list[float]] = {}
LOGIN_ATTEMPT_LOCK = threading.Lock()


def now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def read_config() -> dict[str, str]:
    values = load_env_file(DEFAULT_CONFIG_PATH)
    for key, value in os.environ.items():
        if key.startswith(("DIFY_", "MEETING_")):
            values[key] = value
    return values


def storage_dir() -> Path:
    values = read_config()
    return Path(values.get("MEETING_REVIEW_STORAGE") or str(DEFAULT_STORAGE_DIR)).expanduser()


def output_dir() -> Path:
    values = read_config()
    return Path(values.get("MEETING_FINAL_OUTPUT") or str(DEFAULT_OUTPUT_DIR)).expanduser()


def structured_table_archive_dir() -> Path:
    values = read_config()
    return Path(values.get("MEETING_STRUCTURED_TABLE_ARCHIVE") or str(DEFAULT_STRUCTURED_TABLE_ARCHIVE_DIR)).expanduser()


def planner_status_path() -> Path:
    values = read_config()
    return Path(values.get("MEETING_PLANNER_STATUS_PATH") or str(DEFAULT_PLANNER_STATUS_PATH)).expanduser()


def access_config_path() -> Path:
    values = read_config()
    return Path(values.get("MEETING_ACCESS_CONFIG_PATH") or str(DEFAULT_ACCESS_CONFIG_PATH)).expanduser()


def access_db_path() -> Path:
    values = read_config()
    return Path(values.get("MEETING_ACCESS_DB_PATH") or str(DEFAULT_ACCESS_DB_PATH)).expanduser()


def admin_token_path() -> Path:
    values = read_config()
    return Path(values.get("MEETING_ADMIN_TOKEN_PATH") or str(DEFAULT_ADMIN_TOKEN_PATH)).expanduser()


def admin_password_path() -> Path:
    values = read_config()
    return Path(values.get("MEETING_ADMIN_PASSWORD_PATH") or str(DEFAULT_ADMIN_PASSWORD_PATH)).expanduser()


def access_session_dir() -> Path:
    values = read_config()
    return Path(values.get("MEETING_ACCESS_SESSION_DIR") or str(DEFAULT_ACCESS_SESSION_DIR)).expanduser()


def access_disabled() -> bool:
    values = read_config()
    raw = str(values.get("MEETING_ACCESS_DISABLED") or "").strip().lower()
    return raw in {"1", "true", "yes", "on"}


def access_audit_log_path() -> Path:
    values = read_config()
    return Path(values.get("MEETING_ACCESS_AUDIT_LOG") or str(ACCESS_AUDIT_LOG_PATH)).expanduser()


def token_hash(token: str) -> str:
    return hashlib.sha256(token.encode("utf-8")).hexdigest()


def hash_password(password: str) -> str:
    salt = secrets.token_urlsafe(16)
    digest = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt.encode("utf-8"),
        PASSWORD_HASH_ITERATIONS,
    ).hex()
    return f"pbkdf2_sha256${PASSWORD_HASH_ITERATIONS}${salt}${digest}"


def verify_password(password: str, stored_hash: str) -> bool:
    if not password or not stored_hash:
        return False
    try:
        algorithm, iterations, salt, expected = stored_hash.split("$", 3)
        if algorithm != "pbkdf2_sha256":
            return False
        digest = hashlib.pbkdf2_hmac(
            "sha256",
            password.encode("utf-8"),
            salt.encode("utf-8"),
            int(iterations),
        ).hex()
        return hmac.compare_digest(expected, digest)
    except Exception:
        return False


def generate_access_password() -> str:
    return "Kuma-" + secrets.token_urlsafe(14)


def write_access_user_password(user_id: str, password: str) -> Path:
    if user_id == "admin":
        password_path = admin_password_path()
    else:
        password_dir = access_db_path().parent / "user-passwords"
        password_dir.mkdir(parents=True, exist_ok=True)
        password_path = password_dir / f"{sanitize_title(user_id)}.txt"
    password_path.parent.mkdir(parents=True, exist_ok=True)
    password_path.write_text(password + "\n", encoding="utf-8")
    password_path.chmod(0o600)
    return password_path


def identity_from_access_user(user: dict[str, Any], source: str) -> dict[str, Any]:
    return {
        "authenticated": True,
        "user_id": str(user.get("user_id") or "user"),
        "display_name": str(user.get("display_name") or user.get("user_id") or "user"),
        "role": str(user.get("role") or "member"),
        "groups": normalize_group_list(user.get("groups")),
        "source": source,
    }


def access_db_connect() -> sqlite3.Connection:
    path = access_db_path()
    path.parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(path)
    conn.row_factory = sqlite3.Row
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS access_users (
            user_id TEXT PRIMARY KEY,
            display_name TEXT NOT NULL DEFAULT '',
            role TEXT NOT NULL DEFAULT 'viewer',
            groups_json TEXT NOT NULL DEFAULT '[]',
            password_hash TEXT NOT NULL DEFAULT '',
            token_hash TEXT NOT NULL DEFAULT '',
            disabled INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        )
        """
    )
    conn.commit()
    path.chmod(0o600)
    return conn


def access_user_from_row(row: sqlite3.Row | dict[str, Any]) -> dict[str, Any]:
    raw_groups = row["groups_json"] if isinstance(row, sqlite3.Row) else row.get("groups_json", "[]")
    try:
        groups = json.loads(raw_groups or "[]")
    except Exception:
        groups = []
    return {
        "user_id": str(row["user_id"] if isinstance(row, sqlite3.Row) else row.get("user_id") or ""),
        "display_name": str(row["display_name"] if isinstance(row, sqlite3.Row) else row.get("display_name") or ""),
        "role": str(row["role"] if isinstance(row, sqlite3.Row) else row.get("role") or "viewer"),
        "groups": normalize_group_list(groups),
        "password_hash": str(row["password_hash"] if isinstance(row, sqlite3.Row) else row.get("password_hash") or ""),
        "token_hash": str(row["token_hash"] if isinstance(row, sqlite3.Row) else row.get("token_hash") or ""),
        "disabled": bool(row["disabled"] if isinstance(row, sqlite3.Row) else row.get("disabled")),
        "created_at": str(row["created_at"] if isinstance(row, sqlite3.Row) else row.get("created_at") or ""),
        "updated_at": str(row["updated_at"] if isinstance(row, sqlite3.Row) else row.get("updated_at") or ""),
    }


def insert_or_replace_access_user(conn: sqlite3.Connection, user: dict[str, Any]) -> None:
    timestamp = now_iso()
    created_at = str(user.get("created_at") or timestamp)
    updated_at = str(user.get("updated_at") or timestamp)
    conn.execute(
        """
        INSERT INTO access_users (
            user_id, display_name, role, groups_json, password_hash, token_hash, disabled, created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(user_id) DO UPDATE SET
            display_name=excluded.display_name,
            role=excluded.role,
            groups_json=excluded.groups_json,
            password_hash=excluded.password_hash,
            token_hash=excluded.token_hash,
            disabled=excluded.disabled,
            updated_at=excluded.updated_at
        """,
        (
            str(user.get("user_id") or "").strip(),
            str(user.get("display_name") or user.get("user_id") or ""),
            str(user.get("role") or "viewer"),
            json.dumps(normalize_group_list(user.get("groups")), ensure_ascii=False),
            str(user.get("password_hash") or ""),
            str(user.get("token_hash") or ""),
            1 if user.get("disabled") is True else 0,
            created_at,
            updated_at,
        ),
    )


def migrate_access_config_to_db(conn: sqlite3.Connection) -> None:
    if conn.execute("SELECT COUNT(*) FROM access_users").fetchone()[0]:
        return
    path = access_config_path()
    if not path.exists():
        return
    try:
        loaded = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return
    users = loaded.get("users") if isinstance(loaded, dict) else []
    if not isinstance(users, list):
        return
    for user in users:
        if isinstance(user, dict) and str(user.get("user_id") or "").strip():
            insert_or_replace_access_user(conn, user)
    conn.commit()


def ensure_access_config() -> dict[str, Any]:
    conn = access_db_connect()
    try:
        migrate_access_config_to_db(conn)
        admin_row = conn.execute("SELECT * FROM access_users WHERE user_id = ?", ("admin",)).fetchone()
        if not admin_row:
            token = secrets.token_urlsafe(32)
            password = generate_access_password()
            token_file = admin_token_path()
            token_file.parent.mkdir(parents=True, exist_ok=True)
            token_file.write_text(token + "\n", encoding="utf-8")
            token_file.chmod(0o600)
            write_access_user_password("admin", password)
            insert_or_replace_access_user(
                conn,
                {
                    "user_id": "admin",
                    "display_name": "管理员",
                    "role": "admin",
                    "groups": ["admin", "local"],
                    "token_hash": token_hash(token),
                    "password_hash": hash_password(password),
                    "created_at": now_iso(),
                    "updated_at": now_iso(),
                },
            )
            conn.commit()
        elif not str(admin_row["password_hash"] or "") and not bool(admin_row["disabled"]):
            password = generate_access_password()
            write_access_user_password("admin", password)
            conn.execute(
                "UPDATE access_users SET password_hash = ?, updated_at = ? WHERE user_id = ?",
                (hash_password(password), now_iso(), "admin"),
            )
            conn.commit()
        rows = conn.execute("SELECT * FROM access_users ORDER BY created_at, user_id").fetchall()
        return {"db_path": str(access_db_path()), "users": [access_user_from_row(row) for row in rows]}
    finally:
        conn.close()


def find_access_user(user_id: str) -> dict[str, Any] | None:
    user_id = (user_id or "").strip()
    if not user_id:
        return None
    ensure_access_config()
    conn = access_db_connect()
    try:
        row = conn.execute("SELECT * FROM access_users WHERE user_id = ?", (user_id,)).fetchone()
        return access_user_from_row(row) if row else None
    finally:
        conn.close()


def public_hosts() -> set[str]:
    values = read_config()
    configured = values.get("MEETING_PUBLIC_HOSTS") or ",".join(sorted(PUBLIC_HOSTS))
    return {item.strip().lower() for item in configured.split(",") if item.strip()}


def login_path(next_path: str = "/history") -> str:
    return "/access-login?next=" + urllib.parse.quote(next_path or "/history", safe="")


def normalize_group_list(value: Any) -> list[str]:
    if isinstance(value, str):
        return [item.strip() for item in value.split(",") if item.strip()]
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    return []


def validate_access_token(token: str) -> dict[str, Any] | None:
    if not token:
        return None
    config = ensure_access_config()
    hashed = token_hash(token)
    for user in config.get("users") or []:
        if not isinstance(user, dict):
            continue
        if user.get("disabled") is True:
            continue
        expected = str(user.get("token_hash") or "")
        if expected and hmac.compare_digest(expected, hashed):
            return identity_from_access_user(user, "token")
    return None


def validate_access_password(user_id: str, password: str) -> dict[str, Any] | None:
    user_id = (user_id or "").strip()
    if not user_id or not password:
        return None
    user = find_access_user(user_id)
    if not user or user.get("disabled") is True:
        return None
    if verify_password(password, str(user.get("password_hash") or "")):
        return identity_from_access_user(user, "password")
    return None


def normalize_login_user_id(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_.@-]", "", (value or "").strip())[:80]
    return cleaned


def _login_attempt_keys(client_key: str, user_id: str) -> tuple[str, str]:
    normalized_user = normalize_login_user_id(user_id).lower() or "anonymous"
    return f"ip:{client_key}", f"user:{client_key}:{normalized_user}"


def login_retry_after_seconds(client_key: str, user_id: str) -> int:
    now = time.time()
    cutoff = now - LOGIN_RATE_WINDOW_SECONDS
    ip_key, user_key = _login_attempt_keys(client_key, user_id)
    with LOGIN_ATTEMPT_LOCK:
        for key in list(LOGIN_ATTEMPTS):
            attempts = [stamp for stamp in LOGIN_ATTEMPTS.get(key, []) if stamp >= cutoff]
            if attempts:
                LOGIN_ATTEMPTS[key] = attempts
            else:
                LOGIN_ATTEMPTS.pop(key, None)
        ip_attempts = LOGIN_ATTEMPTS.get(ip_key, [])
        user_attempts = LOGIN_ATTEMPTS.get(user_key, [])
        if len(ip_attempts) >= LOGIN_RATE_MAX_FAILURES_PER_IP or len(user_attempts) >= LOGIN_RATE_MAX_FAILURES_PER_USER:
            oldest = min((ip_attempts + user_attempts) or [now])
            return max(1, int(LOGIN_RATE_LOCK_SECONDS - (now - oldest)))
    return 0


def record_login_failure(client_key: str, user_id: str) -> None:
    now = time.time()
    ip_key, user_key = _login_attempt_keys(client_key, user_id)
    with LOGIN_ATTEMPT_LOCK:
        LOGIN_ATTEMPTS.setdefault(ip_key, []).append(now)
        LOGIN_ATTEMPTS.setdefault(user_key, []).append(now)


def clear_login_failures(client_key: str, user_id: str) -> None:
    ip_key, user_key = _login_attempt_keys(client_key, user_id)
    with LOGIN_ATTEMPT_LOCK:
        LOGIN_ATTEMPTS.pop(user_key, None)
        if len(LOGIN_ATTEMPTS.get(ip_key, [])) <= 2:
            LOGIN_ATTEMPTS.pop(ip_key, None)


def create_access_session(user_id: str) -> str:
    session_token = secrets.token_urlsafe(32)
    session_path = access_session_dir() / f"{token_hash(session_token)}.json"
    session_path.parent.mkdir(parents=True, exist_ok=True)
    payload = {
        "user_id": user_id,
        "created_at": now_iso(),
        "expires_at": (datetime.now() + timedelta(days=30)).isoformat(timespec="seconds"),
    }
    session_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    session_path.chmod(0o600)
    return session_token


def validate_access_session(session_token: str) -> dict[str, Any] | None:
    if not session_token:
        return None
    session_path = access_session_dir() / f"{token_hash(session_token)}.json"
    if not session_path.exists():
        return None
    try:
        payload = json.loads(session_path.read_text(encoding="utf-8"))
        expires_at = datetime.fromisoformat(str(payload.get("expires_at") or ""))
        if expires_at < datetime.now():
            session_path.unlink(missing_ok=True)
            return None
        user = find_access_user(str(payload.get("user_id") or ""))
        if not user or user.get("disabled") is True:
            return None
        return identity_from_access_user(user, "session")
    except Exception:
        return None


def public_access_users() -> list[dict[str, Any]]:
    config = ensure_access_config()
    users: list[dict[str, Any]] = []
    for user in config.get("users") or []:
        if not isinstance(user, dict):
            continue
        users.append(
            {
                "user_id": str(user.get("user_id") or ""),
                "display_name": str(user.get("display_name") or ""),
                "role": str(user.get("role") or ""),
                "groups": normalize_group_list(user.get("groups")),
                "disabled": bool(user.get("disabled")),
                "created_at": str(user.get("created_at") or ""),
                "updated_at": str(user.get("updated_at") or ""),
                "password_present": bool(user.get("password_hash")),
                "token_present": bool(user.get("token_hash")),
            }
        )
    return users


def write_access_user_token(user_id: str, token: str) -> Path:
    token_dir = access_db_path().parent / "user-tokens"
    token_dir.mkdir(parents=True, exist_ok=True)
    token_path = token_dir / f"{sanitize_title(user_id)}.txt"
    token_path.write_text(token + "\n", encoding="utf-8")
    token_path.chmod(0o600)
    return token_path


def mutate_access_user(action: str, payload: dict[str, Any]) -> dict[str, Any]:
    ensure_access_config()
    user_id = str(payload.get("user_id") or "").strip()
    if not user_id:
        raise ValueError("缺少 user_id")
    if action == "rotate":
        action = "rotate_token"
    conn = access_db_connect()
    try:
        row = conn.execute("SELECT * FROM access_users WHERE user_id = ?", (user_id,)).fetchone()
        existing = access_user_from_row(row) if row else None
        if action in {"add", "update", "reset_password", "rotate_token"}:
            if action == "add" and existing:
                raise ValueError(f"用户已存在: {user_id}")
            if action != "add" and not existing:
                raise ValueError(f"用户不存在: {user_id}")
            if not existing:
                existing = {
                    "user_id": user_id,
                    "display_name": user_id,
                    "role": "viewer",
                    "groups": [],
                    "password_hash": "",
                    "token_hash": "",
                    "disabled": False,
                    "created_at": now_iso(),
                }
            role = str(payload.get("role") or existing.get("role") or "viewer")
            if role not in VALID_ACCESS_ROLES:
                raise ValueError(f"无效角色: {role}")
            groups_payload = payload.get("groups")
            if groups_payload in (None, ""):
                groups = existing.get("groups") or []
            else:
                groups = normalize_group_list(groups_payload)
            existing.update(
                {
                    "display_name": str(payload.get("display_name") or existing.get("display_name") or user_id),
                    "role": role,
                    "groups": groups,
                    "disabled": False,
                    "updated_at": now_iso(),
                }
            )
            if action in {"add", "reset_password"}:
                password = str(payload.get("password") or "").strip() or generate_access_password()
                password_path = write_access_user_password(user_id, password)
                existing["password_hash"] = hash_password(password)
                message = f"用户已{'新增' if action == 'add' else '重置密码'}，密码文件：{password_path}"
            elif action == "rotate_token":
                token = secrets.token_urlsafe(32)
                token_path = write_access_user_token(user_id, token)
                existing["token_hash"] = token_hash(token)
                message = f"用户 API 令牌已轮换，令牌文件：{token_path}"
            else:
                message = "用户资料已更新"
            insert_or_replace_access_user(conn, existing)
        elif action == "disable":
            if not existing:
                raise ValueError(f"用户不存在: {user_id}")
            conn.execute("UPDATE access_users SET disabled = 1, updated_at = ? WHERE user_id = ?", (now_iso(), user_id))
            message = "用户已禁用"
        elif action == "remove":
            if not existing:
                raise ValueError(f"用户不存在: {user_id}")
            conn.execute("DELETE FROM access_users WHERE user_id = ?", (user_id,))
            token_path = access_db_path().parent / "user-tokens" / f"{sanitize_title(user_id)}.txt"
            if token_path.exists():
                token_path.unlink()
            password_path = access_db_path().parent / "user-passwords" / f"{sanitize_title(user_id)}.txt"
            if password_path.exists():
                password_path.unlink()
            message = "用户已删除"
        else:
            raise ValueError(f"无效操作: {action}")
        conn.commit()
        return {"ok": True, "action": action, "user_id": user_id, "message": message, "db_path": str(access_db_path())}
    finally:
        conn.close()


def is_public_host(host: str) -> bool:
    hostname = host.split(":", 1)[0].strip().lower()
    return hostname in public_hosts()


def permission_values(value: Any) -> set[str]:
    if isinstance(value, str):
        return {item.strip().lower() for item in re.split(r"[,，\s]+", value) if item.strip()}
    if isinstance(value, list):
        return {str(item).strip().lower() for item in value if str(item).strip()}
    return set()


def can_identity_access(identity: dict[str, Any], permissions: dict[str, Any]) -> bool:
    if identity.get("is_admin"):
        return True
    visibility = str(permissions.get("visibility") or "internal").strip().lower()
    if visibility in {"public", "公开"}:
        return True
    if not identity.get("authenticated"):
        return False
    if visibility in {"internal", "内部", "team", "团队"}:
        return True
    user_id = str(identity.get("user_id") or "").lower()
    groups = {str(item).lower() for item in identity.get("groups") or []}
    owner = str(permissions.get("owner") or "").lower()
    allowed_users = permission_values(permissions.get("allowed_users"))
    allowed_groups = permission_values(permissions.get("allowed_groups"))
    return bool(
        (owner and owner == user_id)
        or (user_id and user_id in allowed_users)
        or (groups and groups.intersection(allowed_groups))
    )


def google_drive_sync_status() -> dict[str, Any]:
    rclone_path = shutil.which("rclone") or "/Users/kumaai/.local/bin/rclone"
    payload: dict[str, Any] = {
        "ok": False,
        "status": "unknown",
        "rclone_path": rclone_path if Path(rclone_path).exists() else "",
        "log_path": str(RCLONE_LOG_PATH),
        "last_exit": "",
        "last_start": "",
        "recent_errors": [],
        "message": "",
    }
    if not Path(rclone_path).exists():
        payload["status"] = "missing_rclone"
        payload["message"] = "未找到 rclone"
        return payload
    if not RCLONE_LOG_PATH.exists():
        payload["status"] = "missing_log"
        payload["message"] = "尚未找到 Google Drive 同步日志"
        return payload
    lines = RCLONE_LOG_PATH.read_text(encoding="utf-8", errors="replace").splitlines()
    tail = lines[-160:]
    starts = [line for line in tail if " sync start:" in line]
    exits = [line for line in tail if "sync exit code:" in line]
    errors = [line for line in tail if re.search(r"\b(ERROR|FATAL|Failed|failed)\b", line)]
    payload["last_start"] = starts[-1] if starts else ""
    payload["last_exit"] = exits[-1] if exits else ""
    payload["recent_errors"] = errors[-5:]
    if "sync exit code: 0" in payload["last_exit"]:
        payload["ok"] = True
        payload["status"] = "updated"
        payload["message"] = "最近一次 Google Drive 同步成功"
    elif payload["last_exit"]:
        payload["status"] = "failed"
        payload["message"] = "最近一次 Google Drive 同步可能失败"
    elif payload["last_start"]:
        payload["status"] = "running_or_incomplete"
        payload["message"] = "检测到同步开始，但暂未看到退出码"
    else:
        payload["status"] = "unknown"
        payload["message"] = "日志中未找到同步开始或退出记录"
    return payload


def dify_mapping_path() -> Path:
    return mapping_path_from_config(DEFAULT_CONFIG_PATH)


def public_base_url() -> str:
    values = read_config()
    port = values.get("MEETING_REVIEW_SERVER_PORT") or str(DEFAULT_PORT)
    return values.get("MEETING_REVIEW_BASE_URL") or f"http://localhost:{port}"


def normalize_local_url(url: str) -> str:
    return re.sub(r"^https?://(?:(?:localhost|127\.0\.0\.1):(?:8767|4396|1000)|kuma\.d91\.global|dify\.d91\.global)", "", url or "") or "/history"


def normalize_meeting_type(value: str) -> str:
    cleaned = re.sub(r"\s+", "", (value or "").strip())
    if not cleaned:
        return "多人复盘会"
    return cleaned if cleaned in MEETING_TYPES else "多人复盘会"


def meeting_type_options(selected: str = "") -> str:
    selected = normalize_meeting_type(selected)
    options = []
    for value in MEETING_TYPES:
        flag = " selected" if value == selected else ""
        options.append(f'<option value="{html.escape(value)}"{flag}>{html.escape(value)}</option>')
    return "".join(options)


def normalize_meeting_series(value: str) -> str:
    cleaned = re.sub(r"\s+", "", (value or "").strip())
    if not cleaned:
        return MEETING_SERIES_DEFAULT
    return cleaned[:60]


def meeting_series_options(selected: str = "") -> str:
    selected = normalize_meeting_series(selected)
    options = []
    for value in MEETING_SERIES:
        flag = " selected" if value == selected else ""
        options.append(f'<option value="{html.escape(value)}"{flag}>{html.escape(value)}</option>')
    if selected and selected not in MEETING_SERIES:
        options.append(f'<option value="{html.escape(selected)}" selected>{html.escape(selected)}</option>')
    options.append(f'<option value="{CUSTOM_CHOICE_VALUE}">新增会议系列...</option>')
    return "".join(options)


def infer_meeting_type_from_text(value: str) -> str:
    text = (value or "").lower()
    if any(keyword in text for keyword in ("专家", "访谈", "产业链", "深访", "电话会", "expert")):
        return "专家交流"
    if any(keyword in text for keyword in ("上市公司", "管理层", "业绩", "财报", "路演", "公司交流", "ir")):
        return "上市公司交流"
    if any(keyword in text for keyword in ("复盘", "晨会", "周会", "例会", "策略会", "多人", "市场回顾")):
        return "多人复盘会"
    return "多人复盘会"


def infer_meeting_series_from_text(value: str, meeting_type: str = "") -> str:
    return MEETING_SERIES_DEFAULT


def infer_meeting_metadata_from_uploads(files: list[dict[str, Any]]) -> tuple[str, str, str]:
    filenames = [str(item.get("filename") or "") for item in files if isinstance(item, dict)]
    text_hints = " ".join(filenames)
    for item in files[:3]:
        if not isinstance(item, dict):
            continue
        suffix = Path(str(item.get("filename") or "")).suffix.lower()
        data = item.get("data")
        if suffix in TEXT_SUFFIXES and isinstance(data, (bytes, bytearray)):
            text_hints += " " + bytes(data[:12000]).decode("utf-8", errors="ignore")
    title = ""
    if filenames:
        stem = Path(filenames[0]).stem
        stem = re.sub(r"20\d{2}[-_.年]?\d{1,2}[-_.月]?\d{0,2}日?", " ", stem)
        stem = re.sub(r"(原文|笔记|转录|录音|会议材料|整理前|导入|副本|copy)", " ", stem, flags=re.I)
        stem = re.sub(r"[_\-—–]+", " ", stem)
        title = sanitize_title(stem.strip())
    meeting_type = infer_meeting_type_from_text(text_hints)
    return title or "投资会议纪要", meeting_type, infer_meeting_series_from_text(text_hints, meeting_type)


def sanitize_title(value: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|]+', "-", value).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned or "投资会议纪要"


def sanitize_filename(value: str, fallback: str = "uploaded-file") -> str:
    name = Path(value or fallback).name
    cleaned = re.sub(r'[\\/:*?"<>|]+', "-", name).strip().strip(".")
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned or fallback


def markdown_title(markdown: str, fallback: str = "投资会议纪要") -> str:
    for line in markdown.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            return sanitize_title(stripped[2:].strip())
    return fallback


def markdown_field(markdown: str, field: str, fallback: str = "") -> str:
    pattern = re.compile(rf"^\*\*{re.escape(field)}\*\*[:：][^\S\r\n]*([^\r\n]+?)[^\S\r\n]*$", re.MULTILINE)
    match = pattern.search(markdown)
    return match.group(1).strip() if match else fallback


PROCESS_METADATA_FIELDS = {"输入来源", "整理说明", "上传文件"}


def strip_process_metadata(markdown: str) -> str:
    """Remove workflow-only metadata from human-facing final notes."""
    lines: list[str] = []
    for line in markdown.splitlines():
        if any(re.match(rf"^\s*\*\*{re.escape(field)}\*\*[:：]", line) for field in PROCESS_METADATA_FIELDS):
            continue
        lines.append(line)
    cleaned = "\n".join(lines)
    cleaned = re.sub(r"(?m)^(【[^】]{1,60}】)\s*A[:：]\s*", r"\1", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def ensure_markdown_metadata(markdown: str, *, meeting_title: str, meeting_type: str, meeting_series: str = "") -> str:
    text = markdown.rstrip()
    additions: list[str] = []
    if meeting_title and not markdown_field(text, "会议标题"):
        additions.append(f"**会议标题**：{meeting_title}")
    if meeting_type and not markdown_field(text, "会议类型"):
        additions.append(f"**会议类型**：{meeting_type}")
    if meeting_series and not markdown_field(text, "会议系列"):
        additions.append(f"**会议系列**：{meeting_series}")
    if not additions:
        return text
    lines = text.splitlines()
    if lines and lines[0].startswith("# "):
        return "\n".join([lines[0], *additions, *lines[1:]]).rstrip()
    return "\n".join(additions + [text]).rstrip()


def extract_targets(markdown: str) -> list[str]:
    targets: set[str] = set()
    for match in re.finditer(r"【[^】]*?｜([^】]+?)】", markdown):
        value = match.group(1).strip()
        if value and value != "-":
            targets.add(value)
    return sorted(targets)


def draft_folder(draft_id: str, meeting_date: str | None = None) -> Path:
    root = storage_dir()
    if meeting_date:
        return root / meeting_date / draft_id
    for candidate in root.glob(f"*/{draft_id}"):
        if candidate.is_dir():
            return candidate
    return root / normalize_meeting_date(None) / draft_id


def meta_path(folder: Path) -> Path:
    return folder / "meta.json"


def load_meta(folder: Path) -> dict[str, Any]:
    path = meta_path(folder)
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def save_meta(folder: Path, meta: dict[str, Any]) -> None:
    folder.mkdir(parents=True, exist_ok=True)
    meta_path(folder).write_text(json.dumps(meta, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def save_version(folder: Path, markdown: str, label: str) -> str:
    versions_dir = folder / "versions"
    versions_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    path = versions_dir / f"{stamp}-{label}.md"
    path.write_text(markdown.rstrip() + "\n", encoding="utf-8")
    return str(path)


def import_job_path(draft_id: str, meeting_date: str) -> Path:
    return draft_folder(draft_id, meeting_date) / IMPORT_JOB_FILENAME


def write_import_job_status(draft_id: str, meeting_date: str, payload: dict[str, Any]) -> None:
    path = import_job_path(draft_id, meeting_date)
    path.parent.mkdir(parents=True, exist_ok=True)
    data = {
        "draft_id": draft_id,
        "meeting_date": meeting_date,
        "updated_at": now_iso(),
        **payload,
    }
    tmp_path = path.with_suffix(".json.tmp")
    tmp_path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    tmp_path.replace(path)


def load_import_job_status(draft_id: str, meeting_date: str) -> dict[str, Any]:
    path = import_job_path(draft_id, meeting_date)
    if not path.exists():
        return {
            "ok": False,
            "status": "not_found",
            "draft_id": draft_id,
            "meeting_date": meeting_date,
            "error": "未找到导入任务。",
        }
    return json.loads(path.read_text(encoding="utf-8"))


def confirm_input_job_path(draft_id: str) -> Path:
    return draft_folder(draft_id) / CONFIRM_INPUT_JOB_FILENAME


def write_confirm_input_job_status(draft_id: str, payload: dict[str, Any]) -> None:
    path = confirm_input_job_path(draft_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    data = {
        "draft_id": draft_id,
        "updated_at": now_iso(),
        **payload,
    }
    tmp_path = path.with_suffix(".json.tmp")
    tmp_path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    tmp_path.replace(path)


def load_confirm_input_job_status(draft_id: str) -> dict[str, Any]:
    path = confirm_input_job_path(draft_id)
    if not draft_id or not path.exists():
        return {
            "ok": False,
            "status": "not_found",
            "draft_id": draft_id,
            "error": "未找到会议整理任务。",
        }
    return json.loads(path.read_text(encoding="utf-8"))


def import_upload_root() -> Path:
    path = storage_dir() / "_import_uploads"
    path.mkdir(parents=True, exist_ok=True)
    return path


def import_upload_dir(upload_id: str) -> Path:
    safe_id = re.sub(r"[^0-9a-fA-F-]", "", upload_id or "")
    if not safe_id:
        raise ValueError("缺少上传任务编号")
    return import_upload_root() / safe_id


def import_upload_manifest_path(upload_id: str) -> Path:
    return import_upload_dir(upload_id) / "manifest.json"


def write_import_upload_manifest(upload_id: str, manifest: dict[str, Any]) -> None:
    folder = import_upload_dir(upload_id)
    folder.mkdir(parents=True, exist_ok=True)
    path = import_upload_manifest_path(upload_id)
    tmp_path = path.with_suffix(".json.tmp")
    tmp_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    tmp_path.replace(path)


def load_import_upload_manifest(upload_id: str) -> dict[str, Any]:
    path = import_upload_manifest_path(upload_id)
    if not path.exists():
        raise FileNotFoundError("未找到上传任务")
    return json.loads(path.read_text(encoding="utf-8"))


def create_import_upload_session(payload: dict[str, Any]) -> dict[str, Any]:
    upload_id = str(uuid4())
    draft_id = str(uuid4())
    files = payload.get("files") if isinstance(payload.get("files"), list) else []
    if not files:
        raise ValueError("请先选择文件")
    file_hints = [{"filename": item.get("name") or item.get("filename") or ""} for item in files]
    inferred_title, inferred_type, inferred_series = infer_meeting_metadata_from_uploads(file_hints)
    title = sanitize_title(str(payload.get("meeting_title") or inferred_title or "投资会议纪要"))
    meeting_date = normalize_meeting_date(str(payload.get("meeting_date") or ""))
    raw_meeting_type = str(payload.get("meeting_type") or "")
    meeting_type = normalize_meeting_type(str(payload.get("custom_meeting_type") or raw_meeting_type or inferred_type or "多人复盘会"))
    meeting_series = normalize_meeting_series(
        str(payload.get("custom_meeting_series") or payload.get("meeting_series") or inferred_series or MEETING_SERIES_DEFAULT)
    )
    folder = import_upload_dir(upload_id)
    folder.mkdir(parents=True, exist_ok=True)
    manifest_files: list[dict[str, Any]] = []
    for index, item in enumerate(files):
        filename = sanitize_filename(str(item.get("name") or item.get("filename") or f"uploaded-{index + 1}"))
        suffix = Path(filename).suffix
        part_path = folder / f"{index:03d}-{uuid4().hex}{suffix}"
        manifest_files.append(
            {
                "index": index,
                "filename": filename,
                "content_type": str(item.get("type") or item.get("content_type") or ""),
                "size": int(item.get("size") or 0),
                "received": 0,
                "path": str(part_path),
            }
        )
    status_path = f"/apps/meeting-minutes/import/status?id={urllib.parse.quote(draft_id)}&date={urllib.parse.quote(meeting_date)}"
    manifest = {
        "upload_id": upload_id,
        "draft_id": draft_id,
        "meeting_date": meeting_date,
        "title": title,
        "meeting_type": meeting_type,
        "meeting_series": meeting_series,
        "operator_notes": str(payload.get("operator_notes") or "").strip(),
        "status_url": status_path,
        "files": manifest_files,
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    write_import_upload_manifest(upload_id, manifest)
    write_import_job_status(
        draft_id,
        meeting_date,
        {
            "ok": True,
            "status": "uploading",
            "stage": "上传会议材料",
            "percent": 1,
            "message": "正在分片上传文件，上传完成后会交给 Dify 工作流。",
            "title": title,
            "meeting_type": meeting_type,
            "meeting_series": meeting_series,
            "file_count": len(manifest_files),
            "status_url": status_path,
        },
    )
    return manifest


def append_import_upload_chunk(upload_id: str, file_index: int, offset: int, data: bytes) -> dict[str, Any]:
    manifest = load_import_upload_manifest(upload_id)
    files = manifest.get("files") if isinstance(manifest.get("files"), list) else []
    if file_index < 0 or file_index >= len(files):
        raise ValueError("文件分片编号无效")
    item = files[file_index]
    path = Path(str(item.get("path") or ""))
    path.parent.mkdir(parents=True, exist_ok=True)
    current_size = path.stat().st_size if path.exists() else 0
    if offset < current_size:
        return {"ok": True, "received": current_size, "duplicate": True}
    if offset > current_size:
        raise ValueError(f"分片偏移不连续：expected {current_size}, got {offset}")
    with path.open("ab") as handle:
        handle.write(data)
    received = current_size + len(data)
    item["received"] = received
    manifest["updated_at"] = now_iso()
    write_import_upload_manifest(upload_id, manifest)
    total_size = sum(int(row.get("size") or 0) for row in files) or 1
    total_received = sum(int(row.get("received") or 0) for row in files)
    percent = min(70, max(2, int(total_received * 70 / total_size)))
    write_import_job_status(
        str(manifest.get("draft_id") or ""),
        normalize_meeting_date(str(manifest.get("meeting_date") or "")),
        {
            "ok": True,
            "status": "uploading",
            "stage": "上传会议材料",
            "percent": percent,
            "message": f"正在上传 {item.get('filename') or '文件'}",
            "title": str(manifest.get("title") or "投资会议纪要"),
            "meeting_type": str(manifest.get("meeting_type") or "多人复盘会"),
            "meeting_series": str(manifest.get("meeting_series") or MEETING_SERIES_DEFAULT),
            "file_count": len(files),
            "status_url": str(manifest.get("status_url") or ""),
        },
    )
    return {"ok": True, "received": received, "percent": percent}


def complete_import_upload_session(upload_id: str) -> dict[str, Any]:
    manifest = load_import_upload_manifest(upload_id)
    files = manifest.get("files") if isinstance(manifest.get("files"), list) else []
    missing = []
    for item in files:
        path = Path(str(item.get("path") or ""))
        expected = int(item.get("size") or 0)
        actual = path.stat().st_size if path.exists() else 0
        if expected and actual != expected:
            missing.append(f"{item.get('filename')}: {actual}/{expected}")
    if missing:
        raise ValueError("文件尚未上传完整：" + "；".join(missing))
    draft_id = str(manifest.get("draft_id") or "")
    meeting_date = normalize_meeting_date(str(manifest.get("meeting_date") or ""))
    write_import_job_status(
        draft_id,
        meeting_date,
        {
            "ok": True,
            "status": "queued",
            "stage": "等待 Dify 工作流",
            "percent": 72,
            "message": "文件上传完成，正在提交给 Dify 工作流。",
            "title": str(manifest.get("title") or "投资会议纪要"),
            "meeting_type": str(manifest.get("meeting_type") or "多人复盘会"),
            "meeting_series": str(manifest.get("meeting_series") or MEETING_SERIES_DEFAULT),
            "file_count": len(files),
            "status_url": str(manifest.get("status_url") or ""),
        },
    )
    job = {
        "draft_id": draft_id,
        "meeting_date": meeting_date,
        "title": str(manifest.get("title") or "投资会议纪要"),
        "meeting_type": str(manifest.get("meeting_type") or "多人复盘会"),
        "meeting_series": str(manifest.get("meeting_series") or MEETING_SERIES_DEFAULT),
        "operator_notes": str(manifest.get("operator_notes") or ""),
        "files": [
            {
                "filename": str(item.get("filename") or Path(str(item.get("path"))).name),
                "content_type": str(item.get("content_type") or ""),
                "path": str(item.get("path") or ""),
            }
            for item in files
        ],
    }
    threading.Thread(target=run_import_job, args=(job,), daemon=True).start()
    return {
        "ok": True,
        "upload_id": upload_id,
        "draft_id": draft_id,
        "meeting_date": meeting_date,
        "status_url": str(manifest.get("status_url") or ""),
    }


def latest_result_path() -> Path:
    return storage_dir() / "latest-result.json"


def ensure_markdown_heading(markdown: str, title: str = "投资会议纪要") -> str:
    text = (markdown or "").strip()
    if not text:
        text = "待校对内容为空。"
    if text.startswith("#"):
        return text
    return f"# {sanitize_title(title)}\n\n{text}"


def save_latest_result(payload: dict[str, Any]) -> None:
    path = latest_result_path()
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def load_latest_result() -> dict[str, Any]:
    path = latest_result_path()
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def validate_post_agent_markdown(markdown: str) -> None:
    text = str(markdown or "").strip()
    placeholder_markers = (
        "未生成任何文本或文件输出",
        "待校对内容为空",
        "Incorrect model credentials",
        "InvokeAuthorizationError",
        "PluginInvokeError",
    )
    if not text:
        raise ValueError("Dify 整理阶段没有返回会议纪要正文，已拒绝创建二次校对草稿。")
    for marker in placeholder_markers:
        if marker in text:
            raise ValueError(f"Dify 整理阶段返回了错误占位内容，命中：{marker}。已拒绝创建二次校对草稿，请检查 Dify 模型节点输出。")
    required_sections = ("## 一、发言整理",)
    missing = [section for section in required_sections if section not in text]
    if missing:
        raise ValueError(f"Dify 整理阶段输出缺少必需章节：{'、'.join(missing)}，已拒绝创建二次校对草稿。")


def create_draft(payload: dict[str, Any]) -> dict[str, Any]:
    raw_markdown = str(payload.get("markdown") or payload.get("draft_markdown") or "").strip()
    title_hint = str(payload.get("meeting_title") or payload.get("title") or "投资会议纪要")
    meeting_type = normalize_meeting_type(
        str(payload.get("custom_meeting_type") or payload.get("meeting_type") or markdown_field(raw_markdown, "会议类型") or "")
    )
    meeting_series = normalize_meeting_series(
        str(payload.get("meeting_series") or markdown_field(raw_markdown, "会议系列") or MEETING_SERIES_DEFAULT)
    )
    markdown = ensure_markdown_heading(raw_markdown, title_hint)
    markdown = ensure_markdown_metadata(markdown, meeting_title=title_hint, meeting_type=meeting_type, meeting_series=meeting_series)
    meeting_date = normalize_meeting_date(str(payload.get("meeting_date") or markdown_field(markdown, "会议日期") or ""))
    title = sanitize_title(str(payload.get("meeting_title") or payload.get("title") or markdown_field(markdown, "会议标题") or markdown_title(markdown)))
    draft_id = str(payload.get("draft_id") or uuid4())
    stage = str(payload.get("stage") or "post_agent").strip() or "post_agent"
    if stage != "pre_agent":
        markdown = strip_process_metadata(markdown)
        validate_post_agent_markdown(markdown)
    status = "input_review" if stage == "pre_agent" else "draft"
    folder = draft_folder(draft_id, meeting_date)
    folder.mkdir(parents=True, exist_ok=True)
    (folder / "draft.md").write_text(markdown.rstrip() + "\n", encoding="utf-8")
    (folder / "current.md").write_text(markdown.rstrip() + "\n", encoding="utf-8")
    version_path = save_version(folder, markdown, "generated")
    created_at = now_iso()
    meta = {
        "id": draft_id,
        "meeting_date": meeting_date,
        "title": title,
        "meeting_type": meeting_type,
        "meeting_series": meeting_series,
        "stage": stage,
        "status": status,
        "created_at": created_at,
        "updated_at": created_at,
        "dify_app_id": str(payload.get("dify_app_id") or payload.get("app_id") or ""),
        "source_input_draft_id": str(payload.get("source_input_draft_id") or ""),
        "input_source": str(payload.get("input_source") or markdown_field(markdown, "输入来源", "未注明")),
        "source_run_id": str(payload.get("workflow_run_id") or ""),
        "versions": [{"label": "generated", "path": version_path, "created_at": created_at}],
        "targets": extract_targets(markdown),
    }
    save_meta(folder, meta)
    base = public_base_url()
    result = {
        "ok": True,
        "draft_id": draft_id,
        "review_url": f"{base}/review?id={urllib.parse.quote(draft_id)}",
        "drafts_url": f"{base}/drafts",
        "history_url": f"{base}/history",
        "obsidian_history_url": f"{base}/history",
        "external_sources_url": f"{base}/external-sources",
        "target_query_url": f"{base}/target-query",
        "sync_status_url": f"{base}/sync-status",
        "result_cache_url": f"{base}/result?id={urllib.parse.quote(draft_id)}",
        "latest_result_url": f"{base}/latest",
        "meeting_date": meeting_date,
        "title": title,
        "meeting_type": meeting_type,
        "meeting_series": meeting_series,
        "stage": stage,
    }
    save_latest_result({**result, "updated_at": now_iso()})
    return result


def _decode_text_bytes(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_docx_text(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"缺少 Word 解析依赖: {exc}") from exc
    document = Document(str(path))
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n".join(chunks).strip()


def resolve_asr_model_choice(choice: str) -> dict[str, str]:
    value = str(choice or "").strip().lower()
    if "sensevoice" in value and "auto" not in value:
        return {
            "engine": "sensevoice",
            "model": "SenseVoiceSmall",
            "sensevoice_model": "iic/SenseVoiceSmall",
            "aux_engine": "paraformer",
        }
    return {
        "engine": "auto",
        "model": "自动：SenseVoiceSmall + Paraformer 辅助校对；禁止降级为 Whisper",
        "sensevoice_model": "iic/SenseVoiceSmall",
        "aux_engine": "paraformer",
    }


def _read_transcript_text(output_dir: Path, stem: str) -> str:
    text_path = output_dir / f"{stem}.txt"
    if not text_path.exists():
        raise RuntimeError(f"音频转录完成但未找到本会话输出文本: {text_path.name}")
    return text_path.read_text(encoding="utf-8", errors="replace").strip()


def _read_transcript_json(output_dir: Path, stem: str) -> dict[str, Any]:
    json_path = output_dir / f"{stem}.json"
    if not json_path.exists():
        return {}
    try:
        data = json.loads(json_path.read_text(encoding="utf-8", errors="replace"))
    except json.JSONDecodeError:
        return {}
    return data if isinstance(data, dict) else {}


def _normalize_asr_timestamp_segments(filename: str, segments: Any, limit: int = 160) -> list[dict[str, str]]:
    if not isinstance(segments, list):
        return []
    normalized: list[dict[str, str]] = []
    for item in segments:
        if not isinstance(item, dict):
            continue
        text = str(item.get("text") or item.get("sentence") or "").strip()
        if not text:
            continue
        start = str(item.get("start") or item.get("start_time") or item.get("begin") or "").strip()
        end = str(item.get("end") or item.get("end_time") or item.get("finish") or "").strip()
        speaker = str(item.get("speaker") or item.get("spk") or "").strip()
        normalized.append(
            {
                "filename": filename,
                "start": start,
                "end": end,
                "speaker": speaker,
                "text": text,
                "source": str(item.get("source") or "sensevoice").strip(),
            }
        )
        if len(normalized) >= limit:
            break
    return normalized


def _format_asr_timestamp_index(segments: list[dict[str, str]], limit: int = 120) -> str:
    lines: list[str] = []
    for item in segments[:limit]:
        start = item.get("start", "")
        end = item.get("end", "")
        speaker = item.get("speaker", "")
        text = item.get("text", "")
        if start or end:
            anchor = f"[{start}-{end}]"
        else:
            anchor = "[时间戳待定位]"
        prefix = f"{speaker}: " if speaker else ""
        lines.append(f"{anchor} {prefix}{text}".strip())
    if len(segments) > limit:
        lines.append(f"...（时间戳索引过长，已截断；共 {len(segments)} 段）")
    return "\n".join(lines).strip()


def transcribe_audio_payload(path: Path, output_dir: Path, *, asr_model_choice: str = "") -> dict[str, Any]:
    output_dir.mkdir(parents=True, exist_ok=True)
    model_config = resolve_asr_model_choice(asr_model_choice)
    primary_dir = output_dir / "primary"
    primary_dir.mkdir(parents=True, exist_ok=True)
    primary_python = (
        DEFAULT_SENSEVOICE_PYTHON
        if DEFAULT_SENSEVOICE_PYTHON and Path(DEFAULT_SENSEVOICE_PYTHON).exists()
        else sys.executable
    )
    command = [
        primary_python,
        str(SCRIPT_DIR / "transcribe_audio.py"),
        str(path),
        "--output-dir",
        str(primary_dir),
        "--engine",
        model_config["engine"],
        "--model",
        model_config["sensevoice_model"],
        "--output-format",
        "all",
        "--language",
        "zh",
    ]
    if DEFAULT_SENSEVOICE_MODEL_CACHE:
        command.extend(["--cache-dir", DEFAULT_SENSEVOICE_MODEL_CACHE])
    completed = subprocess.run(
        command,
        text=True,
        capture_output=True,
        timeout=3600,
        env={**os.environ, "PYTHONUTF8": "1", "PYTHONIOENCODING": "utf-8"},
    )
    if completed.returncode != 0:
        raise RuntimeError((completed.stderr or completed.stdout or "音频转录失败").strip())
    primary_text = _read_transcript_text(primary_dir, path.stem)
    primary_json = _read_transcript_json(primary_dir, path.stem)
    speaker_segments = primary_json.get("sentence_info") if isinstance(primary_json.get("sentence_info"), list) else []
    auxiliary_text = str(primary_json.get("auxiliary_text") or "").strip()
    payload: dict[str, Any] = {
        "ok": True,
        "primary_engine": model_config["engine"],
        "primary_model": model_config["model"],
        "text": primary_text,
        "timestamp_detected": bool(primary_json.get("timestamp_detected") or speaker_segments),
        "timestamp_segments": speaker_segments,
        "timestamp_index": primary_json.get("timestamp_index") or [],
        "timestamp_index_path": primary_json.get("timestamp_index_path") or "",
        "speakers": primary_json.get("speakers") or [],
        "speaker_segments": speaker_segments,
        "sentence_info": speaker_segments,
        "auxiliary_engine": primary_json.get("auxiliary_engine") or "",
        "auxiliary_model": primary_json.get("auxiliary_model") or "",
        "auxiliary_text": auxiliary_text,
        "auxiliary_ok": bool(primary_json.get("auxiliary_ok")),
        "auxiliary_status": primary_json.get("auxiliary_status") or "",
        "asr_comparison_diff": primary_json.get("asr_comparison_diff") or "",
        "auxiliary_transcripts": {"paraformer": auxiliary_text} if auxiliary_text else {},
    }
    return payload


def transcribe_audio_file(path: Path, output_dir: Path, *, asr_model_choice: str = "") -> str:
    return str(transcribe_audio_payload(path, output_dir, asr_model_choice=asr_model_choice).get("text") or "").strip()


def extract_uploaded_files_text(
    files: list[dict[str, Any]],
    *,
    draft_id: str,
    meeting_date: str,
    progress_callback: Any | None = None,
) -> tuple[str, list[str], list[str]]:
    if not files:
        raise ValueError("请先上传文本、Word 文档或音频文件")
    folder = draft_folder(draft_id, meeting_date)
    upload_dir = folder / "uploads"
    transcript_dir = folder / "transcripts"
    upload_dir.mkdir(parents=True, exist_ok=True)
    sections: list[str] = []
    filenames: list[str] = []
    warnings: list[str] = []
    total_files = max(1, len(files))

    def report(stage: str, percent: int, message: str) -> None:
        if callable(progress_callback):
            progress_callback(stage, max(1, min(98, percent)), message)

    for index, file_item in enumerate(files, start=1):
        filename = sanitize_filename(str(file_item.get("filename") or f"uploaded-{index}"))
        data = file_item.get("data") if isinstance(file_item.get("data"), bytes) else b""
        before_percent = 12 + int(70 * (index - 1) / total_files)
        after_percent = 12 + int(70 * index / total_files)
        if not data:
            warnings.append(f"{filename}: 文件为空，已跳过")
            report("读取上传文件", before_percent, f"{filename} 文件为空，已跳过")
            continue
        target = upload_dir / filename
        if target.exists():
            target = upload_dir / f"{target.stem}-{index}{target.suffix}"
        target.write_bytes(data)
        suffix = target.suffix.lower()
        try:
            if suffix in TEXT_SUFFIXES:
                report("读取文本材料", before_percent, f"正在读取 {filename}")
                text = _decode_text_bytes(data).strip()
                source_label = "文本抽取"
            elif suffix in DOCX_SUFFIXES:
                report("解析 Word 文档", before_percent, f"正在解析 {filename}")
                text = extract_docx_text(target)
                source_label = "Word 文档抽取"
            elif suffix in AUDIO_SUFFIXES:
                report("音频转录中", before_percent, f"正在转录 {filename}，录音越长耗时越久")
                asr_payload = transcribe_audio_payload(target, transcript_dir / target.stem)
                text = str(asr_payload.get("text") or "").strip()
                timestamp_segments = _normalize_asr_timestamp_segments(
                    filename,
                    asr_payload.get("timestamp_segments") or asr_payload.get("sentence_info") or [],
                )
                timestamp_index = _format_asr_timestamp_index(timestamp_segments)
                if timestamp_index:
                    text = f"{text}\n\n【SenseVoice 时间戳索引】\n{timestamp_index}".strip()
                source_label = "音频转录"
            else:
                warnings.append(f"{filename}: 暂不支持该格式，已保存原文件但未抽取文本")
                report("检查文件格式", before_percent, f"{filename} 格式暂不支持，已保存原文件")
                continue
        except Exception as exc:
            warnings.append(f"{filename}: {exc}")
            report("处理文件失败", before_percent, f"{filename} 处理失败，继续处理其他文件")
            continue
        if not text:
            warnings.append(f"{filename}: 未抽取到有效文本")
            report("读取上传文件", after_percent, f"{filename} 未抽取到有效文本")
            continue
        filenames.append(filename)
        sections.append(f"### {filename}（{source_label}）\n\n{text.strip()}")
        report("合并上传材料", after_percent, f"已完成 {filename}")
    if not sections:
        detail = "；".join(warnings) if warnings else "未抽取到可校对文本"
        raise ValueError(f"上传文件未生成可校对文本：{detail}")
    return "\n\n".join(sections).strip(), filenames, warnings


def ingest_dify_file_list(files: list[dict[str, Any]], *, mode: str = "all", asr_model_choice: str = "") -> dict[str, Any]:
    """Extract and/or transcribe a Dify file-list upload without requiring type-specific inputs."""
    ingest_id = f"dify-{datetime.now().strftime('%Y%m%d-%H%M%S')}-{uuid4().hex[:8]}"
    ingest_dir = DEFAULT_STORAGE_DIR.parent / "dify-ingest" / ingest_id
    upload_dir = ingest_dir / "uploads"
    transcript_dir = ingest_dir / "transcripts"
    upload_dir.mkdir(parents=True, exist_ok=True)
    transcript_dir.mkdir(parents=True, exist_ok=True)

    document_sections: list[str] = []
    audio_sections: list[str] = []
    comparison_sections: list[str] = []
    timestamp_index_sections: list[str] = []
    all_timestamp_segments: list[dict[str, str]] = []
    filenames: list[str] = []
    warnings: list[str] = []
    saved_files: list[str] = []
    model_config = resolve_asr_model_choice(asr_model_choice)
    extract_documents = mode in {"all", "document"}
    transcribe_audio = mode in {"all", "audio"}

    for index, file_item in enumerate(files, start=1):
        filename = sanitize_filename(str(file_item.get("filename") or f"meeting-file-{index}"))
        data = file_item.get("data") if isinstance(file_item.get("data"), bytes) else b""
        if not data:
            warnings.append(f"{filename}: 文件为空，已跳过")
            continue
        target = upload_dir / filename
        if target.exists():
            target = upload_dir / f"{target.stem}-{index}{target.suffix}"
        target.write_bytes(data)
        saved_files.append(str(target))
        suffix = target.suffix.lower()
        try:
            if suffix in TEXT_SUFFIXES and extract_documents:
                text = _decode_text_bytes(data).strip()
                if text:
                    document_sections.append(f"### {filename}（文本抽取）\n\n{text}")
                    filenames.append(filename)
                else:
                    warnings.append(f"{filename}: 未抽取到有效文本")
            elif suffix in DOCX_SUFFIXES and extract_documents:
                text = extract_docx_text(target).strip()
                if text:
                    document_sections.append(f"### {filename}（Word 文档抽取）\n\n{text}")
                    filenames.append(filename)
                else:
                    warnings.append(f"{filename}: 未抽取到有效文本")
            elif (suffix in AUDIO_SUFFIXES or suffix in VIDEO_SUFFIXES) and transcribe_audio:
                asr_payload = transcribe_audio_payload(target, transcript_dir / target.stem, asr_model_choice=asr_model_choice)
                text = str(asr_payload.get("text") or "").strip()
                if text:
                    audio_sections.append(f"### {filename}（音频/视频转录｜{model_config['model']}）\n\n{text}")
                    filenames.append(filename)
                    timestamp_segments = _normalize_asr_timestamp_segments(
                        filename,
                        asr_payload.get("timestamp_segments") or asr_payload.get("sentence_info") or [],
                    )
                    if timestamp_segments:
                        all_timestamp_segments.extend(timestamp_segments)
                        timestamp_index = _format_asr_timestamp_index(timestamp_segments)
                        if timestamp_index:
                            timestamp_index_sections.append(f"### {filename}（SenseVoice 时间戳索引）\n\n{timestamp_index}")
                    elif model_config["engine"] in {"auto", "sensevoice"}:
                        warnings.append(f"{filename}: SenseVoice 未返回可用时间戳索引，音频存疑项需人工定位")
                    if asr_payload.get("auxiliary_status"):
                        comparison_sections.append(f"- {filename}: {asr_payload.get('auxiliary_status')}")
                else:
                    warnings.append(f"{filename}: 转录完成但未生成有效文本")
            elif suffix == ".pdf" and extract_documents:
                warnings.append(f"{filename}: PDF 已保存，但当前本地抽取端暂未解析 PDF 正文")
            elif mode == "document" and (suffix in AUDIO_SUFFIXES or suffix in VIDEO_SUFFIXES):
                continue
            elif mode == "audio" and (suffix in TEXT_SUFFIXES or suffix in DOCX_SUFFIXES or suffix == ".pdf"):
                continue
            else:
                warnings.append(f"{filename}: 暂不支持该格式，已保存原文件但未抽取文本")
        except Exception as exc:
            warnings.append(f"{filename}: {exc}")

    document_text = "\n\n".join(document_sections).strip()
    audio_text = "\n\n".join(audio_sections).strip()
    asr_comparison_summary = "\n\n".join(comparison_sections).strip()
    sensevoice_timestamp_index = "\n\n".join(timestamp_index_sections).strip()
    timestamp_index_text = f"【SenseVoice 时间戳索引】\n{sensevoice_timestamp_index}" if sensevoice_timestamp_index else ""
    combined_parts = [part for part in [document_text, audio_text, timestamp_index_text] if part]
    combined_text = "\n\n".join(combined_parts).strip()
    if mode == "document":
        input_mode = "文稿抽取"
        summary = "已完成上传文稿/文本抽取；音频/视频会交给独立语音识别节点处理。"
    elif mode == "audio":
        input_mode = "独立语音识别"
        summary = f"已通过独立语音识别节点处理音频/视频；ASR 模型：{model_config['model']}。"
    elif document_text and audio_text:
        input_mode = "文稿+音频"
        summary = "已合并文稿抽取与音频/视频转录内容。初校时请优先校准发言人、关键名词、股票代码、数字和时间。"
    elif audio_text:
        input_mode = "音频/视频"
        summary = "已完成音频/视频转录。初校时请重点补充分段、发言人和专有名词。"
    elif document_text:
        input_mode = "文稿"
        summary = "已完成文稿抽取。初校时可直接修正格式、发言人和关键名词。"
    else:
        input_mode = "未抽取到文件文本"
        summary = "未从上传文件中抽取到可用文本；如已粘贴原文，工作流会继续使用粘贴文本。"
    if warnings:
        summary += "\n\n处理提示：\n" + "\n".join(f"- {item}" for item in warnings)
    return {
        "ok": True,
        "ingest_id": ingest_id,
        "input_mode": input_mode,
        "filenames": filenames,
        "saved_files": saved_files,
        "warnings": warnings,
        "asr_model": model_config["model"],
        "document_text": document_text,
        "text_reference": document_text,
        "audio_text": audio_text,
        "sensevoice_transcript": audio_text,
        "sensevoice_timestamp_index": sensevoice_timestamp_index,
        "sensevoice_timestamp_segments": all_timestamp_segments,
        "timestamp_detected": bool(all_timestamp_segments),
        "auxiliary_transcripts": {},
        "asr_comparison_summary": asr_comparison_summary,
        "combined_text": combined_text,
        "cross_check_summary": summary,
    }


def build_import_review_markdown(
    *,
    title: str,
    meeting_date: str,
    meeting_type: str,
    meeting_series: str,
    source_text: str,
    filenames: list[str],
    warnings: list[str],
    operator_notes: str,
) -> str:
    warning_lines = "\n".join(f"- {item}" for item in warnings) if warnings else "- 未检索到明确记录"
    notes_block = f"\n\n## 三、补充说明\n\n{operator_notes}\n" if operator_notes else ""
    return f"""# {title}

**会议日期**：{meeting_date}
**会议标题**：{title}
**会议类型**：{meeting_type}
**会议系列**：{meeting_series}
**整理时间**：{datetime.now().strftime("%Y-%m-%d")}
**输入来源**：上传文件抽取/转录（整理前人工校对）
**上传文件**：{' / '.join(filenames)}
**整理说明**：本稿是根据上传材料自动抽取或转录生成的输入校对稿。请先人工校对文本，对不同发言人进行标注和分段，关键名词如有转录错误可直接修改；确认后才开始生成会议纪要。本稿不是最终会议纪要。

---

## 一、待输入 Agent 的校对文本

{source_text}

## 二、输入阶段存疑

{warning_lines}
{notes_block}"""


def file_item_path(file_item: dict[str, Any]) -> Path | None:
    raw_path = str(file_item.get("path") or "").strip()
    if raw_path:
        path = Path(raw_path).expanduser()
        return path if path.exists() and path.is_file() else None
    return None


def persist_import_file_for_dify(file_item: dict[str, Any], *, draft_id: str, meeting_date: str, index: int) -> Path:
    filename = sanitize_filename(str(file_item.get("filename") or f"uploaded-{index}"))
    folder = draft_folder(draft_id, meeting_date)
    upload_dir = folder / "uploads"
    upload_dir.mkdir(parents=True, exist_ok=True)
    target = upload_dir / filename
    if target.exists():
        target = upload_dir / f"{target.stem}-{index}{target.suffix}"
    source_path = file_item_path(file_item)
    if source_path:
        shutil.copy2(source_path, target)
    else:
        data = file_item.get("data") if isinstance(file_item.get("data"), bytes) else b""
        if not data:
            raise ValueError(f"{filename}: 文件为空")
        target.write_bytes(data)
    return target


def dify_file_kind(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in AUDIO_SUFFIXES:
        return "audio"
    if suffix in DOCX_SUFFIXES or suffix in TEXT_SUFFIXES or suffix == ".pdf":
        return "document"
    if suffix in VIDEO_SUFFIXES:
        return "video"
    return "custom"


def dify_upload_mime_type(path: Path) -> str:
    suffix = path.suffix.lower()
    overrides = {
        ".ogg": "audio/ogg",
        ".mp3": "audio/mpeg",
        ".m4a": "audio/mp4",
        ".wav": "audio/wav",
        ".aac": "audio/aac",
        ".flac": "audio/flac",
        ".wma": "audio/x-ms-wma",
        ".mp4": "video/mp4",
        ".webm": "video/webm",
        ".mpeg": "video/mpeg",
        ".mpga": "audio/mpeg",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".markdown": "text/markdown",
        ".csv": "text/csv",
        ".tsv": "text/tab-separated-values",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".pdf": "application/pdf",
    }
    return overrides.get(suffix) or mimetypes.guess_type(path.name)[0] or "application/octet-stream"


def extract_document_text_for_dify_bundle(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return _decode_text_bytes(path.read_bytes()).strip()
    if suffix in DOCX_SUFFIXES:
        return extract_docx_text(path).strip()
    raise ValueError(f"{path.name}: 暂不能合并该文稿格式")


def prepare_dify_import_paths(
    files: list[Path],
    *,
    title: str,
    draft_id: str,
    meeting_date: str,
) -> tuple[list[Path], list[str]]:
    prepared = [
        path
        for path in files
        if dify_file_kind(path) in {"document", "audio", "video"}
        or path.suffix.lower() in DOCX_SUFFIXES
        or path.suffix.lower() in TEXT_SUFFIXES
    ]
    notes: list[str] = []

    unsupported = [path.name for path in files if path not in prepared]
    if unsupported:
        notes.append("以下文件已保存到网页端上传目录，但 Dify 工作流暂不处理：%s。" % "、".join(unsupported))
    return prepared, notes


def transcribe_import_media_after_dify_asr(
    files: list[Path],
    *,
    draft_id: str,
    meeting_date: str,
    progress_callback: Any | None = None,
) -> str:
    media_files = [path for path in files if dify_file_kind(path) in {"audio", "video"}]
    if not media_files:
        return ""
    transcript_dir = draft_folder(draft_id, meeting_date) / "transcripts"
    sections: list[str] = []
    for index, path in enumerate(media_files, start=1):
        if callable(progress_callback):
            progress_callback(
                "本地 SenseVoice 转录",
                68 + int(20 * (index - 1) / max(1, len(media_files))),
                f"Dify ASR 未产出文本，正在本地 SenseVoice 转录：{path.name}",
            )
        asr_payload = transcribe_audio_payload(path, transcript_dir / path.stem)
        text = str(asr_payload.get("text") or "").strip()
        if text:
            timestamp_segments = _normalize_asr_timestamp_segments(
                path.name,
                asr_payload.get("timestamp_segments") or asr_payload.get("sentence_info") or [],
            )
            timestamp_index = _format_asr_timestamp_index(timestamp_segments)
            timestamp_block = f"\n\n【SenseVoice 时间戳索引】\n{timestamp_index}" if timestamp_index else ""
            sections.append(f"### {path.name}（Dify ASR无文本后本地 SenseVoice 转录）\n\n{text}{timestamp_block}")
    return "\n\n".join(sections).strip()


def upload_file_to_dify(token: str, path: Path, *, user: str) -> dict[str, Any]:
    file_form = f"file=@{path};type={dify_upload_mime_type(path)}"
    response = subprocess.run(
        [
            "curl",
            "-sS",
            "-X",
            "POST",
            f"{resolved_dify_base_url()}/v1/files/upload",
            "-H",
            f"Authorization: Bearer {token}",
            "-F",
            f"user={user}",
            "-F",
            file_form,
        ],
        text=True,
        capture_output=True,
        timeout=7200,
    )
    if response.returncode != 0:
        raise RuntimeError(response.stderr.strip() or f"Dify 文件上传失败：{path.name}")
    response_text = (response.stdout or "").strip()
    try:
        data = json.loads(response_text or "{}")
    except json.JSONDecodeError as exc:
        detail = response_text[:800] or response.stderr.strip() or "empty response"
        raise RuntimeError(f"Dify 文件上传返回不是 JSON：{path.name}: {detail}") from exc
    if not data.get("id"):
        detail = data.get("message") or data.get("error") or response_text[:800] or "empty response"
        raise RuntimeError(f"Dify 文件上传未返回文件 ID：{path.name}: {detail}")
    return data


def call_dify_workflow_for_import(
    *,
    token: str,
    files: list[Path],
    title: str,
    meeting_date: str,
    meeting_type: str,
    meeting_series: str,
    operator_notes: str,
    draft_id: str,
    progress_callback: Any | None = None,
) -> dict[str, Any]:
    user = f"web-import-{draft_id}"
    inputs: dict[str, Any] = {
        "transcript_text": "",
        "meeting_title": title,
        "meeting_type": meeting_type,
        "custom_meeting_type": "",
        "meeting_series": meeting_series,
        "custom_meeting_series": "",
        "meeting_date": meeting_date,
        "review_date": datetime.now().strftime("%Y-%m-%d"),
        "speakers": "",
        "correction_notes": (
            (operator_notes + "\n" if operator_notes else "")
            + f"网页端导入；会议系列：{meeting_series}；web_import_draft_id={draft_id}。"
        ),
    }
    uploaded: list[dict[str, Any]] = []
    dify_files, preparation_notes = prepare_dify_import_paths(
        files,
        title=title,
        draft_id=draft_id,
        meeting_date=meeting_date,
    )
    if preparation_notes:
        inputs["correction_notes"] += "\n" + "\n".join(preparation_notes)
    for index, path in enumerate(dify_files, start=1):
        if callable(progress_callback):
            progress_callback(
                "上传到 Dify 工作流",
                12 + int(54 * (index - 1) / max(1, len(dify_files))),
                f"正在交给 Dify：{path.name}",
            )
        upload = upload_file_to_dify(token, path, user=user)
        uploaded.append({"path": path, "upload": upload})
    if uploaded:
        inputs["meeting_files"] = [
            {
                "type": dify_file_kind(Path(str(item["path"]))),
                "transfer_method": "local_file",
                "upload_file_id": item["upload"]["id"],
            }
            for item in uploaded
        ]
        first_audio_upload = next(
            (
                item
                for item in uploaded
                if dify_file_kind(Path(str(item["path"]))) == "audio"
            ),
            None,
        )
        if first_audio_upload:
            inputs["asr_audio_file"] = {
                "type": "audio",
                "transfer_method": "local_file",
                "upload_file_id": first_audio_upload["upload"]["id"],
            }
    def run_dify(current_inputs: dict[str, Any], *, retry_user_suffix: str = "") -> dict[str, Any]:
        payload = json.dumps(
            {"inputs": current_inputs, "response_mode": "blocking", "user": user + retry_user_suffix},
            ensure_ascii=False,
        )
        response = subprocess.run(
            [
                "curl",
                "-sS",
                "-X",
                "POST",
                f"{resolved_dify_base_url()}/v1/workflows/run",
                "-H",
                f"Authorization: Bearer {token}",
                "-H",
                "Content-Type: application/json",
                "--data-binary",
                "@-",
            ],
            input=payload,
            text=True,
            capture_output=True,
            timeout=7200,
        )
        if response.returncode != 0:
            raise RuntimeError(response.stderr.strip() or "Dify workflow call failed")
        data = json.loads(response.stdout or "{}")
        if data.get("code") or data.get("status") in {400, "400"}:
            raise RuntimeError(str(data.get("message") or data.get("error") or "Dify workflow failed"))
        status = (data.get("data") or {}).get("status")
        if status not in {"succeeded", "running", None}:
            raise RuntimeError(str((data.get("data") or {}).get("error") or "Dify workflow failed"))
        return data

    def retry_with_local_fallback() -> dict[str, Any]:
        fallback_text = transcribe_import_media_after_dify_asr(
            dify_files,
            draft_id=draft_id,
            meeting_date=meeting_date,
            progress_callback=progress_callback,
        )
        if not fallback_text:
            raise RuntimeError("Dify ASR 未产出文本，本地 SenseVoice 也未生成有效转录文本")
        retry_inputs = dict(inputs)
        retry_inputs["transcript_text"] = "\n\n".join(
            item for item in [str(inputs.get("transcript_text") or "").strip(), fallback_text] if item
        )
        retry_inputs["correction_notes"] = (
            str(retry_inputs.get("correction_notes") or "")
            + "\nDify ASR 已优先尝试但未产出可合并文本；已在 Dify 外部完成本地 SenseVoice 转录后重跑工作流。"
        )
        retry_inputs.pop("asr_audio_file", None)
        document_uploads = [
            item
            for item in uploaded
            if dify_file_kind(Path(str(item["path"]))) == "document"
        ]
        if document_uploads:
            retry_inputs["meeting_files"] = [
                {
                    "type": "document",
                    "transfer_method": "local_file",
                    "upload_file_id": item["upload"]["id"],
                }
                for item in document_uploads
            ]
        else:
            retry_inputs.pop("meeting_files", None)
        if callable(progress_callback):
            progress_callback("重跑 Dify 工作流", 90, "已带入本地 SenseVoice 转录文本，不再让 Dify HTTP 节点长时间等待")
        return run_dify(retry_inputs, retry_user_suffix="-local-fallback")

    if callable(progress_callback):
        progress_callback("运行 Dify 工作流", 66, "Dify 正在优先调用已配置的 ASR 节点")
    try:
        data = run_dify(inputs)
    except RuntimeError as exc:
        if "transcribe-files" not in str(exc):
            raise
        data = retry_with_local_fallback()
    initial_outputs = (data.get("data") or {}).get("outputs") or {}
    initial_markdown = str(initial_outputs.get("result_markdown") or "")
    if any(dify_file_kind(path) in {"audio", "video"} for path in dify_files) and "未检测到有效转录文本" in initial_markdown:
        data = retry_with_local_fallback()
    outputs = (data.get("data") or {}).get("outputs") or {}
    raw_review_url = str(outputs.get("review_url") or "")
    if not raw_review_url:
        detail = (data.get("data") or {}).get("error") or data.get("message") or "Dify workflow 未返回 review_url"
        raise RuntimeError(str(detail))
    review_url = normalize_local_url(raw_review_url)
    if not review_url:
        raise RuntimeError("Dify workflow 未返回 review_url")
    return {
        "ok": True,
        "workflow_run_id": str(data.get("workflow_run_id") or (data.get("data") or {}).get("id") or ""),
        "draft_id": str(outputs.get("draft_id") or draft_id),
        "review_url": review_url,
        "result_cache_url": normalize_local_url(str(outputs.get("result_cache_url") or "")),
        "history_url": normalize_local_url(str(outputs.get("history_url") or "/history")),
        "drafts_url": normalize_local_url(str(outputs.get("drafts_url") or "/drafts")),
        "latest_result_url": normalize_local_url(str(outputs.get("latest_result_url") or "/latest")),
        "review_message": str(outputs.get("review_message") or "Dify 工作流已生成校对页。"),
        "uploaded_files": [item["upload"] for item in uploaded],
    }


def run_import_job(job: dict[str, Any]) -> None:
    draft_id = str(job.get("draft_id") or "")
    meeting_date = normalize_meeting_date(str(job.get("meeting_date") or ""))
    title = sanitize_title(str(job.get("title") or "投资会议纪要"))
    meeting_type = normalize_meeting_type(str(job.get("meeting_type") or "多人复盘会"))
    meeting_series = normalize_meeting_series(str(job.get("meeting_series") or MEETING_SERIES_DEFAULT))
    files = job.get("files") if isinstance(job.get("files"), list) else []
    operator_notes = str(job.get("operator_notes") or "").strip()

    def progress(stage: str, percent: int, message: str) -> None:
        write_import_job_status(
            draft_id,
            meeting_date,
            {
                "ok": True,
                "status": "running",
                "stage": stage,
                "percent": percent,
                "message": message,
                "title": title,
                "meeting_type": meeting_type,
                "meeting_series": meeting_series,
                "file_count": len(files),
            },
        )

    try:
        progress("保存上传材料", 8, "已收到文件，准备交给 Dify 工作流")
        local_files: list[Path] = []
        for index, file_item in enumerate(files, start=1):
            local_files.append(persist_import_file_for_dify(file_item, draft_id=draft_id, meeting_date=meeting_date, index=index))
        token_id, token = create_temp_dify_token()
        try:
            result = call_dify_workflow_for_import(
                token=token,
                files=local_files,
                title=title,
                meeting_date=meeting_date,
                meeting_type=meeting_type,
                meeting_series=meeting_series,
                operator_notes=operator_notes,
                draft_id=draft_id,
                progress_callback=progress,
            )
        finally:
            delete_temp_dify_token(token_id)
        write_import_job_status(
            draft_id,
            meeting_date,
            {
                "ok": True,
                "status": "done",
                "stage": "准备打开人工校对页",
                "percent": 100,
                "message": "Dify 工作流已生成校对页，即将打开人工初校页面",
                "title": title,
                "meeting_type": meeting_type,
                "meeting_series": meeting_series,
                "file_count": len(files),
                "filenames": [path.name for path in local_files],
                "warnings": [],
                "result": result,
            },
        )
    except Exception as exc:
        write_import_job_status(
            draft_id,
            meeting_date,
            {
                "ok": False,
                "status": "failed",
                "stage": "导入失败",
                "percent": 100,
                "message": "导入失败，请检查文件格式或转录服务状态。",
                "error": str(exc),
                "traceback": traceback.format_exc(),
                "title": title,
                "meeting_type": meeting_type,
                "meeting_series": meeting_series,
                "file_count": len(files),
            },
        )


def current_markdown(folder: Path) -> str:
    path = folder / "current.md"
    if path.exists():
        return path.read_text(encoding="utf-8")
    return (folder / "draft.md").read_text(encoding="utf-8")


def save_current(draft_id: str, markdown: str) -> dict[str, Any]:
    folder = draft_folder(draft_id)
    if not folder.exists():
        raise FileNotFoundError("draft not found")
    meta = load_meta(folder)
    if meta.get("status") == "confirmed":
        raise ValueError("confirmed minutes are immutable; create a new revision instead")
    (folder / "current.md").write_text(markdown.rstrip() + "\n", encoding="utf-8")
    version_path = save_version(folder, markdown, "saved")
    meta.setdefault("versions", []).append({"label": "saved", "path": version_path, "created_at": now_iso()})
    meta["updated_at"] = now_iso()
    meta["targets"] = extract_targets(markdown)
    save_meta(folder, meta)
    return {"ok": True, "draft_id": draft_id, "message": "草稿已保存"}


def psql_scalar(sql: str, *, timeout: int = 30) -> str:
    result = subprocess.run(
        ["docker", "exec", DIFY_DB_CONTAINER, "psql", "-U", "postgres", "-d", "dify", "-t", "-A", "-c", sql],
        text=True,
        capture_output=True,
        timeout=timeout,
    )
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or result.stdout.strip() or "psql failed")
    return result.stdout.strip()


def resolved_dify_base_url() -> str:
    values = read_config()
    return str(os.environ.get("DIFY_BASE_URL") or values.get("DIFY_BASE_URL") or DIFY_BASE_URL).rstrip("/")


def resolve_dify_app_id(meta: dict[str, Any] | None = None) -> str:
    meta = meta or {}
    return str(
        meta.get("dify_app_id")
        or meta.get("app_id")
        or os.environ.get("DIFY_WORKFLOW_APP_ID")
        or DIFY_WORKFLOW_APP_ID
    ).strip()


def create_temp_dify_token(app_id: str | None = None) -> tuple[str, str]:
    resolved_app_id = str(app_id or DIFY_WORKFLOW_APP_ID).strip()
    tenant_id = psql_scalar(f"select tenant_id from apps where id='{resolved_app_id}' limit 1;")
    if not tenant_id:
        raise RuntimeError(f"Dify app not found: {resolved_app_id}")
    token = "app-" + secrets.token_urlsafe(48)
    stdout = psql_scalar(
        "insert into api_tokens (app_id, type, token, tenant_id) "
        f"values ('{resolved_app_id}', 'app', '{token}', '{tenant_id}') returning id;"
    )
    match = re.search(r"[0-9a-fA-F-]{36}", stdout)
    if not match:
        raise RuntimeError("temporary Dify token was not created")
    return match.group(0), token


def delete_temp_dify_token(token_id: str) -> None:
    if token_id:
        psql_scalar(f"delete from api_tokens where id='{token_id}';")


def extract_pre_agent_text(markdown: str) -> str:
    body = str(markdown or "").strip()
    start = body.find("## 一、待输入 Agent 的校对文本")
    if start >= 0:
        body = body[start + len("## 一、待输入 Agent 的校对文本") :]
    end = body.find("## 二、输入阶段存疑")
    if end >= 0:
        body = body[:end]
    return body.strip() or str(markdown or "").strip()


def structured_table_path(folder: Path) -> Path:
    return folder / "structured_table.md"


def target_resolution_path(folder: Path) -> Path:
    return folder / "target_resolution.json"


def target_resolution_review_path(folder: Path) -> Path:
    return folder / "target_resolution_review.md"


def load_confirmed_source_text(meta: dict[str, Any]) -> str:
    source_id = str(meta.get("source_input_draft_id") or "").strip()
    if not source_id:
        return ""
    source_folder = draft_folder(source_id)
    for filename in ("input-confirmed.md", "current.md", "draft.md"):
        path = source_folder / filename
        if path.exists():
            return path.read_text(encoding="utf-8")
    return ""


def collect_raw_upload_files(folder: Path, meta: dict[str, Any]) -> list[Path]:
    candidates: list[Path] = []
    source_id = str(meta.get("source_input_draft_id") or "").strip()
    if source_id:
        candidates.append(draft_folder(source_id) / "uploads")
    candidates.append(folder / "uploads")
    files: list[Path] = []
    seen: set[str] = set()
    for upload_dir in candidates:
        if not upload_dir.exists():
            continue
        for path in sorted(upload_dir.iterdir()):
            if not path.is_file():
                continue
            resolved = str(path.resolve())
            if resolved in seen:
                continue
            seen.add(resolved)
            files.append(path)
    return files


def structured_table_archive_path(
    *,
    meeting_date: str,
    meeting_type: str,
    meeting_title: str,
    source_markdown_path: Path,
) -> Path:
    date_part = normalize_meeting_date(meeting_date)
    type_part = sanitize_title(meeting_type or "多人复盘会")
    session_part = sanitize_title(meeting_title or source_markdown_path.stem)
    target_dir = structured_table_archive_dir() / date_part / type_part / session_part
    target_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{date_part} - {session_part} - 结构化表格.md"
    target = target_dir / sanitize_filename(filename, "结构化表格.md")
    if not target.exists():
        return target
    timestamp = datetime.now().strftime("%H%M%S")
    return target.with_name(f"{target.stem}-{timestamp}{target.suffix}")


def generate_structured_table_markdown(final_markdown: str, source_markdown: str, folder: Path) -> tuple[str, dict[str, Any]]:
    skill_output = folder / "structured_table.md"
    json_output = folder / "structured_table.json"
    source_path = folder / "structured_table_source.md"
    confirmed_source_path = folder / "structured_table_confirmed_source.md"
    target_resolution = resolve_targets_for_text(
        "\n\n".join(part for part in [final_markdown, source_markdown] if part),
        max_mentions=120,
    )
    write_target_resolution_artifacts(folder, target_resolution, stage="structured_table")
    enriched_source_markdown = "\n\n".join(
        part for part in [source_markdown.rstrip(), target_resolution_confirmed_markdown(target_resolution)] if part
    )
    source_path.write_text(final_markdown.rstrip() + "\n", encoding="utf-8")
    confirmed_source_path.write_text(enriched_source_markdown.rstrip() + "\n", encoding="utf-8")
    if not STRUCTURED_TABLE_SCRIPT.exists():
        raise FileNotFoundError(f"structured table skill script not found: {STRUCTURED_TABLE_SCRIPT}")
    result = subprocess.run(
        [
            sys.executable,
            str(STRUCTURED_TABLE_SCRIPT),
            "--meeting-markdown",
            str(source_path),
            "--source-markdown",
            str(confirmed_source_path),
            "--output",
            str(skill_output),
            "--json-output",
            str(json_output),
        ],
        text=True,
        capture_output=True,
        timeout=60,
    )
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or result.stdout.strip() or "structured table skill failed")
    table_markdown = skill_output.read_text(encoding="utf-8")
    try:
        rows = json.loads(json_output.read_text(encoding="utf-8")) if json_output.exists() else []
    except Exception:
        rows = []
    meaningful_rows = [
        row for row in rows
        if isinstance(row, dict)
        and str(row.get("target_name") or "").strip() not in {"", "待确认"}
        and not str(row.get("core_viewpoint") or "").strip().startswith("**会议日期**")
    ]
    warning = ""
    if not meaningful_rows:
        warning = (
            "未生成可确认标的行：结构化表格当前只保留股票代码已确认的标的；"
            "本稿未识别到带确认代码的标的，可在表格校对页手工补充后再归档。"
        )
        table_markdown = table_markdown.rstrip() + f"\n\n> {warning}\n"
        skill_output.write_text(table_markdown, encoding="utf-8")
    return table_markdown, {
        "skill_name": "meeting-minutes-structured-table",
        "skill_dir": str(STRUCTURED_TABLE_SKILL_DIR),
        "script": str(STRUCTURED_TABLE_SCRIPT),
        "json_path": str(json_output),
        "row_count": len(rows) if isinstance(rows, list) else 0,
        "meaningful_row_count": len(meaningful_rows),
        "warning": warning,
        "target_resolution_path": str(target_resolution_path(folder)),
        "target_resolution_review_path": str(target_resolution_review_path(folder)),
        "target_resolution_counts": target_resolution.get("counts") or {},
    }


def confirm_structured_table(draft_id: str, table_markdown: str | None = None) -> dict[str, Any]:
    folder = draft_folder(draft_id)
    if not folder.exists():
        raise FileNotFoundError("draft not found")
    meta = load_meta(folder)
    current_table = table_markdown if table_markdown is not None else structured_table_path(folder).read_text(encoding="utf-8")
    current_table = str(current_table or "").strip()
    if not current_table:
        raise ValueError("structured table is empty")
    table_path = structured_table_path(folder)
    table_path.write_text(current_table.rstrip() + "\n", encoding="utf-8")
    confirmed_at = now_iso()
    meta["status"] = "archive_review"
    meta["stage"] = "archive_review"
    meta["structured_table"] = {
        **(meta.get("structured_table") if isinstance(meta.get("structured_table"), dict) else {}),
        "ok": True,
        "path": str(table_path),
        "confirmed_at": confirmed_at,
        "editable_format": "markdown",
    }
    meta["updated_at"] = confirmed_at
    save_meta(folder, meta)
    return {
        "ok": True,
        "draft_id": draft_id,
        "status": meta["status"],
        "archive_confirm_url": f"{public_base_url()}/result?id={urllib.parse.quote(draft_id)}",
        "message": "结构化表格已确认，请在结果页执行最终归档确认。",
    }


def call_dify_workflow_for_minutes(*, draft_id: str, source_markdown: str, meta: dict[str, Any]) -> dict[str, Any]:
    token_id = ""
    try:
        target_resolution = resolve_targets_for_text(source_markdown, max_mentions=120)
        source_folder = draft_folder(draft_id)
        write_target_resolution_artifacts(source_folder, target_resolution, stage="pre_agent_confirmed")
        target_resolution_context = target_resolution_prompt_context(target_resolution)
        token_id, token = create_temp_dify_token(resolve_dify_app_id(meta))
        inputs = {
            "transcript_text": extract_pre_agent_text(source_markdown),
            "meeting_title": str(meta.get("title") or markdown_title(source_markdown)),
            "meeting_type": str(meta.get("meeting_type") or markdown_field(source_markdown, "会议类型", "多人复盘会")),
            "custom_meeting_type": "",
            "meeting_series": str(meta.get("meeting_series") or markdown_field(source_markdown, "会议系列", MEETING_SERIES_DEFAULT)),
            "custom_meeting_series": "",
            "meeting_date": str(meta.get("meeting_date") or markdown_field(source_markdown, "会议日期", "")),
            "review_date": datetime.now().strftime("%Y-%m-%d"),
            "speakers": "",
            "correction_notes": (
                f"__INPUT_REVIEW_CONFIRMED__ source_input_draft_id={draft_id}; "
                f"已完成输入 Agent 前人工校对；来源输入校对稿 draft_id={draft_id}。"
                f"会议系列：{str(meta.get('meeting_series') or markdown_field(source_markdown, '会议系列', MEETING_SERIES_DEFAULT))}。"
                "请基于该人工确认文本整理会议纪要，并在元信息中保留会议系列。"
                + (f"\n\n{target_resolution_context}" if target_resolution_context else "")
            ),
        }
        payload = json.dumps({"inputs": inputs, "response_mode": "blocking", "user": f"input-review-{draft_id}"}, ensure_ascii=False)
        response = subprocess.run(
            [
                "curl",
                "-sS",
                "-X",
                "POST",
                f"{resolved_dify_base_url()}/v1/workflows/run",
                "-H",
                f"Authorization: Bearer {token}",
                "-H",
                "Content-Type: application/json",
                "--data-binary",
                "@-",
            ],
            input=payload,
            text=True,
            capture_output=True,
            timeout=480,
        )
        if response.returncode != 0:
            raise RuntimeError(response.stderr.strip() or "Dify workflow call failed")
        data = json.loads(response.stdout or "{}")
        run_data = data.get("data") or {}
        outputs = run_data.get("outputs") or {}
        if run_data.get("status") not in {"succeeded", "running", None}:
            raise RuntimeError(str(run_data.get("error") or f"Dify workflow failed: {run_data.get('status')}"))
        review_url = normalize_local_url(str(outputs.get("review_url") or ""))
        post_draft_id = str(outputs.get("draft_id") or "")
        if not review_url or not post_draft_id:
            detail_parts = [
                f"status={run_data.get('status') or ''}",
                f"error={run_data.get('error') or ''}",
                f"review_message={outputs.get('review_message') or ''}",
                f"validation_message={outputs.get('validation_message') or outputs.get('review_message') or ''}",
            ]
            raise RuntimeError("Dify workflow did not return post-agent review_url/draft_id; " + "; ".join(detail_parts))
        post_folder = draft_folder(post_draft_id)
        if post_folder.exists():
            write_target_resolution_artifacts(post_folder, target_resolution, stage="post_agent_context")
            post_meta = load_meta(post_folder)
            post_meta["target_resolution"] = {
                "path": str(target_resolution_path(post_folder)),
                "review_path": str(target_resolution_review_path(post_folder)),
                "counts": target_resolution.get("counts") or {},
                "source_input_draft_id": draft_id,
            }
            save_meta(post_folder, post_meta)
        return {
            "ok": True,
            "workflow_run_id": str(data.get("workflow_run_id") or run_data.get("id") or ""),
            "draft_id": post_draft_id,
            "review_url": review_url,
            "result_cache_url": normalize_local_url(str(outputs.get("result_cache_url") or "")),
            "history_url": normalize_local_url(str(outputs.get("history_url") or "/history")),
            "message": "会议纪要整理稿已生成，进入第二次人工校对。",
        }
    finally:
        try:
            delete_temp_dify_token(token_id)
        except Exception:
            pass


def call_dify_workflow_for_confirmation(*, action: str, draft_id: str, markdown: str = "", meta: dict[str, Any] | None = None) -> dict[str, Any]:
    meta = meta or {}
    token_id = ""
    try:
        token_id, token = create_temp_dify_token(resolve_dify_app_id(meta))
        inputs = {
            "workflow_action": action,
            "action_draft_id": draft_id,
            "final_markdown": markdown,
            "transcript_text": "",
            "meeting_title": str(meta.get("title") or "投资会议纪要"),
            "meeting_type": str(meta.get("meeting_type") or "多人复盘会"),
            "custom_meeting_type": "",
            "meeting_series": str(meta.get("meeting_series") or MEETING_SERIES_DEFAULT),
            "custom_meeting_series": "",
            "meeting_date": str(meta.get("meeting_date") or ""),
            "review_date": datetime.now().strftime("%Y-%m-%d"),
            "speakers": "",
            "correction_notes": f"确认动作由 Dify workflow 编排执行；action={action}; draft_id={draft_id}",
            "input_reviewed": "true",
            "source_input_draft_id": draft_id,
        }
        payload = json.dumps({"inputs": inputs, "response_mode": "blocking", "user": f"{action}-{draft_id}"}, ensure_ascii=False)
        response = subprocess.run(
            [
                "curl",
                "-sS",
                "-X",
                "POST",
                f"{resolved_dify_base_url()}/v1/workflows/run",
                "-H",
                f"Authorization: Bearer {token}",
                "-H",
                "Content-Type: application/json",
                "--data-binary",
                "@-",
            ],
            input=payload,
            text=True,
            capture_output=True,
            timeout=7200,
        )
        if response.returncode != 0:
            raise RuntimeError(response.stderr.strip() or "Dify workflow confirmation call failed")
        data = json.loads(response.stdout or "{}")
        run_data = data.get("data") or {}
        outputs = run_data.get("outputs") or {}
        if run_data.get("status") != "succeeded":
            raise RuntimeError(str(run_data.get("error") or f"Dify workflow confirmation failed: {run_data.get('status')}"))
        status = str(outputs.get("status") or ("table_review" if action == "final_review_confirmed" else "confirmed"))
        next_result_url = (
            outputs.get("table_confirm_url")
            or outputs.get("archive_confirm_url")
            or outputs.get("result_cache_url")
            or outputs.get("review_url")
            or f"/result?id={urllib.parse.quote(draft_id)}"
        )
        return {
            "ok": True,
            "draft_id": str(outputs.get("draft_id") or draft_id),
            "status": status,
            "workflow_run_id": str(data.get("workflow_run_id") or run_data.get("id") or ""),
            "review_url": normalize_local_url(str(outputs.get("review_url") or "")),
            "archive_confirm_url": normalize_local_url(str(next_result_url)),
            "table_confirm_url": normalize_local_url(str(outputs.get("table_confirm_url") or "")),
            "history_url": normalize_local_url(str(outputs.get("history_url") or "/history")),
            "drafts_url": normalize_local_url(str(outputs.get("drafts_url") or "/drafts")),
            "message": str(outputs.get("review_message") or "Dify 确认动作已完成。"),
        }
    finally:
        try:
            delete_temp_dify_token(token_id)
        except Exception:
            pass


def execute_confirm_draft(draft_id: str, markdown: str | None = None, *, workflow_run_id: str = "") -> dict[str, Any]:
    folder = draft_folder(draft_id)
    if not folder.exists():
        raise FileNotFoundError("draft not found")
    meta = load_meta(folder)
    final_markdown = strip_process_metadata(markdown if markdown is not None else current_markdown(folder))
    if not final_markdown.startswith("#"):
        raise ValueError("final markdown must start with a heading")
    meeting_date = normalize_meeting_date(str(meta.get("meeting_date") or markdown_field(final_markdown, "会议日期") or ""))
    meeting_type = normalize_meeting_type(str(meta.get("meeting_type") or markdown_field(final_markdown, "会议类型") or "多人复盘会"))
    meeting_series = normalize_meeting_series(str(meta.get("meeting_series") or markdown_field(final_markdown, "会议系列") or ""))
    final_markdown = ensure_markdown_metadata(
        final_markdown,
        meeting_title=str(meta.get("title") or markdown_title(final_markdown)),
        meeting_type=meeting_type,
        meeting_series=meeting_series,
    )
    (folder / "current.md").write_text(final_markdown + "\n", encoding="utf-8")
    final_path = folder / "final.md"
    final_path.write_text(final_markdown + "\n", encoding="utf-8")
    version_path = save_version(folder, final_markdown, "final-confirmed")

    artifacts_payload: dict[str, Any]
    try:
        artifacts_payload = generate_artifacts(final_path, output_dir=folder / "post-review-artifacts", require_final_root=False)
    except Exception as exc:
        artifacts_payload = {"ok": False, "error": str(exc)}

    source_markdown = load_confirmed_source_text(meta)
    table_markdown, table_skill_payload = generate_structured_table_markdown(final_markdown, source_markdown, folder)
    table_path = structured_table_path(folder)

    confirmed_at = now_iso()
    meta["status"] = "table_review"
    meta["stage"] = "table_review"
    if workflow_run_id:
        meta["final_review_workflow_run_id"] = workflow_run_id
    meta["meeting_type"] = meeting_type
    meta["meeting_series"] = meeting_series
    meta["final_confirmed_at"] = confirmed_at
    meta["updated_at"] = confirmed_at
    meta.setdefault("versions", []).append({"label": "final-confirmed", "path": version_path, "created_at": confirmed_at})
    meta["post_review_artifacts"] = artifacts_payload
    meta["structured_table"] = {
        "ok": True,
        "path": str(table_path),
        "source_input_draft_id": str(meta.get("source_input_draft_id") or ""),
        "generated_at": confirmed_at,
        "editable_format": "markdown",
        **table_skill_payload,
    }
    meta["target_resolution"] = {
        "path": table_skill_payload.get("target_resolution_path"),
        "review_path": table_skill_payload.get("target_resolution_review_path"),
        "counts": table_skill_payload.get("target_resolution_counts") or {},
        "stage": "structured_table",
    }
    meta["targets"] = extract_targets(final_markdown)
    save_meta(folder, meta)
    return {
        "ok": bool(artifacts_payload.get("ok")),
        "draft_id": draft_id,
        "status": meta["status"],
        "post_review_artifacts": artifacts_payload,
        "structured_table": meta["structured_table"],
        "table_confirm_url": f"{public_base_url()}/result?id={urllib.parse.quote(draft_id)}",
        "message": "终稿已确认，请在结果页校对并确认结构化表格。",
    }


def confirm_draft(draft_id: str, markdown: str | None = None) -> dict[str, Any]:
    folder = draft_folder(draft_id)
    if not folder.exists():
        raise FileNotFoundError("draft not found")
    meta = load_meta(folder)
    final_markdown = strip_process_metadata(markdown if markdown is not None else current_markdown(folder))
    if not final_markdown.startswith("#"):
        raise ValueError("final markdown must start with a heading")
    if resolve_dify_app_id(meta) == DIFY_WORKFLOW_APP_ID:
        return execute_confirm_draft(draft_id, final_markdown)
    return call_dify_workflow_for_confirmation(
        action="final_review_confirmed",
        draft_id=draft_id,
        markdown=final_markdown,
        meta=meta,
    )


def execute_confirm_archive(draft_id: str, *, workflow_run_id: str = "") -> dict[str, Any]:
    folder = draft_folder(draft_id)
    if not folder.exists():
        raise FileNotFoundError("draft not found")
    meta = load_meta(folder)
    if meta.get("status") == "table_review":
        raise ValueError("structured table must be confirmed before final archive")
    final_path = folder / "final.md"
    if not final_path.exists():
        raise FileNotFoundError("final markdown not found; confirm final draft first")
    final_markdown = strip_process_metadata(final_path.read_text(encoding="utf-8"))
    if not final_markdown.startswith("#"):
        raise ValueError("final markdown must start with a heading")
    meeting_date = normalize_meeting_date(str(meta.get("meeting_date") or markdown_field(final_markdown, "会议日期") or ""))
    meeting_type = normalize_meeting_type(str(meta.get("meeting_type") or markdown_field(final_markdown, "会议类型") or "多人复盘会"))
    meeting_series = normalize_meeting_series(str(meta.get("meeting_series") or markdown_field(final_markdown, "会议系列") or ""))
    final_markdown = ensure_markdown_metadata(
        final_markdown,
        meeting_title=str(meta.get("title") or markdown_title(final_markdown)),
        meeting_type=meeting_type,
        meeting_series=meeting_series,
    )
    final_path.write_text(final_markdown + "\n", encoding="utf-8")
    raw_files = collect_raw_upload_files(folder, meta)
    if raw_files:
        try:
            archived_raw_paths = archive_raw_files(
                raw_files,
                RAW_ARCHIVE_ROOT,
                meeting_date,
                str(meta.get("title") or markdown_title(final_markdown) or ""),
            )
            raw_archive_payload = {
                "ok": True,
                "archived_count": len(archived_raw_paths),
                "archived_paths": [str(path) for path in archived_raw_paths],
            }
        except Exception as exc:
            raw_archive_payload = {"ok": False, "error": str(exc), "source_paths": [str(path) for path in raw_files]}
    else:
        raw_archive_payload = {
            "ok": True,
            "archived_count": 0,
            "archived_paths": [],
            "message": "no uploaded raw files found for this draft",
        }
    result = export_note(final_path, output_dir(), meeting_date)
    structured_table_payload = meta.get("structured_table") if isinstance(meta.get("structured_table"), dict) else {}
    table_source = Path(str(structured_table_payload.get("path") or structured_table_path(folder)))
    if table_source.exists():
        table_archive_path = structured_table_archive_path(
            meeting_date=meeting_date,
            meeting_type=meeting_type,
            meeting_title=str(meta.get("title") or markdown_title(final_markdown) or ""),
            source_markdown_path=Path(result.md_path),
        )
        shutil.copy2(table_source, table_archive_path)
        structured_table_payload = {
            **structured_table_payload,
            "ok": True,
            "path": str(table_source),
            "archive_path": str(table_archive_path),
            "archive_root": str(structured_table_archive_dir()),
            "archive_layout": "date/meeting_type/meeting_session",
            "archived_at": now_iso(),
        }
    else:
        structured_table_payload = {
            **structured_table_payload,
            "ok": False,
            "error": "structured table artifact not found",
        }
    artifacts_payload: dict[str, Any]
    try:
        artifacts_payload = generate_artifacts(Path(result.md_path))
    except Exception as exc:
        artifacts_payload = {"ok": False, "error": str(exc)}
    knowledge_payload: dict[str, Any]
    try:
        dify_config = load_config(DEFAULT_CONFIG_PATH)
        metadata = build_metadata(
            final_markdown,
            source_file=final_path,
            meeting_date=meeting_date,
            meeting_type=meeting_type,
            meeting_series=meeting_series,
            title=str(meta.get("title") or markdown_title(final_markdown)),
            markdown_path=str(result.md_path),
            word_path=str(result.docx_path),
            note_key=draft_id,
        )
        sync_result = sync_markdown_to_dify(
            final_markdown,
            metadata,
            config=dify_config,
            mapping_path=dify_mapping_path(),
        )
        knowledge_payload = {
            "ok": sync_result.ok,
            "action": sync_result.action,
            "dataset_id": sync_result.dataset_id,
            "document_id": sync_result.document_id,
            "document_name": sync_result.document_name,
            "message": sync_result.message,
        }
    except Exception as exc:
        knowledge_payload = {"ok": False, "error": str(exc)}

    meta["status"] = "confirmed"
    meta["stage"] = "final"
    if workflow_run_id:
        meta["archive_workflow_run_id"] = workflow_run_id
    meta["meeting_type"] = meeting_type
    meta["meeting_series"] = meeting_series
    meta["confirmed_at"] = now_iso()
    meta["updated_at"] = meta["confirmed_at"]
    version_path = save_version(folder, final_markdown, "archive-confirmed")
    meta.setdefault("versions", []).append({"label": "archive-confirmed", "path": version_path, "created_at": meta["confirmed_at"]})
    meta["export"] = {
        "markdown_path": str(result.md_path),
        "markdown_created": result.md_created,
        "markdown_message": result.md_message,
        "word_path": str(result.docx_path),
        "word_created": result.docx_created,
        "word_message": result.docx_message,
        "google_drive_synced": result.sync_created,
        "google_drive_message": result.sync_message,
    }
    meta["raw_archive"] = raw_archive_payload
    meta["knowledge"] = knowledge_payload
    meta["post_review_artifacts"] = artifacts_payload
    meta["structured_table"] = structured_table_payload
    meta["targets"] = extract_targets(final_markdown)
    archive_all_ok = (
        bool(result.md_created)
        and bool(result.docx_created)
        and bool(raw_archive_payload.get("ok"))
        and bool(knowledge_payload.get("ok"))
        and bool(structured_table_payload.get("ok"))
        and bool(artifacts_payload.get("ok"))
    )
    archive_message = (
        "归档成功：Markdown、Word、原始记录、结构化表格和知识库同步已完成。"
        if archive_all_ok
        else "归档完成但有异常：部分导出、原始记录归档、结构化表格或知识库同步未完全成功，请查看结果明细。"
    )
    meta["archive_message"] = archive_message
    save_meta(folder, meta)
    return {
        "ok": result.md_created and result.docx_created,
        "draft_id": draft_id,
        "status": meta["status"],
        "message": archive_message,
        "export": meta["export"],
        "raw_archive": raw_archive_payload,
        "knowledge": knowledge_payload,
        "structured_table": structured_table_payload,
        "post_review_artifacts": artifacts_payload,
        "history_url": f"{public_base_url()}/history",
    }


def confirm_archive(draft_id: str) -> dict[str, Any]:
    folder = draft_folder(draft_id)
    if not folder.exists():
        raise FileNotFoundError("draft not found")
    meta = load_meta(folder)
    if resolve_dify_app_id(meta) == DIFY_WORKFLOW_APP_ID:
        return execute_confirm_archive(draft_id)
    markdown = current_markdown(folder) if (folder / "current.md").exists() else ""
    return call_dify_workflow_for_confirmation(
        action="archive_confirmed",
        draft_id=draft_id,
        markdown=markdown,
        meta=meta,
    )


def confirm_input_draft(draft_id: str, markdown: str | None = None) -> dict[str, Any]:
    folder = draft_folder(draft_id)
    if not folder.exists():
        raise FileNotFoundError("draft not found")
    meta = load_meta(folder)
    final_markdown = (markdown if markdown is not None else current_markdown(folder)).strip()
    if not final_markdown.startswith("#"):
        raise ValueError("input markdown must start with a heading")
    (folder / "current.md").write_text(final_markdown + "\n", encoding="utf-8")
    input_path = folder / "input-confirmed.md"
    input_path.write_text(final_markdown + "\n", encoding="utf-8")
    version_path = save_version(folder, final_markdown, "input-confirmed")
    confirmed_at = now_iso()
    meta["status"] = "input_confirmed"
    meta["stage"] = "pre_agent"
    meta["input_confirmed_at"] = confirmed_at
    meta["updated_at"] = confirmed_at
    meta.setdefault("versions", []).append({"label": "input-confirmed", "path": version_path, "created_at": confirmed_at})
    meta["confirmed_input_path"] = str(input_path)
    meta["targets"] = extract_targets(final_markdown)
    continue_payload: dict[str, Any]
    try:
        continue_payload = call_dify_workflow_for_minutes(draft_id=draft_id, source_markdown=final_markdown, meta=meta)
    except Exception as exc:
        continue_payload = {"ok": False, "error": str(exc), "message": "输入稿已确认，但自动进入会议纪要整理失败。"}
    meta["post_agent_run"] = continue_payload
    if continue_payload.get("ok"):
        meta["post_agent_draft_id"] = continue_payload.get("draft_id")
        meta["post_agent_review_url"] = continue_payload.get("review_url")
    save_meta(folder, meta)
    return {
        "ok": True,
        "draft_id": draft_id,
        "status": meta["status"],
        "stage": meta["stage"],
        "confirmed_input_path": str(input_path),
        "post_agent_run": continue_payload,
        "next_review_url": continue_payload.get("review_url") if continue_payload.get("ok") else "",
        "message": continue_payload.get("message") or "输入 Agent 前的人工校对稿已确认。",
        "drafts_url": f"{public_base_url()}/drafts?status=all",
    }


def run_confirm_input_job(draft_id: str, markdown: str | None = None) -> None:
    folder = draft_folder(draft_id)
    title = "投资会议纪要"
    meeting_date = ""
    meeting_type = ""
    meeting_series = MEETING_SERIES_DEFAULT
    if folder.exists():
        try:
            meta = load_meta(folder)
            title = str(meta.get("title") or title)
            meeting_date = str(meta.get("meeting_date") or "")
            meeting_type = str(meta.get("meeting_type") or "")
            meeting_series = str(meta.get("meeting_series") or MEETING_SERIES_DEFAULT)
        except Exception:
            pass

    def write_status(payload: dict[str, Any]) -> None:
        write_confirm_input_job_status(
            draft_id,
            {
                "title": title,
                "meeting_date": meeting_date,
                "meeting_type": meeting_type,
                "meeting_series": meeting_series,
                **payload,
            },
        )

    stop_heartbeat = threading.Event()

    def heartbeat() -> None:
        started_at = time.monotonic()
        while not stop_heartbeat.wait(5):
            elapsed = max(0.0, time.monotonic() - started_at)
            percent = min(92, 35 + int(elapsed / 12))
            write_status(
                {
                    "ok": True,
                    "status": "running",
                    "stage": "Dify 会议整理中",
                    "percent": percent,
                    "message": "Dify 工作流正在根据已确认输入稿生成二次校对稿；长文本可能需要更久。",
                    "status_url": f"/confirm-input/status?id={urllib.parse.quote(draft_id)}",
                }
            )

    heartbeat_thread = threading.Thread(target=heartbeat, daemon=True)
    try:
        write_status(
            {
                "ok": True,
                "status": "running",
                "stage": "保存初校稿",
                "percent": 12,
                "message": "正在保存人工初校后的输入稿。",
                "status_url": f"/confirm-input/status?id={urllib.parse.quote(draft_id)}",
            }
        )
        write_status(
            {
                "ok": True,
                "status": "running",
                "stage": "提交 Dify 工作流",
                "percent": 35,
                "message": "正在提交给 Dify 会议纪要整理工作流。",
                "status_url": f"/confirm-input/status?id={urllib.parse.quote(draft_id)}",
            }
        )
        heartbeat_thread.start()
        result = confirm_input_draft(draft_id, markdown)
        stop_heartbeat.set()
        if result.get("post_agent_run") and not (result.get("post_agent_run") or {}).get("ok"):
            error = str((result.get("post_agent_run") or {}).get("error") or result.get("message") or "Dify 整理失败")
            write_status(
                {
                    "ok": False,
                    "status": "failed",
                    "stage": "整理失败",
                    "percent": 100,
                    "message": "初校稿已保存，但 Dify 工作流整理失败。",
                    "error": error,
                    "result": result,
                    "status_url": f"/confirm-input/status?id={urllib.parse.quote(draft_id)}",
                }
            )
            return
        write_status(
            {
                "ok": True,
                "status": "done",
                "stage": "准备打开二次校对页",
                "percent": 100,
                "message": "会议纪要整理稿已生成，即将打开二次校对页面。",
                "result": result,
                "next_review_url": str(result.get("next_review_url") or ""),
                "status_url": f"/confirm-input/status?id={urllib.parse.quote(draft_id)}",
            }
        )
    except Exception as exc:
        stop_heartbeat.set()
        write_status(
            {
                "ok": False,
                "status": "failed",
                "stage": "整理失败",
                "percent": 100,
                "message": "确认初校后的会议整理失败，请检查 Dify 工作流状态。",
                "error": str(exc),
                "traceback": traceback.format_exc(),
                "status_url": f"/confirm-input/status?id={urllib.parse.quote(draft_id)}",
            }
        )


def start_confirm_input_job(draft_id: str, markdown: str | None = None) -> dict[str, Any]:
    folder = draft_folder(draft_id)
    if not folder.exists():
        raise FileNotFoundError("draft not found")
    existing = load_confirm_input_job_status(draft_id)
    if existing.get("status") in {"queued", "running"}:
        return {
            "ok": True,
            "draft_id": draft_id,
            "status": existing.get("status"),
            "status_url": f"/confirm-input/status?id={urllib.parse.quote(draft_id)}",
            "message": "会议整理任务已经在后台运行。",
        }
    if existing.get("status") == "done":
        return {
            "ok": True,
            "draft_id": draft_id,
            "status": "done",
            "status_url": f"/confirm-input/status?id={urllib.parse.quote(draft_id)}",
            "next_review_url": str(existing.get("next_review_url") or ""),
            "message": "会议整理任务已完成。",
        }
    meta = load_meta(folder)
    status_url = f"/confirm-input/status?id={urllib.parse.quote(draft_id)}"
    write_confirm_input_job_status(
        draft_id,
        {
            "ok": True,
            "status": "queued",
            "stage": "等待 Dify 会议整理",
            "percent": 3,
            "message": "人工初校已提交，会议整理将在后台执行。",
            "title": str(meta.get("title") or "投资会议纪要"),
            "meeting_date": str(meta.get("meeting_date") or ""),
            "meeting_type": str(meta.get("meeting_type") or ""),
            "meeting_series": str(meta.get("meeting_series") or MEETING_SERIES_DEFAULT),
            "status_url": status_url,
        },
    )
    threading.Thread(target=run_confirm_input_job, args=(draft_id, markdown), daemon=True).start()
    return {
        "ok": True,
        "draft_id": draft_id,
        "status": "queued",
        "status_url": status_url,
        "message": "人工初校已提交，会议整理正在后台运行。",
    }


def knowledge_index_by_markdown_path() -> dict[str, dict[str, Any]]:
    index: dict[str, dict[str, Any]] = {}
    mapping = load_mapping(dify_mapping_path())
    for key, value in mapping.items():
        if not isinstance(value, dict):
            continue
        if key.startswith("backfill:"):
            index[key.removeprefix("backfill:")] = {
                "ok": True,
                "document_id": value.get("document_id") or "",
                "document_name": value.get("document_name") or "",
                "dataset_id": value.get("dataset_id") or "",
                "meeting_type": value.get("meeting_type") or "",
                "meeting_series": value.get("meeting_series") or "",
                "updated_at": value.get("updated_at") or "",
                "source": "mapping",
            }

    for meta_file in storage_dir().glob("*/*/meta.json"):
        try:
            meta = json.loads(meta_file.read_text(encoding="utf-8"))
        except Exception:
            continue
        export = meta.get("export") or {}
        markdown_path = str(export.get("markdown_path") or "")
        if not markdown_path:
            continue
        knowledge = meta.get("knowledge") or {}
        index[markdown_path] = {
            "ok": bool(knowledge.get("ok")),
            "document_id": knowledge.get("document_id") or "",
            "document_name": knowledge.get("document_name") or "",
            "dataset_id": knowledge.get("dataset_id") or "",
            "meeting_type": meta.get("meeting_type") or "",
            "meeting_series": meta.get("meeting_series") or "",
            "message": knowledge.get("message") or knowledge.get("error") or "",
            "updated_at": meta.get("updated_at") or "",
            "source": "review_meta",
        }
    return index


def permission_meta_for_note(path: Path, markdown: str) -> dict[str, Any]:
    sidecar = path.with_suffix(".permissions.json")
    permissions: dict[str, Any] = {
        "visibility": "internal",
        "owner": "local",
        "allowed_users": [],
        "allowed_groups": [],
        "source": "default",
    }
    if sidecar.exists():
        try:
            loaded = json.loads(sidecar.read_text(encoding="utf-8"))
            if isinstance(loaded, dict):
                for key in ("visibility", "owner", "allowed_users", "allowed_groups"):
                    if key in loaded:
                        permissions[key] = loaded[key]
                permissions["source"] = "sidecar"
        except Exception:
            permissions["source"] = "sidecar_error"

    visibility = markdown_field(markdown, "可见范围") or markdown_field(markdown, "权限")
    owner = markdown_field(markdown, "负责人") or markdown_field(markdown, "owner")
    if visibility:
        permissions["visibility"] = visibility
        permissions["source"] = "markdown"
    if owner:
        permissions["owner"] = owner
        permissions["source"] = "markdown"
    return permissions


def scan_history() -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    final_root = output_dir()
    knowledge_index = knowledge_index_by_markdown_path()
    if final_root.exists():
        for path in sorted(final_root.rglob("*.md")):
            if ".derived" in path.parts:
                continue
            try:
                content = path.read_text(encoding="utf-8")
            except Exception:
                continue
            meeting_date = normalize_meeting_date(markdown_field(content, "会议日期") or _date_from_path(path))
            title = markdown_title(content, sanitize_title(re.sub(r"^20\d{2}-\d{2}-\d{2}\s*-\s*", "", path.stem)))
            word_path = path.with_suffix(".docx")
            items.append(
                {
                    "id": "file:" + str(path),
                    "meeting_date": meeting_date,
                    "title": title,
                    "meeting_type": normalize_meeting_type(markdown_field(content, "会议类型", "多人复盘会")),
                    "meeting_series": normalize_meeting_series(markdown_field(content, "会议系列", MEETING_SERIES_DEFAULT)),
                    "status": "confirmed",
                    "updated_at": datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds"),
                    "targets": extract_targets(content),
                    "review_url": f"{public_base_url()}/view?path={urllib.parse.quote(str(path))}",
                    "markdown_path": str(path),
                    "word_path": str(word_path) if word_path.exists() else "",
                    "knowledge": knowledge_index.get(str(path), {"ok": False, "message": "未找到知识库映射"}),
                    "permissions": permission_meta_for_note(path, content),
                    "content": content,
                }
            )
    items.sort(key=lambda item: (str(item.get("meeting_date") or ""), str(item.get("updated_at") or "")), reverse=True)
    return items


def external_root_for_path(path: Path) -> tuple[str, str, Path] | None:
    try:
        resolved = path.resolve()
    except Exception:
        return None
    for kind, label, root in EXTERNAL_SOURCE_ROOTS:
        try:
            root_resolved = root.resolve()
        except Exception:
            continue
        if resolved == root_resolved or root_resolved in resolved.parents:
            return kind, label, root_resolved
    return None


def scan_external_sources() -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for kind, label, root in EXTERNAL_SOURCE_ROOTS:
        if not root.exists():
            continue
        for path in sorted(root.rglob("*.md")):
            try:
                content = path.read_text(encoding="utf-8")
            except Exception:
                continue
            raw_path = markdown_field(content, "原始归档")
            item = {
                "id": f"{kind}:{path}",
                "kind": kind,
                "kind_label": label,
                "date": normalize_meeting_date(markdown_field(content, "资料日期") or _date_from_path(path)),
                "title": markdown_title(content, sanitize_title(re.sub(r"^20\d{2}-\d{2}-\d{2}\s*-\s*", "", path.stem))),
                "type": markdown_field(content, "资料类型", label),
                "source_link": markdown_field(content, "原文链接"),
                "raw_path": raw_path,
                "markdown_path": str(path),
                "permissions": permission_meta_for_note(path, content),
                "updated_at": datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds"),
                "view_path": f"/external-view?path={urllib.parse.quote(str(path))}",
                "view_url": f"{public_base_url()}/external-view?path={urllib.parse.quote(str(path))}",
                "content": content,
            }
            items.append(item)
    items.sort(key=lambda item: (str(item.get("date") or ""), str(item.get("updated_at") or "")), reverse=True)
    deduped: dict[str, dict[str, Any]] = {}
    for item in items:
        source_link = str(item.get("source_link") or "")
        key = f"{item.get('kind')}:{source_link}" if source_link else str(item.get("id") or item.get("markdown_path"))
        if key in deduped:
            deduped[key]["version_count"] = int(deduped[key].get("version_count") or 1) + 1
            continue
        item["version_count"] = 1
        deduped[key] = item
    return list(deduped.values())


def filter_external_sources(items: list[dict[str, Any]], query: dict[str, list[str]]) -> list[dict[str, Any]]:
    date = (query.get("date") or [""])[0].strip()
    kind = (query.get("kind") or [""])[0].strip()
    keyword = (query.get("q") or [""])[0].strip().lower()
    visibility = (query.get("visibility") or [""])[0].strip().lower()
    owner = (query.get("owner") or [""])[0].strip().lower()
    filtered = []
    for item in items:
        haystack = "\n".join(
            [
                str(item.get("title") or ""),
                str(item.get("type") or ""),
                str(item.get("source_link") or ""),
                str(item.get("raw_path") or ""),
                str(item.get("markdown_path") or ""),
                str(item.get("content") or ""),
            ]
        ).lower()
        if date and str(item.get("date") or "") != date:
            continue
        if kind and str(item.get("kind") or "") != kind:
            continue
        if keyword and keyword not in haystack:
            continue
        permissions = item.get("permissions") or {}
        if visibility and visibility != str(permissions.get("visibility") or "").lower():
            continue
        if owner and owner not in str(permissions.get("owner") or "").lower():
            continue
        filtered.append(item)
    return filtered


def filter_accessible_items(items: list[dict[str, Any]], identity: dict[str, Any]) -> list[dict[str, Any]]:
    return [
        item
        for item in items
        if can_identity_access(identity, item.get("permissions") if isinstance(item.get("permissions"), dict) else {})
    ]


def _date_from_path(path: Path) -> str:
    match = re.search(r"20\d{2}-\d{2}-\d{2}", str(path))
    return match.group(0) if match else ""


def filter_history(items: list[dict[str, Any]], query: dict[str, list[str]]) -> list[dict[str, Any]]:
    date = (query.get("date") or [""])[0].strip()
    meeting_type = (query.get("meeting_type") or [""])[0].strip()
    meeting_series = (query.get("meeting_series") or [""])[0].strip()
    keyword = (query.get("q") or [""])[0].strip().lower()
    symbol = (query.get("symbol") or [""])[0].strip().lower()
    visibility = (query.get("visibility") or [""])[0].strip().lower()
    owner = (query.get("owner") or [""])[0].strip().lower()
    filtered = []
    for item in items:
        haystack = "\n".join(
            [
                str(item.get("title") or ""),
                str(item.get("meeting_type") or ""),
                str(item.get("meeting_series") or ""),
                str(item.get("meeting_date") or ""),
                " ".join(str(target) for target in item.get("targets") or []),
                str(item.get("markdown_path") or ""),
                str(item.get("word_path") or ""),
                json.dumps(item.get("knowledge") or {}, ensure_ascii=False),
                json.dumps(item.get("permissions") or {}, ensure_ascii=False),
                str(item.get("content") or ""),
            ]
        ).lower()
        if date and str(item.get("meeting_date") or "") != date:
            continue
        if meeting_type and str(item.get("meeting_type") or "") != meeting_type:
            continue
        if meeting_series and str(item.get("meeting_series") or "") != meeting_series:
            continue
        if keyword and keyword not in haystack:
            continue
        if symbol and symbol not in haystack:
            continue
        permissions = item.get("permissions") or {}
        if visibility and visibility != str(permissions.get("visibility") or "").lower():
            continue
        if owner and owner not in str(permissions.get("owner") or "").lower():
            continue
        filtered.append(item)
    return filtered


def snippet_for_terms(text: str, terms: list[str], *, window: int = 120) -> str:
    lowered = text.lower()
    for term in terms:
        if not term:
            continue
        index = lowered.find(term.lower())
        if index >= 0:
            start = max(0, index - window)
            end = min(len(text), index + len(term) + window)
            prefix = "..." if start else ""
            suffix = "..." if end < len(text) else ""
            return prefix + re.sub(r"\s+", " ", text[start:end]).strip() + suffix
    return re.sub(r"\s+", " ", text[: window * 2]).strip()


def original_text_for_search(markdown: str) -> str:
    start = markdown.find("## 一、发言整理")
    if start < 0:
        start = 0
    end_candidates = [
        position
        for marker in ("## 二、存疑与待确认", "## 二、AI结构化总结", "## 三、标的汇总表", "## 四、存疑与待确认")
        for position in [markdown.find(marker, start + 1)]
        if position >= 0
    ]
    end = min(end_candidates) if end_candidates else len(markdown)
    return markdown[start:end]


def external_text_for_search(markdown: str) -> str:
    sections = [
        ("## 一、正文", ("## 二、待人工确认",)),
        ("## 四、按时间整理的原文", ("## 五、待人工确认",)),
    ]
    for start_marker, end_markers in sections:
        start = markdown.find(start_marker)
        if start < 0:
            continue
        end_candidates = [
            position
            for marker in end_markers
            for position in [markdown.find(marker, start + 1)]
            if position >= 0
        ]
        end = min(end_candidates) if end_candidates else len(markdown)
        return markdown[start:end]
    return markdown


def target_query(query_text: str, *, limit: int = 20, identity: dict[str, Any] | None = None) -> dict[str, Any]:
    query_text = query_text.strip()
    if not query_text:
        return {"ok": False, "error": "缺少查询标的", "query": query_text, "symbol_lookup": None, "results": []}
    symbol_lookup = query_symbols(
        query_text,
        market="all",
        limit=8,
        root=DEFAULT_SYMBOL_ROOT,
        alias_path=DEFAULT_ALIAS_PATH,
    )
    terms = {query_text}
    for candidate in symbol_lookup.get("candidates", []):
        if isinstance(candidate, dict):
            terms.add(str(candidate.get("symbol") or ""))
            terms.add(str(candidate.get("name") or ""))
    terms = {term for term in terms if term}

    results: list[dict[str, Any]] = []
    access_identity = identity or {
        "authenticated": True,
        "user_id": "local",
        "role": "admin",
        "groups": ["admin", "local"],
        "is_admin": True,
    }
    for item in filter_accessible_items(scan_history(), access_identity):
        content = str(item.get("content") or "")
        evidence_source = original_text_for_search(content)
        haystack = "\n".join(
            [
                str(item.get("title") or ""),
                " ".join(str(target) for target in item.get("targets") or []),
                evidence_source,
            ]
        ).lower()
        matched_terms = sorted(term for term in terms if term.lower() in haystack)
        if not matched_terms:
            continue
        result = {key: value for key, value in item.items() if key != "content"}
        result["source_type"] = "meeting_minutes"
        result["source_label"] = "会议纪要"
        result["matched_terms"] = matched_terms
        result["evidence_snippet"] = snippet_for_terms(evidence_source, matched_terms)
        results.append(result)
    for item in filter_accessible_items(scan_external_sources(), access_identity):
        content = str(item.get("content") or "")
        evidence_source = external_text_for_search(content)
        haystack = "\n".join(
            [
                str(item.get("title") or ""),
                str(item.get("type") or ""),
                str(item.get("source_link") or ""),
                evidence_source,
            ]
        ).lower()
        matched_terms = sorted(term for term in terms if term.lower() in haystack)
        if not matched_terms:
            continue
        result = {key: value for key, value in item.items() if key != "content"}
        result["source_type"] = "external_source"
        result["source_label"] = str(item.get("kind_label") or "外部资料")
        result["meeting_date"] = item.get("date") or ""
        result["review_url"] = item.get("view_url") or item.get("view_path") or ""
        result["word_path"] = ""
        result["matched_terms"] = matched_terms
        result["evidence_snippet"] = snippet_for_terms(evidence_source, matched_terms)
        results.append(result)
    results.sort(key=lambda item: (str(item.get("meeting_date") or item.get("date") or ""), str(item.get("updated_at") or "")), reverse=True)
    results = results[:limit]
    source_counts = {
        "meeting_minutes": sum(1 for item in results if item.get("source_type") == "meeting_minutes"),
        "external_source": sum(1 for item in results if item.get("source_type") == "external_source"),
    }
    return {
        "ok": True,
        "query": query_text,
        "symbol_lookup": symbol_lookup,
        "result_count": len(results),
        "source_counts": source_counts,
        "results": results,
        "note": "这是基于本地正式纪要和外部资料归档的可追溯检索结果，不是新的投资结论。",
    }


TARGET_RESOLUTION_CONFIRMED_THRESHOLD = 0.9
TARGET_RESOLUTION_REVIEW_THRESHOLD = 0.65
TARGET_RESOLUTION_STOPWORDS = {
    "会议纪要",
    "会议资料",
    "结构化表格",
    "投资判断",
    "核心观点",
    "人工确认",
    "存疑词",
    "待确认",
    "有限公司",
    "股份公司",
}
TARGET_NAME_SUFFIX_RE = re.compile(
    r"[\u4e00-\u9fffA-Za-z0-9·]{1,14}"
    r"(?:科技|股份|新材|材料|精工|电气|数控|能源|药业|生物|电子|光电|智能|实业|集团|化学|化工|矿业|传媒|证券|银行|南油)"
)
STOCK_PROFILE_CACHE: dict[str, dict[str, Any]] = {}


def clean_target_mention_name(value: str) -> str:
    text = str(value or "").strip()
    text = re.sub(r"[`*_#>\[\]【】]", "", text)
    text = re.sub(r"[（(]\s*(?:\d{6}(?:\.[A-Z]{2})?|[A-Z]{2,8}(?:\.[A-Z]{1,4})?)\s*[）)]", "", text)
    text = re.sub(r"^.*?(看好|关注|强推|买入|加仓|增持|提到|提及|标的是|公司是|股票是)", "", text).strip()
    text = text.strip(" ：:，,。；;、/\\|")
    return text[:40]


def target_context_sentence(text: str, start: int, end: int) -> str:
    left = max(text.rfind(marker, 0, start) for marker in ("\n", "。", "！", "？", ";", "；"))
    right_candidates = [text.find(marker, end) for marker in ("\n", "。", "！", "？", ";", "；")]
    right_candidates = [position for position in right_candidates if position >= 0]
    right = min(right_candidates) if right_candidates else min(len(text), end + 140)
    snippet = text[left + 1 : right + 1].strip()
    snippet = re.sub(r"\s+", " ", snippet)
    return snippet[:260]


def is_reasonable_target_mention(name: str) -> bool:
    if not name or name in TARGET_RESOLUTION_STOPWORDS:
        return False
    if len(name) < 2 or len(name) > 28:
        return False
    if any(stop in name for stop in TARGET_RESOLUTION_STOPWORDS):
        return False
    if re.fullmatch(r"\d{1,6}", name):
        return False
    if re.search(r"(会议|纪要|表格|观点|判断|确认|问题|来源|模型|生成|结构化|工作流|服务器|上下文)", name):
        return False
    return True


def add_target_mention(
    mentions: list[dict[str, Any]],
    seen: set[str],
    *,
    raw_name: str,
    context: str = "",
    source: str = "auto",
    speaker: str = "",
    code_hint: str = "",
) -> None:
    cleaned = clean_target_mention_name(raw_name)
    if not is_reasonable_target_mention(cleaned):
        return
    if re.search(r"[/／、和与及,，;；]", cleaned):
        for part in split_target_query_names(cleaned):
            if part == cleaned:
                continue
            add_target_mention(
                mentions,
                seen,
                raw_name=part,
                context=context,
                source=source,
                speaker=speaker,
                code_hint=code_hint,
            )
        return
    key = cleaned.lower()
    if key in seen:
        return
    seen.add(key)
    mentions.append(
        {
            "raw_name": cleaned,
            "context": context.strip()[:260],
            "source": source,
            "speaker": speaker.strip(),
            "code_hint": code_hint.strip(),
        }
    )


def extract_target_mentions_from_text(text: str, *, limit: int = 120) -> list[dict[str, Any]]:
    body = str(text or "")
    mentions: list[dict[str, Any]] = []
    seen: set[str] = set()
    code_pattern = r"(?:\d{6}(?:\.[A-Z]{2})?|[A-Z]{2,8}(?:\.[A-Z]{1,4})?)"
    for match in re.finditer(rf"([\u4e00-\u9fffA-Za-z0-9·]{{2,22}})\s*[（(]({code_pattern})[）)]", body):
        add_target_mention(
            mentions,
            seen,
            raw_name=match.group(1),
            context=target_context_sentence(body, match.start(), match.end()),
            source="code_pair",
            code_hint=match.group(2),
        )
        if len(mentions) >= limit:
            return mentions
    for match in re.finditer(r"\*\*([^*\n]{2,40})\*\*", body):
        add_target_mention(
            mentions,
            seen,
            raw_name=match.group(1),
            context=target_context_sentence(body, match.start(), match.end()),
            source="bold",
        )
        if len(mentions) >= limit:
            return mentions
    for match in TARGET_NAME_SUFFIX_RE.finditer(body):
        add_target_mention(
            mentions,
            seen,
            raw_name=match.group(0),
            context=target_context_sentence(body, match.start(), match.end()),
            source="suffix",
        )
        if len(mentions) >= limit:
            return mentions
    return mentions


def split_target_query_names(raw_name: str) -> list[str]:
    has_separator = bool(re.search(r"[/／、和与及,，;；]", raw_name))
    pieces = [piece.strip() for piece in re.split(r"[/／、和与及,，;；]+", raw_name) if piece.strip()]
    cleaned = [clean_target_mention_name(piece) for piece in pieces]
    cleaned = [piece for piece in cleaned if is_reasonable_target_mention(piece)]
    if cleaned:
        return cleaned
    if has_separator:
        return []
    fallback = clean_target_mention_name(raw_name)
    return [fallback] if is_reasonable_target_mention(fallback) else []


def symbol_code_value(symbol: str) -> str:
    match = re.search(r"\d{6}", str(symbol or ""))
    return match.group(0) if match else ""


def eastmoney_market_id(code: str) -> str:
    digits = symbol_code_value(code)
    if digits.startswith(("6", "9")):
        return "1"
    return "0"


def stock_profile_for_symbol(symbol: str) -> dict[str, Any]:
    digits = symbol_code_value(symbol)
    if not digits:
        return {}
    cached = STOCK_PROFILE_CACHE.get(digits)
    if cached is not None:
        return cached
    url = "https://push2.eastmoney.com/api/qt/stock/get"
    params = {
        "fltt": "2",
        "invt": "2",
        "fields": "f57,f58,f84,f85,f127,f116,f117,f189,f43",
        "secid": f"{eastmoney_market_id(digits)}.{digits}",
    }
    full_url = f"{url}?{urllib.parse.urlencode(params)}"
    profile: dict[str, Any] = {
        "code": digits,
        "symbol": symbol,
        "name": "",
        "industry": "",
        "source": "eastmoney_push2",
        "ok": False,
    }
    try:
        request = urllib.request.Request(full_url, headers={"User-Agent": "meeting-minutes-mcp-bridge/1.0"})
        with urllib.request.urlopen(request, timeout=6) as response:
            payload = json.loads(response.read().decode("utf-8", errors="replace"))
        data = payload.get("data") if isinstance(payload, dict) else {}
        if isinstance(data, dict):
            profile.update(
                {
                    "ok": True,
                    "code": str(data.get("f57") or digits),
                    "name": str(data.get("f58") or ""),
                    "industry": str(data.get("f127") or ""),
                    "total_shares": data.get("f84"),
                    "float_shares": data.get("f85"),
                    "mcap": data.get("f116"),
                    "float_mcap": data.get("f117"),
                    "list_date": str(data.get("f189") or ""),
                    "price": data.get("f43"),
                }
            )
    except Exception as exc:
        profile["error"] = str(exc)
    STOCK_PROFILE_CACHE[digits] = profile
    return profile


def a_stock_stock_profile(arguments: dict[str, Any]) -> dict[str, Any]:
    raw_codes = arguments.get("codes")
    if raw_codes is None:
        raw_codes = arguments.get("code") or arguments.get("symbols") or arguments.get("symbol")
    if isinstance(raw_codes, str):
        codes = [part.strip() for part in re.split(r"[,，\s]+", raw_codes) if part.strip()]
    elif isinstance(raw_codes, list):
        codes = [str(part).strip() for part in raw_codes if str(part).strip()]
    else:
        codes = []
    profiles = [stock_profile_for_symbol(code) for code in codes[:50]]
    return {
        "ok": True,
        "tool": "stock_profile",
        "profiles": profiles,
        "source": "eastmoney_push2",
    }


def resolve_single_target_mention(mention: dict[str, Any]) -> dict[str, Any]:
    raw_name = clean_target_mention_name(str(mention.get("raw_name") or mention.get("name") or ""))
    code_hint = str(mention.get("code_hint") or "").strip()
    query_names = split_target_query_names(raw_name)
    candidates_by_symbol: dict[str, dict[str, Any]] = {}
    lookups: list[dict[str, Any]] = []
    for query_name in query_names[:4]:
        if not query_name:
            continue
        lookup = query_symbols(
            query_name,
            market="all",
            limit=5,
            root=DEFAULT_SYMBOL_ROOT,
            alias_path=DEFAULT_ALIAS_PATH,
        )
        lookups.append({"query": query_name, **lookup})
        for candidate in lookup.get("candidates", []):
            if not isinstance(candidate, dict):
                continue
            symbol = str(candidate.get("symbol") or "")
            if not symbol:
                continue
            current = candidates_by_symbol.get(symbol)
            if current is None or float(candidate.get("confidence") or 0) > float(current.get("confidence") or 0):
                candidates_by_symbol[symbol] = candidate
    candidates = sorted(candidates_by_symbol.values(), key=lambda item: float(item.get("confidence") or 0), reverse=True)
    best = candidates[0] if candidates else {}
    confidence = float(best.get("confidence") or 0)
    reason = str(best.get("match_type") or "")
    hint_code = symbol_code_value(code_hint)
    best_code = symbol_code_value(str(best.get("symbol") or ""))
    if hint_code and best_code and hint_code == best_code:
        confidence = max(confidence, 0.98)
        reason = f"{reason or 'symbol_lookup'}+source_code_hint"
    if hint_code and not candidates:
        by_code = query_symbols(
            hint_code,
            market="all",
            limit=5,
            root=DEFAULT_SYMBOL_ROOT,
            alias_path=DEFAULT_ALIAS_PATH,
        )
        lookups.append({"query": hint_code, **by_code})
        candidates = [candidate for candidate in by_code.get("candidates", []) if isinstance(candidate, dict)]
        best = candidates[0] if candidates else {}
        confidence = max(float(best.get("confidence") or 0), 0.96 if best else 0)
        reason = str(best.get("match_type") or "source_code_hint")
    if confidence >= TARGET_RESOLUTION_CONFIRMED_THRESHOLD:
        decision = "confirmed"
    elif confidence >= TARGET_RESOLUTION_REVIEW_THRESHOLD:
        decision = "review"
    else:
        decision = "rejected"
    stock_profile = stock_profile_for_symbol(str(best.get("symbol") or "")) if best else {}
    return {
        "raw_name": raw_name,
        "query_names": query_names,
        "context": str(mention.get("context") or "")[:260],
        "speaker": str(mention.get("speaker") or ""),
        "source": str(mention.get("source") or ""),
        "code_hint": code_hint,
        "canonical_name": str(best.get("name") or ""),
        "stock_code": str(best.get("symbol") or ""),
        "sector_name": str(stock_profile.get("industry") or ""),
        "profile": stock_profile,
        "market": str(best.get("market") or ""),
        "confidence": round(confidence, 4),
        "decision": decision,
        "reason": reason,
        "alternatives": candidates[:5],
    }


def normalize_target_mentions_argument(arguments: dict[str, Any]) -> list[dict[str, Any]]:
    raw_mentions = arguments.get("mentions")
    mentions: list[dict[str, Any]] = []
    seen: set[str] = set()
    if isinstance(raw_mentions, list):
        for item in raw_mentions:
            if isinstance(item, dict):
                add_target_mention(
                    mentions,
                    seen,
                    raw_name=str(item.get("raw_name") or item.get("name") or item.get("target") or ""),
                    context=str(item.get("context") or ""),
                    source=str(item.get("source") or "provided"),
                    speaker=str(item.get("speaker") or ""),
                    code_hint=str(item.get("code_hint") or item.get("stock_code") or ""),
                )
            else:
                add_target_mention(mentions, seen, raw_name=str(item), source="provided")
    text = "\n\n".join(
        str(arguments.get(key) or "")
        for key in ("text", "meeting_markdown", "source_markdown", "transcript_text")
        if arguments.get(key)
    )
    auto_extract = arguments.get("auto_extract", True)
    if text and (auto_extract or not mentions):
        try:
            max_mentions = max(1, min(int(arguments.get("max_mentions") or 120), 300))
        except (TypeError, ValueError):
            max_mentions = 120
        for item in extract_target_mentions_from_text(text, limit=max_mentions):
            add_target_mention(
                mentions,
                seen,
                raw_name=str(item.get("raw_name") or ""),
                context=str(item.get("context") or ""),
                source=str(item.get("source") or "auto"),
                speaker=str(item.get("speaker") or ""),
                code_hint=str(item.get("code_hint") or ""),
            )
    return mentions


def a_stock_resolve_targets(arguments: dict[str, Any]) -> dict[str, Any]:
    mentions = normalize_target_mentions_argument(arguments)
    resolved = [resolve_single_target_mention(mention) for mention in mentions]
    counts = {
        "confirmed": sum(1 for item in resolved if item.get("decision") == "confirmed"),
        "review": sum(1 for item in resolved if item.get("decision") == "review"),
        "rejected": sum(1 for item in resolved if item.get("decision") == "rejected"),
        "total": len(resolved),
    }
    return {
        "ok": True,
        "tool": "resolve_targets",
        "mentions": resolved,
        "confirmed": [item for item in resolved if item.get("decision") == "confirmed"],
        "review": [item for item in resolved if item.get("decision") == "review"],
        "rejected": [item for item in resolved if item.get("decision") == "rejected"],
        "counts": counts,
        "policy": {
            "confirmed_threshold": TARGET_RESOLUTION_CONFIRMED_THRESHOLD,
            "review_threshold": TARGET_RESOLUTION_REVIEW_THRESHOLD,
            "confirmed_usage": "可在会议纪要和结构化表格中自动使用 canonical_name 与 stock_code。",
            "review_usage": "只进入存疑与待确认，不自动写入正式代码。",
        },
    }


def resolve_targets_for_text(text: str, *, max_mentions: int = 120) -> dict[str, Any]:
    return a_stock_resolve_targets({"text": text, "max_mentions": max_mentions, "auto_extract": True})


def target_resolution_confirmed_markdown(resolution: dict[str, Any]) -> str:
    confirmed = resolution.get("confirmed") if isinstance(resolution.get("confirmed"), list) else []
    lines = ["## MCP确认标的代码"]
    for item in confirmed:
        if not isinstance(item, dict):
            continue
        name = str(item.get("canonical_name") or "").strip()
        code = str(item.get("stock_code") or "").strip()
        sector = str(item.get("sector_name") or "").strip()
        raw_name = str(item.get("raw_name") or "").strip()
        if name and code:
            alias = f"；原文称：{raw_name}" if raw_name and raw_name != name else ""
            sector_text = f"；板块：{sector}" if sector else ""
            lines.append(f"- {name}（{code}）{sector_text}{alias}")
    return "\n".join(lines) if len(lines) > 1 else ""


def target_resolution_prompt_context(resolution: dict[str, Any]) -> str:
    confirmed = [item for item in resolution.get("confirmed") or [] if isinstance(item, dict)]
    review = [item for item in resolution.get("review") or [] if isinstance(item, dict)]
    if not confirmed and not review:
        return ""
    lines = [
        "标的名称与代码核验上下文（来自本地 a-stock-data MCP；不得编造未列出的代码）：",
    ]
    if confirmed:
        lines.append("可直接使用的高置信标的：")
        for item in confirmed[:40]:
            raw_name = str(item.get("raw_name") or "")
            name = str(item.get("canonical_name") or "")
            code = str(item.get("stock_code") or "")
            sector = str(item.get("sector_name") or "")
            confidence = str(item.get("confidence") or "")
            if name and code:
                sector_text = f"，板块 {sector}" if sector else ""
                lines.append(f"- 原文“{raw_name}” => {name}（{code}）{sector_text}，置信度 {confidence}")
    if review:
        lines.append("需放入存疑与待确认、不得自动写入代码的候选：")
        for item in review[:30]:
            raw_name = str(item.get("raw_name") or "")
            name = str(item.get("canonical_name") or "")
            code = str(item.get("stock_code") or "")
            confidence = str(item.get("confidence") or "")
            if name and code:
                lines.append(f"- 原文“{raw_name}” 可能是 {name}（{code}），置信度 {confidence}")
    lines.append("生成会议纪要时：高置信标的可写标准名称和代码；待确认候选只写入存疑，不得改写原文判断。")
    return "\n".join(lines)


def target_resolution_review_markdown(resolution: dict[str, Any]) -> str:
    lines = ["# 标的名称与代码核验", ""]
    counts = resolution.get("counts") if isinstance(resolution.get("counts"), dict) else {}
    lines.append(
        f"确认 {counts.get('confirmed', 0)} 个；待确认 {counts.get('review', 0)} 个；未采用 {counts.get('rejected', 0)} 个。"
    )
    for title, key in (("## 可自动使用", "confirmed"), ("## 待人工确认", "review"), ("## 未采用", "rejected")):
        rows = [item for item in resolution.get(key) or [] if isinstance(item, dict)]
        if not rows:
            continue
        lines.extend(["", title, "", "| 原文 | 标准名称 | 股票代码 | 板块 | 置信度 | 上下文 |", "| --- | --- | --- | --- | --- | --- |"])
        for item in rows:
            lines.append(
                "| "
                + " | ".join(
                    [
                        str(item.get("raw_name") or "").replace("|", "/"),
                        str(item.get("canonical_name") or "").replace("|", "/"),
                        str(item.get("stock_code") or "").replace("|", "/"),
                        str(item.get("sector_name") or "").replace("|", "/"),
                        str(item.get("confidence") or ""),
                        str(item.get("context") or "").replace("|", "/")[:120],
                    ]
                )
                + " |"
            )
    return "\n".join(lines).rstrip() + "\n"


def write_target_resolution_artifacts(folder: Path, resolution: dict[str, Any], *, stage: str) -> None:
    folder.mkdir(parents=True, exist_ok=True)
    payload = {**resolution, "stage": stage, "updated_at": now_iso()}
    target_resolution_path(folder).write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    target_resolution_review_path(folder).write_text(target_resolution_review_markdown(payload), encoding="utf-8")


def default_mcp_config() -> dict[str, Any]:
    return {
        "version": 1,
        "updated_at": now_iso(),
        "servers": [
            {
                "id": "a-stock-data",
                "name": "A股全栈数据工具包",
                "type": "local_skill_bridge",
                "enabled": True,
                "description": "会议纪要工作流通过本地 MCP bridge 调用 A 股数据能力；当前开放代码候选和实时行情。",
                "skill_path": "/Users/kumaai/.codex/skills/a-stock-data/SKILL.md",
                "call_url": "http://host.docker.internal:8767/api/mcp/call",
                "tools": [
                    {
                        "name": "symbol_lookup",
                        "description": "按公司名、简称或股票代码查询本地 A 股代码候选。",
                        "arguments_schema": {
                            "query": "标的名称、简称或代码",
                            "limit": "返回候选数，默认 8",
                        },
                    },
                    {
                        "name": "quote",
                        "description": "通过腾讯行情接口查询 A 股实时行情快照。",
                        "arguments_schema": {
                            "codes": "股票代码列表，例如 ['600519', '300750']",
                        },
                    },
                    {
                        "name": "resolve_targets",
                        "description": "从会议原文上下文批量解析标的名称，并用本地 A 股代码库核验标准名称和股票代码。",
                        "arguments_schema": {
                            "mentions": "可选，模型抽取的标的候选列表，每项可含 raw_name/context/speaker/code_hint",
                            "text": "可选，会议原文或终稿 Markdown；未提供 mentions 时自动抽取候选",
                            "max_mentions": "最多解析的候选数，默认 120",
                        },
                    },
                    {
                        "name": "stock_profile",
                        "description": "按已确认股票代码查询公司基础信息和行业，用于结构化表格板块兜底补全。",
                        "arguments_schema": {
                            "codes": "股票代码列表，例如 ['300750.SZ', '688300.SH']",
                        },
                    },
                ],
            }
        ],
    }


def merge_mcp_config_defaults(config: dict[str, Any]) -> dict[str, Any]:
    defaults = default_mcp_config()
    existing_servers = config.get("servers") if isinstance(config.get("servers"), list) else []
    existing_by_id = {str(server.get("id") or ""): server for server in existing_servers if isinstance(server, dict)}
    merged_servers: list[dict[str, Any]] = []
    for default_server in defaults.get("servers") or []:
        if not isinstance(default_server, dict):
            continue
        server_id = str(default_server.get("id") or "")
        server = existing_by_id.pop(server_id, None)
        if server is None:
            merged_servers.append(default_server)
            continue
        for key in ("name", "type", "description", "skill_path", "call_url"):
            server.setdefault(key, default_server.get(key))
        existing_tools = server.get("tools") if isinstance(server.get("tools"), list) else []
        existing_tool_names = {str(tool.get("name") or "") for tool in existing_tools if isinstance(tool, dict)}
        for tool in default_server.get("tools") or []:
            if isinstance(tool, dict) and str(tool.get("name") or "") not in existing_tool_names:
                existing_tools.append(tool)
        server["tools"] = existing_tools
        merged_servers.append(server)
    merged_servers.extend(server for server in existing_by_id.values() if isinstance(server, dict))
    return {**defaults, **config, "servers": merged_servers}


def load_mcp_config() -> dict[str, Any]:
    if not DEFAULT_MCP_CONFIG_PATH.exists():
        config = default_mcp_config()
        DEFAULT_MCP_CONFIG_PATH.parent.mkdir(parents=True, exist_ok=True)
        tmp_path = DEFAULT_MCP_CONFIG_PATH.with_suffix(".json.tmp")
        tmp_path.write_text(json.dumps(config, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        tmp_path.replace(DEFAULT_MCP_CONFIG_PATH)
        return config
    try:
        config = json.loads(DEFAULT_MCP_CONFIG_PATH.read_text(encoding="utf-8"))
    except Exception:
        config = default_mcp_config()
    if not isinstance(config, dict):
        config = default_mcp_config()
    servers = config.get("servers")
    if not isinstance(servers, list):
        config["servers"] = default_mcp_config()["servers"]
    merged = merge_mcp_config_defaults(config)
    if merged != config:
        try:
            DEFAULT_MCP_CONFIG_PATH.parent.mkdir(parents=True, exist_ok=True)
            tmp_path = DEFAULT_MCP_CONFIG_PATH.with_suffix(".json.tmp")
            tmp_path.write_text(json.dumps(merged, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
            tmp_path.replace(DEFAULT_MCP_CONFIG_PATH)
        except Exception:
            pass
    return merged


def enabled_mcp_server(server_id: str) -> dict[str, Any] | None:
    for server in load_mcp_config().get("servers") or []:
        if not isinstance(server, dict):
            continue
        if server.get("id") == server_id and server.get("enabled", True):
            return server
    return None


def normalize_a_stock_code(value: str) -> str:
    code = value.strip().lower()
    for prefix in ("sh", "sz", "bj"):
        if code.startswith(prefix):
            code = code[len(prefix):]
            break
    code = "".join(ch for ch in code if ch.isdigit())
    if not code:
        return ""
    if code.startswith("6"):
        return f"sh{code}"
    if code.startswith(("0", "3")):
        return f"sz{code}"
    if code.startswith(("4", "8")):
        return f"bj{code}"
    return code


def a_stock_symbol_lookup(arguments: dict[str, Any]) -> dict[str, Any]:
    query_text = str(arguments.get("query") or arguments.get("q") or "").strip()
    limit_value = arguments.get("limit")
    try:
        limit = max(1, min(int(limit_value or 8), 50))
    except (TypeError, ValueError):
        limit = 8
    if not query_text:
        return {"ok": False, "error": "缺少 query 参数", "candidates": []}
    result = query_symbols(
        query_text,
        market="all",
        limit=limit,
        root=DEFAULT_SYMBOL_ROOT,
        alias_path=DEFAULT_ALIAS_PATH,
    )
    return {"ok": True, "tool": "symbol_lookup", "data": result}


def a_stock_tencent_quote(arguments: dict[str, Any]) -> dict[str, Any]:
    raw_codes = arguments.get("codes")
    if raw_codes is None:
        raw_codes = arguments.get("code") or arguments.get("symbols") or arguments.get("symbol")
    if isinstance(raw_codes, str):
        codes = [part.strip() for part in re.split(r"[,，\s]+", raw_codes) if part.strip()]
    elif isinstance(raw_codes, list):
        codes = [str(part).strip() for part in raw_codes if str(part).strip()]
    else:
        codes = []
    normalized_codes = [normalize_a_stock_code(code) for code in codes]
    normalized_codes = [code for code in normalized_codes if code]
    if not normalized_codes:
        return {"ok": False, "error": "缺少 codes 参数", "quotes": []}
    query = ",".join(normalized_codes[:30])
    url = f"https://qt.gtimg.cn/q={urllib.parse.quote(query)}"
    try:
        request = urllib.request.Request(url, headers={"User-Agent": "meeting-minutes-mcp-bridge/1.0"})
        with urllib.request.urlopen(request, timeout=8) as response:
            text = response.read().decode("gbk", errors="replace")
    except Exception as exc:
        return {"ok": False, "error": f"腾讯行情查询失败：{exc}", "quotes": [], "url": url}
    quotes: list[dict[str, Any]] = []
    for line in text.splitlines():
        if not line.strip() or "=" not in line:
            continue
        symbol_part, value_part = line.split("=", 1)
        symbol = symbol_part.rsplit("_", 1)[-1]
        payload = value_part.strip().strip('";')
        fields = payload.split("~") if payload else []
        quotes.append(
            {
                "symbol": symbol,
                "name": fields[1] if len(fields) > 1 else "",
                "code": fields[2] if len(fields) > 2 else "",
                "price": fields[3] if len(fields) > 3 else "",
                "previous_close": fields[4] if len(fields) > 4 else "",
                "open": fields[5] if len(fields) > 5 else "",
                "change": fields[31] if len(fields) > 31 else "",
                "change_percent": fields[32] if len(fields) > 32 else "",
                "datetime": fields[30] if len(fields) > 30 else "",
                "raw_fields": fields,
            }
        )
    return {"ok": True, "tool": "quote", "quotes": quotes, "source": "tencent"}


def mcp_bridge_call(payload: dict[str, Any]) -> dict[str, Any]:
    server_id = str(payload.get("server") or payload.get("server_id") or "").strip()
    tool_name = str(payload.get("tool") or payload.get("name") or "").strip()
    arguments = payload.get("arguments")
    if not isinstance(arguments, dict):
        arguments = {}
    if server_id != "a-stock-data":
        return {"ok": False, "error": f"未知或未支持的 MCP server：{server_id or '-'}"}
    server = enabled_mcp_server(server_id)
    if server is None:
        return {"ok": False, "error": f"MCP server 未启用：{server_id}"}
    if tool_name == "symbol_lookup":
        result = a_stock_symbol_lookup(arguments)
    elif tool_name == "quote":
        result = a_stock_tencent_quote(arguments)
    elif tool_name == "resolve_targets":
        result = a_stock_resolve_targets(arguments)
    elif tool_name == "stock_profile":
        result = a_stock_stock_profile(arguments)
    else:
        return {"ok": False, "error": f"a-stock-data 不支持该 tool：{tool_name or '-'}"}
    return {
        "ok": bool(result.get("ok")),
        "server": server_id,
        "server_name": server.get("name"),
        "tool": tool_name,
        "result": result,
    }


def knowledge_answer(query_text: str, *, identity: dict[str, Any]) -> dict[str, Any]:
    query_text = query_text.strip()
    if not query_text:
        return {"ok": False, "error": "请输入问题", "answer": "", "sources": []}
    result = target_query(query_text, limit=8, identity=identity)
    sources = result.get("results") or []
    if not sources:
        return {
            "ok": True,
            "answer": "未检索到明确记录。请换一个标的、简称、代码或关键词再查。",
            "sources": [],
            "mode": "permission_filtered_rag",
        }
    lines = [
        "基于当前账号可访问的已归档资料，检索到以下证据：",
    ]
    for index, item in enumerate(sources[:5], start=1):
        source_label = str(item.get("source_label") or "资料")
        title = str(item.get("title") or "未命名资料")
        date = str(item.get("meeting_date") or item.get("date") or "")
        snippet = str(item.get("evidence_snippet") or "").strip()
        lines.append(f"{index}. {source_label}《{title}》{date}：{snippet or '未抽取到可展示片段'}")
    lines.append("以上回答只使用当前账号权限内资料；权限外内容不会进入检索结果。")
    return {
        "ok": True,
        "answer": "\n".join(lines),
        "sources": sources[:8],
        "mode": "permission_filtered_rag",
        "source_counts": result.get("source_counts") or {},
    }


def load_planner_status() -> dict[str, Any]:
    path = planner_status_path()
    if path.exists():
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception as exc:
            return {
                "name": "Dify 工作流规划 Agent",
                "status": "blocked",
                "updated_at": now_iso(),
                "current_task": "读取规划状态失败",
                "acceptance": ["状态文件必须是 UTF-8 JSON"],
                "completed": [],
                "blocked": [f"{path}: {exc}"],
                "next_task": "修复 planner status JSON",
            }
    return {
        "name": "Dify 工作流规划 Agent",
        "status": "active",
        "updated_at": now_iso(),
        "current_task": "建立规划 agent 工作状态可见化入口",
        "acceptance": [
            "本机有持久状态文件",
            "页面能显示当前任务、验收标准、最近 run id、已完成事项和阻塞项",
            "Dify 同源入口可访问",
        ],
        "completed": [
            "恢复 Dify Skill Agent 生产链路",
            "修复人工校对页默认可直接编辑",
            "完成 Skill Agent 长文档回归与清洗节点修复",
        ],
        "recent_runs": [
            {
                "name": "Skill Agent 长文档回归",
                "run_id": "f2f746fb-1ae3-4d8a-8f5f-2a435bd11c90",
                "status": "succeeded",
            },
            {
                "name": "Skill Agent 清洗节点回归",
                "run_id": "8f7008c7-27eb-405a-b0d2-99926491637a",
                "status": "succeeded",
            },
        ],
        "blocked": [],
        "next_task": "选择下一个北极星 P0 任务",
    }


def planner_status_html() -> str:
    status = load_planner_status()

    def list_items(values: Any, empty: str = "无") -> str:
        if not isinstance(values, list) or not values:
            return f"<li>{html.escape(empty)}</li>"
        return "".join(f"<li>{html.escape(str(item))}</li>" for item in values)

    runs = status.get("recent_runs") or []
    if isinstance(runs, list) and runs:
        run_rows = []
        for run in runs:
            if not isinstance(run, dict):
                continue
            run_rows.append(
                "<tr>"
                f"<td>{html.escape(str(run.get('name') or '-'))}</td>"
                f"<td><code>{html.escape(str(run.get('run_id') or '-'))}</code></td>"
                f"<td>{html.escape(str(run.get('status') or '-'))}</td>"
                "</tr>"
            )
        runs_html = (
            "<table><thead><tr><th>名称</th><th>Run ID</th><th>状态</th></tr></thead><tbody>"
            + "".join(run_rows)
            + "</tbody></table>"
        )
    else:
        runs_html = "<p class=\"muted\">暂无记录。</p>"
    blocked = status.get("blocked") or []
    blocked_class = "error" if blocked else "muted"
    path = planner_status_path()
    return f"""
<div class="meta">
  <strong>{html.escape(str(status.get("name") or "Dify 工作流规划 Agent"))}</strong>
  <span class="status">{html.escape(str(status.get("status") or "active"))}</span>
  <div class="muted">更新时间：{html.escape(str(status.get("updated_at") or ""))}</div>
  <div class="muted">状态文件：{html.escape(str(path))}</div>
</div>
<div class="item">
  <h2>当前任务</h2>
  <p>{html.escape(str(status.get("current_task") or "-"))}</p>
</div>
<div class="item">
  <h2>验收标准</h2>
  <ul>{list_items(status.get("acceptance"))}</ul>
</div>
<div class="item">
  <h2>最近运行</h2>
  {runs_html}
</div>
<div class="item">
  <h2>已完成</h2>
  <ul>{list_items(status.get("completed"))}</ul>
</div>
<div class="item">
  <h2>阻塞项</h2>
  <ul class="{blocked_class}">{list_items(blocked)}</ul>
</div>
<div class="item">
  <h2>下一任务</h2>
  <p>{html.escape(str(status.get("next_task") or "-"))}</p>
</div>
"""


def local_health_report() -> dict[str, Any]:
    script = SCRIPT_DIR / "check_investment_workflow_health.py"
    try:
        result = subprocess.run(
            [sys.executable, str(script), "--json"],
            text=True,
            capture_output=True,
            timeout=45,
            env={**os.environ, "PYTHONUTF8": "1", "PYTHONIOENCODING": "utf-8"},
        )
    except subprocess.TimeoutExpired:
        return {"ok": False, "error": "health check timeout", "summary": {"overall": "error"}, "checks": []}
    except Exception as exc:
        return {"ok": False, "error": str(exc), "summary": {"overall": "error"}, "checks": []}
    try:
        payload = json.loads(result.stdout)
    except json.JSONDecodeError:
        return {
            "ok": False,
            "error": "health check returned non-json output",
            "stdout": result.stdout[-2000:],
            "stderr": result.stderr[-2000:],
            "summary": {"overall": "error"},
            "checks": [],
        }
    payload["ok"] = payload.get("summary", {}).get("overall") != "error"
    return payload


def load_access_audit_events(limit: int = 100) -> list[dict[str, Any]]:
    path = access_audit_log_path()
    if not path.exists():
        return []
    events: list[dict[str, Any]] = []
    try:
        lines = path.read_text(encoding="utf-8", errors="replace").splitlines()[-max(1, limit):]
    except Exception:
        return []
    for line in lines:
        try:
            item = json.loads(line)
        except json.JSONDecodeError:
            continue
        if isinstance(item, dict):
            events.append(item)
    events.reverse()
    return events


def load_mapping_audit_report() -> dict[str, Any]:
    script = SCRIPT_DIR / "audit_dify_dataset_mapping.py"
    if not script.exists():
        return {"ok": False, "error": "mapping audit script missing", "issues": []}
    try:
        result = subprocess.run(
            [sys.executable, str(script), "--json"],
            text=True,
            capture_output=True,
            timeout=30,
        )
        payload = json.loads(result.stdout or "{}")
    except Exception as exc:
        return {"ok": False, "error": str(exc), "issues": []}
    mapping = load_mapping(dify_mapping_path())
    for issue in payload.get("issues") or []:
        if not isinstance(issue, dict):
            continue
        key = str(issue.get("key") or "")
        mapped = mapping.get(key) if isinstance(mapping.get(key), dict) else {}
        issue["document_id"] = mapped.get("document_id") or issue.get("document_id") or ""
        issue["document_name"] = mapped.get("document_name") or issue.get("document_name") or ""
        issue["dataset_id"] = mapped.get("dataset_id") or issue.get("dataset_id") or ""
        issue["updated_at"] = mapped.get("updated_at") or issue.get("updated_at") or ""
    return payload


def scan_drafts() -> list[dict[str, Any]]:
    drafts: list[dict[str, Any]] = []
    root = storage_dir()
    if not root.exists():
        return drafts
    for meta_path in root.glob("*/*/meta.json"):
        try:
            meta = json.loads(meta_path.read_text(encoding="utf-8"))
        except Exception:
            continue
        if not isinstance(meta, dict):
            continue
        draft_id = str(meta.get("id") or meta_path.parent.name)
        targets = meta.get("targets") if isinstance(meta.get("targets"), list) else []
        drafts.append(
            {
                "id": draft_id,
                "meeting_date": str(meta.get("meeting_date") or meta_path.parent.parent.name),
                "title": str(meta.get("title") or "投资会议纪要"),
                "meeting_type": normalize_meeting_type(str(meta.get("meeting_type") or "多人复盘会")),
                "meeting_series": normalize_meeting_series(str(meta.get("meeting_series") or MEETING_SERIES_DEFAULT)),
                "stage": str(meta.get("stage") or ""),
                "status": str(meta.get("status") or "draft"),
                "created_at": str(meta.get("created_at") or ""),
                "updated_at": str(meta.get("updated_at") or ""),
                "input_source": str(meta.get("input_source") or ""),
                "source_run_id": str(meta.get("source_run_id") or ""),
                "target_count": len(targets),
                "review_url": f"/review?id={urllib.parse.quote(draft_id)}",
            }
        )
    drafts.sort(key=lambda item: (item.get("updated_at") or item.get("created_at") or ""), reverse=True)
    return drafts


def filter_drafts(drafts: list[dict[str, Any]], query: dict[str, list[str]]) -> list[dict[str, Any]]:
    status_filter = (query.get("status") or ["open"])[0].strip().lower() or "open"
    meeting_type = (query.get("meeting_type") or [""])[0].strip()
    meeting_series = (query.get("meeting_series") or [""])[0].strip()
    if status_filter == "all":
        filtered = drafts
    if status_filter == "open":
        filtered = [item for item in drafts if str(item.get("status") or "").lower() not in {"confirmed", "archived"}]
    elif status_filter != "all":
        filtered = [item for item in drafts if str(item.get("status") or "").lower() == status_filter]
    if meeting_type:
        filtered = [item for item in filtered if str(item.get("meeting_type") or "") == meeting_type]
    if meeting_series:
        filtered = [item for item in filtered if str(item.get("meeting_series") or "") == meeting_series]
    return filtered


def logic_tags_for_text(text: str) -> list[str]:
    value = str(text or "")
    rules = [
        ("催化", ("催化", "发布", "大会", "上市", "政策", "落地", "回收", "招标")),
        ("业绩", ("利润", "收入", "毛利", "业绩", "指引", "PE", "估值", "市值")),
        ("订单", ("订单", "客户", "供货", "产能", "扩产", "排期", "交付", "量产")),
        ("风险", ("风险", "不确定", "压力", "腐败", "监管", "退市", "立案", "担心")),
        ("存疑", ("**", "待确认", "无法确认", "存疑", "听说", "可能")),
    ]
    tags = [label for label, keywords in rules if any(keyword in value for keyword in keywords)]
    return tags or ["观点"]


def dashboard_data(date: str | None = None, identity: dict[str, Any] | None = None) -> dict[str, Any]:
    selected_date = normalize_meeting_date(date or datetime.now().strftime("%Y-%m-%d"))
    access_identity = identity or {
        "authenticated": True,
        "user_id": "local",
        "role": "admin",
        "groups": ["admin", "local"],
        "is_admin": True,
    }
    items = [
        item
        for item in filter_accessible_items(scan_history(), access_identity)
        if str(item.get("meeting_date") or "") == selected_date
    ]
    deduped: dict[str, dict[str, Any]] = {}
    for item in items:
        key = str(item.get("markdown_path") or item.get("id") or item.get("title"))
        deduped[key] = item
    items = list(deduped.values())
    target_counter: Counter[str] = Counter()
    target_logic: dict[str, list[dict[str, str]]] = defaultdict(list)
    speaker_views: dict[str, list[dict[str, str]]] = defaultdict(list)
    meeting_types: Counter[str] = Counter()
    for item in items:
        content = str(item.get("content") or "")
        title = str(item.get("title") or "")
        meeting_type = str(item.get("meeting_type") or "多人复盘会")
        meeting_types[meeting_type] += 1
        for target in item.get("targets") or []:
            target_name = str(target)
            if not target_name or target_name == "-":
                continue
            target_counter[target_name] += 1
            target_logic[target_name].append(
                {
                    "title": title,
                    "meeting_type": meeting_type,
                    "snippet": snippet_for_terms(original_text_for_search(content), [target_name], window=90),
                    "logic_tags": logic_tags_for_text(snippet_for_terms(original_text_for_search(content), [target_name], window=90)),
                    "url": str(item.get("review_url") or ""),
                }
            )
        current_speaker = ""
        for raw_line in original_text_for_search(content).splitlines():
            line = raw_line.strip()
            if line.startswith("### "):
                current_speaker = line[4:].strip()
                continue
            if current_speaker and line.startswith("【") and "】" in line:
                speaker_views[current_speaker].append(
                    {
                        "title": title,
                        "meeting_type": meeting_type,
                        "heading": line,
                        "url": str(item.get("review_url") or ""),
                    }
                )
    return {
        "date": selected_date,
        "meeting_count": len(items),
        "meeting_types": dict(meeting_types),
        "deduped_minutes": [
            {key: value for key, value in item.items() if key != "content"}
            for item in sorted(items, key=lambda row: str(row.get("updated_at") or ""), reverse=True)
        ],
        "high_frequency_targets": [
            {"target": target, "count": count, "logic": target_logic.get(target, [])[:5]}
            for target, count in target_counter.most_common(20)
        ],
        "speaker_rollup": {
            speaker: values[:10]
            for speaker, values in speaker_views.items()
            if len(values) > 1
        },
    }


PORTAL_SECTIONS: list[dict[str, str]] = [
    {
        "key": "tools",
        "title": "工具",
        "description": "OpenFinance 工具入口，以 OpenMinute 会议纪要为当前主工具。",
    },
    {
        "key": "knowledge",
        "title": "知识库",
        "description": "会议纪要 RAG 问答、外部资料和标的证据入口。",
    },
    {
        "key": "dashboard",
        "title": "看板",
        "description": "基于已确认资料生成的当日视图和运营视图。",
    },
    {
        "key": "data",
        "title": "数据",
        "description": "最高级别账号可访问的历史数据、审计和映射状态。",
    },
    {
        "key": "account",
        "title": "账号",
        "description": "账号信息、角色、权限和后台管理入口。",
    },
]


PORTAL_MODULES: list[dict[str, str]] = [
    {
        "section": "tools",
        "title": "OpenMinute-会议纪要",
        "href": "/apps/meeting-minutes",
        "description": "导入材料、输入校对、自动整理、终稿校对和确认归档。",
        "status": "已上线",
    },
    {
        "section": "tools",
        "title": "导入会议材料",
        "href": "/apps/meeting-minutes/import",
        "description": "录入会议标题、会议日期、会议类型和会议系列。",
        "status": "已上线",
    },
    {
        "section": "knowledge",
        "title": "外部资料归档",
        "href": "/external-sources",
        "description": "微信公众号文章、微信聊天记录等外部资料的归档与查看。",
        "status": "已上线",
    },
    {
        "section": "knowledge",
        "title": "知识库问答",
        "href": "/knowledge",
        "description": "基于账号权限过滤后的会议纪要与外部资料进行证据问答。",
        "status": "已上线",
    },
    {
        "section": "tools",
        "title": "待校对草稿",
        "href": "/drafts",
        "description": "初校草稿、整理后草稿和待确认终稿集中处理。",
        "status": "已上线",
    },
    {
        "section": "knowledge",
        "title": "历史纪要",
        "href": "/history",
        "description": "按日期、会议类型和标题检索人工确认后的正式纪要。",
        "status": "已上线",
    },
    {
        "section": "knowledge",
        "title": "标的证据查询",
        "href": "/target-query",
        "description": "按股票、公司或关键词查看纪要与外部资料中的证据来源。",
        "status": "已上线",
    },
    {
        "section": "tools",
        "title": "MCP 服务器",
        "href": "/mcp-servers",
        "description": "查看会议纪要工作流可调用的本地 MCP bridge server 和 tool。",
        "status": "已上线",
    },
    {
        "section": "dashboard",
        "title": "公众号文章看板",
        "href": "/dashboards",
        "description": "实时读取 WeRSS 公众号文章库，轮播展示并站内阅读全文。",
        "status": "已上线",
    },
    {
        "section": "dashboard",
        "title": "投研数据看板",
        "href": "/dashboard",
        "description": "查看当日高频股票、逻辑摘录、去重纪要和同发言人观点汇总。",
        "status": "已上线",
    },
    {
        "section": "dashboard",
        "title": "规划 Agent 状态",
        "href": "/agent-status",
        "description": "查看需求拆解、任务流转和当前工作流执行状态。",
        "status": "已上线",
    },
    {
        "section": "account",
        "title": "访问用户",
        "href": "/access-users",
        "description": "后台维护账号、角色、密码、API 令牌和停用状态；不开放自主注册。",
        "status": "已上线",
    },
    {
        "section": "data",
        "title": "访问审计",
        "href": "/access-audit",
        "description": "查看公网访问、登录、下载和资料查看审计记录。",
        "status": "已上线",
    },
    {
        "section": "data",
        "title": "映射审计",
        "href": "/mapping-audit",
        "description": "检查本地归档、Dify 知识库和历史索引之间的映射状态。",
        "status": "已上线",
    },
    {
        "section": "data",
        "title": "同步状态",
        "href": "/sync-status",
        "description": "查看 Google Drive 同步和归档上传结果。",
        "status": "已上线",
    },
    {
        "section": "data",
        "title": "健康检查",
        "href": "/health-report",
        "description": "检查服务、代理、知识库、公开访问和本地配置状态。",
        "status": "已上线",
    },
    {
        "section": "tools",
        "title": "OpenFinance 工具预留",
        "href": "/workflows",
        "description": "后续新增投研自动化、报告生成或资料处理工作流时接入这里。",
        "status": "预留",
    },
]


def portal_modules(section: str) -> list[dict[str, str]]:
    return [item for item in PORTAL_MODULES if item.get("section") == section]


def module_icon(title: str) -> str:
    mapping = {
        "OpenMinute-会议纪要": "M",
        "导入会议材料": "+",
        "外部资料归档": "A",
        "知识库问答": "AI",
        "待校对草稿": "R",
        "历史纪要": "H",
        "标的证据查询": "Q",
        "投研数据看板": "D",
        "规划 Agent 状态": "P",
        "访问用户": "U",
        "访问审计": "L",
        "映射审计": "K",
        "同步状态": "S",
        "健康检查": "OK",
    }
    return mapping.get(title, "·")


def portal_module_rows(section: str, *, compact: bool = False, identity: dict[str, Any] | None = None) -> str:
    rows = []
    items = portal_modules(section)
    if compact:
        items = [item for item in items if item.get("status") != "预留"]
        if section == "dashboard":
            items = items[:1]
        else:
            items = items[:3]
    for item in items:
        href = html.escape(item["href"])
        login_attrs = ""
        if section != "ops" and identity is not None and not identity.get("authenticated"):
            href = "#login"
            login_attrs = f' data-protected-link data-next="{html.escape(item["href"])}"'
        rows.append(
            f"""<a class="module-row reveal" href="{href}"{login_attrs}>
  <span class="module-icon" aria-hidden="true">{html.escape(module_icon(item['title']))}</span>
  <span class="module-main">
    <strong>{html.escape(item['title'])}</strong>
    <span>{html.escape(item['description'])}</span>
  </span>
</a>"""
        )
    return "".join(rows)


def portal_status_cards(identity: dict[str, Any]) -> str:
    if not identity.get("authenticated"):
        return """
<div class="metric-row">
  <div class="metric"><strong>未登录</strong><span>登录后查看草稿、归档和看板数据。</span></div>
  <div class="metric"><strong>个人研究</strong><span>本网站所有信息和功能仅供个人研究使用，不对外开放。</span></div>
</div>
"""
    drafts = scan_drafts()
    today = datetime.now().strftime("%Y-%m-%d")
    open_drafts = sum(1 for item in drafts if str(item.get("status") or "").lower() not in {"confirmed", "archived"})
    today_confirmed = sum(
        1
        for item in drafts
        if str(item.get("status") or "").lower() == "confirmed" and str(item.get("meeting_date") or "") == today
    )
    meeting_types = Counter(str(item.get("meeting_type") or "多人复盘会") for item in drafts if item.get("meeting_type"))
    sync = google_drive_sync_status()
    sync_label = "正常" if sync.get("ok") else "待检查"
    type_label = " / ".join(f"{key} {value}" for key, value in meeting_types.most_common(3)) or "暂无"
    return f"""
<div class="metric-row">
  <div class="metric"><strong>{open_drafts}</strong><span>待处理草稿</span></div>
  <div class="metric"><strong>{today_confirmed}</strong><span>今日正式纪要</span></div>
  <div class="metric"><strong>{html.escape(type_label)}</strong><span>主要会议类型</span></div>
  <div class="metric"><strong>{html.escape(sync_label)}</strong><span>归档同步</span></div>
</div>
"""


def workflow_stage_header(current_stage: str, description: str = "") -> str:
    desc = f'<div class="muted">{html.escape(description)}</div>' if description else ""
    return f"""
<div class="workflow-title">
  <h1 class="page-title">OpenMinute-会议纪要</h1>
  <div class="stage-banner">
    <span class="stage-kicker">当前环节</span>
    <strong>{html.escape(current_stage)}</strong>
    {desc}
  </div>
</div>
"""


def html_page(title: str, body: str, identity: dict[str, Any] | None = None) -> bytes:
    is_authenticated = bool(identity.get("authenticated")) if isinstance(identity, dict) else True
    def nav_href(path: str) -> str:
        return path if is_authenticated else "#login"

    def nav_attrs(path: str) -> str:
        return f' data-protected-link data-next="{html.escape(path)}"'

    nav_auth_link = (
        '<a href="/access-logout">退出</a>'
        if is_authenticated
        else '<a href="/access-login?next=/tools" data-login-open data-next="/tools">登录</a>'
    )
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1, viewport-fit=cover">
  <title>{html.escape(title)}</title>
  <style>
    :root {{ --bg:#f7f8f6; --surface:#ffffff; --surface-2:#eef3ef; --control-bg:#ffffff; --paper-bg:#ffffff; --paper-ink:#17201b; --code-bg:#ffffff; --table-bg:rgba(255,255,255,.75); --table-head-bg:rgba(238,243,239,.78); --table-hover-bg:rgba(255,255,255,.86); --ink:#17201b; --muted:#66736c; --line:#dde3de; --accent:#176b5b; --accent-2:#2e6f9e; --warn:#b56a1d; --danger:#b42318; --shadow:0 22px 60px rgba(23,32,27,.08); font-family: -apple-system, BlinkMacSystemFont, "PingFang SC", "Microsoft YaHei", "Noto Sans CJK SC", "Helvetica Neue", Arial, sans-serif; color: var(--ink); background: var(--bg); }}
    * {{ box-sizing: border-box; }}
    html {{ scroll-behavior: smooth; width: 100%; overflow-x: hidden; }}
    body {{ margin: 0; width: 100%; min-width: 0; min-height: 100vh; overflow-x: hidden; background: radial-gradient(circle at 14% 0%, var(--accent-soft), transparent 30rem), linear-gradient(180deg, var(--surface) 0%, var(--bg) 48%, var(--surface-2) 100%); }}
    body::selection {{ background: var(--accent-soft); }}
    header {{ position: sticky; top: 0; z-index: 20; background: rgba(247,248,246,.84); border-bottom: 1px solid rgba(221,227,222,.76); backdrop-filter: blur(18px); }}
    .topbar {{ max-width: 1220px; margin: 0 auto; padding: 14px 24px; display: flex; justify-content: space-between; gap: 18px; align-items: center; }}
    .brand {{ display: flex; flex-direction: column; gap: 2px; min-width: 190px; color: var(--ink); }}
    .brand strong {{ font-size: 16px; letter-spacing: 0; }}
    .brand span {{ color: var(--muted); font-size: 12px; }}
    main {{ max-width: 1220px; margin: 0 auto; padding: 28px 24px 56px; animation: page-in .28s ease-out both; }}
    a {{ color: var(--accent); text-decoration: none; transition: color .16s ease, background .16s ease, transform .16s ease, border-color .16s ease; }}
    a:hover {{ color: var(--accent-2); }}
    button, input, select, textarea {{ font: inherit; }}
    input, select, textarea {{ max-width: 100%; border: 1px solid var(--line); background: var(--control-bg); color: var(--ink); border-radius: 10px; outline: none; transition: border-color .16s ease, box-shadow .16s ease, background .16s ease; }}
    input, select {{ min-height: 42px; padding: 9px 12px; }}
    textarea {{ width: 100%; min-height: 56vh; padding: 16px; resize: vertical; font-size: 15px; line-height: 1.68; }}
    input:focus, select:focus, textarea:focus, .paper:focus, .direct-editor:focus {{ border-color: color-mix(in srgb, var(--accent) 62%, transparent); box-shadow: 0 0 0 4px var(--accent-soft); }}
    button, .button, .bar a {{ min-height: 40px; border: 1px solid color-mix(in srgb, var(--accent) 32%, transparent); background: var(--accent); color: #fff; border-radius: 999px; padding: 9px 14px; cursor: pointer; display: inline-flex; align-items: center; justify-content: center; gap: 8px; line-height: 1; }}
    button:hover, .button:hover, .bar a:hover {{ transform: translateY(-1px); background: var(--accent-2); color: #fff; }}
    button:active, .button:active, .bar a:active {{ transform: translateY(0) scale(.98); }}
    button.secondary, .bar a.secondary, .bar a[href="/history"], .bar a[href="/drafts"], .bar a[href="/dashboard"], .bar a[href="/apps/meeting-minutes"], .bar a[href="/external-sources"], .bar a[href="/target-query"] {{ background: rgba(255,255,255,.8); color: var(--accent); border-color: var(--line); }}
    button.tool {{ background: color-mix(in srgb, var(--surface) 84%, transparent); color: var(--ink); border-color: var(--line); padding: 7px 10px; min-width: 36px; border-radius: 10px; }}
    .bar {{ display: flex; gap: 10px; align-items: center; margin: 14px 0; flex-wrap: wrap; }}
    .meta, .item {{ background: rgba(255,255,255,.78); border: 1px solid rgba(221,227,222,.9); border-radius: 14px; padding: 16px; margin: 12px 0; box-shadow: 0 1px 0 rgba(23,32,27,.03); }}
    .item {{ transition: transform .16s ease, border-color .16s ease, background .16s ease; }}
    .item:hover {{ transform: translateY(-1px); border-color: color-mix(in srgb, var(--accent) 28%, transparent); }}
    .muted {{ color: var(--muted); font-size: 13px; line-height: 1.58; }}
    .filters {{ display: flex; gap: 10px; align-items: center; flex-wrap: wrap; padding: 12px 0; }}
    .filters input {{ min-width: 160px; }}
    .status {{ display: inline-flex; align-items: center; gap: 5px; padding: 3px 9px; border-radius: 999px; background: var(--accent-soft); color: var(--accent); font-size: 12px; border: 1px solid color-mix(in srgb, var(--accent) 18%, transparent); }}
    .error {{ color: var(--danger); }}
    .nav {{ display: flex; gap: 4px; align-items: center; flex-wrap: wrap; justify-content: flex-end; }}
    .nav a {{ color: #33433b; padding: 8px 10px; border-radius: 999px; font-size: 14px; }}
    .nav a:hover {{ background: var(--surface-2); color: var(--accent); }}
    .portal-hero {{ width: 100vw; min-height: min(680px, calc(100svh - 82px)); margin: -28px 0 30px calc(50% - 50vw); padding: 86px max(24px, calc((100vw - 1220px) / 2 + 24px)) 72px; display: grid; place-items: end start; background: linear-gradient(90deg, rgba(13,32,27,.88), rgba(13,32,27,.55) 45%, rgba(13,32,27,.16)), url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='1600' height='900' viewBox='0 0 1600 900'%3E%3Cdefs%3E%3ClinearGradient id='g' x1='0' y1='0' x2='1' y2='1'%3E%3Cstop stop-color='%2334473e'/%3E%3Cstop offset='.58' stop-color='%23899b8e'/%3E%3Cstop offset='1' stop-color='%23d8ded5'/%3E%3C/linearGradient%3E%3C/defs%3E%3Crect width='1600' height='900' fill='url(%23g)'/%3E%3Cg fill='none' stroke='%23ffffff' stroke-opacity='.18'%3E%3Cpath d='M120 690 C420 520 610 690 910 510 S1300 360 1510 480' stroke-width='4'/%3E%3Cpath d='M160 750 C460 590 690 730 980 590 S1340 520 1510 610' stroke-width='2'/%3E%3C/g%3E%3Cg fill='%23ffffff' fill-opacity='.18'%3E%3Ccircle cx='1160' cy='310' r='90'/%3E%3Ccircle cx='1320' cy='430' r='34'/%3E%3Crect x='1000' y='560' width='360' height='10' rx='5'/%3E%3Crect x='1040' y='600' width='260' height='10' rx='5'/%3E%3Crect x='1080' y='640' width='320' height='10' rx='5'/%3E%3C/g%3E%3C/svg%3E") center/cover; color: #fff; }}
    .portal-hero > div:first-child {{ max-width: 760px; }}
    .portal-hero h1 {{ margin: 0 0 16px; font-size: clamp(40px, 5.6vw, 72px); line-height: 1.04; letter-spacing: 0; word-break: keep-all; overflow-wrap: normal; }}
    .portal-hero p {{ margin: 0; max-width: 560px; color: rgba(255,255,255,.82); line-height: 1.78; font-size: 16px; }}
    .portal-hero .bar a, .portal-hero .bar button {{ background: #fff; color: #13352d; border-color: rgba(255,255,255,.7); }}
    .portal-hero .bar a:hover {{ background: rgba(255,255,255,.88); }}
    .notice {{ display: inline-flex; margin-top: 18px; padding: 9px 12px; border: 1px solid rgba(255,255,255,.22); border-radius: 999px; background: rgba(255,255,255,.12); color: rgba(255,255,255,.92); backdrop-filter: blur(10px); }}
    .metric-row {{ display: grid; grid-template-columns: repeat(4, minmax(0, 1fr)); gap: 10px; }}
    .metric {{ background: rgba(255,255,255,.82); border: 1px solid rgba(221,227,222,.9); border-radius: 14px; padding: 14px; min-height: 78px; box-shadow: 0 10px 34px rgba(23,32,27,.06); }}
    .metric strong {{ display: block; color: var(--ink); font-size: 21px; line-height: 1.25; overflow-wrap: anywhere; }}
    .metric span {{ display: block; color: var(--muted); font-size: 12px; margin-top: 6px; line-height: 1.4; }}
    .section-head {{ display: flex; justify-content: space-between; gap: 14px; align-items: flex-end; margin-top: 30px; margin-bottom: 10px; }}
    .section-head h2 {{ margin: 0; font-size: 21px; letter-spacing: 0; }}
    .section-head p {{ margin: 5px 0 0; color: var(--muted); font-size: 13px; }}
    .module-list {{ display: grid; gap: 2px; border-top: 1px solid var(--line); }}
    .module-row {{ position: relative; padding: 18px 0; display: grid; grid-template-columns: 44px minmax(0, 1fr) auto; gap: 16px; align-items: center; color: var(--ink); border-bottom: 1px solid var(--line); }}
    .module-row:hover {{ transform: translateX(4px); }}
    .module-row::after {{ content: "→"; color: var(--accent); opacity: 0; transform: translateX(-4px); transition: opacity .16s ease, transform .16s ease; }}
    .module-row:hover::after {{ opacity: 1; transform: translateX(0); }}
    .module-icon {{ width: 36px; height: 36px; border-radius: 12px; background: var(--surface-2); color: var(--accent); display: inline-flex; align-items: center; justify-content: center; font-weight: 700; font-size: 12px; }}
    .module-main {{ display: grid; gap: 4px; min-width: 0; }}
    .module-main strong {{ color: var(--ink); font-size: 16px; }}
    .module-main span {{ color: var(--muted); font-size: 13px; line-height: 1.45; }}
    .module-status {{ flex: 0 0 auto; border: 1px solid var(--line); border-radius: 999px; color: var(--muted); padding: 4px 9px; font-size: 12px; background: rgba(255,255,255,.7); }}
    .split {{ display: grid; grid-template-columns: 1fr 1fr; gap: 28px; align-items: start; }}
    .page-title {{ margin: 0 0 14px; font-size: clamp(28px, 4vw, 48px); letter-spacing: 0; line-height: 1.05; }}
    .stepper {{ display: grid; grid-template-columns: repeat(4, minmax(0, 1fr)); gap: 10px; margin: 18px 0; }}
    .step {{ padding: 12px 0; border-top: 2px solid var(--line); color: var(--muted); }}
    .step strong {{ display:block; color: var(--ink); margin-bottom: 4px; }}
    .step.active {{ border-color: var(--accent); color: var(--ink); }}
    .app-shell {{ display: grid; grid-template-columns: 220px minmax(0, 1fr); gap: 22px; align-items: start; }}
    .side-rail {{ position: sticky; top: 86px; display: grid; gap: 8px; }}
    .side-rail a {{ color: var(--muted); padding: 9px 0; border-bottom: 1px solid var(--line); }}
    .side-rail a:hover {{ color: var(--accent); transform: translateX(3px); }}
    .editor-shell {{ background: color-mix(in srgb, var(--surface-2) 72%, transparent); border: 1px solid var(--line); border-radius: 18px; padding: 18px; }}
    .toolbar {{ position: sticky; top: 66px; z-index: 2; background: color-mix(in srgb, var(--surface) 88%, transparent); border: 1px solid var(--line); border-radius: 14px; padding: 9px; margin-bottom: 14px; display: flex; gap: 8px; align-items: center; flex-wrap: wrap; backdrop-filter: blur(12px); }}
    .direct-edit-shell {{ background: color-mix(in srgb, var(--surface-2) 72%, transparent); border: 1px solid var(--line); border-radius: 18px; padding: 18px; }}
    .direct-editor {{ min-height: 68vh; font-family: "SFMono-Regular", Menlo, Consolas, "PingFang SC", monospace; line-height: 1.68; }}
    .review-hint {{ margin: 0 0 12px; padding: 10px 12px; border: 1px solid color-mix(in srgb, var(--accent) 20%, transparent); border-radius: 12px; background: var(--accent-soft); color: var(--ink); font-size: 14px; line-height: 1.58; }}
    .review-actions {{ position: sticky; bottom: 12px; z-index: 5; background: color-mix(in srgb, var(--surface) 88%, transparent); border: 1px solid var(--line); border-radius: 18px; padding: 12px; box-shadow: var(--shadow); backdrop-filter: blur(14px); }}
    .review-actions button:not(.secondary) {{ min-height: 46px; padding-inline: 18px; font-weight: 700; }}
    .tool.active {{ border-color: color-mix(in srgb, var(--accent) 48%, transparent); color: var(--accent); font-weight: 700; background: var(--accent-soft); }}
    .paper {{ background: var(--paper-bg); color: var(--paper-ink); max-width: 900px; min-height: 68vh; margin: 0 auto; padding: 48px 54px; border: 1px solid var(--line); border-radius: 8px; box-shadow: var(--shadow); line-height: 1.78; font-size: 15px; outline: none; }}
    .paper h1 {{ font-size: 27px; text-align: center; margin: 0 0 24px; }}
    .paper h2 {{ font-size: 20px; border-bottom: 1px solid var(--line); padding-bottom: 8px; margin-top: 30px; }}
    .paper h3 {{ font-size: 17px; margin-top: 24px; }}
    .paper p {{ margin: 8px 0; }}
    .paper ul, .paper ol {{ padding-left: 24px; }}
    table {{ width: 100%; border-collapse: collapse; margin: 14px 0; background: var(--table-bg); border-radius: 12px; overflow: hidden; }}
    th, td {{ border-bottom: 1px solid var(--line); padding: 10px 12px; vertical-align: top; text-align: left; }}
    th {{ background: var(--table-head-bg); color: var(--paper-ink); font-size: 13px; }}
    tr:hover td {{ background: var(--table-hover-bg); }}
    .source-panel {{ display: none; }}
    .editor-shell {{ display: none; }}
    .preview-mode .direct-edit-shell {{ display: none; }}
    .preview-mode .editor-shell {{ display: block; }}
    .source-mode .editor-shell {{ display: none; }}
    .source-mode .direct-edit-shell {{ display: none; }}
    .source-mode .source-panel {{ display: block; }}
    .markdown-field {{ display: none; }}
    details.source-details {{ max-width: 966px; margin: 14px auto 0; }}
    details.source-details summary {{ cursor: pointer; color: var(--accent); }}
    pre.markdown-source {{ white-space: pre-wrap; overflow-wrap: anywhere; background: var(--code-bg); color: var(--paper-ink); border: 1px solid var(--line); border-radius: 12px; padding: 14px; line-height: 1.55; }}
    :root {{ --bg:#f6f8fc; --surface:#ffffff; --surface-2:#edf3ff; --control-bg:#ffffff; --paper-bg:#ffffff; --paper-ink:#0b1020; --code-bg:#ffffff; --table-bg:rgba(255,255,255,.75); --table-head-bg:rgba(237,243,255,.88); --table-hover-bg:rgba(255,255,255,.86); --ink:#0b1020; --muted:#667085; --line:#d9e2f2; --accent:#2563eb; --accent-2:#06b6d4; --accent-soft:rgba(37,99,235,.12); --warn:#b76e00; --danger:#dc2626; --shadow:0 18px 48px rgba(15,23,42,.10); color-scheme: light; }}
    :root[data-theme="dark"] {{ --bg:#080b12; --surface:#111827; --surface-2:#172033; --control-bg:#0f172a; --paper-bg:#0f172a; --paper-ink:#e8eef8; --code-bg:#0b1220; --table-bg:#111827; --table-head-bg:#1b2537; --table-hover-bg:#172033; --ink:#eef4ff; --muted:#b2bdd0; --line:#344158; --accent:#7bb2ff; --accent-2:#52d7e2; --accent-soft:rgba(123,178,255,.18); --warn:#fbbf24; --danger:#fb7185; --shadow:0 22px 70px rgba(0,0,0,.36); color-scheme: dark; }}
    @media (prefers-color-scheme: dark) {{
      :root:not([data-theme="light"]) {{ --bg:#080b12; --surface:#111827; --surface-2:#172033; --control-bg:#0f172a; --paper-bg:#0f172a; --paper-ink:#e8eef8; --code-bg:#0b1220; --table-bg:#111827; --table-head-bg:#1b2537; --table-hover-bg:#172033; --ink:#eef4ff; --muted:#b2bdd0; --line:#344158; --accent:#7bb2ff; --accent-2:#52d7e2; --accent-soft:rgba(123,178,255,.18); --warn:#fbbf24; --danger:#fb7185; --shadow:0 22px 70px rgba(0,0,0,.36); color-scheme: dark; }}
    }}
    body {{ background:
      linear-gradient(180deg, color-mix(in srgb, var(--surface) 92%, transparent) 0%, color-mix(in srgb, var(--bg) 96%, transparent) 42%, var(--bg) 100%),
      radial-gradient(900px 520px at var(--mx, 70%) var(--my, 18%), var(--accent-soft), transparent 58%);
    }}
    header {{ background: color-mix(in srgb, var(--bg) 78%, transparent); border-bottom: 1px solid color-mix(in srgb, var(--line) 78%, transparent); }}
    .topbar {{ padding: 12px 24px; }}
    .brand {{ flex-direction: row; align-items: center; gap: 10px; min-width: 230px; }}
    .brand::before {{ content: ""; width: 26px; height: 26px; border-radius: 8px; background: linear-gradient(135deg, #0b1020, var(--accent), var(--accent-2)); box-shadow: inset 0 0 0 1px rgba(255,255,255,.24), 0 0 26px var(--accent-soft); }}
    .brand span {{ padding-left: 10px; border-left: 1px solid var(--line); }}
    .nav {{ gap: 2px; }}
    .nav a {{ border-radius: 8px; color: var(--muted); padding: 8px 11px; }}
    .nav a:hover {{ background: var(--surface-2); color: var(--ink); transform: translateY(-1px); }}
    button, .button, .bar a {{ border-radius: 8px; min-height: 42px; padding: 10px 14px; box-shadow: 0 10px 24px color-mix(in srgb, var(--accent) 18%, transparent); }}
    button.secondary, .bar a.secondary, .bar a[href="/history"], .bar a[href="/drafts"], .bar a[href="/dashboard"], .bar a[href="/apps/meeting-minutes"], .bar a[href="/external-sources"], .bar a[href="/target-query"] {{ box-shadow: none; background: color-mix(in srgb, var(--surface) 72%, transparent); color: var(--ink); }}
    .portal-hero {{ position: relative; isolation: isolate; overflow: hidden; grid-template-columns: minmax(0, 560px) minmax(360px, 1fr); gap: 56px; align-items: end; place-items: initial; min-height: min(620px, calc(100svh - 70px)); padding-top: 74px; background:
      linear-gradient(115deg, rgba(5,10,22,.98) 0%, rgba(12,26,54,.94) 48%, rgba(15,23,42,.78) 100%),
      linear-gradient(135deg, #080b12, #123a63 48%, #0f766e);
    }}
    .portal-hero::before {{ content: ""; position: absolute; inset: 0; z-index: -2; background-image:
      linear-gradient(rgba(255,255,255,.055) 1px, transparent 1px),
      linear-gradient(90deg, rgba(255,255,255,.055) 1px, transparent 1px);
      background-size: 44px 44px; mask-image: linear-gradient(90deg, rgba(0,0,0,.85), rgba(0,0,0,.2)); }}
    .portal-hero::after {{ content: ""; position: absolute; right: -14vw; top: 10%; width: 62vw; height: 62vw; z-index: -1; background: conic-gradient(from 190deg, rgba(6,182,212,.46), rgba(37,99,235,.36), rgba(255,255,255,.08), rgba(20,184,166,.34)); filter: blur(42px); opacity: .48; transform: rotate(-10deg); }}
    .hero-copy {{ padding-bottom: 4px; animation: slide-up .48s ease both; }}
    .eyebrow {{ margin: 0 0 18px; color: rgba(255,255,255,.62); font-size: 13px; letter-spacing: .08em; text-transform: uppercase; }}
    .portal-hero h1 {{ font-size: clamp(46px, 7.2vw, 92px); line-height: .98; max-width: 680px; margin-bottom: 18px; }}
    .portal-hero p {{ font-size: 17px; line-height: 1.72; color: rgba(255,255,255,.74); }}
    .portal-hero .bar {{ margin-top: 28px; }}
    .portal-hero .bar a {{ background: #f4fffb; color: #07110f; border: 0; box-shadow: 0 16px 36px rgba(0,0,0,.2); }}
    .portal-hero .bar a.secondary {{ background: rgba(255,255,255,.10); color: #fff; border: 1px solid rgba(255,255,255,.20); box-shadow: none; }}
    .notice {{ display: block; width: fit-content; margin-top: 18px; border-radius: 8px; background: rgba(255,255,255,.08); border-color: rgba(255,255,255,.16); color: rgba(255,255,255,.78); }}
    .hero-panel {{ align-self: center; justify-self: end; width: min(520px, 100%); padding: 18px; border: 1px solid rgba(255,255,255,.16); border-radius: 18px; background: rgba(246,250,248,.10); color: #fff; backdrop-filter: blur(16px); box-shadow: 0 30px 90px rgba(0,0,0,.28); animation: float-in .62s ease both .08s; }}
    .hero-panel-top {{ display: flex; justify-content: space-between; align-items: center; padding-bottom: 14px; border-bottom: 1px solid rgba(255,255,255,.14); color: rgba(255,255,255,.72); font-size: 13px; }}
    .flow-line {{ display: grid; grid-template-columns: repeat(5, 1fr); gap: 8px; margin: 18px 0; }}
    .flow-line span {{ height: 7px; border-radius: 99px; background: rgba(255,255,255,.18); overflow: hidden; }}
    .flow-line span:nth-child(-n+3) {{ background: linear-gradient(90deg, #36d1dc, #8fc6ff); }}
    .hero-list {{ display: grid; gap: 10px; }}
    .hero-task {{ display: grid; grid-template-columns: 24px minmax(0,1fr) auto; gap: 10px; align-items: center; padding: 12px 0; border-bottom: 1px solid rgba(255,255,255,.10); }}
    .hero-task:last-child {{ border-bottom: 0; }}
    .hero-task b {{ font-size: 14px; color: #fff; }}
    .hero-task small {{ color: rgba(255,255,255,.58); }}
    .hero-dot {{ width: 10px; height: 10px; border-radius: 99px; background: #36d1dc; box-shadow: 0 0 0 6px rgba(54,209,220,.14); }}
    .quick-strip {{ display: grid; grid-template-columns: 1.1fr .9fr; gap: 24px; margin: 28px 0 8px; }}
    .work-surface {{ padding: 26px 0; border-top: 1px solid var(--line); border-bottom: 1px solid var(--line); }}
    .surface-title {{ display: flex; justify-content: space-between; align-items: end; gap: 18px; margin-bottom: 12px; }}
    .surface-title h2 {{ margin: 0; font-size: 24px; }}
    .surface-title p {{ margin: 4px 0 0; color: var(--muted); }}
    .module-list {{ border-top: 1px solid rgba(217,224,221,.9); }}
    .module-row {{ padding: 19px 0; grid-template-columns: 36px minmax(0, 1fr) 22px; }}
    .module-row::after {{ opacity: .38; transform: none; }}
    .module-row:hover {{ transform: translateX(6px); }}
    .module-icon {{ border-radius: 8px; background: var(--surface-2); color: var(--accent); }}
    .module-status {{ border-radius: 7px; background: var(--surface); }}
    .meta, .item {{ border-radius: 10px; background: color-mix(in srgb, var(--surface) 74%, transparent); box-shadow: none; }}
    .metric-row {{ gap: 0; border-top: 1px solid var(--line); border-left: 1px solid var(--line); }}
    .metric {{ border-radius: 0; border: 0; border-right: 1px solid var(--line); border-bottom: 1px solid var(--line); background: color-mix(in srgb, var(--surface) 76%, transparent); box-shadow: none; }}
    .metric strong {{ font-size: 22px; }}
    .page-title {{ font-size: clamp(34px, 5vw, 62px); max-width: 780px; margin-bottom: 22px; }}
    .app-shell {{ grid-template-columns: 260px minmax(0, 1fr); gap: 36px; }}
    .side-rail {{ top: 78px; padding-top: 5px; }}
    .side-rail a {{ display: flex; justify-content: space-between; align-items: center; border-bottom-color: rgba(217,224,221,.86); }}
    .side-rail a::after {{ content: "↗"; opacity: .34; }}
    .stepper {{ gap: 0; border-top: 1px solid var(--line); border-left: 1px solid var(--line); }}
    .step {{ border-top: 0; border-right: 1px solid var(--line); border-bottom: 1px solid var(--line); padding: 14px; background: color-mix(in srgb, var(--surface) 68%, transparent); }}
    .step.active {{ background: var(--surface-2); box-shadow: inset 0 2px 0 var(--accent); }}
    .import-layout {{ display: grid; grid-template-columns: minmax(0, 1fr) 310px; gap: 24px; align-items: start; }}
    .process-panel {{ position: sticky; top: 84px; padding: 4px 0; }}
    .process-panel ol {{ list-style: none; margin: 0; padding: 0; border-top: 1px solid var(--line); }}
    .process-panel li {{ display: grid; grid-template-columns: 28px 1fr; gap: 10px; padding: 14px 0; border-bottom: 1px solid var(--line); color: var(--muted); }}
    .process-panel b {{ color: var(--ink); }}
    .process-index {{ width: 24px; height: 24px; border-radius: 7px; display: inline-flex; align-items: center; justify-content: center; background: var(--surface-2); color: var(--accent); font-size: 12px; }}
    .input-grid {{ display: grid; grid-template-columns: minmax(190px, 1.3fr) minmax(142px, .74fr) minmax(150px, .95fr) minmax(150px, .95fr); gap: 10px; margin-bottom: 12px; }}
    .data-window {{ min-height: 220px; width: 100%; resize: vertical; font-family: "SFMono-Regular", Menlo, Consolas, "PingFang SC", monospace; }}
    .import-layout textarea {{ min-height: 58vh; border-radius: 10px; background: var(--control-bg); color: var(--ink); }}
    .import-layout .field textarea {{ min-height: 112px; }}
    .field {{ display: grid; gap: 7px; }}
    .field > span {{ color: var(--muted); font-size: 13px; font-weight: 700; }}
    .field input, .field select {{ width: 100%; }}
    .field input[hidden] {{ display: none !important; }}
    .upload-zone {{ margin: 14px 0; padding: 20px; border: 1px dashed color-mix(in srgb, var(--accent) 42%, transparent); border-radius: 12px; background: linear-gradient(180deg, var(--surface-2), color-mix(in srgb, var(--surface) 72%, transparent)); }}
    .upload-zone.is-dragover {{ border-color: var(--accent); background: color-mix(in srgb, var(--accent) 10%, var(--surface)); }}
    .upload-zone input[type="file"] {{ width: 100%; min-height: 48px; padding: 10px; background: var(--surface); border-style: solid; }}
    .upload-zone strong {{ display: block; margin-bottom: 4px; }}
    .upload-zone .muted {{ margin: 0; }}
    .file-list {{ display: grid; gap: 8px; margin-top: 12px; }}
    .file-item {{ display: flex; align-items: center; justify-content: space-between; gap: 12px; padding: 9px 11px; border: 1px solid var(--line); border-radius: 8px; background: var(--surface); color: var(--ink); }}
    .file-item span:first-child {{ min-width: 0; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }}
    .file-meta {{ flex: 0 0 auto; color: var(--muted); font-size: 12px; }}
    .file-remove {{ flex: 0 0 auto; min-height: 30px; padding: 6px 10px; border-radius: 7px; background: transparent; color: var(--danger); border-color: color-mix(in srgb, var(--danger) 38%, transparent); }}
    .file-remove:hover {{ background: color-mix(in srgb, var(--danger) 10%, transparent); color: var(--danger); }}
    .login-overlay {{ position: fixed; inset: 0; z-index: 60; display: none; align-items: center; justify-content: center; padding: 24px; background: rgba(9,16,14,.48); backdrop-filter: blur(10px); }}
    .login-overlay.open {{ display: flex; animation: fade-in .16s ease both; }}
    .login-dialog {{ width: min(420px, 100%); background: var(--surface); border: 1px solid var(--line); border-radius: 16px; padding: 24px; box-shadow: 0 28px 90px rgba(0,0,0,.28); }}
    .login-dialog h2 {{ margin: 0 0 8px; font-size: 24px; }}
    .login-dialog p {{ margin: 0 0 18px; color: var(--muted); line-height: 1.58; }}
    .login-dialog input {{ width: 100%; margin-bottom: 12px; }}
    .login-close {{ position: absolute; top: 18px; right: 18px; width: 38px; min-width: 38px; padding: 0; border-radius: 999px; background: color-mix(in srgb, var(--surface) 92%, transparent); color: var(--ink); box-shadow: none; }}
    input::placeholder, textarea::placeholder {{ color: color-mix(in srgb, var(--muted) 82%, transparent); }}
    .theme-switcher {{ display: inline-flex; gap: 2px; padding: 2px; border: 1px solid var(--line); border-radius: 10px; background: color-mix(in srgb, var(--surface) 68%, transparent); }}
    .theme-switcher button {{ min-height: 30px; padding: 5px 8px; border: 0; border-radius: 7px; background: transparent; color: var(--muted); box-shadow: none; font-size: 12px; }}
    .theme-switcher button.active {{ background: var(--surface-2); color: var(--ink); }}
    .workflow-title {{ display: grid; gap: 12px; margin-bottom: 18px; }}
    .stage-banner {{ display: grid; gap: 4px; max-width: 760px; padding: 14px 16px; border-left: 3px solid var(--accent); background: color-mix(in srgb, var(--surface-2) 72%, transparent); border-radius: 10px; }}
    .stage-kicker {{ color: var(--accent); font-size: 12px; font-weight: 700; letter-spacing: .08em; }}
    .processing-overlay {{ position: fixed; inset: 0; z-index: 70; display: none; align-items: center; justify-content: center; padding: 24px; background: rgba(4,8,18,.58); backdrop-filter: blur(12px); }}
    .processing-overlay.open {{ display: flex; animation: fade-in .16s ease both; }}
    .processing-dialog {{ width: min(440px, 100%); padding: 26px; border: 1px solid color-mix(in srgb, var(--accent) 28%, var(--line)); border-radius: 18px; background: var(--surface); box-shadow: 0 30px 100px rgba(0,0,0,.34); }}
    .processing-spinner {{ width: 42px; height: 42px; border-radius: 50%; border: 3px solid var(--line); border-top-color: var(--accent-2); animation: spin .8s linear infinite; margin-bottom: 16px; }}
    .processing-dialog h2 {{ margin: 0 0 8px; font-size: 22px; }}
    .processing-dialog p {{ margin: 0; color: var(--muted); line-height: 1.58; }}
    .processing-progress {{ display: none; gap: 10px; margin-top: 18px; }}
    .processing-progress.show {{ display: grid; }}
    .progress-head {{ display: flex; align-items: center; justify-content: space-between; gap: 12px; color: var(--muted); font-size: 13px; }}
    .progress-track {{ height: 9px; overflow: hidden; border-radius: 999px; background: var(--surface-2); border: 1px solid var(--line); }}
    .progress-fill {{ height: 100%; width: 0%; border-radius: inherit; background: linear-gradient(90deg, var(--accent), var(--accent-2)); transition: width .42s ease; }}
    .progress-steps {{ display: grid; gap: 7px; margin-top: 2px; color: var(--muted); font-size: 13px; }}
    .progress-step {{ display: flex; align-items: center; gap: 8px; }}
    .progress-step::before {{ content: ""; width: 8px; height: 8px; border-radius: 50%; background: var(--line); }}
    .progress-step.active {{ color: var(--ink); font-weight: 700; }}
    .progress-step.active::before {{ background: var(--accent); box-shadow: 0 0 0 4px color-mix(in srgb, var(--accent) 16%, transparent); }}
    .progress-step.done::before {{ background: var(--accent-2); }}
    .form-error {{ display: none; margin: 0 0 12px; color: var(--danger); }}
    .form-error.show {{ display: block; }}
    @keyframes fade-in {{ from {{ opacity: 0; }} to {{ opacity: 1; }} }}
    @keyframes spin {{ to {{ transform: rotate(360deg); }} }}
    @keyframes slide-up {{ from {{ opacity: 0; transform: translateY(18px); }} to {{ opacity: 1; transform: translateY(0); }} }}
    @keyframes float-in {{ from {{ opacity: 0; transform: translateY(22px) scale(.98); }} to {{ opacity: 1; transform: translateY(0) scale(1); }} }}
    .reveal {{ opacity: 0; transform: translateY(12px); animation: reveal .42s ease forwards; }}
    .reveal:nth-child(2) {{ animation-delay: .04s; }}
    .reveal:nth-child(3) {{ animation-delay: .08s; }}
    .reveal:nth-child(4) {{ animation-delay: .12s; }}
    @keyframes page-in {{ from {{ opacity: 0; transform: translateY(8px); }} to {{ opacity: 1; transform: translateY(0); }} }}
    @keyframes reveal {{ to {{ opacity: 1; transform: translateY(0); }} }}
    @media (max-width: 760px) {{
      header {{ position: sticky; top: 0; }}
      .topbar {{ align-items: stretch; flex-direction: column; gap: 10px; padding: 10px 14px 9px; }}
      .brand {{ min-width: 0; width: 100%; align-items: center; gap: 8px; flex-wrap: wrap; overflow: visible; }}
      .brand span {{ min-width: 0; overflow: visible; text-overflow: clip; white-space: normal; line-height: 1.3; }}
      .nav {{ justify-content: flex-start; flex-wrap: nowrap; gap: 6px; overflow-x: auto; padding: 2px 0 6px; -webkit-overflow-scrolling: touch; scrollbar-width: none; }}
      .nav::-webkit-scrollbar {{ display: none; }}
      .nav a {{ flex: 0 0 auto; min-height: 36px; display: inline-flex; align-items: center; padding: 8px 11px; }}
      main {{ padding: 18px 14px 44px; }}
      .portal-hero, .split {{ grid-template-columns: minmax(0, 1fr); }}
      .quick-strip, .import-layout, .input-grid {{ grid-template-columns: minmax(0, 1fr); gap: 14px; }}
      .portal-hero {{ width: 100%; max-width: 100%; min-height: auto; margin: -18px 0 22px; padding: 44px 16px 34px; border-radius: 0 0 18px 18px; }}
      .portal-hero > * {{ min-width: 0; max-width: 100%; }}
      .hero-copy, .portal-hero > div:first-child {{ width: 100%; max-width: 360px; }}
      .portal-hero h1 {{ font-size: 40px; line-height: 1.04; white-space: normal; word-break: normal; overflow-wrap: anywhere; }}
      .portal-hero p {{ font-size: 15px; line-height: 1.62; }}
      .notice {{ width: auto; max-width: 100%; white-space: normal; overflow-wrap: anywhere; }}
      .portal-hero .bar {{ align-items: stretch; }}
      .portal-hero .bar a, .portal-hero .bar button {{ flex: 1 1 100%; width: 100%; }}
      .hero-panel {{ display: none; }}
      .hero-task {{ grid-template-columns: 20px minmax(0, 1fr); }}
      .hero-task small {{ grid-column: 2; justify-self: start; }}
      .metric-row {{ grid-template-columns: repeat(2, minmax(0, 1fr)); }}
      .metric {{ min-width: 0; padding: 12px; }}
      .section-head, .surface-title {{ align-items: stretch; flex-direction: column; gap: 8px; }}
      .module-row {{ grid-template-columns: 34px minmax(0, 1fr); gap: 10px; padding: 15px 0; }}
      .module-row:hover {{ transform: none; }}
      .module-row::after {{ display: none; }}
      .module-status {{ grid-column: 2; width: fit-content; max-width: 100%; }}
      .app-shell, .stepper {{ grid-template-columns: minmax(0, 1fr); }}
      .side-rail {{ position: static; }}
      .side-rail a {{ padding: 10px 0; }}
      .bar, .filters {{ align-items: stretch; gap: 8px; }}
      .filters input, .filters select, .filters button, .filters .button {{ width: 100%; min-width: 0; }}
      .bar button, .bar .button, .bar a {{ min-width: 0; min-height: 44px; white-space: normal; text-align: center; }}
      .meta, .item {{ padding: 14px; border-radius: 10px; }}
      .page-title {{ font-size: 34px; line-height: 1.08; margin-bottom: 16px; overflow-wrap: anywhere; }}
      .paper {{ max-width: 100%; min-height: 58vh; padding: 24px 18px; box-shadow: none; overflow-wrap: anywhere; }}
      .direct-edit-shell, .editor-shell {{ padding: 12px; border-radius: 12px; }}
      .direct-editor, textarea {{ min-height: 48vh; font-size: 15px; }}
      .toolbar {{ position: static; overflow-x: auto; flex-wrap: nowrap; -webkit-overflow-scrolling: touch; }}
      .toolbar .tool {{ flex: 0 0 auto; }}
      table {{ display: block; width: 100%; max-width: 100%; overflow-x: auto; white-space: nowrap; -webkit-overflow-scrolling: touch; border-radius: 10px; }}
      th, td {{ padding: 9px 10px; }}
      pre.markdown-source {{ max-width: 100%; overflow-x: auto; }}
      .login-overlay, .processing-overlay {{ align-items: flex-end; padding: 14px; }}
      .login-dialog, .processing-dialog {{ width: 100%; border-radius: 16px; padding: 20px; }}
      .process-panel {{ position: static; }}
      .upload-zone {{ padding: 14px; }}
      .upload-zone input[type="file"] {{ min-width: 0; }}
    }}
    @media (max-width: 430px) {{
      .topbar {{ padding-inline: 12px; }}
      main {{ padding-inline: 12px; }}
      .brand strong {{ font-size: 15px; }}
      .brand span {{ flex: 0 0 100%; padding-left: 34px; border-left: 0; font-size: 11px; }}
      .portal-hero {{ padding: 38px 14px 30px; }}
      .portal-hero h1 {{ font-size: 34px; }}
      .metric-row {{ grid-template-columns: minmax(0, 1fr); }}
      .bar button, .bar .button, .bar a {{ width: 100%; }}
      input, select {{ width: 100%; min-width: 0; }}
    }}
    @media (prefers-reduced-motion: reduce) {{
      *, *::before, *::after {{ animation-duration: .001ms !important; transition-duration: .001ms !important; scroll-behavior: auto !important; }}
    }}
  </style>
</head>
<body data-authenticated="{str(is_authenticated).lower()}">
  <header>
    <div class="topbar">
      <a class="brand" href="/"><strong>D91 AI投研平台</strong><span>OpenFinance · KumaAI</span></a>
      <nav class="nav">
        <a href="/">首页</a>
        <a href="{nav_href('/dashboard')}"{nav_attrs('/dashboard')}>看板</a>
        <a href="{nav_href('/knowledge')}"{nav_attrs('/knowledge')}>知识库</a>
        <a href="{nav_href('/tools')}"{nav_attrs('/tools')}>工具</a>
        <a href="{nav_href('/data')}"{nav_attrs('/data')}>数据</a>
        <a href="{nav_href('/account')}"{nav_attrs('/account')}>账号</a>
        <span class="theme-switcher" aria-label="主题">
          <button type="button" data-theme-choice="system">跟随</button>
          <button type="button" data-theme-choice="light">浅色</button>
          <button type="button" data-theme-choice="dark">深色</button>
        </span>
        {nav_auth_link}
      </nav>
    </div>
  </header>
  <main>{body}</main>
  <div class="login-overlay" id="login-overlay" aria-hidden="true">
    <button class="login-close" type="button" data-login-close aria-label="关闭登录">×</button>
    <form class="login-dialog" method="post" action="/access-login">
      <h2>登录后继续</h2>
      <p>请输入用户名和密码。登录成功后仍停留首页，再由你选择具体工作流。</p>
      <p class="form-error" id="login-error"></p>
      <input type="hidden" name="next" id="login-next" value="/tools">
      <input name="user_id" placeholder="用户名" autocomplete="username" required>
      <input name="password" type="password" placeholder="密码" autocomplete="current-password" required>
      <button type="submit" data-keep-label="1">登录</button>
    </form>
  </div>
  <div class="processing-overlay" id="processing-overlay" aria-hidden="true">
    <div class="processing-dialog" role="status" aria-live="polite">
      <div class="processing-spinner" aria-hidden="true"></div>
      <h2 id="processing-title">正在处理</h2>
      <p id="processing-message">系统正在执行，请不要重复提交。</p>
      <div class="processing-progress" id="processing-progress" aria-hidden="true">
        <div class="progress-head"><span id="processing-stage">准备中</span><strong id="processing-percent">0%</strong></div>
        <div class="progress-track" aria-hidden="true"><div class="progress-fill" id="processing-fill"></div></div>
        <div class="progress-steps" id="processing-steps"></div>
      </div>
    </div>
  </div>
  <script>
    const themeButtons = Array.from(document.querySelectorAll("[data-theme-choice]"));
    const savedTheme = localStorage.getItem("kuma-theme") || "system";
    const applyTheme = (mode) => {{
      const theme = mode || "system";
      if (theme === "system") {{
        document.documentElement.removeAttribute("data-theme");
      }} else {{
        document.documentElement.dataset.theme = theme;
      }}
      themeButtons.forEach((button) => button.classList.toggle("active", button.dataset.themeChoice === theme));
    }};
    applyTheme(savedTheme);
    themeButtons.forEach((button) => {{
      button.addEventListener("click", () => {{
        localStorage.setItem("kuma-theme", button.dataset.themeChoice || "system");
        applyTheme(button.dataset.themeChoice || "system");
      }});
    }});
    const isAuthenticated = document.body.dataset.authenticated === "true";
    const loginOverlay = document.getElementById("login-overlay");
    const loginNext = document.getElementById("login-next");
    const loginError = document.getElementById("login-error");
    const processingOverlay = document.getElementById("processing-overlay");
    const processingTitle = document.getElementById("processing-title");
    const processingMessage = document.getElementById("processing-message");
    const processingProgress = document.getElementById("processing-progress");
    const processingStage = document.getElementById("processing-stage");
    const processingPercent = document.getElementById("processing-percent");
    const processingFill = document.getElementById("processing-fill");
    const processingSteps = document.getElementById("processing-steps");
    let processingTimer = null;
    const openLogin = (nextPath) => {{
      if (!loginOverlay) return;
      loginNext.value = nextPath || "/tools";
      if (loginError) {{
        loginError.textContent = "";
        loginError.classList.remove("show");
      }}
      loginOverlay.classList.add("open");
      loginOverlay.setAttribute("aria-hidden", "false");
      const input = loginOverlay.querySelector('input[name="user_id"]');
      if (input) input.focus();
    }};
    const setProcessingProgress = (percent, stageText, stages = [], activeIndex = 0) => {{
      if (!processingProgress) return;
      const safePercent = Math.max(0, Math.min(98, Math.round(percent || 0)));
      processingProgress.classList.add("show");
      processingProgress.setAttribute("aria-hidden", "false");
      if (processingStage) processingStage.textContent = stageText || "处理中";
      if (processingPercent) processingPercent.textContent = `${{safePercent}}%`;
      if (processingFill) processingFill.style.width = `${{safePercent}}%`;
      if (processingSteps) {{
        processingSteps.innerHTML = "";
        stages.forEach((stage, index) => {{
          const row = document.createElement("div");
          row.className = "progress-step" + (index < activeIndex ? " done" : index === activeIndex ? " active" : "");
          row.textContent = stage.label || stage;
          processingSteps.appendChild(row);
        }});
      }}
    }};
    const hideProcessingProgress = () => {{
      if (processingTimer) {{
        window.clearInterval(processingTimer);
        processingTimer = null;
      }}
      processingProgress?.classList.remove("show");
      processingProgress?.setAttribute("aria-hidden", "true");
      if (processingFill) processingFill.style.width = "0%";
      if (processingPercent) processingPercent.textContent = "0%";
      if (processingSteps) processingSteps.innerHTML = "";
    }};
    const showProcessing = (title, message) => {{
      if (!processingOverlay) return;
      hideProcessingProgress();
      processingTitle.textContent = title || "正在处理";
      processingMessage.textContent = message || "系统正在执行，请不要重复提交。";
      processingOverlay.classList.add("open");
      processingOverlay.setAttribute("aria-hidden", "false");
    }};
    const importProgressStages = (files) => {{
      const names = files.map((file) => file.name || "");
      const hasText = names.some((name) => /\\.(txt|md|markdown|csv|tsv)$/i.test(name));
      const hasDocx = names.some((name) => /\\.docx$/i.test(name));
      const hasAudio = names.some((name) => /\\.(mp3|m4a|wav|aac|flac|ogg|wma|mp4)$/i.test(name));
      const stages = [
        {{ label: "上传会议材料", target: 18, duration: 900 }},
      ];
      if (hasText) stages.push({{ label: "读取文本材料", target: 34, duration: 1100 }});
      if (hasDocx) stages.push({{ label: "解析 Word 文档", target: hasAudio ? 46 : 58, duration: 1600 }});
      if (hasAudio) stages.push({{ label: "音频转录中，时间取决于录音长度", target: 78, duration: 9000 }});
      stages.push({{ label: "合并材料并生成输入校对稿", target: 92, duration: 2600 }});
      stages.push({{ label: "准备打开人工校对页", target: 97, duration: 6000 }});
      return stages;
    }};
    const startImportProgress = (form) => {{
      const input = form.querySelector("#meeting-files-input");
      const files = Array.from(input?.files || []);
      const fileText = files.length ? `${{files.length}} 个文件` : "已选择的文件";
      const stages = importProgressStages(files);
      let stageIndex = 0;
      let startedAt = Date.now();
      let previousTarget = 0;
      showProcessing("正在导入会议材料", `正在处理 ${{fileText}}，页面完成后会自动进入人工校对。`);
      const tick = () => {{
        const stage = stages[stageIndex] || stages[stages.length - 1];
        const elapsed = Date.now() - startedAt;
        const ratio = Math.min(1, elapsed / Math.max(stage.duration || 1200, 300));
        const percent = previousTarget + (stage.target - previousTarget) * ratio;
        setProcessingProgress(percent, stage.label, stages, stageIndex);
        processingMessage.textContent = stage.label === "音频转录中，时间取决于录音长度"
          ? `正在转录音频。长录音需要更久，请保持本页打开；完成后会自动跳转。`
          : `正在处理 ${{fileText}}：${{stage.label}}。`;
        if (ratio >= 1 && stageIndex < stages.length - 1) {{
          previousTarget = stage.target;
          stageIndex += 1;
          startedAt = Date.now();
        }}
      }};
      tick();
      processingTimer = window.setInterval(tick, 650);
    }};
    document.querySelectorAll("[data-protected-link], [data-login-open]").forEach((link) => {{
      link.addEventListener("click", (event) => {{
        if (isAuthenticated && !link.hasAttribute("data-login-open")) return;
        event.preventDefault();
        openLogin(link.dataset.next || link.getAttribute("href") || "/tools");
      }});
    }});
    document.querySelectorAll("[data-login-close]").forEach((button) => {{
      button.addEventListener("click", () => {{
        loginOverlay.classList.remove("open");
        loginOverlay.setAttribute("aria-hidden", "true");
      }});
    }});
    loginOverlay?.addEventListener("click", (event) => {{
      if (event.target === loginOverlay) {{
        loginOverlay.classList.remove("open");
        loginOverlay.setAttribute("aria-hidden", "true");
      }}
    }});
    loginOverlay?.querySelector("form")?.addEventListener("submit", async (event) => {{
      event.preventDefault();
      const form = event.currentTarget;
      const button = form.querySelector('button[type="submit"]');
      const originalText = button?.textContent || "登录";
      if (button) {{
        button.disabled = true;
        button.textContent = "登录中...";
      }}
      if (loginError) {{
        loginError.textContent = "";
        loginError.classList.remove("show");
      }}
      try {{
        const response = await fetch(form.action, {{
          method: "POST",
          body: new FormData(form),
          headers: {{ "Accept": "application/json" }},
          credentials: "same-origin",
        }});
        const payload = await response.json().catch(() => ({{ ok: false, error: "登录失败，请稍后重试。" }}));
        if (!response.ok || !payload.ok) {{
          if (loginError) {{
            loginError.textContent = payload.error || "登录失败，请检查用户名和密码。";
            loginError.classList.add("show");
          }}
          return;
        }}
        window.location.href = payload.redirect || "/";
      }} catch (error) {{
        if (loginError) {{
          loginError.textContent = "登录失败，请检查网络后重试。";
          loginError.classList.add("show");
        }}
      }} finally {{
        if (button) {{
          button.disabled = false;
          button.textContent = originalText;
        }}
      }}
    }});
    document.querySelectorAll("form").forEach((form) => {{
      if (form.closest(".login-overlay")) return;
      form.addEventListener("submit", (event) => {{
        const button = event.submitter || form.querySelector('button[type="submit"]');
        if (button && !button.dataset.keepLabel) {{
          button.dataset.originalText = button.textContent;
          button.textContent = "处理中...";
          button.setAttribute("aria-busy", "true");
        }}
        const hasProcessingCopy = button?.dataset.processingTitle || button?.dataset.processingMessage || form.dataset.processingTitle || form.dataset.processingMessage;
        if (form.dataset.processing !== "off" && hasProcessingCopy) {{
          const actionPath = new URL(form.action || window.location.href, window.location.href).pathname;
          if (actionPath === "/apps/meeting-minutes/import") {{
            startImportProgress(form);
          }} else {{
            showProcessing(button?.dataset.processingTitle || form.dataset.processingTitle || "正在处理", button?.dataset.processingMessage || form.dataset.processingMessage || "系统正在执行，请不要重复提交。");
          }}
          Array.from(form.elements).forEach((element) => {{
            if (element.tagName === "BUTTON" || element.tagName === "INPUT" || element.tagName === "SELECT" || element.tagName === "TEXTAREA") {{
              if (element.type !== "hidden") element.setAttribute("aria-disabled", "true");
            }}
          }});
        }}
      }});
    }});
    const filesInput = document.getElementById("meeting-files-input");
    const titleInput = document.getElementById("meeting-title-input");
    const typeSelect = document.getElementById("meeting-type-select");
    const seriesSelect = document.getElementById("meeting-series-select");
    const dateInput = document.getElementById("meeting-date-input");
    const smartHint = document.getElementById("upload-smart-hint");
    const uploadZone = document.querySelector(".upload-zone");
    const fileList = document.getElementById("meeting-files-list");
    const selectedFiles = [];
    const setupCustomChoice = (select) => {{
      if (!select) return;
      const input = document.getElementById(select.dataset.customInput || "");
      const customValue = select.dataset.customValue || "__custom__";
      const update = (focusInput = false) => {{
        const isCustom = select.value === customValue;
        if (input) {{
          input.hidden = !isCustom;
          input.required = isCustom;
          if (!isCustom) input.value = "";
          if (isCustom && focusInput) input.focus();
        }}
      }};
      select.addEventListener("change", () => update(true));
      update(false);
    }};
    setupCustomChoice(typeSelect);
    setupCustomChoice(seriesSelect);
    const cleanMeetingTitle = (name) => {{
      return (name || "投资会议纪要")
        .replace(/\\.[^.]+$/, "")
        .replace(/20\\d{{2}}[-_.年]?\\d{{1,2}}[-_.月]?\\d{{0,2}}日?/g, " ")
        .replace(/(原文|笔记|转录|录音|会议材料|整理前|导入|副本|copy)/gi, " ")
        .replace(/[_\\-—–]+/g, " ")
        .replace(/\\s+/g, " ")
        .trim() || "投资会议纪要";
    }};
    const inferType = (text) => {{
      const value = (text || "").toLowerCase();
      if (/(专家|访谈|产业链|深访|电话会|expert)/.test(value)) return "专家交流";
      if (/(上市公司|管理层|业绩|财报|路演|公司交流|ir)/.test(value)) return "上市公司交流";
      if (/(复盘|晨会|周会|例会|策略会|多人|市场回顾)/.test(value)) return "多人复盘会";
      return "多人复盘会";
    }};
    const inferSeries = (text, type) => {{
      return "研究所周会";
    }};
    const inferDate = (text) => {{
      const value = text || "";
      const match = value.match(/(20\\d{{2}})[-_.年]?(\\d{{1,2}})[-_.月]?(\\d{{1,2}})?/);
      if (!match || !match[3]) return "";
      return `${{match[1]}}-${{match[2].padStart(2, "0")}}-${{match[3].padStart(2, "0")}}`;
    }};
    const formatFileSize = (size) => {{
      if (!Number.isFinite(size)) return "";
      if (size >= 1024 * 1024) return `${{(size / 1024 / 1024).toFixed(1)}} MB`;
      if (size >= 1024) return `${{(size / 1024).toFixed(1)}} KB`;
      return `${{size}} B`;
    }};
    const fileKey = (file) => `${{file.name}}::${{file.size}}::${{file.lastModified}}`;
    const syncInputFiles = () => {{
      if (!filesInput || typeof DataTransfer === "undefined") return;
      const dataTransfer = new DataTransfer();
      selectedFiles.forEach((file) => dataTransfer.items.add(file));
      filesInput.files = dataTransfer.files;
    }};
    const renderFileList = (files) => {{
      if (!fileList) return;
      fileList.innerHTML = "";
      files.forEach((file, index) => {{
        const row = document.createElement("div");
        row.className = "file-item";
        const name = document.createElement("span");
        name.textContent = file.name;
        const size = document.createElement("span");
        size.className = "file-meta";
        size.textContent = formatFileSize(file.size);
        const remove = document.createElement("button");
        remove.type = "button";
        remove.className = "file-remove";
        remove.textContent = "删除";
        remove.setAttribute("aria-label", `删除 ${{file.name}}`);
        remove.addEventListener("click", () => {{
          selectedFiles.splice(index, 1);
          syncInputFiles();
          refreshUploadSelection();
        }});
        row.append(name, size, remove);
        fileList.appendChild(row);
      }});
    }};
    const addFiles = (incomingFiles) => {{
      const existing = new Set(selectedFiles.map(fileKey));
      incomingFiles.forEach((file) => {{
        const key = fileKey(file);
        if (!existing.has(key)) {{
          selectedFiles.push(file);
          existing.add(key);
        }}
      }});
      syncInputFiles();
      refreshUploadSelection();
    }};
    const refreshUploadSelection = async () => {{
      const files = selectedFiles;
      if (!files.length) {{
        if (smartHint) smartHint.textContent = "尚未选择文件。";
        renderFileList([]);
        return;
      }}
      renderFileList(files);
      const names = files.map((file) => file.name).join(" ");
      let sampleText = names;
      for (const file of files.slice(0, 3)) {{
        if (/\\.(txt|md|markdown|csv|tsv)$/i.test(file.name)) {{
          try {{
            sampleText += " " + (await file.slice(0, 12000).text());
          }} catch (error) {{}}
        }}
      }}
      const inferredTitle = cleanMeetingTitle(files[0].name);
      const inferredType = inferType(sampleText);
      const inferredSeries = inferSeries(sampleText, inferredType);
      const inferredDate = inferDate(names);
      if (titleInput && (!titleInput.value || titleInput.dataset.autoFilled === "1")) {{
        titleInput.value = inferredTitle;
        titleInput.dataset.autoFilled = "1";
      }}
      if (typeSelect && inferredType) {{
        typeSelect.value = inferredType;
        typeSelect.dispatchEvent(new Event("change"));
      }}
      if (seriesSelect && inferredSeries) {{
        seriesSelect.value = inferredSeries;
        seriesSelect.dispatchEvent(new Event("change"));
      }}
      if (dateInput && inferredDate) {{
        dateInput.value = inferredDate;
      }}
      if (smartHint) {{
        const totalSize = files.reduce((sum, file) => sum + file.size, 0);
        smartHint.textContent = `已选择 ${{files.length}} 个文件，合计 ${{formatFileSize(totalSize)}}。已预填：${{inferredTitle}} / ${{inferredType}} / ${{inferredSeries}}${{inferredDate ? " / " + inferredDate : ""}}`;
      }}
    }};
    filesInput?.addEventListener("change", () => {{
      addFiles(Array.from(filesInput.files || []));
    }});
    if (filesInput && uploadZone) {{
      ["dragenter", "dragover"].forEach((eventName) => {{
        uploadZone.addEventListener(eventName, (event) => {{
          event.preventDefault();
          uploadZone.classList.add("is-dragover");
        }});
      }});
      ["dragleave", "drop"].forEach((eventName) => {{
        uploadZone.addEventListener(eventName, () => uploadZone.classList.remove("is-dragover"));
      }});
      uploadZone.addEventListener("drop", (event) => {{
        event.preventDefault();
        const droppedFiles = Array.from(event.dataTransfer?.files || []);
        if (!droppedFiles.length) return;
        addFiles(droppedFiles);
      }});
    }}
    const uploadImportInChunks = async (form) => {{
      const files = selectedFiles.length ? selectedFiles : Array.from(filesInput?.files || []);
      if (!files.length) {{
        window.alert("请先选择要上传的会议材料。");
        return;
      }}
      const formData = new FormData(form);
      const metadata = {{
        meeting_title: formData.get("meeting_title") || "",
        meeting_date: formData.get("meeting_date") || "",
        meeting_type: formData.get("meeting_type") || "",
        custom_meeting_type: formData.get("custom_meeting_type") || "",
        meeting_series: formData.get("meeting_series") || "",
        custom_meeting_series: formData.get("custom_meeting_series") || "",
        operator_notes: formData.get("operator_notes") || "",
        files: files.map((file) => ({{
          name: file.name,
          size: file.size,
          type: file.type || "",
          last_modified: file.lastModified || 0,
        }})),
      }};
      showProcessing("正在上传会议材料", "正在分片上传文件；上传完成后会交给 Dify 工作流处理。");
      const uploadStages = [
        {{ label: "分片上传会议材料" }},
        {{ label: "提交给 Dify 工作流" }},
        {{ label: "Dify 抽取/转录/生成校对稿" }},
        {{ label: "打开人工校对页" }},
      ];
      const startResponse = await fetch("/api/import-upload/start", {{
        method: "POST",
        headers: {{ "Content-Type": "application/json", "Accept": "application/json" }},
        body: JSON.stringify(metadata),
        credentials: "same-origin",
      }});
      const startPayload = await startResponse.json().catch(() => ({{ ok: false, error: "上传任务创建失败" }}));
      if (!startResponse.ok || !startPayload.ok) {{
        throw new Error(startPayload.error || "上传任务创建失败");
      }}
      const chunkSize = Number(startPayload.chunk_size || (4 * 1024 * 1024));
      const totalBytes = files.reduce((sum, file) => sum + file.size, 0) || 1;
      let uploadedBytes = 0;
      for (let fileIndex = 0; fileIndex < files.length; fileIndex += 1) {{
        const file = files[fileIndex];
        let offset = 0;
        while (offset < file.size) {{
          const chunkStart = offset;
          const chunkEnd = Math.min(chunkStart + chunkSize, file.size);
          const chunk = file.slice(chunkStart, chunkEnd);
          let payload = null;
          let response = null;
          for (let attempt = 1; attempt <= 3; attempt += 1) {{
            const url = `/api/import-upload/chunk?upload_id=${{encodeURIComponent(startPayload.upload_id)}}&file_index=${{fileIndex}}&offset=${{chunkStart}}`;
            try {{
              response = await fetch(url, {{
                method: "POST",
                headers: {{ "Content-Type": "application/octet-stream", "Accept": "application/json" }},
                body: chunk,
                credentials: "same-origin",
              }});
              payload = await response.json().catch(() => ({{ ok: false, error: "分片上传失败" }}));
            }} catch (error) {{
              payload = {{ ok: false, error: error.message || "分片上传失败" }};
            }}
            if (response?.ok && payload?.ok) break;
            if (attempt >= 3) {{
              throw new Error(payload?.error || `分片上传失败：${{file.name}}`);
            }}
            await new Promise((resolve) => window.setTimeout(resolve, 800 * attempt));
          }}
          const received = Math.min(file.size, Math.max(chunkStart, Number(payload.received || 0)));
          if (received <= chunkStart) {{
            throw new Error(`分片上传未写入：${{file.name}}`);
          }}
          uploadedBytes += received - chunkStart;
          offset = received;
          const percent = Math.min(70, Math.max(2, Math.round((uploadedBytes / totalBytes) * 70)));
          setProcessingProgress(percent, `正在上传 ${{file.name}}`, uploadStages, 0);
          if (processingMessage) {{
            processingMessage.textContent = `正在上传第 ${{fileIndex + 1}}/${{files.length}} 个文件，已完成 ${{percent}}%。`;
          }}
        }}
      }}
      setProcessingProgress(74, "提交给 Dify 工作流", uploadStages, 1);
      if (processingMessage) processingMessage.textContent = "文件已上传完成，正在提交给 Dify 工作流。";
      const completeResponse = await fetch("/api/import-upload/complete", {{
        method: "POST",
        headers: {{ "Content-Type": "application/json", "Accept": "application/json" }},
        body: JSON.stringify({{ upload_id: startPayload.upload_id }}),
        credentials: "same-origin",
      }});
      const completePayload = await completeResponse.json().catch(() => ({{ ok: false, error: "提交 Dify 工作流失败" }}));
      if (!completeResponse.ok || !completePayload.ok) {{
        throw new Error(completePayload.error || "提交 Dify 工作流失败");
      }}
      setProcessingProgress(78, "Dify 工作流处理中", uploadStages, 2);
      window.location.href = completePayload.status_url || startPayload.status_url || "/drafts";
    }};
    const importForm = filesInput?.closest("form");
    importForm?.addEventListener("submit", async (event) => {{
      event.preventDefault();
      if (importForm.dataset.uploading === "1") return;
      importForm.dataset.uploading = "1";
      try {{
        await uploadImportInChunks(importForm);
      }} catch (error) {{
        importForm.dataset.uploading = "0";
        if (processingMessage) processingMessage.textContent = error.message || "上传失败，请稍后重试。";
        if (processingTitle) processingTitle.textContent = "上传失败";
        setProcessingProgress(0, "上传失败", [{{ label: error.message || "上传失败" }}], 0);
      }}
    }});
    titleInput?.addEventListener("input", () => {{
      titleInput.dataset.autoFilled = "0";
    }});
    window.addEventListener("pointermove", (event) => {{
      document.documentElement.style.setProperty("--mx", `${{event.clientX}}px`);
      document.documentElement.style.setProperty("--my", `${{event.clientY}}px`);
    }}, {{ passive: true }});
  </script>
</body>
</html>""".encode("utf-8")


def inline_markdown_to_html(text: str) -> str:
    escaped = html.escape(text)
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)
    escaped = re.sub(r"`([^`]+?)`", r"<code>\1</code>", escaped)
    return escaped


def markdown_to_readable_html(markdown: str) -> str:
    parts: list[str] = []
    lines = markdown.replace("\r\n", "\n").splitlines()
    list_type = ""
    table_lines: list[str] = []

    def close_list() -> None:
        nonlocal list_type
        if list_type:
            parts.append(f"</{list_type}>")
            list_type = ""

    def flush_table() -> None:
        nonlocal table_lines
        if not table_lines:
            return
        close_list()
        rows = []
        for raw in table_lines:
            stripped = raw.strip()
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            if cells and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                continue
            rows.append(cells)
        parts.append("<table><tbody>")
        for row_index, cells in enumerate(rows):
            cell_tag = "th" if row_index == 0 else "td"
            cells_html = "".join(f"<{cell_tag}>{inline_markdown_to_html(cell)}</{cell_tag}>" for cell in cells)
            parts.append(f"<tr>{cells_html}</tr>")
        parts.append("</tbody></table>")
        table_lines = []

    for raw_line in lines:
        stripped = raw_line.strip()
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines.append(raw_line)
            continue
        flush_table()
        if not stripped or stripped == "---":
            close_list()
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            close_list()
            level = len(heading.group(1))
            parts.append(f"<h{level}>{inline_markdown_to_html(heading.group(2))}</h{level}>")
            continue
        unordered = re.match(r"^[-*]\s+(.+)$", stripped)
        if unordered:
            if list_type != "ul":
                close_list()
                parts.append("<ul>")
                list_type = "ul"
            parts.append(f"<li>{inline_markdown_to_html(unordered.group(1))}</li>")
            continue
        ordered = re.match(r"^\d+[.]\s+(.+)$", stripped)
        if ordered:
            if list_type != "ol":
                close_list()
                parts.append("<ol>")
                list_type = "ol"
            parts.append(f"<li>{inline_markdown_to_html(ordered.group(1))}</li>")
            continue
        close_list()
        parts.append(f"<p>{inline_markdown_to_html(stripped)}</p>")
    flush_table()
    close_list()
    return "\n".join(parts)


def review_editor_html(
    markdown: str,
    draft_id: str,
    *,
    confirm_action: str = "/confirm",
    confirm_label: str = "确认归档",
    helper_text: str = "直接在下方文本框里修改，完成后点击确认按钮。",
) -> str:
    encoded_markdown = html.escape(markdown)
    quoted_draft_id = urllib.parse.quote(draft_id)
    quoted_action = html.escape(f"{confirm_action}?id={quoted_draft_id}")
    return f"""
<form id="review-form" method="post" action="/save?id={quoted_draft_id}" data-processing-title="正在保存校对结果" data-processing-message="系统正在保存当前校对内容。">
  <textarea class="markdown-field" name="markdown" id="markdown-field">{encoded_markdown}</textarea>
  <div class="direct-edit-shell">
    <div class="review-hint">{html.escape(helper_text)}</div>
    <div class="toolbar" role="toolbar" aria-label="校对工具栏">
      <button class="tool active" type="button" id="show-direct" title="直接编辑正文">编辑</button>
      <button class="tool" type="button" id="show-preview" title="查看排版预览">预览</button>
    </div>
    <textarea id="direct-editor" class="direct-editor" spellcheck="false">{encoded_markdown}</textarea>
  </div>
  <div class="editor-shell">
    <div class="toolbar" role="toolbar" aria-label="校对工具栏">
      <select id="block-format" title="段落格式">
        <option value="p">正文</option>
        <option value="h1">标题 1</option>
        <option value="h2">标题 2</option>
        <option value="h3">标题 3</option>
      </select>
      <button class="tool" type="button" data-command="bold" title="加粗"><strong>B</strong></button>
      <button class="tool" type="button" data-command="insertUnorderedList" title="项目符号">•</button>
      <button class="tool" type="button" data-command="insertOrderedList" title="编号">1.</button>
      <button class="tool" type="button" id="toggle-source" title="查看 Markdown 源文">MD</button>
    </div>
    <div id="rich-editor" class="paper" contenteditable="true" spellcheck="false"></div>
  </div>
  <div class="source-panel">
    <textarea id="source-editor">{encoded_markdown}</textarea>
    <div class="bar"><button class="secondary" type="button" id="back-to-rich">返回富文本校对</button></div>
  </div>
  <div class="bar review-actions">
    <button type="submit" class="secondary" data-processing-title="正在保存草稿" data-processing-message="系统正在保存当前校对内容。">保存草稿</button>
    <button type="submit" formaction="{quoted_action}" data-processing-title="正在运行工作流" data-processing-message="系统正在执行下一步整理或生成，完成前请不要重复提交。">{html.escape(confirm_label)}</button>
    <a class="secondary" href="/history">返回历史</a>
  </div>
</form>
<script>
const initialMarkdown = document.getElementById("markdown-field").value;
const richEditor = document.getElementById("rich-editor");
const sourceEditor = document.getElementById("source-editor");
const directEditor = document.getElementById("direct-editor");
const form = document.getElementById("review-form");
const markdownField = document.getElementById("markdown-field");

function escapeHtml(value) {{
  return value.replace(/&/g, "&amp;").replace(/</g, "&lt;").replace(/>/g, "&gt;");
}}

function inlineMarkdownToHtml(value) {{
  return escapeHtml(value).replace(/\\*\\*(.+?)\\*\\*/g, "<strong>$1</strong>");
}}

function markdownToHtml(markdown) {{
  const lines = markdown.replace(/\\r\\n/g, "\\n").split("\\n");
  const htmlParts = [];
  let listType = "";
  let tableLines = [];
  const closeList = () => {{
    if (listType) {{
      htmlParts.push(`</${{listType}}>`);
      listType = "";
    }}
  }};
  const flushTable = () => {{
    if (!tableLines.length) return;
    closeList();
    const rows = tableLines.filter((line) => !/^\\s*\\|?\\s*:?-{{3,}}/.test(line.replace(/\\|/g, " | ")));
    htmlParts.push("<table><tbody>");
    rows.forEach((line, index) => {{
      const cells = line.trim().replace(/^\\|/, "").replace(/\\|$/, "").split("|");
      const tag = index === 0 ? "th" : "td";
      htmlParts.push("<tr>" + cells.map((cell) => `<${{tag}}>${{inlineMarkdownToHtml(cell.trim())}}</${{tag}}>`).join("") + "</tr>");
    }});
    htmlParts.push("</tbody></table>");
    tableLines = [];
  }};
  lines.forEach((line) => {{
    const trimmed = line.trim();
    if (trimmed.includes("|") && /^\\s*\\|?.+\\|.+/.test(trimmed)) {{
      tableLines.push(line);
      return;
    }}
    flushTable();
    if (!trimmed) {{
      closeList();
      return;
    }}
    const heading = trimmed.match(/^(#{{1,3}})\\s+(.+)$/);
    if (heading) {{
      closeList();
      const level = heading[1].length;
      htmlParts.push(`<h${{level}}>${{inlineMarkdownToHtml(heading[2])}}</h${{level}}>`);
      return;
    }}
    const unordered = trimmed.match(/^[-*]\\s+(.+)$/);
    if (unordered) {{
      if (listType !== "ul") {{
        closeList();
        htmlParts.push("<ul>");
        listType = "ul";
      }}
      htmlParts.push(`<li>${{inlineMarkdownToHtml(unordered[1])}}</li>`);
      return;
    }}
    const ordered = trimmed.match(/^\\d+[.]\\s+(.+)$/);
    if (ordered) {{
      if (listType !== "ol") {{
        closeList();
        htmlParts.push("<ol>");
        listType = "ol";
      }}
      htmlParts.push(`<li>${{inlineMarkdownToHtml(ordered[1])}}</li>`);
      return;
    }}
    closeList();
    htmlParts.push(`<p>${{inlineMarkdownToHtml(trimmed)}}</p>`);
  }});
  flushTable();
  closeList();
  return htmlParts.join("\\n");
}}

function textWithInlineMarkdown(node) {{
  let output = "";
  node.childNodes.forEach((child) => {{
    if (child.nodeType === Node.TEXT_NODE) {{
      output += child.textContent;
    }} else if (child.nodeName === "BR") {{
      output += "\\n";
    }} else if (child.nodeName === "STRONG" || child.nodeName === "B") {{
      output += "**" + textWithInlineMarkdown(child).trim() + "**";
    }} else {{
      output += textWithInlineMarkdown(child);
    }}
  }});
  return output.replace(/\\u00a0/g, " ").trim();
}}

function htmlToMarkdown(root) {{
  const blocks = [];
  root.childNodes.forEach((node) => {{
    if (node.nodeType === Node.TEXT_NODE) {{
      const text = node.textContent.trim();
      if (text) blocks.push(text);
      return;
    }}
    const tag = node.nodeName.toLowerCase();
    if (tag === "h1" || tag === "h2" || tag === "h3") {{
      blocks.push("#".repeat(Number(tag.slice(1))) + " " + textWithInlineMarkdown(node));
    }} else if (tag === "ul" || tag === "ol") {{
      Array.from(node.children).forEach((li, index) => {{
        const prefix = tag === "ol" ? `${{index + 1}}. ` : "- ";
        blocks.push(prefix + textWithInlineMarkdown(li));
      }});
    }} else if (tag === "table") {{
      const rows = Array.from(node.querySelectorAll("tr")).map((row) =>
        Array.from(row.children).map((cell) => textWithInlineMarkdown(cell))
      ).filter((cells) => cells.length);
      rows.forEach((cells, index) => {{
        blocks.push("| " + cells.join(" | ") + " |");
        if (index === 0) blocks.push("| " + cells.map(() => "---").join(" | ") + " |");
      }});
    }} else if (tag === "div" || tag === "p") {{
      const text = textWithInlineMarkdown(node);
      if (text) blocks.push(text);
    }}
  }});
  return blocks.join("\\n\\n").replace(/\\n{{3,}}/g, "\\n\\n").trim() + "\\n";
}}

function syncFromRich() {{
  markdownField.value = htmlToMarkdown(richEditor);
  sourceEditor.value = markdownField.value;
  directEditor.value = markdownField.value;
}}

richEditor.innerHTML = markdownToHtml(initialMarkdown);

function syncFromDirect() {{
  markdownField.value = directEditor.value;
  sourceEditor.value = directEditor.value;
  richEditor.innerHTML = markdownToHtml(directEditor.value);
}}

directEditor.addEventListener("input", () => {{
  markdownField.value = directEditor.value;
  sourceEditor.value = directEditor.value;
}});

document.querySelectorAll("[data-command]").forEach((button) => {{
  button.addEventListener("click", () => {{
    document.execCommand(button.dataset.command, false, null);
    richEditor.focus();
  }});
}});

document.getElementById("block-format").addEventListener("change", (event) => {{
  document.execCommand("formatBlock", false, event.target.value);
  richEditor.focus();
}});

document.getElementById("toggle-source").addEventListener("click", () => {{
  syncFromDirect();
  document.body.classList.add("source-mode");
}});

document.getElementById("back-to-rich").addEventListener("click", () => {{
  markdownField.value = sourceEditor.value;
  directEditor.value = sourceEditor.value;
  richEditor.innerHTML = markdownToHtml(sourceEditor.value);
  document.body.classList.remove("source-mode");
}});

document.getElementById("show-preview").addEventListener("click", () => {{
  syncFromDirect();
  document.body.classList.add("preview-mode");
}});

document.getElementById("show-direct").addEventListener("click", () => {{
  if (document.body.classList.contains("source-mode")) {{
    directEditor.value = sourceEditor.value;
  }} else {{
    directEditor.value = markdownField.value;
  }}
  document.body.classList.remove("preview-mode");
  document.body.classList.remove("source-mode");
  directEditor.focus();
}});

form.addEventListener("submit", () => {{
  if (document.body.classList.contains("source-mode")) {{
    markdownField.value = sourceEditor.value;
  }} else if (document.body.classList.contains("preview-mode")) {{
    syncFromRich();
  }} else {{
    markdownField.value = directEditor.value;
  }}
}});

window.addEventListener("load", () => {{
  directEditor.focus();
  directEditor.setSelectionRange(0, 0);
}});
</script>
"""


class ReviewHandler(BaseHTTPRequestHandler):
    server_version = "MeetingMinutesReview/1.0"

    def log_message(self, fmt: str, *args) -> None:
        LOG_DIR.mkdir(parents=True, exist_ok=True)
        with (LOG_DIR / "meeting-minutes-review-server.log").open("a", encoding="utf-8") as handle:
            handle.write("%s - %s\n" % (self.log_date_time_string(), fmt % args))

    def _security_headers(self) -> None:
        self.send_header("X-Content-Type-Options", "nosniff")
        self.send_header("X-Frame-Options", "DENY")
        self.send_header("Referrer-Policy", "strict-origin-when-cross-origin")
        self.send_header("Permissions-Policy", "camera=(), microphone=(), geolocation=(), payment=()")
        self.send_header(
            "Content-Security-Policy",
            "default-src 'self'; img-src 'self' data:; style-src 'self' 'unsafe-inline'; "
            "script-src 'self' 'unsafe-inline'; connect-src 'self'; frame-ancestors 'none'; "
            "base-uri 'self'; form-action 'self'",
        )
        self.send_header("Cross-Origin-Opener-Policy", "same-origin")

    def _send(self, status: int, body: bytes, content_type: str) -> None:
        self.send_response(status)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(body)))
        self._security_headers()
        self.end_headers()
        self.wfile.write(body)

    def _json(self, status: int, payload: dict[str, Any]) -> None:
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self._send(status, data, "application/json; charset=utf-8")

    def _redirect(self, location: str) -> None:
        self.send_response(303)
        self.send_header("Location", location)
        self.send_header("Content-Length", "0")
        self._security_headers()
        self.end_headers()

    def _is_platform_proxy_path(self, path: str) -> bool:
        if path in PLATFORM_PAGE_PATHS:
            return True
        return any(path.startswith(prefix) for prefix in (*PLATFORM_API_PREFIXES, *PLATFORM_ASSET_PREFIXES))

    def _platform_proxy_requires_auth(self, path: str) -> bool:
        if path in PUBLIC_PLATFORM_PAGE_PATHS:
            return False
        return not any(path.startswith(prefix) for prefix in PLATFORM_ASSET_PREFIXES)

    def _proxy_platform(self, identity: dict[str, Any] | None = None) -> None:
        parsed = urllib.parse.urlparse(self.path)
        target = f"{PLATFORM_PROXY_BASE_URL}{parsed.path}"
        if parsed.query:
            target = f"{target}?{parsed.query}"
        headers = {
            "Accept": self.headers.get("Accept", "*/*"),
            "User-Agent": self.headers.get("User-Agent", "D91 public proxy"),
            "X-Forwarded-Host": self.headers.get("Host", ""),
            "X-Forwarded-Proto": "https" if is_public_host(self.headers.get("Host", "")) else "http",
        }
        if identity:
            headers["X-Forwarded-Email"] = str(identity.get("user_id") or "authenticated")
        try:
            req = urllib.request.Request(target, headers=headers, method="GET")
            with urllib.request.urlopen(req, timeout=20) as response:
                body = response.read()
                status = response.status
                response_headers = response.headers
        except urllib.error.HTTPError as exc:
            body = exc.read()
            status = exc.code
            response_headers = exc.headers
        except Exception as exc:
            self._json(502, {"ok": False, "error": "platform_proxy_failed", "message": str(exc)})
            return
        self.send_response(status)
        for key, value in response_headers.items():
            lower_key = key.lower()
            if lower_key in {"connection", "content-length", "keep-alive", "proxy-authenticate", "proxy-authorization", "transfer-encoding"}:
                continue
            self.send_header(key, value)
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _wants_json_response(self) -> bool:
        content_type = self.headers.get("Content-Type", "")
        accept = self.headers.get("Accept", "")
        return "application/json" in content_type or (
            "application/json" in accept and "text/html" not in accept
        )

    def _client_key(self) -> str:
        forwarded = (self.headers.get("X-Forwarded-For") or "").split(",")[0].strip()
        remote = self.client_address[0] if self.client_address else "unknown"
        return forwarded or remote or "unknown"

    def _read_payload(self) -> dict[str, Any]:
        length = int(self.headers.get("Content-Length", "0") or "0")
        if length > MAX_POST_BYTES:
            raise ValueError("request body too large")
        content_type = self.headers.get("Content-Type", "")
        if "multipart/form-data" in content_type:
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": content_type,
                    "CONTENT_LENGTH": str(length),
                },
                keep_blank_values=True,
            )
            payload: dict[str, Any] = {"_files": []}
            keys = form.keys() if hasattr(form, "keys") else []
            for key in keys:
                values = form[key]
                fields = values if isinstance(values, list) else [values]
                for field in fields:
                    if getattr(field, "filename", None):
                        payload["_files"].append(
                            {
                                "field": key,
                                "filename": field.filename,
                                "content_type": getattr(field, "type", "") or "",
                                "data": field.file.read(),
                            }
                        )
                    else:
                        payload[key] = field.value
            return payload
        raw = self.rfile.read(length).decode("utf-8") if length else ""
        if "application/json" in content_type:
            return json.loads(raw or "{}")
        parsed = urllib.parse.parse_qs(raw, keep_blank_values=True)
        return {key: values[-1] for key, values in parsed.items()}

    def _read_binary_body(self, max_bytes: int = MAX_IMPORT_CHUNK_BYTES) -> tuple[bytes, dict[str, Any]]:
        length = int(self.headers.get("Content-Length", "0") or "0")
        if length > max_bytes:
            raise ValueError("上传分片过大")
        started_at = time.monotonic()
        data = self.rfile.read(length) if length else b""
        elapsed = time.monotonic() - started_at
        return data, {
            "declared_bytes": length,
            "received_bytes": len(data),
            "read_seconds": round(elapsed, 3),
        }

    def _request_identity(self) -> dict[str, Any]:
        host = self.headers.get("X-Forwarded-Host") or self.headers.get("Host") or ""
        if access_disabled():
            return {
                "authenticated": True,
                "user_id": "direct",
                "display_name": "直接访问用户",
                "role": "admin",
                "groups": ["admin", "direct"],
                "is_admin": True,
                "source": "access_disabled",
                "host": host,
            }
        if not is_public_host(host):
            return {
                "authenticated": True,
                "user_id": "local",
                "display_name": "本机管理员",
                "role": "admin",
                "groups": ["admin", "local"],
                "is_admin": True,
                "source": "local",
                "host": host,
            }
        token = self.headers.get("X-Workflow-Access-Token") or ""
        auth_header = self.headers.get("Authorization") or ""
        if not token and auth_header.lower().startswith("bearer "):
            token = auth_header.split(" ", 1)[1].strip()
        if token:
            token_identity = validate_access_token(token)
            if token_identity:
                token_identity["is_admin"] = str(token_identity.get("role") or "").lower() == "admin"
                token_identity["host"] = host
                return token_identity
        else:
            cookie_header = self.headers.get("Cookie") or ""
            parsed_cookies = cookies.SimpleCookie()
            try:
                parsed_cookies.load(cookie_header)
                cookie_value = parsed_cookies.get(AUTH_COOKIE_NAME).value if parsed_cookies.get(AUTH_COOKIE_NAME) else ""
            except Exception:
                cookie_value = ""
            cookie_identity = validate_access_session(cookie_value) or validate_access_token(cookie_value)
            if cookie_identity:
                cookie_identity["is_admin"] = str(cookie_identity.get("role") or "").lower() == "admin"
                cookie_identity["host"] = host
                return cookie_identity
        return {
            "authenticated": False,
            "user_id": "anonymous",
            "display_name": "未登录用户",
            "role": "anonymous",
            "groups": [],
            "is_admin": False,
            "source": "anonymous",
            "host": host,
        }

    def _require_authenticated(self, query: dict[str, list[str]] | None = None) -> dict[str, Any] | None:
        identity = self._request_identity()
        if identity.get("authenticated"):
            return identity
        self._audit("auth_required", identity, "blocked")
        parsed = urllib.parse.urlparse(self.path)
        next_path = parsed.path + (("?" + parsed.query) if parsed.query else "")
        if self._wants_json_response() or parsed.path.startswith("/api/"):
            self._json(401, {"ok": False, "error": "authentication_required", "login_url": login_path(next_path)})
        else:
            self.send_response(302)
            self.send_header("Location", login_path(next_path))
            self.send_header("Content-Length", "0")
            self._security_headers()
            self.end_headers()
        return None

    def _require_role(self, roles: set[str], message: str) -> dict[str, Any] | None:
        identity = self._require_authenticated()
        if identity is None:
            return None
        role = str(identity.get("role") or "").lower()
        if identity.get("is_admin") or role in roles:
            return identity
        self._deny_access(message)
        return None

    def _require_editor(self) -> dict[str, Any] | None:
        return self._require_role({"editor"}, "只有编辑或管理员可以上传、校对或归档会议纪要。")

    def _require_admin(self) -> dict[str, Any] | None:
        return self._require_role({"admin"}, "只有管理员可以访问后台管理功能。")

    def _deny_access(self, message: str = "当前身份无权查看该资料。") -> None:
        self._audit("forbidden", self._request_identity(), "blocked", message=message)
        if self._wants_json_response():
            self._json(403, {"ok": False, "error": "forbidden", "message": message})
            return
        body = f"""
<div class="meta">
  <strong>无权访问</strong>
  <div class="muted">{html.escape(message)}</div>
</div>
<div class="bar"><a href="/history">返回历史纪要</a><a href="/external-sources">返回资料源</a></div>
"""
        self._send(403, html_page("无权访问", body), "text/html; charset=utf-8")

    def _check_post_origin(self) -> bool:
        host = self.headers.get("X-Forwarded-Host") or self.headers.get("Host") or ""
        if not is_public_host(host):
            return True
        origin = self.headers.get("Origin") or ""
        referer = self.headers.get("Referer") or ""
        source = origin or referer
        if not source:
            return True
        source_host = urllib.parse.urlparse(source).netloc
        expected_host = host.split(",", 1)[0].strip()
        if source_host == expected_host:
            return True
        self._audit("blocked_cross_site_post", self._request_identity(), "blocked", origin=origin, referer=referer)
        self._send(403, html_page("请求被拦截", "<p class=\"error\">请求来源不匹配，请刷新页面后重试。</p>"), "text/html; charset=utf-8")
        return False

    def _set_access_cookie(self, token: str) -> None:
        secure = "; Secure" if is_public_host(self.headers.get("X-Forwarded-Host") or self.headers.get("Host") or "") else ""
        self.send_header(
            "Set-Cookie",
            f"{AUTH_COOKIE_NAME}={urllib.parse.quote(token)}; Path=/; HttpOnly; SameSite=Lax{secure}; Max-Age=2592000",
        )

    def _clear_access_cookie(self) -> None:
        secure = "; Secure" if is_public_host(self.headers.get("X-Forwarded-Host") or self.headers.get("Host") or "") else ""
        self.send_header("Set-Cookie", f"{AUTH_COOKIE_NAME}=; Path=/; HttpOnly; SameSite=Lax{secure}; Max-Age=0")

    def _audit(self, event: str, identity: dict[str, Any] | None = None, status: str = "ok", **details: Any) -> None:
        try:
            actor = identity or self._request_identity()
            payload = {
                "ts": now_iso(),
                "event": event,
                "status": status,
                "user_id": actor.get("user_id"),
                "role": actor.get("role"),
                "groups": actor.get("groups") or [],
                "identity_source": actor.get("source"),
                "host": self.headers.get("Host") or "",
                "x_forwarded_host": self.headers.get("X-Forwarded-Host") or "",
                "x_forwarded_for": self.headers.get("X-Forwarded-For") or "",
                "client": self.client_address[0] if self.client_address else "",
                "method": self.command,
                "path": urllib.parse.urlparse(self.path).path,
                "details": details,
            }
            log_path = access_audit_log_path()
            log_path.parent.mkdir(parents=True, exist_ok=True)
            with log_path.open("a", encoding="utf-8") as handle:
                handle.write(json.dumps(payload, ensure_ascii=False, sort_keys=True) + "\n")
        except Exception as exc:
            self.log_message("access audit failed: %s", exc)

    def _handle_home(self) -> None:
        identity = self._request_identity()
        login_attrs = "" if identity.get("authenticated") else ' data-protected-link data-next="/apps/meeting-minutes/import"'
        tools_attrs = "" if identity.get("authenticated") else ' data-protected-link data-next="/tools"'
        body = f"""
<section class="portal-hero">
  <div class="hero-copy">
    <p class="eyebrow">OpenFinance · Developed by KumaAI</p>
    <h1>D91 AI投研平台</h1>
    <p>首页优先展示投研看板、正式归档和知识库状态；OpenMinute 等工具入口放在下方，减少日常使用路径干扰。</p>
    <div class="notice">本网站所有信息和功能仅供个人研究使用，不对外开放。</div>
    <div class="bar">
      <a href="{html.escape('/apps/meeting-minutes/import' if identity.get('authenticated') else '#login')}"{login_attrs}>导入会议材料</a>
      <a class="secondary" href="{html.escape('/dashboards' if identity.get('authenticated') else '#login')}"{'' if identity.get('authenticated') else ' data-protected-link data-next="/dashboards"'}>公众号文章看板</a>
    </div>
  </div>
  <div class="hero-panel" aria-label="会议纪要流程概览">
    <div class="hero-panel-top"><span>OpenMinute</span><span>默认研究所周会</span></div>
    <div class="flow-line"><span></span><span></span><span></span><span></span><span></span></div>
    <div class="hero-list">
      <div class="hero-task"><span class="hero-dot"></span><div><b>导入材料</b><br><small>标题、日期、类型、系列</small></div><small>01</small></div>
      <div class="hero-task"><span class="hero-dot"></span><div><b>人工初校</b><br><small>标注发言人、分段并修正关键名词</small></div><small>02</small></div>
      <div class="hero-task"><span class="hero-dot"></span><div><b>按类型整理</b><br><small>多人复盘 / 上市公司 / 专家交流</small></div><small>03</small></div>
      <div class="hero-task"><span class="hero-dot"></span><div><b>二次校对与归档</b><br><small>确认发言人、标的和关键名词无误</small></div><small>04</small></div>
    </div>
  </div>
</section>
<div class="quick-strip">
  <section class="work-surface">
    <div class="surface-title"><div><h2>公众号文章看板</h2><p>实时展示 WeRSS 公众号文章，并支持站内阅读全文。</p></div><a href="{html.escape('/dashboards' if identity.get('authenticated') else '#login')}"{'' if identity.get('authenticated') else ' data-protected-link data-next="/dashboards"'}>打开看板</a></div>
    <div class="module-list">{portal_module_rows("dashboard", compact=True, identity=identity)}</div>
  </section>
  <section class="work-surface">
    <div class="surface-title"><div><h2>知识库</h2><p>按账号权限过滤后检索会议纪要和外部资料。</p></div><a href="{html.escape('/knowledge' if identity.get('authenticated') else '#login')}"{'' if identity.get('authenticated') else ' data-protected-link data-next="/knowledge"'}>进入知识库</a></div>
    <div class="module-list">{portal_module_rows("knowledge", compact=True, identity=identity)}</div>
  </section>
</div>
<div class="quick-strip">
  <section class="work-surface">
    <div class="surface-title"><div><h2>工具入口</h2><p>OpenFinance 工具从这里进入。</p></div><a href="{html.escape('/tools' if identity.get('authenticated') else '#login')}"{tools_attrs}>全部工具</a></div>
    <div class="module-list">{portal_module_rows("tools", identity=identity)}</div>
  </section>
</div>
"""
        self._send(200, html_page(SITE_NAME, body, identity), "text/html; charset=utf-8")

    def _handle_apps(self, identity: dict[str, Any]) -> None:
        section = next(item for item in PORTAL_SECTIONS if item["key"] == "tools")
        link_attr = lambda path: "" if identity.get("authenticated") else f' data-protected-link data-next="{html.escape(path)}"'
        href_for = lambda path: html.escape(path if identity.get("authenticated") else "#login")
        body = f"""
<h1 class="page-title">工具</h1>
<div class="meta">
  <strong>{html.escape(section["title"])}</strong>
  <div class="muted">OpenFinance 工具入口。当前主工具是 OpenMinute-会议纪要，后续工具继续在这里接入。</div>
</div>
<div class="app-shell">
  <aside class="side-rail">
    <a href="{href_for('/apps/meeting-minutes')}"{link_attr('/apps/meeting-minutes')}>OpenMinute</a>
    <a href="{href_for('/apps/meeting-minutes/import')}"{link_attr('/apps/meeting-minutes/import')}>导入会议材料</a>
    <a href="{href_for('/drafts')}"{link_attr('/drafts')}>待校对草稿</a>
    <a href="{href_for('/dashboard')}"{link_attr('/dashboard')}>数据看板</a>
  </aside>
  <section>
    <div class="module-list">{portal_module_rows("tools", identity=identity)}</div>
  </section>
</div>
"""
        self._send(200, html_page("工具", body, identity), "text/html; charset=utf-8")

    def _handle_knowledge(self, identity: dict[str, Any], query: dict[str, list[str]] | None = None) -> None:
        query = query or {}
        question = (query.get("q") or [""])[0].strip()
        answer_block = ""
        if question:
            answer = knowledge_answer(question, identity=identity)
            source_rows = []
            for item in answer.get("sources") or []:
                url = normalize_local_url(str(item.get("review_url") or item.get("view_url") or item.get("view_path") or ""))
                source_rows.append(
                    f"""<div class="item">
  <div><a href="{html.escape(url)}"><strong>{html.escape(str(item.get('title') or '资料'))}</strong></a> <span class="status">{html.escape(str(item.get('source_label') or '资料'))}</span></div>
  <div class="muted">日期：{html.escape(str(item.get('meeting_date') or item.get('date') or '-'))}　命中：{html.escape(' / '.join(str(term) for term in item.get('matched_terms') or []) or '-')}</div>
  <div>{html.escape(str(item.get('evidence_snippet') or ''))}</div>
</div>"""
                )
            answer_block = f"""
<div class="item">
  <strong>AI 问答结果</strong>
  <p class="muted">当前实现先按账号权限做 RAG 检索和证据汇总；如果接入独立问答模型，仍必须只使用这些权限过滤后的证据。</p>
  <pre class="markdown-source">{html.escape(str(answer.get('answer') or answer.get('error') or ''))}</pre>
</div>
<div class="section-head"><div><h2>引用证据</h2><p>权限外资料不会显示。</p></div></div>
{''.join(source_rows) if source_rows else '<p class="muted">暂无引用证据。</p>'}
"""
        body = f"""
<h1 class="page-title">知识库</h1>
<div class="meta">
  <strong>会议纪要 RAG 问答</strong>
  <div class="muted">知识库只使用当前账号有权访问的正式会议纪要和外部资料。不同账号级别会先做权限过滤，再进入问答和证据展示。</div>
</div>
<form class="filters bar" method="get" action="/knowledge">
  <input name="q" placeholder="输入标的、公司、主题或问题" value="{html.escape(question)}">
  <button type="submit">提问</button>
  <a class="secondary" href="/target-query">标的证据查询</a>
</form>
{answer_block}
<div class="module-list">
  {portal_module_rows("knowledge", identity=identity)}
</div>
"""
        self._send(200, html_page("知识库", body, identity), "text/html; charset=utf-8")

    def _handle_operations(self) -> None:
        body = f"""
<h1 class="page-title">数据</h1>
<div class="meta">
  <strong>历史数据与后台状态</strong>
  <div class="muted">只有最高级别账号可以访问。这里集中展示历史数据、访问审计、知识库映射、同步状态和健康检查。</div>
</div>
<div class="module-list">{portal_module_rows("data")}</div>
<div class="item">
  <strong>可编辑数据展示窗口</strong>
  <p class="muted">用于临时粘贴、比对或整理从看板/API 导出的数据；内容仅保存在当前浏览器页面，不写入服务器。</p>
  <textarea class="data-window" id="data-window" placeholder="粘贴需要临时查看或整理的数据"></textarea>
</div>
<script>
const dataWindow = document.getElementById("data-window");
if (dataWindow) {{
  dataWindow.value = localStorage.getItem("d91-data-window") || "";
  dataWindow.addEventListener("input", () => localStorage.setItem("d91-data-window", dataWindow.value));
}}
</script>
"""
        self._send(200, html_page("数据", body), "text/html; charset=utf-8")

    def _handle_account(self, identity: dict[str, Any]) -> None:
        groups = ", ".join(str(item) for item in identity.get("groups") or []) or "-"
        admin_links = (
            '<a href="/access-users">管理访问用户</a><a href="/access-audit">访问审计</a>'
            if identity.get("is_admin")
            else ""
        )
        body = f"""
<h1 class="page-title">账号</h1>
<div class="meta">
  <strong>{html.escape(str(identity.get('display_name') or identity.get('user_id') or '未登录'))}</strong>
  <div class="muted">账号信息只展示当前登录身份和权限。公网账号由管理员维护，不开放自助注册。</div>
</div>
<table>
  <tbody>
    <tr><th>用户 ID</th><td>{html.escape(str(identity.get('user_id') or '-'))}</td></tr>
    <tr><th>角色</th><td>{html.escape(str(identity.get('role') or '-'))}</td></tr>
    <tr><th>分组</th><td>{html.escape(groups)}</td></tr>
    <tr><th>身份来源</th><td>{html.escape(str(identity.get('source') or '-'))}</td></tr>
  </tbody>
</table>
<div class="bar">
  {admin_links}
  <a href="/access-logout">退出登录</a>
</div>
"""
        self._send(200, html_page("账号", body, identity), "text/html; charset=utf-8")

    def _handle_meeting_app(self) -> None:
        body = """
""" + workflow_stage_header("OpenMinute-会议纪要", "从导入材料开始，也可以继续处理待校对草稿或查看历史纪要。") + """
<div class="meta">
  <strong>从材料到归档，按步骤处理</strong>
  <div class="muted">系统会先让你校对转录文本，再按会议类型生成纪要；当前默认会议系列为研究所周会，后续可按需要新增系列，最终归档前还会再次确认。</div>
</div>
<div class="stepper">
  <div class="step active"><strong>1 上传材料</strong><span class="muted">录音、文稿或粘贴文本</span></div>
  <div class="step"><strong>2 初步校对</strong><span class="muted">标注发言人、分段并修正关键名词</span></div>
  <div class="step"><strong>3 二次校对</strong><span class="muted">确认发言人、标的、关键名词无误</span></div>
  <div class="step"><strong>4 归档确认</strong><span class="muted">生成 Word、知识库和看板数据</span></div>
</div>
<div class="bar">
  <a href="/apps/meeting-minutes/import">导入材料</a>
  <a class="secondary" href="/drafts">待校对</a>
  <a href="/history">历史纪要</a>
  <a href="/dashboard">数据看板</a>
</div>
<div class="item">
  <strong>会议类型</strong>
  <p class="muted">默认类型：多人复盘会；仅当材料明确是单家公司专场或专家问答时，选择上市公司交流或专家交流。当前预置会议系列只有“研究所周会”，也支持新增系列；类型和系列都会进入草稿 metadata、历史筛选和知识库归档 metadata。</p>
</div>
"""
        self._send(200, html_page("OpenMinute-会议纪要", body), "text/html; charset=utf-8")

    def _handle_dashboard(self, query: dict[str, list[str]], identity: dict[str, Any]) -> None:
        selected_date = (query.get("date") or [datetime.now().strftime("%Y-%m-%d")])[0]
        data = dashboard_data(selected_date, identity)
        target_rows = []
        for item in data.get("high_frequency_targets") or []:
            logic = item.get("logic") or []
            snippets = "".join(
                f"<li><a href=\"{html.escape(normalize_local_url(str(row.get('url') or '')))}\">{html.escape(str(row.get('title') or ''))}</a> <span class=\"status\">{html.escape('/'.join(row.get('logic_tags') or ['观点']))}</span>：{html.escape(str(row.get('snippet') or ''))}</li>"
                for row in logic
            )
            target_rows.append(
                f"""<div class="item">
  <strong>{html.escape(str(item.get('target') or '-'))}</strong> <span class="status">{html.escape(str(item.get('count') or 0))} 次</span>
  {f'<ul>{snippets}</ul>' if snippets else '<p class="muted">未抽取到逻辑摘录。</p>'}
</div>"""
            )
        minute_rows = []
        for item in data.get("deduped_minutes") or []:
            minute_rows.append(
                f"""<tr>
  <td>{html.escape(str(item.get('meeting_date') or ''))}</td>
  <td><span class="status">{html.escape(str(item.get('meeting_type') or '多人复盘会'))}</span></td>
  <td><a href="{html.escape(normalize_local_url(str(item.get('review_url') or '')))}">{html.escape(str(item.get('title') or '投资会议纪要'))}</a></td>
  <td>{html.escape(' / '.join(str(target) for target in item.get('targets') or []) or '-')}</td>
</tr>"""
            )
        speaker_blocks = []
        for speaker, values in (data.get("speaker_rollup") or {}).items():
            rows = "".join(
                f"<li><a href=\"{html.escape(normalize_local_url(str(row.get('url') or '')))}\">{html.escape(str(row.get('title') or ''))}</a> <span class=\"status\">{html.escape(str(row.get('meeting_type') or '多人复盘会'))}</span> {html.escape(str(row.get('heading') or ''))}</li>"
                for row in values
            )
            speaker_blocks.append(f"<div class=\"item\"><strong>{html.escape(str(speaker))}</strong><ul>{rows}</ul></div>")
        type_text = " / ".join(f"{key}: {value}" for key, value in (data.get("meeting_types") or {}).items()) or "-"
        body = f"""
<h1 class="page-title">投研数据看板</h1>
<div class="meta">
  <strong>多看板视图，只读取已归档终稿</strong>
  <div class="muted">基于人工确认后的正式纪要生成，包含高频股票、发言人观点、当日去重纪要和可编辑数据展示窗口；不生成新的投资结论。</div>
</div>
<form class="filters bar" method="get" action="/dashboard">
  <input name="date" placeholder="日期 YYYY-MM-DD" value="{html.escape(str(data.get('date') or selected_date))}">
  <button type="submit">刷新</button>
</form>
<div class="metric-row">
  <div class="metric"><strong>{html.escape(str(data.get('meeting_count') or 0))}</strong><span>去重会议纪要</span></div>
  <div class="metric"><strong>{html.escape(str(len(data.get('high_frequency_targets') or [])))}</strong><span>高频股票</span></div>
  <div class="metric"><strong>{html.escape(str(len(data.get('speaker_rollup') or {})))}</strong><span>多会议发言人</span></div>
  <div class="metric"><strong>{html.escape(type_text)}</strong><span>会议类型</span></div>
</div>
<div class="split">
  <section>
    <div class="section-head"><div><h2>当日高频股票与逻辑</h2><p>点击来源可回到正式纪要。</p></div></div>
    {''.join(target_rows) if target_rows else '<div class="item"><strong>暂无高频股票</strong><p class="muted">当日暂无确认归档纪要，或纪要中未抽取到标的。</p><div class="bar"><a href="/apps/meeting-minutes/import">去整理会议</a></div></div>'}
  </section>
  <section>
    <div class="section-head"><div><h2>同发言人观点汇总</h2><p>同一发言人跨多场会议的观点线索。</p></div></div>
    {''.join(speaker_blocks) if speaker_blocks else '<div class="item"><strong>暂无多会议汇总</strong><p class="muted">当日未发现同发言人多条观点。</p></div>'}
  </section>
</div>
<div class="section-head"><div><h2>当日去重会议纪要</h2><p>只展示正式归档版本。</p></div></div>
<table><thead><tr><th>日期</th><th>类型</th><th>标题</th><th>标的</th></tr></thead><tbody>{''.join(minute_rows) if minute_rows else '<tr><td colspan="4" class="muted">暂无正式纪要。</td></tr>'}</tbody></table>
<div class="section-head"><div><h2>可编辑数据展示窗口</h2><p>用于临时整理看板摘录，内容仅保存在浏览器本地。</p></div></div>
<textarea class="data-window" id="dashboard-data-window" placeholder="粘贴或整理看板数据"></textarea>
<script>
const dashboardDataWindow = document.getElementById("dashboard-data-window");
if (dashboardDataWindow) {{
  dashboardDataWindow.value = localStorage.getItem("d91-dashboard-data-window") || "";
  dashboardDataWindow.addEventListener("input", () => localStorage.setItem("d91-dashboard-data-window", dashboardDataWindow.value));
}}
</script>
"""
        self._send(200, html_page("投研数据看板", body), "text/html; charset=utf-8")

    def _handle_import_form(self) -> None:
        today = datetime.now().strftime("%Y-%m-%d")
        body = f"""
{workflow_stage_header("导入材料", "上传材料后先生成输入校对稿，不直接写入正式纪要。")}
<div class="stepper">
  <div class="step active"><strong>1 填写信息</strong><span class="muted">标题、日期、类型、系列</span></div>
  <div class="step"><strong>2 上传材料</strong><span class="muted">文本、文档或音频</span></div>
  <div class="step"><strong>3 初步校对</strong><span class="muted">标注发言人、分段、修正关键名词</span></div>
  <div class="step"><strong>4 二次校对</strong><span class="muted">确认发言人、标的和关键名词</span></div>
</div>
<div class="meta">
  <strong>先建立输入校对稿</strong>
  <div class="muted">第一步只上传文本、Word 文档或音频。系统会先抽取或转录成文字，生成可人工校对的输入稿；初步校对要标注不同发言人、按自然主题分段，关键名词转录错误也可以直接修改。</div>
</div>
<div class="import-layout">
<form method="post" action="/apps/meeting-minutes/import" enctype="multipart/form-data" class="item" data-processing-title="正在导入会议材料" data-processing-message="系统正在抽取文档或转录音频，并生成可校对输入稿。请不要重复提交。">
  <div class="input-grid">
    <label class="field"><span>会议标题</span><input name="meeting_title" id="meeting-title-input" placeholder="选择文件后自动预填，也可以手动修改"></label>
    <label class="field"><span>会议日期</span><input name="meeting_date" id="meeting-date-input" type="date" value="{html.escape(today)}"></label>
    <label class="field"><span>会议类型</span><select name="meeting_type" id="meeting-type-select">{meeting_type_options("多人复盘会")}</select></label>
    <label class="field"><span>会议系列</span><select name="meeting_series" id="meeting-series-select" data-custom-input="custom-meeting-series-input" data-custom-value="{CUSTOM_CHOICE_VALUE}">{meeting_series_options(MEETING_SERIES_DEFAULT)}</select><input name="custom_meeting_series" id="custom-meeting-series-input" placeholder="填写新增会议系列" hidden></label>
  </div>
  <div class="upload-zone">
    <strong>上传会议材料</strong>
    <p class="muted">支持一次选择或拖入多个文本、Markdown、Word 文档和常见音频文件。系统会逐个保存并抽取或转录，合并生成一份输入校对稿。</p>
    <input type="file" name="meeting_files" id="meeting-files-input" accept=".txt,.md,.markdown,.docx,.mp3,.m4a,.wav,.aac,.flac,.ogg,.wma,.mp4" multiple required>
    <p class="muted" id="upload-smart-hint">尚未选择文件。</p>
    <div class="file-list" id="meeting-files-list" aria-live="polite"></div>
  </div>
  <label class="field">
    <span>补充说明</span>
    <textarea name="operator_notes" placeholder="可选：写给校对阶段看的提示，例如发言人名单、需要重点核对的公司名、股票代码、数字或转录易错词。"></textarea>
  </label>
  <div class="bar"><button type="submit">上传并生成初步校对稿</button><a class="secondary" href="/apps/meeting-minutes">返回</a></div>
</form>
<aside class="process-panel">
  <ol>
    <li><span class="process-index">1</span><div><b>上传材料</b><br><span>文本、Word、音频都从这里进入。</span></div></li>
    <li><span class="process-index">2</span><div><b>抽取或转录</b><br><span>系统生成可校对文本，不直接产出纪要。</span></div></li>
    <li><span class="process-index">3</span><div><b>人工初校</b><br><span>标注不同发言人，拆分段落，并修正关键名词。</span></div></li>
    <li><span class="process-index">4</span><div><b>整理与二校</b><br><span>确认发言人、标的、关键名词等信息无误后再归档。</span></div></li>
  </ol>
</aside>
</div>
"""
        self._send(200, html_page("OpenMinute-会议纪要", body), "text/html; charset=utf-8")

    def _handle_import_post(self) -> None:
        payload = self._read_payload()
        upload_files = payload.get("_files") if isinstance(payload.get("_files"), list) else []
        if not upload_files:
            raise ValueError("请先上传文本、Word 文档或音频文件")
        inferred_title, inferred_type, inferred_series = infer_meeting_metadata_from_uploads(upload_files)
        title = sanitize_title(str(payload.get("meeting_title") or inferred_title or "投资会议纪要"))
        meeting_date = normalize_meeting_date(str(payload.get("meeting_date") or ""))
        selected_meeting_type_raw = str(payload.get("meeting_type") or "")
        selected_meeting_type = normalize_meeting_type(selected_meeting_type_raw)
        if selected_meeting_type in {"", "多人复盘会"} and inferred_type != "多人复盘会":
            meeting_type = inferred_type
        else:
            meeting_type = normalize_meeting_type(selected_meeting_type or inferred_type or "多人复盘会")
        selected_meeting_series = str(payload.get("meeting_series") or "")
        if selected_meeting_series == CUSTOM_CHOICE_VALUE:
            selected_meeting_series = ""
        meeting_series = normalize_meeting_series(
            str(payload.get("custom_meeting_series") or selected_meeting_series or inferred_series or MEETING_SERIES_DEFAULT)
        )
        draft_id = str(uuid4())
        operator_notes = str(payload.get("operator_notes") or "").strip()
        status_path = f"/apps/meeting-minutes/import/status?id={urllib.parse.quote(draft_id)}&date={urllib.parse.quote(meeting_date)}"
        write_import_job_status(
            draft_id,
            meeting_date,
            {
                "ok": True,
                "status": "queued",
                "stage": "等待导入",
                "percent": 3,
                "message": "文件已收到，正在准备后台导入任务",
                "title": title,
                "meeting_type": meeting_type,
                "meeting_series": meeting_series,
                "file_count": len(upload_files),
                "status_url": status_path,
            },
        )
        job = {
            "draft_id": draft_id,
            "meeting_date": meeting_date,
            "title": title,
            "meeting_type": meeting_type,
            "meeting_series": meeting_series,
            "files": upload_files,
            "operator_notes": operator_notes,
        }
        thread = threading.Thread(target=run_import_job, args=(job,), daemon=True)
        thread.start()
        if self._wants_json_response():
            self._json(
                202,
                {
                    "ok": True,
                    "draft_id": draft_id,
                    "meeting_date": meeting_date,
                    "status": "queued",
                    "status_url": status_path,
                },
            )
        else:
            self._redirect(status_path)

    def _handle_import_status_page(self, query: dict[str, list[str]]) -> None:
        draft_id = (query.get("id") or [""])[0]
        meeting_date = normalize_meeting_date((query.get("date") or [""])[0])
        if not draft_id:
            self._send(400, html_page("导入状态", '<p class="error">缺少导入任务编号。</p>'), "text/html; charset=utf-8")
            return
        safe_id = html.escape(draft_id)
        body = f"""
{workflow_stage_header("导入进度", "上传请求已完成，抽取、转录和文本整理在后台继续执行。")}
<div class="stepper">
  <div class="step active"><strong>1 上传完成</strong><span class="muted">文件已进入后台任务</span></div>
  <div class="step active"><strong>2 抽取/转录</strong><span class="muted">按文件逐个处理</span></div>
  <div class="step"><strong>3 文本整理</strong><span class="muted">合并为输入校对稿</span></div>
  <div class="step"><strong>4 初步校对</strong><span class="muted">完成后自动打开</span></div>
</div>
<section class="item">
  <h2 id="import-stage">正在准备导入</h2>
  <p class="muted" id="import-message">文件已经上传完成，后台任务正在启动。</p>
  <div class="progress-head"><span id="import-percent">3%</span><span id="import-status">等待导入</span></div>
  <div class="progress-track" aria-hidden="true"><div class="progress-fill" id="import-fill" style="width: 3%"></div></div>
  <div class="progress-steps" id="import-steps">
    <div class="progress-step active">保存上传材料</div>
    <div class="progress-step">读取文档或音频转录</div>
    <div class="progress-step">合并生成输入校对稿</div>
    <div class="progress-step">打开人工校对页</div>
  </div>
  <p class="muted">任务编号：{safe_id}</p>
  <div class="bar"><a class="secondary" href="/drafts">查看待校对草稿</a></div>
</section>
<script>
(function() {{
  const draftId = {json.dumps(draft_id, ensure_ascii=False)};
  const meetingDate = {json.dumps(meeting_date, ensure_ascii=False)};
  const stageEl = document.getElementById("import-stage");
  const messageEl = document.getElementById("import-message");
  const percentEl = document.getElementById("import-percent");
  const statusEl = document.getElementById("import-status");
  const fillEl = document.getElementById("import-fill");
  const steps = Array.from(document.querySelectorAll("#import-steps .progress-step"));

  function setSteps(percent) {{
    const activeIndex = percent >= 96 ? 3 : percent >= 84 ? 2 : percent >= 12 ? 1 : 0;
    steps.forEach((step, index) => {{
      step.className = "progress-step" + (index < activeIndex ? " done" : index === activeIndex ? " active" : "");
    }});
  }}

  async function poll() {{
    try {{
      const url = "/api/import-status?id=" + encodeURIComponent(draftId) + "&date=" + encodeURIComponent(meetingDate);
      const response = await fetch(url, {{headers: {{"Accept": "application/json"}}}});
      const payload = await response.json();
      const data = payload.data || payload;
      const percent = Math.max(1, Math.min(100, Number(data.percent || 3)));
      stageEl.textContent = data.stage || "导入中";
      messageEl.textContent = data.message || "后台任务正在处理。";
      statusEl.textContent = data.status === "done" ? "已完成" : data.status === "failed" ? "失败" : "处理中";
      percentEl.textContent = Math.round(percent) + "%";
      fillEl.style.width = percent + "%";
      setSteps(percent);
      if (data.status === "done") {{
        const reviewUrl = data.result && data.result.review_url ? data.result.review_url : "/review?id=" + encodeURIComponent(draftId);
        window.setTimeout(() => {{ window.location.href = reviewUrl; }}, 800);
        return;
      }}
      if (data.status === "failed") {{
        messageEl.textContent = (data.error || data.message || "导入失败，请返回上传页重新尝试。");
        return;
      }}
    }} catch (error) {{
      messageEl.textContent = "暂时读取不到进度，后台任务可能仍在处理。";
    }}
    window.setTimeout(poll, 1600);
  }}
  poll();
}})();
</script>
"""
        self._send(200, html_page("OpenMinute-导入进度", body), "text/html; charset=utf-8")

    def _handle_confirm_input_status_page(self, query: dict[str, list[str]]) -> None:
        draft_id = (query.get("id") or [""])[0]
        if not draft_id:
            self._send(400, html_page("整理状态", '<p class="error">缺少整理任务编号。</p>'), "text/html; charset=utf-8")
            return
        safe_id = html.escape(draft_id)
        body = f"""
{workflow_stage_header("会议整理进度", "人工初校已提交，Dify 会议纪要整理在后台继续执行。")}
<div class="stepper">
  <div class="step active"><strong>1 初校已提交</strong><span class="muted">输入稿已进入后台任务</span></div>
  <div class="step active"><strong>2 Dify 整理</strong><span class="muted">生成会议纪要草稿</span></div>
  <div class="step"><strong>3 二次校对</strong><span class="muted">完成后自动打开</span></div>
  <div class="step"><strong>4 归档确认</strong><span class="muted">二校后执行</span></div>
</div>
<section class="item">
  <h2 id="confirm-stage">正在准备整理</h2>
  <p class="muted" id="confirm-message">会议整理任务已经提交，后台任务正在启动。</p>
  <div class="progress-head"><span id="confirm-percent">3%</span><span id="confirm-status">等待整理</span></div>
  <div class="progress-track" aria-hidden="true"><div class="progress-fill" id="confirm-fill" style="width: 3%"></div></div>
  <div class="progress-steps" id="confirm-steps">
    <div class="progress-step active">保存初校稿</div>
    <div class="progress-step">提交 Dify 工作流</div>
    <div class="progress-step">生成二次校对稿</div>
    <div class="progress-step">打开二次校对页</div>
  </div>
  <p class="muted">任务编号：{safe_id}</p>
  <div class="bar"><a class="secondary" href="/drafts">查看待校对草稿</a></div>
</section>
<script>
(function() {{
  const draftId = {json.dumps(draft_id, ensure_ascii=False)};
  const stageEl = document.getElementById("confirm-stage");
  const messageEl = document.getElementById("confirm-message");
  const percentEl = document.getElementById("confirm-percent");
  const statusEl = document.getElementById("confirm-status");
  const fillEl = document.getElementById("confirm-fill");
  const steps = Array.from(document.querySelectorAll("#confirm-steps .progress-step"));

  function setSteps(percent) {{
    const activeIndex = percent >= 96 ? 3 : percent >= 35 ? 2 : percent >= 12 ? 1 : 0;
    steps.forEach((step, index) => {{
      step.className = "progress-step" + (index < activeIndex ? " done" : index === activeIndex ? " active" : "");
    }});
  }}

  async function poll() {{
    try {{
      const url = "/api/confirm-input-status?id=" + encodeURIComponent(draftId);
      const response = await fetch(url, {{headers: {{"Accept": "application/json"}}}});
      const payload = await response.json();
      const data = payload.data || payload;
      const percent = Math.max(1, Math.min(100, Number(data.percent || 3)));
      stageEl.textContent = data.stage || "整理中";
      messageEl.textContent = data.message || "Dify 工作流正在处理。";
      statusEl.textContent = data.status === "done" ? "已完成" : data.status === "failed" ? "失败" : "处理中";
      percentEl.textContent = Math.round(percent) + "%";
      fillEl.style.width = percent + "%";
      setSteps(percent);
      if (data.status === "done") {{
        const reviewUrl = data.next_review_url || (data.result && data.result.next_review_url) || "";
        window.setTimeout(() => {{ window.location.href = reviewUrl || "/drafts"; }}, 800);
        return;
      }}
      if (data.status === "failed") {{
        messageEl.textContent = data.error || data.message || "整理失败，请返回草稿页检查。";
        return;
      }}
    }} catch (error) {{
      messageEl.textContent = "暂时读取不到进度，后台任务可能仍在处理。";
    }}
    window.setTimeout(poll, 1600);
  }}
  poll();
}})();
</script>
"""
        self._send(200, html_page("OpenMinute-整理进度", body), "text/html; charset=utf-8")

    def _handle_import_status_api(self, query: dict[str, list[str]]) -> None:
        draft_id = (query.get("id") or [""])[0]
        meeting_date = normalize_meeting_date((query.get("date") or [""])[0])
        if not draft_id:
            self._json(400, {"ok": False, "error": "缺少导入任务编号"})
            return
        status = load_import_job_status(draft_id, meeting_date)
        self._json(200 if status.get("ok", True) else 500, {"ok": bool(status.get("ok", True)), "data": status})

    def _handle_confirm_input_status_api(self, query: dict[str, list[str]]) -> None:
        draft_id = (query.get("id") or [""])[0]
        if not draft_id:
            self._json(400, {"ok": False, "error": "缺少整理任务编号"})
            return
        status = load_confirm_input_job_status(draft_id)
        self._json(200 if status.get("ok", True) else 500, {"ok": bool(status.get("ok", True)), "data": status})

    def _handle_import_upload_start(self) -> None:
        started_at = time.monotonic()
        payload = self._read_payload()
        manifest = create_import_upload_session(payload)
        files = manifest.get("files") if isinstance(manifest.get("files"), list) else []
        total_size = sum(int(item.get("size") or 0) for item in files)
        self.log_message(
            "import_upload_start upload_id=%s draft_id=%s files=%s total_bytes=%s elapsed=%.3fs",
            manifest.get("upload_id"),
            manifest.get("draft_id"),
            len(files),
            total_size,
            time.monotonic() - started_at,
        )
        self._json(
            202,
            {
                "ok": True,
                "upload_id": manifest.get("upload_id"),
                "draft_id": manifest.get("draft_id"),
                "meeting_date": manifest.get("meeting_date"),
                "status_url": manifest.get("status_url"),
                "chunk_size": MAX_IMPORT_CHUNK_BYTES,
            },
        )

    def _handle_import_upload_chunk(self, query: dict[str, list[str]]) -> None:
        started_at = time.monotonic()
        upload_id = (query.get("upload_id") or [""])[0]
        file_index = int((query.get("file_index") or ["0"])[0] or "0")
        offset = int((query.get("offset") or ["0"])[0] or "0")
        data, read_stats = self._read_binary_body()
        result = append_import_upload_chunk(upload_id, file_index, offset, data)
        self.log_message(
            "import_upload_chunk upload_id=%s file_index=%s offset=%s declared_bytes=%s received_bytes=%s read_seconds=%s stored_received=%s duplicate=%s elapsed=%.3fs",
            upload_id,
            file_index,
            offset,
            read_stats["declared_bytes"],
            read_stats["received_bytes"],
            read_stats["read_seconds"],
            result.get("received"),
            bool(result.get("duplicate")),
            time.monotonic() - started_at,
        )
        self._json(200, result)

    def _handle_import_upload_complete(self) -> None:
        started_at = time.monotonic()
        payload = self._read_payload()
        upload_id = str(payload.get("upload_id") or "")
        result = complete_import_upload_session(upload_id)
        self.log_message(
            "import_upload_complete upload_id=%s draft_id=%s elapsed=%.3fs",
            upload_id,
            result.get("draft_id"),
            time.monotonic() - started_at,
        )
        self._json(202, result)

    def _handle_access_login(self, query: dict[str, list[str]], error: str = "") -> None:
        ensure_access_config()
        next_path = (query.get("next") or ["/history"])[0] or "/history"
        if not next_path.startswith("/"):
            next_path = "/history"
        error_html = f'<p class="error">{html.escape(error)}</p>' if error else ""
        body = f"""
<h1 class="page-title">访问登录</h1>
<div class="meta">
  <strong>请输入用户名和密码</strong>
  <div class="muted">公网查看历史纪要、外部资料、下载文件或标的证据时需要登录。本机 localhost 访问默认视为管理员。</div>
</div>
{error_html}
<form class="filters item" method="post" action="/access-login">
  <input type="hidden" name="next" value="{html.escape(next_path)}">
  <input name="user_id" placeholder="用户名" autocomplete="username" required>
  <input name="password" type="password" placeholder="密码" autocomplete="current-password" required>
  <button type="submit">登录</button>
</form>
"""
        self._send(200, html_page("访问登录", body, {"authenticated": False}), "text/html; charset=utf-8")

    def _handle_access_login_post(self) -> None:
        payload = self._read_payload()
        user_id = normalize_login_user_id(str(payload.get("user_id") or ""))
        password = str(payload.get("password") or "")
        token = str(payload.get("token") or "").strip()
        next_path = str(payload.get("next") or "/")
        if not next_path.startswith("/"):
            next_path = "/"
        client_key = self._client_key()
        retry_after = login_retry_after_seconds(client_key, user_id)
        if retry_after:
            message = f"登录失败次数过多，请 {max(1, retry_after // 60)} 分钟后再试。"
            self._audit("login", {"authenticated": False, "user_id": user_id or "anonymous", "role": "anonymous", "groups": [], "source": "anonymous"}, "rate_limited")
            if self._wants_json_response():
                self._json(429, {"ok": False, "error": message, "retry_after": retry_after})
            else:
                self._handle_access_login({"next": ["/"]}, message)
            return
        cookie_value = ""
        identity = validate_access_password(user_id, password)
        if identity:
            cookie_value = create_access_session(str(identity.get("user_id") or user_id))
        elif token:
            identity = validate_access_token(token)
            cookie_value = token if identity else ""
        if not identity:
            record_login_failure(client_key, user_id)
            self._audit("login", {"authenticated": False, "user_id": user_id or "anonymous", "role": "anonymous", "groups": [], "source": "anonymous"}, "failed")
            if self._wants_json_response():
                self._json(401, {"ok": False, "error": "登录失败，请检查用户名和密码。"})
            else:
                self._handle_access_login({"next": ["/"]}, "登录失败，请检查用户名和密码。")
            return
        clear_login_failures(client_key, user_id)
        self._audit("login", identity, "ok")
        if self._wants_json_response():
            self.send_response(200)
            self._set_access_cookie(cookie_value)
            data = json.dumps({"ok": True, "redirect": "/"}, ensure_ascii=False).encode("utf-8")
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.send_header("Content-Length", str(len(data)))
            self._security_headers()
            self.end_headers()
            self.wfile.write(data)
            return
        self.send_response(303)
        self._set_access_cookie(cookie_value)
        self.send_header("Location", "/")
        self.send_header("Content-Length", "0")
        self._security_headers()
        self.end_headers()

    def _handle_access_logout(self) -> None:
        self._audit("logout", self._request_identity(), "ok")
        self.send_response(303)
        self._clear_access_cookie()
        self.send_header("Location", "/access-login")
        self.send_header("Content-Length", "0")
        self._security_headers()
        self.end_headers()

    def do_GET(self) -> None:
        parsed = urllib.parse.urlparse(self.path)
        query = urllib.parse.parse_qs(parsed.query)
        if parsed.path == "/dashboards":
            self._redirect("/mp")
            return
        if self._is_platform_proxy_path(parsed.path):
            identity = None
            if self._platform_proxy_requires_auth(parsed.path):
                identity = self._require_authenticated(query)
                if identity is None:
                    return
            self._proxy_platform(identity)
            return
        if parsed.path in ("", "/", "/portal"):
            self._handle_home()
            return
        if parsed.path in ("/apps", "/tools", "/workflows"):
            self._handle_apps(self._request_identity())
            return
        if parsed.path == "/knowledge":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            self._handle_knowledge(identity, query)
            return
        if parsed.path in ("/data", "/operations", "/admin"):
            if self._require_admin() is None:
                return
            self._handle_operations()
            return
        if parsed.path == "/account":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            self._handle_account(identity)
            return
        if parsed.path == "/apps/meeting-minutes":
            if self._require_authenticated(query) is None:
                return
            self._handle_meeting_app()
            return
        if parsed.path == "/apps/meeting-minutes/import":
            if self._require_editor() is None:
                return
            self._handle_import_form()
            return
        if parsed.path == "/apps/meeting-minutes/import/status":
            if self._require_editor() is None:
                return
            self._handle_import_status_page(query)
            return
        if parsed.path == "/confirm-input/status":
            if self._require_editor() is None:
                return
            self._handle_confirm_input_status_page(query)
            return
        if parsed.path == "/mp":
            self._redirect("/external-sources")
            return
        if parsed.path == "/health":
            self._json(200, {"ok": True, "service": "meeting-minutes-review"})
            return
        if parsed.path == "/access-login":
            self._handle_access_login(query)
            return
        if parsed.path == "/access-logout":
            self._handle_access_logout()
            return
        if parsed.path == "/review":
            if self._require_editor() is None:
                return
            self._handle_review(query)
            return
        if parsed.path == "/history":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            self._handle_history(query, identity)
            return
        if parsed.path == "/drafts":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            self._handle_drafts(query, identity)
            return
        if parsed.path == "/external-sources":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            self._handle_external_sources(query, identity)
            return
        if parsed.path == "/external-view":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            self._handle_external_view(query, identity)
            return
        if parsed.path == "/agent-status":
            if self._require_authenticated(query) is None:
                return
            self._handle_agent_status()
            return
        if parsed.path == "/sync-status":
            if self._require_authenticated(query) is None:
                return
            self._handle_sync_status()
            return
        if parsed.path == "/access-audit":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            self._handle_access_audit(query, identity)
            return
        if parsed.path == "/access-users":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            self._handle_access_users(identity)
            return
        if parsed.path == "/mapping-audit":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            self._handle_mapping_audit(identity)
            return
        if parsed.path == "/health-report":
            if self._require_authenticated(query) is None:
                return
            self._handle_health_report()
            return
        if parsed.path == "/dashboard":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            self._handle_dashboard(query, identity)
            return
        if parsed.path == "/target-query":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            self._handle_target_query(query, identity)
            return
        if parsed.path == "/mcp-servers":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            self._handle_mcp_servers(identity)
            return
        if parsed.path == "/view":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            self._handle_view(query, identity)
            return
        if parsed.path == "/result":
            if self._require_editor() is None:
                return
            self._handle_result(query)
            return
        if parsed.path == "/latest":
            if self._require_authenticated(query) is None:
                return
            self._handle_latest()
            return
        if parsed.path == "/download":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            self._handle_download(query, identity)
            return
        if parsed.path == "/api/history":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            items = filter_history(filter_accessible_items(scan_history(), identity), query)
            for item in items:
                item.pop("content", None)
            self._audit("api_history", identity, "ok", result_count=len(items), filters={k: v[-1] for k, v in query.items() if v})
            self._json(200, {"ok": True, "data": items})
            return
        if parsed.path == "/api/drafts":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            role = str(identity.get("role") or "").lower()
            if role not in {"admin", "editor"}:
                self._deny_access("只有管理员或编辑可以查看待校对草稿。")
                return
            drafts = filter_drafts(scan_drafts(), query)
            status_filter = (query.get("status") or ["open"])[0].strip().lower() or "open"
            self._audit("api_drafts", identity, "ok", result_count=len(drafts), status_filter=status_filter)
            self._json(200, {"ok": True, "data": drafts})
            return
        if parsed.path == "/api/import-status":
            if self._require_editor() is None:
                return
            self._handle_import_status_api(query)
            return
        if parsed.path == "/api/confirm-input-status":
            if self._require_editor() is None:
                return
            self._handle_confirm_input_status_api(query)
            return
        if parsed.path == "/api/external-sources":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            items = filter_external_sources(filter_accessible_items(scan_external_sources(), identity), query)
            for item in items:
                item.pop("content", None)
            self._audit("api_external_sources", identity, "ok", result_count=len(items), filters={k: v[-1] for k, v in query.items() if v})
            self._json(200, {"ok": True, "data": items})
            return
        if parsed.path == "/api/agent-status":
            if self._require_authenticated(query) is None:
                return
            self._json(200, {"ok": True, "data": load_planner_status()})
            return
        if parsed.path == "/api/sync-status":
            if self._require_authenticated(query) is None:
                return
            self._json(200, {"ok": True, "data": google_drive_sync_status()})
            return
        if parsed.path == "/api/access-audit":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            if not identity.get("is_admin"):
                self._deny_access("只有管理员可以查看访问审计日志。")
                return
            limit = int((query.get("limit") or ["100"])[0] or "100")
            events = load_access_audit_events(limit=max(1, min(limit, 500)))
            self._audit("api_access_audit", identity, "ok", result_count=len(events))
            self._json(200, {"ok": True, "data": events})
            return
        if parsed.path == "/api/access-users":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            if not identity.get("is_admin"):
                self._deny_access("只有管理员可以查看访问用户。")
                return
            users = public_access_users()
            self._audit("api_access_users", identity, "ok", user_count=len(users))
            self._json(200, {"ok": True, "data": users})
            return
        if parsed.path == "/api/mapping-audit":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            if not identity.get("is_admin"):
                self._deny_access("只有管理员可以查看知识库映射审计。")
                return
            report = load_mapping_audit_report()
            self._audit("api_mapping_audit", identity, "ok", issue_count=report.get("issue_count"))
            self._json(200, {"ok": True, "data": report})
            return
        if parsed.path == "/api/health-report":
            if self._require_authenticated(query) is None:
                return
            self._json(200, {"ok": True, "data": local_health_report()})
            return
        if parsed.path == "/api/dashboard":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            selected_date = (query.get("date") or [datetime.now().strftime("%Y-%m-%d")])[0]
            self._json(200, {"ok": True, "data": dashboard_data(selected_date, identity)})
            return
        if parsed.path == "/api/target-query":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            query_text = (query.get("q") or [""])[0]
            result = target_query(query_text, identity=identity)
            self._audit("api_target_query", identity, "ok", query=query_text, result_count=result.get("result_count"))
            self._json(200, {"ok": True, "data": result})
            return
        if parsed.path == "/api/mcp-servers":
            identity = self._require_authenticated(query)
            if identity is None:
                return
            self._audit("api_mcp_servers", identity, "ok")
            self._json(
                200,
                {
                    "ok": True,
                    "data": load_mcp_config(),
                    "config_path": str(DEFAULT_MCP_CONFIG_PATH),
                    "call_url": "http://host.docker.internal:8767/api/mcp/call",
                },
            )
            return
        self._json(404, {"ok": False, "error": "not found"})

    def do_POST(self) -> None:
        parsed = urllib.parse.urlparse(self.path)
        query = urllib.parse.parse_qs(parsed.query)
        try:
            if not self._check_post_origin():
                return
            if parsed.path in ("/draft", "/prepare-review"):
                if is_public_host(self.headers.get("X-Forwarded-Host") or self.headers.get("Host") or ""):
                    self._deny_access("公网入口不接受直接创建草稿请求。")
                    return
                self._json(200, create_draft(self._read_payload()))
                return
            if parsed.path == "/access-login":
                self._handle_access_login_post()
                return
            if parsed.path == "/access-users":
                identity = self._require_authenticated(query)
                if identity is None:
                    return
                self._handle_access_users_post(identity)
                return
            if parsed.path == "/apps/meeting-minutes/import":
                if self._require_editor() is None:
                    return
                self._handle_import_post()
                return
            if parsed.path == "/api/import-upload/start":
                if self._require_editor() is None:
                    return
                self._handle_import_upload_start()
                return
            if parsed.path == "/api/import-upload/chunk":
                if self._require_editor() is None:
                    return
                self._handle_import_upload_chunk(query)
                return
            if parsed.path == "/api/import-upload/complete":
                if self._require_editor() is None:
                    return
                self._handle_import_upload_complete()
                return
            if parsed.path == "/dify/ingest-files":
                if is_public_host(self.headers.get("X-Forwarded-Host") or self.headers.get("Host") or ""):
                    self._deny_access("公网入口不接受直接执行 Dify 文件抽取请求。")
                    return
                payload = self._read_payload()
                files = payload.get("_files") if isinstance(payload.get("_files"), list) else []
                self._json(200, ingest_dify_file_list(files, asr_model_choice=str(payload.get("asr_model") or "")))
                return
            if parsed.path == "/dify/extract-files":
                if is_public_host(self.headers.get("X-Forwarded-Host") or self.headers.get("Host") or ""):
                    self._deny_access("公网入口不接受直接执行 Dify 文件抽取请求。")
                    return
                payload = self._read_payload()
                files = payload.get("_files") if isinstance(payload.get("_files"), list) else []
                self._json(200, ingest_dify_file_list(files, mode="document"))
                return
            if parsed.path == "/dify/transcribe-files":
                if is_public_host(self.headers.get("X-Forwarded-Host") or self.headers.get("Host") or ""):
                    self._deny_access("公网入口不接受直接执行 Dify 语音识别请求。")
                    return
                payload = self._read_payload()
                files = payload.get("_files") if isinstance(payload.get("_files"), list) else []
                self._json(
                    200,
                    ingest_dify_file_list(
                        files,
                        mode="audio",
                        asr_model_choice=str(payload.get("asr_model") or ""),
                    ),
                )
                return
            if parsed.path == "/api/mcp/call":
                if is_public_host(self.headers.get("X-Forwarded-Host") or self.headers.get("Host") or ""):
                    self._deny_access("公网入口不接受直接执行 MCP bridge 请求。")
                    return
                payload = self._read_payload()
                if not isinstance(payload, dict):
                    payload = {}
                self._json(200, mcp_bridge_call(payload))
                return
            if parsed.path == "/dify/dispatch":
                if is_public_host(self.headers.get("X-Forwarded-Host") or self.headers.get("Host") or ""):
                    self._deny_access("公网入口不接受直接执行 Dify 回调请求。")
                    return
                payload = self._read_payload()
                action = str(payload.get("action") or "").strip()
                if action == "final_review_confirmed":
                    draft_id = str(payload.get("draft_id") or "").strip()
                    markdown = str(payload.get("markdown") or "")
                    result = execute_confirm_draft(
                        draft_id,
                        markdown,
                        workflow_run_id=str(payload.get("workflow_run_id") or ""),
                    )
                    self._json(200, result)
                    return
                if action == "archive_confirmed":
                    draft_id = str(payload.get("draft_id") or "").strip()
                    result = execute_confirm_archive(
                        draft_id,
                        workflow_run_id=str(payload.get("workflow_run_id") or ""),
                    )
                    self._json(200, result)
                    return
                try:
                    self._json(200, create_draft(payload))
                except ValueError as exc:
                    self._json(
                        200,
                        {
                            "ok": False,
                            "error": str(exc),
                            "message": str(exc),
                            "draft_id": str(payload.get("draft_id") or ""),
                            "workflow_run_id": str(payload.get("workflow_run_id") or ""),
                        },
                    )
                return
            if parsed.path == "/save":
                if self._require_editor() is None:
                    return
                draft_id = (query.get("id") or [""])[0]
                payload = self._read_payload()
                self._json(200, save_current(draft_id, str(payload.get("markdown") or "")))
                return
            if parsed.path == "/confirm":
                identity = self._require_editor()
                if identity is None:
                    return
                draft_id = (query.get("id") or [""])[0]
                payload = self._read_payload()
                markdown = str(payload.get("markdown") or "") or None
                result = confirm_draft(draft_id, markdown)
                self._audit("confirm_draft", identity, "ok" if result.get("ok") else "failed", draft_id=draft_id, draft_status=result.get("status"))
                if self._wants_json_response():
                    self._json(200, result)
                else:
                    self._redirect(f"/result?id={urllib.parse.quote(draft_id)}")
                return
            if parsed.path == "/confirm-archive":
                identity = self._require_editor()
                if identity is None:
                    return
                draft_id = (query.get("id") or [""])[0]
                result = confirm_archive(draft_id)
                self._audit("confirm_archive", identity, "ok" if result.get("ok") else "failed", draft_id=draft_id, draft_status=result.get("status"))
                if self._wants_json_response():
                    self._json(200, result)
                else:
                    self._redirect(f"/result?id={urllib.parse.quote(draft_id)}")
                return
            if parsed.path == "/confirm-table":
                identity = self._require_editor()
                if identity is None:
                    return
                draft_id = (query.get("id") or [""])[0]
                payload = self._read_payload()
                result = confirm_structured_table(draft_id, str(payload.get("structured_table") or ""))
                self._audit("confirm_structured_table", identity, "ok" if result.get("ok") else "failed", draft_id=draft_id, draft_status=result.get("status"))
                if self._wants_json_response():
                    self._json(200, result)
                else:
                    self._redirect(f"/result?id={urllib.parse.quote(draft_id)}")
                return
            if parsed.path == "/confirm-input":
                identity = self._require_editor()
                if identity is None:
                    return
                draft_id = (query.get("id") or [""])[0]
                payload = self._read_payload()
                markdown = str(payload.get("markdown") or "") or None
                result = start_confirm_input_job(draft_id, markdown)
                self._audit("confirm_input_draft", identity, "ok" if result.get("ok") else "failed", draft_id=draft_id, draft_status=result.get("status"))
                if self._wants_json_response():
                    self._json(202 if result.get("status") in {"queued", "running"} else 200, result)
                else:
                    self._redirect(normalize_local_url(str(result.get("status_url") or "")) or f"/confirm-input/status?id={urllib.parse.quote(draft_id)}")
                return
            self._json(404, {"ok": False, "error": "not found"})
        except Exception as exc:
            self.log_message("request failed: %s\n%s", exc, traceback.format_exc())
            self._json(500, {"ok": False, "error": str(exc)})

    def _handle_review(self, query: dict[str, list[str]]) -> None:
        draft_id = (query.get("id") or [""])[0]
        folder = draft_folder(draft_id)
        if not draft_id or not folder.exists():
            self._send(404, html_page("未找到校对稿", "<p class=\"error\">未找到校对稿。</p>"), "text/html; charset=utf-8")
            return
        meta = load_meta(folder)
        markdown = current_markdown(folder)
        stage = str(meta.get("stage") or "post_agent")
        meeting_type = html.escape(str(meta.get("meeting_type") or markdown_field(markdown, "会议类型", "多人复盘会")))
        confirm_action = "/confirm-input" if stage == "pre_agent" else "/confirm"
        confirm_label = "确认文本，开始生成会议纪要" if stage == "pre_agent" else "确认终稿，进入归档确认"
        gate_text = (
            "这是根据上传材料抽取/转录出的初校文本。可以直接编辑正文，修正发言人、段落、关键名词、股票代码和数字；确认后才会开始生成会议纪要。"
            if stage == "pre_agent"
            else "这是整理后的终稿校对页。可以直接改正文；确认后进入归档确认页，最终写入正式库仍需要再次确认。"
        )
        helper_text = (
            "直接在这里改初校文本。改完点击“确认文本，开始生成会议纪要”，不需要填写任何确认字段。"
            if stage == "pre_agent"
            else "直接在这里改终稿。改完点击“确认终稿，进入归档确认”。"
        )
        stepper = (
            """<div class="stepper">
  <div class="step active"><strong>1 初步校对</strong><span class="muted">检查转录文本</span></div>
  <div class="step"><strong>2 会议整理</strong><span class="muted">确认后自动执行</span></div>
  <div class="step"><strong>3 二次校对</strong><span class="muted">稍后确认终稿</span></div>
  <div class="step"><strong>4 归档确认</strong><span class="muted">最后写入正式库</span></div>
</div>"""
            if stage == "pre_agent"
            else """<div class="stepper">
  <div class="step"><strong>1 初步校对</strong><span class="muted">已完成</span></div>
  <div class="step"><strong>2 会议整理</strong><span class="muted">已生成草稿</span></div>
  <div class="step active"><strong>3 二次校对</strong><span class="muted">确认终稿</span></div>
  <div class="step"><strong>4 归档确认</strong><span class="muted">下一步</span></div>
</div>"""
        )
        body = f"""
{workflow_stage_header('人工初步校对' if stage == 'pre_agent' else '人工二次校对', gate_text)}
{stepper}
<div class="meta">
  <div><strong>{html.escape(str(meta.get("title") or "投资会议纪要"))}</strong></div>
  <div class="muted">会议日期：{html.escape(str(meta.get("meeting_date") or ""))}　会议类型：{meeting_type}　会议系列：{html.escape(str(meta.get("meeting_series") or markdown_field(markdown, "会议系列", MEETING_SERIES_DEFAULT)))}　更新时间：{html.escape(str(meta.get("updated_at") or ""))}</div>
</div>
{review_editor_html(markdown, draft_id, confirm_action=confirm_action, confirm_label=confirm_label, helper_text=helper_text)}
"""
        self._send(200, html_page("OpenMinute-会议纪要", body), "text/html; charset=utf-8")

    def _handle_agent_status(self) -> None:
        self._send(200, html_page("规划 Agent 工作状态", planner_status_html()), "text/html; charset=utf-8")

    def _handle_sync_status(self) -> None:
        status = google_drive_sync_status()
        recent_errors = status.get("recent_errors") or []
        error_rows = "".join(f"<li>{html.escape(str(line))}</li>" for line in recent_errors)
        body = f"""
<div class="meta">
  <strong>Google Drive 同步状态</strong>
  <div class="muted">这里读取本机同步日志，只展示最近一次同步结果，不会触发新的同步。</div>
</div>
<div class="item">
  <div>状态：<span class="status">{html.escape(str(status.get("status") or "unknown"))}</span></div>
  <div class="muted">说明：{html.escape(str(status.get("message") or ""))}</div>
  <div class="muted">最近开始：{html.escape(str(status.get("last_start") or "-"))}</div>
  <div class="muted">最近退出：{html.escape(str(status.get("last_exit") or "-"))}</div>
  <div class="muted">rclone：{html.escape(str(status.get("rclone_path") or "-"))}</div>
  <div class="muted">日志：{html.escape(str(status.get("log_path") or "-"))}</div>
</div>
<div class="item">
  <strong>最近错误</strong>
  {f'<ul>{error_rows}</ul>' if error_rows else '<p class="muted">未检测到最近错误。</p>'}
</div>
"""
        self._send(200, html_page("Google Drive 同步状态", body), "text/html; charset=utf-8")

    def _handle_access_audit(self, query: dict[str, list[str]], identity: dict[str, Any]) -> None:
        if not identity.get("is_admin"):
            self._deny_access("只有管理员可以查看访问审计日志。")
            return
        limit = int((query.get("limit") or ["100"])[0] or "100")
        limit = max(1, min(limit, 500))
        events = load_access_audit_events(limit=limit)
        self._audit("access_audit_view", identity, "ok", result_count=len(events))
        rows = []
        for item in events:
            details = item.get("details") if isinstance(item.get("details"), dict) else {}
            rows.append(
                "<tr>"
                f"<td>{html.escape(str(item.get('ts') or ''))}</td>"
                f"<td><span class=\"status\">{html.escape(str(item.get('event') or ''))}</span></td>"
                f"<td>{html.escape(str(item.get('status') or ''))}</td>"
                f"<td>{html.escape(str(item.get('user_id') or ''))}</td>"
                f"<td>{html.escape(str(item.get('host') or item.get('x_forwarded_host') or ''))}</td>"
                f"<td>{html.escape(str(item.get('path') or ''))}</td>"
                f"<td class=\"muted\">{html.escape(json.dumps(details, ensure_ascii=False, sort_keys=True))}</td>"
                "</tr>"
            )
        body = f"""
<div class="meta">
  <strong>公网访问审计</strong>
  <div class="muted">最近 {html.escape(str(limit))} 条访问、查询和下载记录。日志路径：{html.escape(str(access_audit_log_path()))}</div>
</div>
<form class="filters bar" method="get" action="/access-audit">
  <input name="limit" placeholder="显示条数" value="{html.escape(str(limit))}">
  <button type="submit">刷新</button>
</form>
<table>
  <thead><tr><th>时间</th><th>事件</th><th>状态</th><th>用户</th><th>Host</th><th>路径</th><th>详情</th></tr></thead>
  <tbody>{''.join(rows) if rows else '<tr><td colspan="7" class="muted">暂无审计记录。</td></tr>'}</tbody>
</table>
"""
        self._send(200, html_page("公网访问审计", body), "text/html; charset=utf-8")

    def _handle_access_users(self, identity: dict[str, Any]) -> None:
        if not identity.get("is_admin"):
            self._deny_access("只有管理员可以查看访问用户。")
            return
        users = public_access_users()
        self._audit("access_users_view", identity, "ok", user_count=len(users))
        rows = []
        for user in users:
            status = "禁用" if user.get("disabled") else "启用"
            rows.append(
                "<tr>"
                f"<td>{html.escape(str(user.get('user_id') or ''))}</td>"
                f"<td>{html.escape(str(user.get('display_name') or ''))}</td>"
                f"<td><span class=\"status\">{html.escape(str(user.get('role') or ''))}</span></td>"
                f"<td>{html.escape(', '.join(user.get('groups') or []))}</td>"
                f"<td>{html.escape(status)}</td>"
                f"<td>{'是' if user.get('password_present') else '否'}</td>"
                f"<td>{'是' if user.get('token_present') else '否'}</td>"
                f"<td class=\"muted\">{html.escape(str(user.get('created_at') or ''))}</td>"
                f"<td class=\"muted\">{html.escape(str(user.get('updated_at') or ''))}</td>"
                "</tr>"
            )
        body = f"""
<div class="meta">
  <strong>访问用户</strong>
  <div class="muted">后台直接写入本机 SQLite 用户数据库；不开放注册，不显示密码明文。新增或重置密码时，密码只写入本机私有 password 文件；API 令牌仅用于程序调用兼容。</div>
</div>
<form class="item" method="post" action="/access-users">
  <strong>管理员操作</strong>
  <div class="bar">
    <select name="action">
      <option value="add">新增</option>
      <option value="update">修改资料/分组</option>
      <option value="reset_password">重置密码</option>
      <option value="rotate_token">轮换 API 令牌</option>
      <option value="disable">禁用</option>
      <option value="remove">删除</option>
    </select>
    <input name="user_id" placeholder="用户 ID" required>
    <input name="display_name" placeholder="显示名">
    <select name="role">
      <option value="viewer">viewer</option>
      <option value="editor">editor</option>
      <option value="admin">admin</option>
      <option value="external_viewer">external_viewer</option>
    </select>
    <input name="groups" placeholder="分组，逗号分隔">
    <input name="password" type="password" placeholder="密码；留空则自动生成" autocomplete="new-password">
    <button type="submit">执行</button>
  </div>
</form>
<table>
  <thead><tr><th>用户 ID</th><th>显示名</th><th>角色</th><th>分组</th><th>状态</th><th>密码</th><th>API 令牌</th><th>创建时间</th><th>更新时间</th></tr></thead>
  <tbody>{''.join(rows) if rows else '<tr><td colspan="9" class="muted">暂无访问用户。</td></tr>'}</tbody>
</table>
<div class="item">
  <strong>本机管理命令</strong>
  <p><code>python3 /Users/kumaai/.codex/skills/投资会议纪要整理/scripts/manage_access_users.py add --user-id USER --role viewer --groups research</code></p>
  <p class="muted">可用子命令：<code>list</code>、<code>add</code>、<code>update</code>、<code>reset-password</code>、<code>rotate-token</code>、<code>disable</code>、<code>remove</code>。脚本只输出本机私有文件路径，不把密码或令牌明文写入页面。</p>
</div>
"""
        self._send(200, html_page("访问用户", body), "text/html; charset=utf-8")

    def _handle_access_users_post(self, identity: dict[str, Any]) -> None:
        if not identity.get("is_admin"):
            self._deny_access("只有管理员可以管理访问用户。")
            return
        payload = self._read_payload()
        action = str(payload.get("action") or "").strip()
        try:
            result = mutate_access_user(action, payload)
        except ValueError as exc:
            self._audit("access_user_mutation", identity, "failed", action=action, error=str(exc))
            if self._wants_json_response():
                self._json(400, {"ok": False, "error": str(exc)})
            else:
                body = f'<p class="error">{html.escape(str(exc))}</p><div class="bar"><a href="/access-users">返回访问用户</a></div>'
                self._send(400, html_page("访问用户", body), "text/html; charset=utf-8")
            return
        self._audit("access_user_mutation", identity, "ok", action=action, user_id=result.get("user_id"))
        if self._wants_json_response():
            self._json(200, result)
        else:
            self._redirect("/access-users")

    def _handle_drafts(self, query: dict[str, list[str]], identity: dict[str, Any]) -> None:
        role = str(identity.get("role") or "").lower()
        if role not in {"admin", "editor"}:
            self._deny_access("只有管理员或编辑可以查看待校对草稿。")
            return
        status_filter = (query.get("status") or ["open"])[0].strip().lower() or "open"
        selected_type = (query.get("meeting_type") or [""])[0].strip()
        selected_series = (query.get("meeting_series") or [""])[0].strip()
        drafts = filter_drafts(scan_drafts(), query)
        self._audit("drafts_view", identity, "ok", result_count=len(drafts), status_filter=status_filter)
        rows = []
        for item in drafts[:300]:
            rows.append(
                "<tr>"
                f"<td>{html.escape(str(item.get('meeting_date') or ''))}</td>"
                f"<td><a href=\"{html.escape(str(item.get('review_url') or ''))}\">{html.escape(str(item.get('title') or ''))}</a></td>"
                f"<td>{html.escape(str(item.get('meeting_type') or '多人复盘会'))}</td>"
                f"<td>{html.escape(str(item.get('meeting_series') or MEETING_SERIES_DEFAULT))}</td>"
                f"<td>{html.escape(str(item.get('stage') or '-'))}</td>"
                f"<td><span class=\"status\">{html.escape(str(item.get('status') or ''))}</span></td>"
                f"<td>{html.escape(str(item.get('input_source') or ''))}</td>"
                f"<td>{html.escape(str(item.get('target_count') or 0))}</td>"
                f"<td class=\"muted\">{html.escape(str(item.get('updated_at') or item.get('created_at') or ''))}</td>"
                f"<td class=\"muted\">{html.escape(str(item.get('source_run_id') or ''))}</td>"
                "</tr>"
            )
        body = f"""
<h1 class="page-title">待校对草稿</h1>
<div class="meta">
  <strong>只处理还没完成的稿件</strong>
  <div class="muted">草稿不是最终会议纪要；完成二次校对和归档确认后，才会进入历史纪要、Obsidian 正式目录和 Dify 知识库。</div>
</div>
<form class="filters bar" method="get" action="/drafts">
  <select name="status">
    <option value="open" {'selected' if status_filter == 'open' else ''}>未确认</option>
    <option value="draft" {'selected' if status_filter == 'draft' else ''}>draft</option>
    <option value="confirmed" {'selected' if status_filter == 'confirmed' else ''}>confirmed</option>
    <option value="all" {'selected' if status_filter == 'all' else ''}>全部</option>
  </select>
  <select name="meeting_type">
    <option value="">全部类型</option>
    {meeting_type_options(selected_type)}
  </select>
  <select name="meeting_series">
    <option value="">全部系列</option>
    {meeting_series_options(selected_series)}
  </select>
  <button type="submit">筛选</button>
</form>
<table>
  <thead><tr><th>会议日期</th><th>草稿</th><th>类型</th><th>系列</th><th>阶段</th><th>状态</th><th>输入来源</th><th>标的数</th><th>更新时间</th><th>Run ID</th></tr></thead>
  <tbody>{''.join(rows) if rows else '<tr><td colspan="10" class="muted">暂无草稿。</td></tr>'}</tbody>
</table>
"""
        self._send(200, html_page("待校对草稿", body), "text/html; charset=utf-8")

    def _handle_mapping_audit(self, identity: dict[str, Any]) -> None:
        if not identity.get("is_admin"):
            self._deny_access("只有管理员可以查看知识库映射审计。")
            return
        report = load_mapping_audit_report()
        self._audit("mapping_audit_view", identity, "ok", issue_count=report.get("issue_count"))
        rows = []
        for issue in report.get("issues") or []:
            if not isinstance(issue, dict):
                continue
            rows.append(
                "<tr>"
                f"<td><span class=\"status\">{html.escape(str(issue.get('severity') or ''))}</span></td>"
                f"<td>{html.escape(str(issue.get('type') or ''))}</td>"
                f"<td>{html.escape(str(issue.get('path') or issue.get('key') or ''))}</td>"
                f"<td>{html.escape(str(issue.get('document_name') or ''))}</td>"
                f"<td>{html.escape(str(issue.get('document_id') or ''))}</td>"
                f"<td class=\"muted\">{html.escape(str(issue.get('dataset_id') or ''))}</td>"
                "</tr>"
            )
        body = f"""
<div class="meta">
  <strong>Dify 知识库映射审计</strong>
  <div class="muted">只读核对入口，不会删除或修改 Dify 文档。映射条目：{html.escape(str(report.get('mapping_count') or 0))}，正式纪要：{html.escape(str(report.get('minutes_count') or 0))}，问题：{html.escape(str(report.get('issue_count') or 0))}</div>
</div>
<table>
  <thead><tr><th>级别</th><th>类型</th><th>本地路径 / 映射键</th><th>Dify 文档名</th><th>Dify 文档 ID</th><th>Dataset</th></tr></thead>
  <tbody>{''.join(rows) if rows else '<tr><td colspan="6" class="muted">未发现映射问题。</td></tr>'}</tbody>
</table>
<div class="item">
  <strong>处理原则</strong>
  <p class="muted">当前剩余问题是旧 backfill 的本地源文件缺失。这里仅展示核对信息；确认 Dify 文档不是唯一版本之前，不自动删除、不自动合并。</p>
</div>
"""
        self._send(200, html_page("知识库映射审计", body), "text/html; charset=utf-8")

    def _handle_health_report(self) -> None:
        report = local_health_report()
        summary = report.get("summary") or {}
        counts = summary.get("counts") or {}
        rows = []
        for item in report.get("checks") or []:
            status = str(item.get("status") or "warning")
            detail = item.get("details") or {}
            rows.append(
                "<tr>"
                f"<td><span class=\"status\">{html.escape(status.upper())}</span></td>"
                f"<td>{html.escape(str(item.get('name') or '-'))}</td>"
                f"<td>{html.escape(str(item.get('message') or '-'))}</td>"
                f"<td class=\"muted\">{html.escape(str(detail.get('path') or detail.get('url') or detail.get('log') or ''))}</td>"
                "</tr>"
            )
        body = f"""
<div class="meta">
  <strong>工作流健康检查</strong>
  <div class="muted">生成时间：{html.escape(str(report.get('generated_at') or now_iso()))}</div>
</div>
<div class="item">
  <div>整体状态：<span class="status">{html.escape(str(summary.get('overall') or 'unknown').upper())}</span></div>
  <div class="muted">OK {html.escape(str(counts.get('ok') or 0))} / WARNING {html.escape(str(counts.get('warning') or 0))} / ERROR {html.escape(str(counts.get('error') or 0))}</div>
</div>
<table>
  <thead><tr><th>状态</th><th>项目</th><th>说明</th><th>路径 / URL</th></tr></thead>
  <tbody>{''.join(rows)}</tbody>
</table>
"""
        self._send(200, html_page("工作流健康检查", body), "text/html; charset=utf-8")

    def _handle_mcp_servers(self, identity: dict[str, Any]) -> None:
        config = load_mcp_config()
        rows: list[str] = []
        for server in config.get("servers") or []:
            if not isinstance(server, dict):
                continue
            tools = server.get("tools") if isinstance(server.get("tools"), list) else []
            tool_rows = "".join(
                f"<li><strong>{html.escape(str(tool.get('name') or '-'))}</strong>：{html.escape(str(tool.get('description') or ''))}</li>"
                for tool in tools
                if isinstance(tool, dict)
            )
            enabled = "启用" if server.get("enabled", True) else "停用"
            rows.append(
                f"""<div class="item">
  <div><strong>{html.escape(str(server.get('name') or server.get('id') or '-'))}</strong> <span class="status">{html.escape(enabled)}</span> <span class="status">{html.escape(str(server.get('type') or ''))}</span></div>
  <div class="muted">ID：{html.escape(str(server.get('id') or '-'))}</div>
  <div>{html.escape(str(server.get('description') or ''))}</div>
  <div class="muted">调用地址：{html.escape(str(server.get('call_url') or 'http://host.docker.internal:8767/api/mcp/call'))}</div>
  <div class="muted">Skill：{html.escape(str(server.get('skill_path') or ''))}</div>
  {f'<ul>{tool_rows}</ul>' if tool_rows else '<p class="muted">未配置工具。</p>'}
</div>"""
            )
        example_payload = html.escape(
            'POST /api/mcp/call\n'
            '{"server":"a-stock-data","tool":"symbol_lookup","arguments":{"query":"宁德时代","limit":3}}\n'
            '{"server":"a-stock-data","tool":"quote","arguments":{"codes":["300750","600519"]}}\n'
            '{"server":"a-stock-data","tool":"resolve_targets","arguments":{"text":"看好宁德时代和联瑞新材","max_mentions":5}}'
        )
        body = f"""
<div class="meta">
  <strong>MCP 服务器</strong>
  <div class="muted">这里展示会议纪要工作流通过本地 bridge 可调用的 MCP server。Dify 工作流中用 HTTP 节点调用，不修改 Dify 本体源码。</div>
</div>
<div class="item">
  <div><strong>查看接口</strong>：<code>/api/mcp-servers</code></div>
  <div><strong>Dify 调用地址</strong>：<code>http://host.docker.internal:8767/api/mcp/call</code></div>
  <div><strong>本机调试地址</strong>：<code>http://127.0.0.1:8767/api/mcp/call</code></div>
  <div><strong>配置文件</strong>：<code>{html.escape(str(DEFAULT_MCP_CONFIG_PATH))}</code></div>
</div>
<div class="item">
  <strong>调用示例</strong>
  <pre>{example_payload}</pre>
</div>
{''.join(rows) if rows else '<p class="muted">还没有配置 MCP server。</p>'}
"""
        self._audit("mcp_servers", identity, "ok", server_count=len(rows))
        self._send(200, html_page("MCP 服务器", body), "text/html; charset=utf-8")

    def _handle_target_query(self, query: dict[str, list[str]], identity: dict[str, Any]) -> None:
        query_text = (query.get("q") or [""])[0].strip()
        escaped_query = html.escape(query_text)
        result_html = ""
        if query_text:
            result = target_query(query_text, identity=identity)
            self._audit("target_query", identity, "ok", query=query_text, result_count=result.get("result_count"))
            lookup = result.get("symbol_lookup") or {}
            candidates = lookup.get("candidates") or []
            candidate_rows = "".join(
                f"<li>{html.escape(str(candidate.get('symbol') or ''))} / {html.escape(str(candidate.get('name') or ''))} / {html.escape(str(candidate.get('market') or ''))} / {html.escape(str(candidate.get('confidence') or ''))}</li>"
                for candidate in candidates
                if isinstance(candidate, dict)
            )
            rows = []
            for item in result.get("results") or []:
                review_url = str(item.get("review_url") or "")
                md_path = str(item.get("markdown_path") or "")
                word_path = str(item.get("word_path") or "")
                is_meeting = item.get("source_type") == "meeting_minutes"
                md_download = f"/download?path={urllib.parse.quote(md_path)}" if is_meeting and md_path else ""
                word_download = f"/download?path={urllib.parse.quote(word_path)}" if is_meeting and word_path else ""
                source_label = str(item.get("source_label") or "资料")
                rows.append(
                    f"""<div class="item">
  <div><a href="{html.escape(review_url)}"><strong>{html.escape(str(item.get('title') or '投资会议纪要'))}</strong></a> <span class="status">{html.escape(source_label)}</span> <span class="status">{html.escape(str(item.get('meeting_date') or item.get('date') or ''))}</span></div>
  <div class="muted">命中：{html.escape(' / '.join(str(term) for term in item.get('matched_terms') or []))}</div>
  <div>{html.escape(str(item.get('evidence_snippet') or ''))}</div>
  <div class="bar"><a href="{html.escape(review_url)}">查看原文</a>{f' <a href="{html.escape(md_download)}">下载 Markdown</a>' if md_download else ''}{f' <a href="{html.escape(word_download)}">下载 Word</a>' if word_download else ''}</div>
</div>"""
                )
            counts = result.get("source_counts") or {}
            result_html = f"""
<div class="item">
  <strong>代码候选</strong>
  <div class="muted">状态：{html.escape(str(lookup.get('status') or ''))}。仅 confirmed 候选可直接写入；其他情况应保留存疑。</div>
  {f'<ul>{candidate_rows}</ul>' if candidate_rows else '<p class="muted">未找到代码候选。</p>'}
</div>
<div class="muted">共 {html.escape(str(result.get('result_count') or 0))} 条证据；会议纪要 {html.escape(str(counts.get('meeting_minutes') or 0))} 条，外部资料 {html.escape(str(counts.get('external_source') or 0))} 条。结果是证据检索，不是新的投资结论。</div>
{''.join(rows) if rows else '<p class="muted">未检索到相关证据。</p>'}
"""
        body = f"""
<div class="meta">
  <strong>标的历史观点查询</strong>
  <div class="muted">输入标的名称、简称或代码，系统检索本机 Obsidian 正式纪要和外部资料归档并返回可追溯证据片段。</div>
</div>
<form class="filters bar" method="get" action="/target-query">
  <input name="q" placeholder="标的 / 代码 / 关键词" value="{escaped_query}">
  <button type="submit">查询</button>
</form>
{result_html}
"""
        self._send(200, html_page("标的历史观点查询", body), "text/html; charset=utf-8")

    def _handle_history(self, query: dict[str, list[str]], identity: dict[str, Any]) -> None:
        items = filter_history(filter_accessible_items(scan_history(), identity), query)
        self._audit("history_list", identity, "ok", result_count=len(items), filters={k: v[-1] for k, v in query.items() if v})
        date = html.escape((query.get("date") or [""])[0])
        selected_type = (query.get("meeting_type") or [""])[0]
        selected_series = (query.get("meeting_series") or [""])[0]
        keyword = html.escape((query.get("q") or [""])[0])
        symbol = html.escape((query.get("symbol") or [""])[0])
        visibility = html.escape((query.get("visibility") or [""])[0])
        owner = html.escape((query.get("owner") or [""])[0])
        rows = []
        for item in items:
            export_path = item.get("markdown_path") or ""
            word_path = item.get("word_path") or ""
            targets = " / ".join(str(target) for target in item.get("targets") or [])
            knowledge = item.get("knowledge") or {}
            knowledge_ok = bool(knowledge.get("ok"))
            knowledge_text = "已写入知识库" if knowledge_ok else "未匹配知识库"
            document_id = str(knowledge.get("document_id") or "")
            document_name = str(knowledge.get("document_name") or "")
            knowledge_detail = document_name or document_id or str(knowledge.get("message") or "")
            permissions = item.get("permissions") or {}
            permission_text = f"{permissions.get('visibility') or '-'} / {permissions.get('owner') or '-'}"
            md_download = f"/download?path={urllib.parse.quote(str(export_path))}" if export_path else ""
            word_download = f"/download?path={urllib.parse.quote(str(word_path))}" if word_path else ""
            rows.append(
                f"""<div class="item">
  <div><a href="{html.escape(str(item.get('review_url')))}"><strong>{html.escape(str(item.get('title') or '投资会议纪要'))}</strong></a> <span class="status">{html.escape(str(item.get('status') or 'draft'))}</span></div>
  <div class="muted">会议日期：{html.escape(str(item.get('meeting_date') or ''))}　会议类型：{html.escape(str(item.get('meeting_type') or '多人复盘会'))}　会议系列：{html.escape(str(item.get('meeting_series') or MEETING_SERIES_DEFAULT))}　来源：Obsidian 正式归档</div>
  <div class="muted">知识库：<span class="status">{html.escape(knowledge_text)}</span>{f'　{html.escape(knowledge_detail)}' if knowledge_detail else ''}</div>
  <div class="muted">权限：{html.escape(permission_text)}</div>
  <div class="muted">标的：{html.escape(targets or '-')}</div>
  <div class="muted">Markdown：{html.escape(str(export_path or '-'))}</div>
  <div class="muted">Word：{html.escape(str(word_path or '-'))}</div>
  <div class="bar"><a href="{html.escape(str(item.get('review_url')))}">查看原始格式</a>{f' <a href="{html.escape(md_download)}">下载 Markdown</a>' if md_download else ''}{f' <a href="{html.escape(word_download)}">下载 Word</a>' if word_download else ''}</div>
</div>"""
            )
        body = f"""
<h1 class="page-title">历史会议纪要</h1>
<div class="meta">
  <strong>正式归档，只看终稿</strong>
  <div class="muted">这里读取本机 Obsidian 正式归档目录，用于阅读、检索和下载 Markdown / Word。当前身份：{html.escape(str(identity.get('display_name') or identity.get('user_id') or '-'))}</div>
</div>
<form class="filters bar" method="get" action="/history">
  <input name="date" placeholder="日期 YYYY-MM-DD" value="{date}">
  <select name="meeting_type"><option value="">全部类型</option>{meeting_type_options(selected_type)}</select>
  <select name="meeting_series"><option value="">全部系列</option>{meeting_series_options(selected_series)}</select>
  <input name="q" placeholder="标题 / 关键词" value="{keyword}">
  <input name="symbol" placeholder="标的关键词" value="{symbol}">
  <input name="visibility" placeholder="可见范围" value="{visibility}">
  <input name="owner" placeholder="负责人" value="{owner}">
  <button type="submit">筛选</button>
  <a href="/history">清空</a>
</form>
<div class="section-head"><div><h2>检索结果</h2><p>共 {len(items)} 条。</p></div></div>
{''.join(rows) if rows else '<p class="muted">暂无匹配纪要。</p>'}
"""
        self._send(200, html_page("历史会议纪要", body), "text/html; charset=utf-8")

    def _handle_external_sources(self, query: dict[str, list[str]], identity: dict[str, Any]) -> None:
        items = filter_external_sources(filter_accessible_items(scan_external_sources(), identity), query)
        self._audit("external_sources_list", identity, "ok", result_count=len(items), filters={k: v[-1] for k, v in query.items() if v})
        date = html.escape((query.get("date") or [""])[0])
        keyword = html.escape((query.get("q") or [""])[0])
        selected_kind = html.escape((query.get("kind") or [""])[0])
        visibility = html.escape((query.get("visibility") or [""])[0])
        owner = html.escape((query.get("owner") or [""])[0])
        options = ['<option value="">全部资料</option>']
        for kind, label, _root in EXTERNAL_SOURCE_ROOTS:
            selected = " selected" if selected_kind == kind else ""
            options.append(f'<option value="{html.escape(kind)}"{selected}>{html.escape(label)}</option>')
        rows = []
        for item in items:
            view_url = str(item.get("view_url") or "")
            source_link = str(item.get("source_link") or "")
            raw_path = str(item.get("raw_path") or "")
            markdown_path = str(item.get("markdown_path") or "")
            version_count = int(item.get("version_count") or 1)
            permissions = item.get("permissions") or {}
            permission_text = f"{permissions.get('visibility') or '-'} / {permissions.get('owner') or '-'}"
            rows.append(
                f"""<div class="item">
  <div><a href="{html.escape(normalize_local_url(view_url))}"><strong>{html.escape(str(item.get('title') or '外部资料'))}</strong></a> <span class="status">{html.escape(str(item.get('kind_label') or '资料'))}</span></div>
  <div class="muted">资料日期：{html.escape(str(item.get('date') or ''))}　类型：{html.escape(str(item.get('type') or ''))}</div>
  <div class="muted">版本：{html.escape(str(version_count))}</div>
  <div class="muted">权限：{html.escape(permission_text)}</div>
  <div class="muted">原文链接：{f'<a href="{html.escape(source_link)}">{html.escape(source_link)}</a>' if source_link else '-'}</div>
  <div class="muted">原始归档：{html.escape(raw_path or '-')}</div>
  <div class="muted">Markdown：{html.escape(markdown_path or '-')}</div>
  <div class="bar"><a href="{html.escape(normalize_local_url(view_url))}">查看正文</a></div>
</div>"""
            )
        body = f"""
<div class="meta">
  <strong>外部资料归档</strong>
  <div class="muted">这里读取微信公众号文章和微信聊天分析的 Obsidian 归档。它们是投研资料源，不等同于人工确认后的会议纪要。当前身份：{html.escape(str(identity.get('display_name') or identity.get('user_id') or '-'))}</div>
</div>
<form class="filters bar" method="get" action="/external-sources">
  <input name="date" placeholder="日期 YYYY-MM-DD" value="{date}">
  <input name="q" placeholder="标题 / 关键词 / 标的" value="{keyword}">
  <select name="kind">{''.join(options)}</select>
  <input name="visibility" placeholder="可见范围" value="{visibility}">
  <input name="owner" placeholder="负责人" value="{owner}">
  <button type="submit">筛选</button>
  <a href="/external-sources">清空</a>
</form>
<div class="muted">共 {len(items)} 条</div>
{''.join(rows) if rows else '<p class="muted">暂无匹配资料。</p>'}
"""
        self._send(200, html_page("外部资料归档", body), "text/html; charset=utf-8")

    def _handle_external_view(self, query: dict[str, list[str]], identity: dict[str, Any]) -> None:
        raw_path = (query.get("path") or [""])[0]
        path = Path(raw_path)
        root_info = external_root_for_path(path)
        if not root_info:
            body = "<p class=\"error\">无法读取资料：路径不在外部资料归档目录内。</p>"
            self._send(404, html_page("外部资料归档", body), "text/html; charset=utf-8")
            return
        try:
            resolved = path.resolve()
            content = resolved.read_text(encoding="utf-8")
        except Exception as exc:
            body = f"<p class=\"error\">无法读取资料：{html.escape(str(exc))}</p>"
            self._send(404, html_page("外部资料归档", body), "text/html; charset=utf-8")
            return
        if not can_identity_access(identity, permission_meta_for_note(resolved, content)):
            self._deny_access("当前身份无权查看该外部资料。")
            return
        kind, label, _root = root_info
        self._audit("external_view", identity, "ok", path=str(resolved), kind=kind)
        body = f"""
<div class="bar"><a href="/external-sources">返回资料源</a></div>
<div class="meta">
  <strong>{html.escape(label)}</strong>
  <div class="muted">{html.escape(str(resolved))}</div>
</div>
<div class="editor-shell">
  <article class="paper">{markdown_to_readable_html(content)}</article>
</div>
<details class="source-details">
  <summary>查看 Markdown 源文</summary>
  <pre class="markdown-source">{html.escape(content)}</pre>
</details>
"""
        self._send(200, html_page(markdown_title(content), body), "text/html; charset=utf-8")

    def _handle_view(self, query: dict[str, list[str]], identity: dict[str, Any]) -> None:
        raw_path = (query.get("path") or [""])[0]
        path = Path(raw_path)
        final_root = output_dir().resolve()
        try:
            resolved = path.resolve()
            if final_root not in resolved.parents and resolved != final_root:
                raise ValueError("path is outside final output directory")
            content = resolved.read_text(encoding="utf-8")
        except Exception as exc:
            body = f"<p class=\"error\">无法读取纪要：{html.escape(str(exc))}</p>"
            self._send(404, html_page("历史会议纪要", body), "text/html; charset=utf-8")
            return
        if not can_identity_access(identity, permission_meta_for_note(resolved, content)):
            self._deny_access("当前身份无权查看该会议纪要。")
            return
        self._audit("minute_view", identity, "ok", path=str(resolved))
        body = f"""
<div class="bar"><a href="/history">返回历史</a></div>
<div class="meta"><div class="muted">{html.escape(str(resolved))}</div></div>
<div class="bar">
  <a href="/download?path={urllib.parse.quote(str(resolved))}">下载 Markdown</a>
  {f'<a href="/download?path={urllib.parse.quote(str(resolved.with_suffix(".docx")))}">下载 Word</a>' if resolved.with_suffix(".docx").exists() else ''}
</div>
<div class="editor-shell">
  <article class="paper">{markdown_to_readable_html(content)}</article>
</div>
<details class="source-details">
  <summary>查看 Markdown 源文</summary>
  <pre class="markdown-source">{html.escape(content)}</pre>
</details>
"""
        self._send(200, html_page(markdown_title(content), body), "text/html; charset=utf-8")

    def _handle_result(self, query: dict[str, list[str]]) -> None:
        draft_id = (query.get("id") or [""])[0]
        folder = draft_folder(draft_id)
        if not draft_id or not folder.exists():
            self._send(404, html_page("工作流结果缓存", "<p class=\"error\">未找到该工作流结果缓存。</p>"), "text/html; charset=utf-8")
            return
        meta = load_meta(folder)
        markdown = current_markdown(folder)
        review_url = f"/review?id={urllib.parse.quote(draft_id)}"
        export = meta.get("export") if isinstance(meta.get("export"), dict) else {}
        raw_archive = meta.get("raw_archive") if isinstance(meta.get("raw_archive"), dict) else {}
        knowledge = meta.get("knowledge") if isinstance(meta.get("knowledge"), dict) else {}
        artifacts = meta.get("post_review_artifacts") if isinstance(meta.get("post_review_artifacts"), dict) else {}
        structured_table = meta.get("structured_table") if isinstance(meta.get("structured_table"), dict) else {}
        table_path = Path(str(structured_table.get("path") or structured_table_path(folder)))
        table_markdown = table_path.read_text(encoding="utf-8") if table_path.exists() else ""
        table_preview_html = markdown_to_readable_html(table_markdown or "# 结构化表格\n\n暂无表格内容。")

        def state(ok: Any) -> str:
            return "ok" if ok else "failed"

        result_blocks = ""
        if meta.get("status") == "table_review":
            result_blocks = f"""
<div class="item">
  <strong>表格确认</strong>
  <div class="muted">终稿已确认。请校对结构化表格，确认后才会进入正式归档确认页。</div>
  <form method="post" action="/confirm-table?id={html.escape(urllib.parse.quote(draft_id))}" data-processing-title="正在确认结构化表格" data-processing-message="系统正在保存表格并进入最终归档确认。">
    <textarea name="structured_table" rows="18" spellcheck="false">{html.escape(table_markdown)}</textarea>
    <div class="bar">
      <button type="submit">确认结构化表格</button>
      <a href="{review_url}">返回二次校对</a>
    </div>
  </form>
  <details>
    <summary>表格预览</summary>
    <article class="paper">{table_preview_html}</article>
  </details>
</div>
"""
        elif meta.get("status") == "archive_review":
            artifact_md = str(artifacts.get("markdown_path") or "")
            artifact_json = str(artifacts.get("json_path") or "")
            result_blocks = f"""
<div class="item">
  <strong>终稿待归档确认</strong>
  <div class="muted">终稿与结构化表格已确认。确认无误后再写入正式归档、Word、Dify 知识库和 Google Drive。</div>
  <h3>Obsidian 归档</h3>
  <table>
    <tbody>
      <tr><th>原始记录</th><td><span class="status">pending</span></td><td class="muted">最终确认后写入 00 Inbox/会议原始记录（如本次有上传文件）</td></tr>
      <tr><th>最终 Markdown / Word</th><td><span class="status">pending</span></td><td class="muted">最终确认后写入 01 Projects/会议纪要</td></tr>
      <tr><th>结构化表格</th><td><span class="status">{html.escape(state(structured_table.get('ok')))}</span></td><td class="muted">最终确认后单独归档到会议结构化表格目录，按日期 / 会议类型 / 会议场次分类</td></tr>
      <tr><th>终稿校验文件</th><td><span class="status">{html.escape(state(artifacts.get('ok')))}</span></td><td class="muted">{html.escape(artifact_md or artifact_json or str(artifacts.get('error') or '-'))}</td></tr>
    </tbody>
  </table>
  <h3>Dify 知识库入库</h3>
  <table>
    <tbody>
      <tr><th>最终 Markdown</th><td><span class="status">pending</span></td><td class="muted">最终确认后同步到会议纪要知识库</td></tr>
      <tr><th>结构化表格</th><td><span class="status">excluded</span></td><td class="muted">不写入会议纪要正文，不进入知识库正文；单独进入会议结构化表格归档</td></tr>
    </tbody>
  </table>
  <h3>不入库内容</h3>
  <table>
    <tbody>
      <tr><th>原始音频</th><td><span class="status">excluded</span></td><td class="muted">只归档为原始证据，不进入知识库正文</td></tr>
      <tr><th>差异明细/人工校对日志</th><td><span class="status">excluded</span></td><td class="muted">只作为过程记录，不进入知识库正文</td></tr>
    </tbody>
  </table>
  <form method="post" action="/confirm-archive?id={html.escape(urllib.parse.quote(draft_id))}" class="bar" data-processing-title="正在确认归档" data-processing-message="系统正在写入正式归档、生成文件并同步知识库。请不要重复提交。">
    <button type="submit">确认归档</button>
    <a href="{review_url}">返回二次校对</a>
  </form>
</div>
"""
        elif meta.get("status") == "confirmed":
            artifact_md = str(artifacts.get("markdown_path") or "")
            artifact_json = str(artifacts.get("json_path") or "")
            exported_markdown_path = str(export.get("markdown_path") or "")
            exported_word_path = str(export.get("word_path") or "")
            exported_markdown_download = (
                f"/download?path={urllib.parse.quote(exported_markdown_path)}"
                if export.get("markdown_created") and exported_markdown_path
                else ""
            )
            exported_word_download = (
                f"/download?path={urllib.parse.quote(exported_word_path)}"
                if export.get("word_created") and exported_word_path
                else ""
            )
            download_buttons = "".join(
                [
                    f'<a href="{html.escape(exported_word_download)}">下载 Word 纪要</a>' if exported_word_download else "",
                    f'<a href="{html.escape(exported_markdown_download)}">下载 Markdown 纪要</a>' if exported_markdown_download else "",
                ]
            )
            archive_message = str(meta.get("archive_message") or "").strip()
            archive_all_ok = (
                bool(export.get("markdown_created"))
                and bool(export.get("word_created"))
                and bool(raw_archive.get("ok"))
                and bool(knowledge.get("ok"))
                and bool(structured_table.get("ok"))
                and bool(artifacts.get("ok"))
            )
            archive_title = "归档成功" if archive_all_ok else "归档完成但有异常"
            archive_message = archive_message or (
                "Markdown、Word、原始记录、结构化表格和知识库同步已完成。"
                if archive_all_ok
                else "部分导出、原始记录归档、结构化表格或知识库同步未完全成功，请查看下方明细。"
            )
            result_blocks = f"""
<div class="item">
  <strong>{html.escape(archive_title)}</strong>
  <div class="muted">{html.escape(archive_message)}</div>
</div>
<div class="item">
  <strong>下载生成的纪要</strong>
  <div class="muted">全部确认完成后，可在这里下载本次生成的正式会议纪要。</div>
  <div class="bar">
    {download_buttons or '<span class="muted">暂未生成可下载的 Markdown 或 Word 文件，请查看下方导出状态。</span>'}
  </div>
</div>
<div class="item">
  <strong>确认归档结果</strong>
  <table>
    <tbody>
      <tr><th>Markdown 导出</th><td><span class="status">{html.escape(state(export.get('markdown_created')))}</span></td><td class="muted">{html.escape(str(export.get('markdown_path') or export.get('markdown_message') or '-'))}</td></tr>
      <tr><th>Word 导出</th><td><span class="status">{html.escape(state(export.get('word_created')))}</span></td><td class="muted">{html.escape(str(export.get('word_path') or export.get('word_message') or '-'))}</td></tr>
      <tr><th>原始记录归档</th><td><span class="status">{html.escape(state(raw_archive.get('ok')))}</span></td><td class="muted">{html.escape(str(raw_archive.get('archived_count') if 'archived_count' in raw_archive else raw_archive.get('error') or '-'))}</td></tr>
      <tr><th>Google Drive 同步</th><td><span class="status">{html.escape(state(export.get('google_drive_synced')))}</span></td><td class="muted">{html.escape(str(export.get('google_drive_message') or '-'))}</td></tr>
      <tr><th>Dify 知识库</th><td><span class="status">{html.escape(state(knowledge.get('ok')))}</span></td><td class="muted">{html.escape(str(knowledge.get('message') or knowledge.get('error') or knowledge.get('document_id') or '-'))}</td></tr>
      <tr><th>结构化表格</th><td><span class="status">{html.escape(state(structured_table.get('ok')))}</span></td><td class="muted">{html.escape(str(structured_table.get('archive_path') or structured_table.get('path') or structured_table.get('error') or '-'))}</td></tr>
      <tr><th>终稿校验文件</th><td><span class="status">{html.escape(state(artifacts.get('ok')))}</span></td><td class="muted">{html.escape(artifact_md or artifact_json or str(artifacts.get('error') or '-'))}</td></tr>
    </tbody>
  </table>
</div>
"""
        stage_name = "归档确认" if meta.get("status") == "archive_review" else "结果查看"
        body = f"""
{workflow_stage_header(stage_name, "查看本次工作流结果；需要归档时请在下方确认。")}
<div class="meta">
  <strong>{html.escape(str(meta.get("title") or markdown_title(markdown)))}</strong>
  <div class="muted">会议日期：{html.escape(str(meta.get("meeting_date") or ""))}　会议类型：{html.escape(str(meta.get("meeting_type") or "多人复盘会"))}　会议系列：{html.escape(str(meta.get("meeting_series") or MEETING_SERIES_DEFAULT))}　更新时间：{html.escape(str(meta.get("updated_at") or ""))}</div>
</div>
<div class="bar">
  <a href="{review_url}">打开人工校对</a>
  <a href="/drafts">查看待校对草稿</a>
  <a href="/history">查看 Obsidian 历史会议纪要</a>
  <a href="/external-sources">查看外部资料归档</a>
  <a href="/target-query">标的证据检索</a>
</div>
{result_blocks}
<div class="editor-shell">
  <article class="paper">{markdown_to_readable_html(markdown)}</article>
</div>
"""
        self._send(200, html_page("OpenMinute-会议纪要", body), "text/html; charset=utf-8")

    def _handle_latest(self) -> None:
        latest = load_latest_result()
        result_url = str(latest.get("result_cache_url") or "")
        if result_url:
            self.send_response(302)
            self.send_header("Location", normalize_local_url(result_url))
            self._security_headers()
            self.end_headers()
            return
        self._send(404, html_page("最近一次工作流结果", "<p class=\"error\">还没有可恢复的工作流结果。</p>"), "text/html; charset=utf-8")

    def _handle_download(self, query: dict[str, list[str]], identity: dict[str, Any]) -> None:
        raw_path = (query.get("path") or [""])[0]
        path = Path(raw_path)
        final_root = output_dir().resolve()
        try:
            resolved = path.resolve()
            if final_root not in resolved.parents and resolved != final_root:
                raise ValueError("path is outside final output directory")
            if resolved.suffix.lower() not in {".md", ".docx"}:
                raise ValueError("only Markdown and Word files can be downloaded")
            permission_source = resolved if resolved.suffix.lower() == ".md" else resolved.with_suffix(".md")
            if permission_source.exists():
                permission_markdown = permission_source.read_text(encoding="utf-8")
                if not can_identity_access(identity, permission_meta_for_note(permission_source, permission_markdown)):
                    self._deny_access("当前身份无权下载该会议纪要文件。")
                    return
            data = resolved.read_bytes()
        except Exception as exc:
            body = f"<p class=\"error\">无法下载文件：{html.escape(str(exc))}</p>"
            self._send(404, html_page("下载失败", body), "text/html; charset=utf-8")
            return
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if resolved.suffix.lower() == ".docx" else "text/markdown; charset=utf-8"
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{urllib.parse.quote(resolved.name)}")
        self.send_header("Content-Length", str(len(data)))
        self._security_headers()
        self.end_headers()
        self.wfile.write(data)
        self._audit("download", identity, "ok", path=str(resolved), suffix=resolved.suffix.lower(), bytes=len(data))


def main() -> int:
    values = read_config()
    host = os.environ.get("MEETING_REVIEW_SERVER_HOST", DEFAULT_HOST)
    port = int(os.environ.get("MEETING_REVIEW_SERVER_PORT") or values.get("MEETING_REVIEW_SERVER_PORT") or DEFAULT_PORT)
    storage_dir().mkdir(parents=True, exist_ok=True)
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    server = ThreadingHTTPServer((host, port), ReviewHandler)
    print(f"Meeting minutes review server listening on http://{host}:{port}", flush=True)
    server.serve_forever()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
