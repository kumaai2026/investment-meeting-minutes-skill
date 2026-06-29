#!/usr/bin/env python3
"""Validate the current meeting-minutes Markdown output contract."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

REQUIRED_SECTIONS = [
    "## 一、发言整理",
]

REQUIRED_METADATA_FIELDS = ["会议日期", "整理时间", "会议标题", "会议类型", "会议系列"]
METADATA_BOLD_LABELS = {
    "会议日期",
    "整理时间",
    "会议标题",
    "会议类型",
    "会议系列",
    "会议标的",
}

FORBIDDEN_PATTERNS = [
    r"AI结构化总结",
    r"标的汇总表",
    r"(?m)^##+\s*(?:摘要|投研摘要|整理者总结|投资结论|核心结论)\s*$",
    r"##\s*三[、.．]",
    r"##\s*四[、.．]",
    r"处理说明",
    r"(?m)^\*\*输入来源\*\*[:：]",
    r"(?m)^\*\*整理说明\*\*[:：]",
    r"Tool Call",
    r"skill_name",
    r"Traceback",
    r"Observation:",
    r"(?m)^###\s*问答\s*\d+",
    r"(?im)^\s*A[:：]",
]
FORBIDDEN_ESCAPE_HEADINGS = {
    "发言片段",
    "未归类",
    "主题整理",
    "内容摘要",
    "观点汇总",
    "核心观点",
    "整体判断",
    "整理者总结",
    "投研摘要",
    "投资结论",
}
REVIEW_SPEAKER_HEADING_TOPIC_TERMS = {
    "AI硬件",
    "半导体",
    "存储",
    "核心观点",
    "科技",
}
REVIEW_SPEAKER_ROLE_TERMS = {
    "CEO",
    "CFO",
    "嘉宾",
    "专家",
    "分析师",
    "主持人",
    "研究员",
    "管理层",
}

DOCUMENT_ONLY_SOURCE_MODES = {"document", "document-only", "text", "text-only", "文稿", "纯文本"}
AUDIO_SOURCE_MODES = {"audio", "audio-only", "audio-text", "audio+text", "音频", "音频转写", "文稿+音频"}
AUDIO_AMBIGUITY_HEADER = ["时间戳", "原始表述", "当前判断", "候选项", "人工确认"]
DOCUMENT_AMBIGUITY_HEADER = ["原始表述", "当前判断", "候选项", "人工确认"]
PLACEHOLDER_VALUES = {"", "-", "无", "暂无", "无存疑", "暂无存疑", "none", "n/a"}
REQUIRED_WORD_TEXT = ["一、发言整理"]
FORBIDDEN_WORD_TEXT = [
    "AI结构化总结",
    "标的汇总表",
    "三、处理说明",
    "三、标的汇总表",
    "四、存疑与待确认",
    "输入来源",
    "整理说明",
    "#### ",
    "##### ",
]
MEETING_TYPES = {"多人复盘会", "上市公司交流", "专家交流"}
QUESTION_LINE_RE = re.compile(r"^\*\*【[^】\n]+】\*\*\s*$", re.MULTILINE)
SINGLE_TIMESTAMP_RE = re.compile(r"^(?:\d{1,2}:[0-5]\d|\d{1,2}:[0-5]\d:[0-5]\d)$")
VERIFICATION_REQUIRED_FIELDS = [
    "原始表述",
    "存疑类型",
    "当前判断",
    "候选项",
    "是否需要 sidecar",
    "上下文依据",
    "检索/证据路径",
    "最终处理",
]
TIMESTAMP_INDEX_REQUIRED_FIELDS = ["source", "precision", "text", "start", "end", "start_ms", "end_ms"]
RELIABLE_TIMESTAMP_PRECISIONS = {"sentence", "phrase"}
LOW_TIMESTAMP_PRECISIONS = {"segment", "chunk", "unavailable"}
MAX_RELIABLE_VAD_SEGMENT_MS = 10000


def body_section(markdown: str) -> str:
    body_match = re.search(
        r"## 一、发言整理(?P<body>.*?)(?=^## 二、存疑与待确认|\Z)",
        markdown,
        re.S | re.M,
    )
    return body_match.group("body") if body_match else ""


def body_subheadings(markdown: str) -> list[str]:
    return re.findall(r"^###(?!#)\s*(.+?)\s*$", body_section(markdown), re.MULTILINE)


def bold_question_lines(markdown: str) -> list[re.Match[str]]:
    return list(QUESTION_LINE_RE.finditer(body_section(markdown)))


def expert_questions_without_answers(markdown: str) -> list[str]:
    body = body_section(markdown)
    questions = bold_question_lines(markdown)
    findings: list[str] = []
    for index, match in enumerate(questions):
        next_start = questions[index + 1].start() if index + 1 < len(questions) else len(body)
        answer_block = body[match.end() : next_start]
        answer_lines = [
            line.strip()
            for line in answer_block.splitlines()
            if line.strip() and not line.strip().startswith("###")
        ]
        if not answer_lines:
            findings.append(match.group(0).strip())
    return findings


def validate_meeting_type_reference(markdown: str, meeting_type: str) -> list[str]:
    errors: list[str] = []
    if not meeting_type:
        return errors
    if meeting_type not in MEETING_TYPES:
        errors.append(f"会议类型只能是: {' / '.join(sorted(MEETING_TYPES))}")
        return errors

    subheadings = body_subheadings(markdown)
    questions = bold_question_lines(markdown)
    body = body_section(markdown)
    if meeting_type == "多人复盘会":
        if not subheadings:
            errors.append("多人复盘会必须按发言轮次使用三级发言人标题")
        topic_headings = review_speaker_heading_topic_findings(subheadings)
        if topic_headings:
            preview = "；".join(topic_headings[:4])
            errors.append(f"发言人标题不是标的或主题标题: {preview}")
        if "标的：" in body:
            errors.append("多人复盘会小段标题不使用 标的： 前缀，应使用【标的(代码)】标题行")
        if "后半段" in body:
            errors.append("多人复盘会同一发言人再次发言时保留同名标题，不使用（后半段）")
    elif meeting_type == "上市公司交流":
        title = markdown_field(markdown, "会议标题")
        if title and not title.endswith("公司交流会议"):
            errors.append("上市公司交流的会议标题应使用 XX公司交流会议")
        if not metadata_field_present(markdown, "会议标的"):
            errors.append("上市公司交流必须包含会议元信息字段: 会议标的")
        if not subheadings:
            errors.append("上市公司交流必须按实际会议环节使用三级标题")
        if re.search(r"(?m)^\s*【[^】]+】", body):
            errors.append("上市公司交流正文不使用【公司经营】【管理层回应】等额外主题标签")
        if re.search(r"(?m)^\*\*(?:Q|q|提问|问)[:：]", body):
            errors.append("上市公司交流问题格式应为 **【问题】**，不使用 Q：或提问：")
        if "标的：" in body:
            errors.append("上市公司交流正文不重复出现 标的： 行，标的信息只放在元信息")
    elif meeting_type == "专家交流":
        if not questions:
            errors.append("专家交流必须使用加粗问题格式，例如 **【问题原文】**")
        if re.search(r"(?m)^\*\*(?:Q|q|提问|问)[:：]", body):
            errors.append("专家交流问题格式应为 **【问题】**，不使用 Q：或提问：")
        missing_answers = expert_questions_without_answers(markdown)
        if missing_answers:
            preview = "；".join(missing_answers[:4])
            errors.append(f"专家交流每个问题后必须保留对应回答: {preview}")
    return errors


def review_speaker_heading_topic_findings(subheadings: list[str]) -> list[str]:
    findings: list[str] = []
    for heading in subheadings:
        normalized = re.sub(r"[\s｜|、/／：:（）()【】]+", "", heading)
        if not normalized:
            continue
        if any(role in normalized for role in REVIEW_SPEAKER_ROLE_TERMS):
            continue
        if any(term in normalized for term in REVIEW_SPEAKER_HEADING_TOPIC_TERMS):
            findings.append(f"### {heading}")
    return findings


def forbidden_escape_heading_findings(markdown: str) -> list[str]:
    findings: list[str] = []
    for match in re.finditer(r"(?m)^(#{2,6})\s*(.+?)\s*$", markdown):
        heading = match.group(2).strip()
        normalized = re.sub(r"\s+", "", heading.strip(" #"))
        if normalized in FORBIDDEN_ESCAPE_HEADINGS:
            findings.append(match.group(0).strip())
    return findings


def review_meeting_heading_warnings(markdown: str, meeting_type: str) -> list[str]:
    if meeting_type != "多人复盘会":
        return []
    body = body_section(markdown)
    if not body:
        return []

    warnings: list[str] = []
    current_speaker = ""
    has_target_heading = False
    has_sector_heading = False
    for raw_line in body.splitlines():
        line = raw_line.strip()
        if not line or line == "---" or line.startswith("|"):
            continue
        if line.startswith("### "):
            current_speaker = line.removeprefix("### ").strip()
            has_target_heading = False
            has_sector_heading = False
            continue
        if line.startswith("#### "):
            has_target_heading = True
            has_sector_heading = False
            continue
        if line.startswith("##### "):
            has_sector_heading = True
            continue
        if line.startswith("#"):
            continue
        speaker_hint = f"{current_speaker}：" if current_speaker else ""
        preview = line[:50]
        if not has_target_heading:
            warnings.append(f"多人复盘会正文段落缺少相邻标的标题，建议复核: {speaker_hint}{preview}")
        elif not has_sector_heading:
            warnings.append(f"多人复盘会正文段落缺少相邻板块标题，建议复核: {speaker_hint}{preview}")
        if len(warnings) >= 4:
            break
    return warnings


def metadata_field_present(markdown: str, field: str) -> bool:
    return bool(re.search(rf"^\*\*{re.escape(field)}\*\*[:：]\s*\S+", markdown, re.MULTILINE))


def markdown_field(markdown: str, field: str) -> str:
    match = re.search(rf"^\*\*{re.escape(field)}\*\*[:：]\s*(.+?)\s*$", markdown, re.MULTILINE)
    return match.group(1).strip() if match else ""


def body_bold_terms_missing_timestamps(markdown: str) -> list[str]:
    body = body_section(markdown)
    if not body:
        return []
    missing: list[str] = []
    for line in body.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r"^\*\*(Q|q|提问|问)[:：]", stripped) or QUESTION_LINE_RE.match(stripped):
            continue
        for match in re.finditer(r"\*\*([^*\n]{1,120})\*\*", stripped):
            term = match.group(1).strip()
            if not term or term in METADATA_BOLD_LABELS:
                continue
            following = stripped[match.end() : match.end() + 40]
            if not re.match(r"^（存疑时间戳：[^）]+）", following):
                missing.append(term)
    return missing


def body_bold_terms(markdown: str) -> list[str]:
    body = body_section(markdown)
    if not body:
        return []
    terms: list[str] = []
    for line in body.splitlines():
        stripped = line.strip()
        if not stripped or re.match(r"^\*\*(Q|q|提问|问)[:：]", stripped) or QUESTION_LINE_RE.match(stripped):
            continue
        for match in re.finditer(r"\*\*([^*\n]{1,120})\*\*", stripped):
            term = match.group(1).strip()
            if term and term not in METADATA_BOLD_LABELS:
                terms.append(term)
    return terms


def markdown_bold_terms(markdown: str) -> list[str]:
    terms: list[str] = []
    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        metadata_match = re.match(r"^\*\*([^*\n：:]{1,20})\*\*[:：]", stripped)
        if metadata_match and metadata_match.group(1).strip() in METADATA_BOLD_LABELS:
            continue
        if re.match(r"^\*\*(Q|q|提问|问)[:：]", stripped) or QUESTION_LINE_RE.match(stripped):
            continue
        for match in re.finditer(r"\*\*([^*\n]{1,80})\*\*", stripped):
            term = match.group(1).strip()
            if term and term not in METADATA_BOLD_LABELS:
                terms.append(term)
    return terms


def docx_paragraph_text(path: Path) -> tuple[str, list[dict[str, Any]], list[list[list[str]]]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，无法校验 Word 文件") from exc

    document = Document(str(path))
    paragraphs: list[str] = []
    runs: list[dict[str, Any]] = []
    tables: list[list[list[str]]] = []
    for paragraph in document.paragraphs:
        paragraphs.append(paragraph.text)
        for run in paragraph.runs:
            if run.text:
                runs.append({"text": run.text, "bold": bool(run.bold), "underline": bool(run.underline)})
    for table in document.tables:
        table_rows: list[list[str]] = []
        for row in table.rows:
            row_values: list[str] = []
            for cell in row.cells:
                row_values.append(cell.text.strip())
                paragraphs.append(cell.text)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            runs.append({"text": run.text, "bold": bool(run.bold), "underline": bool(run.underline)})
            table_rows.append(row_values)
        tables.append(table_rows)
    return "\n".join(paragraphs), runs, tables


def table_contains_fields(table: list[list[str]], fields: list[str]) -> bool:
    cells = [cell.strip().rstrip("：:") for row in table for cell in row]
    return all(field in cells for field in fields)


def first_row_matches(table: list[list[str]], headers: list[str]) -> bool:
    return bool(table) and [cell.strip() for cell in table[0]] == headers


def markdown_ambiguity_headers(markdown: str) -> list[str]:
    table_lines = ambiguity_table_lines(markdown)
    if not table_lines:
        return []
    return [cell.strip() for cell in table_lines[0].strip("|").split("|")]


def validate_word_contract(docx_path: Path, markdown: str | None = None) -> dict[str, Any]:
    errors: list[str] = []
    warnings: list[str] = []
    if not docx_path.exists():
        return {"ok": False, "errors": [f"Word 文件不存在: {docx_path}"], "warnings": []}

    try:
        text, runs, tables = docx_paragraph_text(docx_path)
    except Exception as exc:  # noqa: BLE001
        return {"ok": False, "errors": [f"Word 文件无法打开: {exc}"], "warnings": []}

    for required in REQUIRED_WORD_TEXT:
        if required not in text:
            errors.append(f"Word 缺少必需章节: {required}")
    for forbidden in FORBIDDEN_WORD_TEXT:
        if forbidden in text:
            errors.append(f"Word 包含禁止内容: {forbidden}")
    if not re.search(r"[\u4e00-\u9fff]", text):
        errors.append("Word 未检测到中文内容")

    if markdown:
        if not any(table_contains_fields(table, REQUIRED_METADATA_FIELDS) for table in tables):
            errors.append("Word 元信息必须使用真正表格，并包含会议日期、整理时间、会议标题、会议类型、会议系列")
        ambiguity_headers = markdown_ambiguity_headers(markdown)
        if ambiguity_headers and not any(first_row_matches(table, ambiguity_headers) for table in tables):
            errors.append(f"Word 存疑表必须使用真正表格，并包含表头: {' | '.join(ambiguity_headers)}")
        for term in markdown_bold_terms(markdown):
            matched = any(term in run["text"] and run["bold"] and run["underline"] for run in runs)
            if not matched:
                errors.append(f"Markdown 加粗存疑词未在 Word 中检测到加粗+下划线: {term}")

    return {
        "ok": not errors,
        "errors": errors,
        "warnings": warnings,
        "paragraph_count": text.count("\n") + 1 if text else 0,
        "run_count": len(runs),
        "table_count": len(tables),
    }


def read_verification_records(path: Path) -> list[dict[str, Any]]:
    if path.suffix == ".jsonl":
        records: list[dict[str, Any]] = []
        for line in path.read_text(encoding="utf-8").splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            item = json.loads(stripped)
            if not isinstance(item, dict):
                raise ValueError("JSONL 每行必须是 object")
            records.append(item)
        return records

    payload = json.loads(path.read_text(encoding="utf-8"))
    if isinstance(payload, list):
        records = payload
    elif isinstance(payload, dict):
        records = payload.get("records")
    else:
        raise ValueError("verification JSON 顶层必须是 object 或 list")
    if not isinstance(records, list):
        raise ValueError("verification JSON 必须包含 records 数组，或顶层直接为数组")
    if not all(isinstance(item, dict) for item in records):
        raise ValueError("verification records 中每条记录必须是 object")
    return records


def validate_verification_sidecar(verification_path: Path | None, *, require_verification: bool = False) -> dict[str, Any]:
    errors: list[str] = []
    warnings: list[str] = []
    if verification_path is None:
        if require_verification:
            errors.append("已要求 verification sidecar，但未传入 --verification")
        return {"ok": not errors, "errors": errors, "warnings": warnings, "record_count": 0}
    if not verification_path.exists():
        return {"ok": False, "errors": [f"verification sidecar 不存在: {verification_path}"], "warnings": [], "record_count": 0}

    try:
        records = read_verification_records(verification_path)
    except (OSError, UnicodeDecodeError, json.JSONDecodeError, ValueError) as exc:
        return {"ok": False, "errors": [f"verification sidecar 无法读取或解析: {exc}"], "warnings": [], "record_count": 0}

    if require_verification and not records:
        errors.append("已要求 verification sidecar，但 records 为空")
    for index, record in enumerate(records, start=1):
        missing = [field for field in VERIFICATION_REQUIRED_FIELDS if field not in record]
        if missing:
            errors.append(f"verification 第 {index} 条缺少字段: {', '.join(missing)}")
    return {"ok": not errors, "errors": errors, "warnings": warnings, "record_count": len(records)}


def read_timestamp_index_records(path: Path) -> list[dict[str, Any]]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    if isinstance(payload, list):
        records = payload
    elif isinstance(payload, dict):
        records = payload.get("timestamp_index") or payload.get("records")
    else:
        raise ValueError("timestamp_index JSON 顶层必须是 object 或 list")
    if not isinstance(records, list):
        raise ValueError("timestamp_index JSON 必须包含 timestamp_index/records 数组，或顶层直接为数组")
    if not all(isinstance(item, dict) for item in records):
        raise ValueError("timestamp_index 中每条记录必须是 object")
    return records


def _timestamp_index_record_is_used_for_doubtful(record: dict[str, Any]) -> bool:
    for field in ("used_for_doubtful", "selected_for_ambiguity", "selected_for_doubtful"):
        if bool(record.get(field)):
            return True
    return bool(str(record.get("doubtful_term") or "").strip())


def _timestamp_record_ms(record: dict[str, Any], field: str) -> int | None:
    try:
        return int(round(float(record.get(field))))
    except (TypeError, ValueError):
        return None


def _timestamp_record_is_reliable_for_doubtful(record: dict[str, Any]) -> bool:
    precision = str(record.get("precision") or "").strip().lower()
    source = str(record.get("source") or "").strip().lower()
    if precision in RELIABLE_TIMESTAMP_PRECISIONS:
        return True
    if precision == "segment" and source == "sensevoice_vad_segment":
        start_ms = _timestamp_record_ms(record, "start_ms")
        end_ms = _timestamp_record_ms(record, "end_ms")
        if start_ms is None or end_ms is None:
            return False
        duration_ms = end_ms - start_ms
        return 0 < duration_ms <= MAX_RELIABLE_VAD_SEGMENT_MS
    return False


def validate_timestamp_index_records(
    records: list[dict[str, Any]],
    *,
    require_reliable: bool = False,
) -> dict[str, Any]:
    errors: list[str] = []
    warnings: list[str] = []
    reliable_count = 0
    selected_count = 0
    for index, record in enumerate(records, start=1):
        missing = [field for field in TIMESTAMP_INDEX_REQUIRED_FIELDS if field not in record]
        if missing:
            errors.append(f"timestamp_index 第 {index} 条缺少字段: {', '.join(missing)}")
        precision = str(record.get("precision") or "").strip().lower()
        used_for_doubtful = _timestamp_index_record_is_used_for_doubtful(record)
        if used_for_doubtful:
            selected_count += 1
        if _timestamp_record_is_reliable_for_doubtful(record):
            reliable_count += 1
        if precision in RELIABLE_TIMESTAMP_PRECISIONS:
            for field in ("start", "end", "start_ms", "end_ms"):
                if str(record.get(field) or "").strip() == "":
                    errors.append(f"timestamp_index 第 {index} 条句级/短句级 anchor 缺少 {field}")
        elif precision == "segment":
            source = str(record.get("source") or "").strip().lower()
            if source != "sensevoice_vad_segment":
                errors.append(f"timestamp_index 第 {index} 条低精度 segment 不得作为可靠存疑时间戳；只有短 VAD segment 可用")
            else:
                start_ms = _timestamp_record_ms(record, "start_ms")
                end_ms = _timestamp_record_ms(record, "end_ms")
                if start_ms is None or end_ms is None:
                    errors.append(f"timestamp_index 第 {index} 条短 VAD segment start_ms/end_ms 必须为数字")
                else:
                    duration_ms = end_ms - start_ms
                    if duration_ms <= 0:
                        errors.append(f"timestamp_index 第 {index} 条短 VAD segment duration_ms 必须大于 0")
                    elif duration_ms > MAX_RELIABLE_VAD_SEGMENT_MS:
                        errors.append(
                            f"timestamp_index 第 {index} 条短 VAD segment duration_ms 超过 10000: {duration_ms}"
                        )
        elif precision in LOW_TIMESTAMP_PRECISIONS:
            if used_for_doubtful:
                errors.append(f"timestamp_index 第 {index} 条低精度 {precision} 不得作为可靠存疑时间戳")
        elif precision:
            warnings.append(f"timestamp_index 第 {index} 条 precision 未知，建议复核: {precision}")
        elif require_reliable:
            errors.append(f"timestamp_index 第 {index} 条缺少 precision")

        if require_reliable and used_for_doubtful:
            if not _timestamp_record_is_reliable_for_doubtful(record):
                errors.append(
                    f"timestamp_index 第 {index} 条被标记用于存疑，但不是 sentence/phrase 或短 VAD segment: {precision or '空'}"
                )

    if require_reliable and records and reliable_count == 0:
        errors.append("要求可靠时间戳时，timestamp_index 至少应包含 sentence/phrase anchor 或短 VAD segment")
    if require_reliable and selected_count == 0:
        warnings.append("要求可靠时间戳时，timestamp_index 未标记 used_for_doubtful/selected_for_ambiguity；仅完成字段和精度基础检查")
    return {
        "ok": not errors,
        "errors": errors,
        "warnings": warnings,
        "record_count": len(records),
        "reliable_count": reliable_count,
    }


def validate_timestamp_index_file(
    timestamp_index_path: Path | None,
    *,
    require_reliable: bool = False,
) -> dict[str, Any]:
    if timestamp_index_path is None:
        return {"ok": True, "errors": [], "warnings": [], "record_count": 0, "reliable_count": 0}
    if not timestamp_index_path.exists():
        return {
            "ok": False,
            "errors": [f"timestamp_index 不存在: {timestamp_index_path}"],
            "warnings": [],
            "record_count": 0,
            "reliable_count": 0,
        }
    try:
        records = read_timestamp_index_records(timestamp_index_path)
    except (OSError, UnicodeDecodeError, json.JSONDecodeError, ValueError) as exc:
        return {
            "ok": False,
            "errors": [f"timestamp_index 无法读取或解析: {exc}"],
            "warnings": [],
            "record_count": 0,
            "reliable_count": 0,
        }
    return validate_timestamp_index_records(records, require_reliable=require_reliable)


def audio_timestamp_fallbacks(markdown: str) -> list[str]:
    findings: list[str] = []
    for match in re.finditer(r"存疑时间戳：未提供", markdown):
        start = max(0, match.start() - 40)
        end = min(len(markdown), match.end() + 40)
        findings.append(re.sub(r"\s+", " ", markdown[start:end]).strip())
    for match in re.finditer(r"(?m)^\|\s*未提供\s*\|", markdown):
        start = max(0, match.start() - 40)
        end = min(len(markdown), match.end() + 80)
        findings.append(re.sub(r"\s+", " ", markdown[start:end]).strip())

    return findings


def is_valid_timestamp_value(value: str) -> bool:
    normalized = re.sub(r"\s+", "", value.strip())
    if not normalized:
        return False
    parts = normalized.split("-")
    if len(parts) not in {1, 2}:
        return False
    return all(SINGLE_TIMESTAMP_RE.match(part) for part in parts)


def invalid_inline_timestamp_markers(markdown: str) -> list[str]:
    findings: list[str] = []
    for match in re.finditer(r"存疑时间戳[:：]([^）\n]+)", markdown):
        value = match.group(1).strip()
        if is_valid_timestamp_value(value):
            continue
        start = max(0, match.start() - 40)
        end = min(len(markdown), match.end() + 60)
        findings.append(re.sub(r"\s+", " ", markdown[start:end]).strip())
    return findings


def invalid_table_timestamp_values(markdown: str, headers: list[str]) -> list[str]:
    if "时间戳" not in headers:
        return []
    timestamp_index = headers.index("时间戳")
    original_index = headers.index("原始表述") if "原始表述" in headers else 0
    findings: list[str] = []
    for line in ambiguity_table_lines(markdown)[2:]:
        if _is_separator_line(line):
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if _is_placeholder_ambiguity_row(cells, original_index):
            continue
        value = cells[timestamp_index].strip() if timestamp_index < len(cells) else ""
        if not is_valid_timestamp_value(value):
            findings.append(line)
    return findings


def contains_timestamp_markers(markdown: str) -> list[str]:
    findings: list[str] = []
    for pattern in (
        r"存疑时间戳[:：]",
        r"(?m)^\|\s*时间戳\s*\|",
        r"(?m)^\|\s*(?:\d{1,2}:)?\d{1,2}:\d{2}(?:\s*-\s*(?:\d{1,2}:)?\d{1,2}:\d{2})?\s*\|",
        r"(?m)^\|\s*未提供\s*\|",
    ):
        for match in re.finditer(pattern, markdown):
            start = max(0, match.start() - 40)
            end = min(len(markdown), match.end() + 60)
            findings.append(re.sub(r"\s+", " ", markdown[start:end]).strip())
            if len(findings) >= 8:
                return findings
    return findings


def normalize_source_mode(source_mode: str | None) -> str:
    value = str(source_mode or "").strip().lower()
    if value in DOCUMENT_ONLY_SOURCE_MODES:
        return "document"
    if value in AUDIO_SOURCE_MODES:
        return "audio"
    return "auto"


def normalize_timestamp_mode(timestamp_mode: str | None, require_audio_timestamps: bool) -> str:
    if require_audio_timestamps:
        return "reliable"
    value = str(timestamp_mode or "auto").strip().lower()
    if value in {"reliable", "unavailable", "auto"}:
        return value
    return "auto"


def expected_ambiguity_headers(normalized_source_mode: str, timestamp_mode: str) -> list[list[str]]:
    if timestamp_mode == "reliable":
        return [AUDIO_AMBIGUITY_HEADER]
    if timestamp_mode == "unavailable" or normalized_source_mode == "document":
        return [DOCUMENT_AMBIGUITY_HEADER]
    return [AUDIO_AMBIGUITY_HEADER, DOCUMENT_AMBIGUITY_HEADER]


def format_allowed_headers(headers: list[list[str]]) -> str:
    return " 或 ".join(" | ".join(header) for header in headers)


def has_ambiguity_section(markdown: str) -> bool:
    return bool(re.search(r"^## 二、存疑与待确认\s*$", markdown, re.MULTILINE))


def ambiguity_table_lines(markdown: str) -> list[str]:
    ambiguity_match = re.search(r"## 二、存疑与待确认(?P<body>.*)$", markdown, re.S)
    if not ambiguity_match:
        return []
    return [
        line.strip()
        for line in ambiguity_match.group("body").splitlines()
        if line.strip().startswith("|")
    ]


def _is_separator_line(line: str) -> bool:
    return bool(re.match(r"^\|\s*[-:| ]+\s*$", line.strip()))


def _is_placeholder_ambiguity_row(cells: list[str], original_index: int) -> bool:
    if not cells or original_index >= len(cells):
        return True
    original = cells[original_index].strip().lower()
    return original in PLACEHOLDER_VALUES


def ambiguity_rows_with_manual_confirmation(markdown: str, headers: list[str]) -> list[str]:
    if "人工确认" not in headers:
        return ["存疑表缺少强制列: 人工确认"]
    if headers[-1] != "人工确认":
        return ["人工确认列必须是最后一列"]
    manual_index = headers.index("人工确认")
    original_index = headers.index("原始表述") if "原始表述" in headers else 0
    findings: list[str] = []
    table_lines = ambiguity_table_lines(markdown)
    for line in table_lines[2:]:
        if _is_separator_line(line):
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if _is_placeholder_ambiguity_row(cells, original_index):
            continue
        if manual_index >= len(cells):
            findings.append(line)
            continue
        if cells[manual_index].strip():
            findings.append(line)
    return findings


def real_ambiguity_row_count(markdown: str, headers: list[str]) -> int:
    original_index = headers.index("原始表述") if "原始表述" in headers else 0
    count = 0
    for line in ambiguity_table_lines(markdown)[2:]:
        if _is_separator_line(line):
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if not _is_placeholder_ambiguity_row(cells, original_index):
            count += 1
    return count


def validate_contract(
    markdown: str,
    *,
    required_terms: list[str] | None = None,
    forbidden_terms: list[str] | None = None,
    source_mode: str | None = None,
    require_audio_timestamps: bool = False,
    timestamp_mode: str | None = "auto",
) -> dict[str, Any]:
    errors: list[str] = []
    warnings: list[str] = []
    normalized_source_mode = normalize_source_mode(source_mode)
    normalized_timestamp_mode = normalize_timestamp_mode(timestamp_mode, require_audio_timestamps)
    ambiguity_headers: list[str] = []

    stripped = markdown.lstrip()
    if not stripped.startswith("# "):
        errors.append("Markdown 必须以一级标题开头")
    elif not re.search(r"(?m)^# 投资会议纪要(?:\s|｜|$)", stripped):
        errors.append("一级标题必须以 # 投资会议纪要 开头")

    for field in REQUIRED_METADATA_FIELDS:
        if not metadata_field_present(markdown, field):
            errors.append(f"缺少会议元信息字段: {field}")

    meeting_type = markdown_field(markdown, "会议类型")
    errors.extend(validate_meeting_type_reference(markdown, meeting_type))
    warnings.extend(review_meeting_heading_warnings(markdown, meeting_type))

    body_position = markdown.find("## 一、发言整理")
    if body_position < 0:
        errors.append("缺少必需章节: ## 一、发言整理")
    ambiguity_position = markdown.find("## 二、存疑与待确认")
    if body_position >= 0 and ambiguity_position >= 0 and body_position > ambiguity_position:
        errors.append("必需章节顺序错误")

    for pattern in FORBIDDEN_PATTERNS:
        if re.search(pattern, markdown):
            errors.append(f"包含禁止输出内容: {pattern}")
    escape_headings = forbidden_escape_heading_findings(markdown)
    if escape_headings:
        preview = "；".join(escape_headings[:6])
        errors.append(f"包含契约外逃逸标题: {preview}")

    has_body_section = "## 一、发言整理" in markdown
    has_subheading = bool(re.search(r"^###(?!#)\s+", markdown, re.MULTILINE))
    has_bold_question = bool(bold_question_lines(markdown))
    if has_body_section and not has_subheading and not has_bold_question:
        warnings.append("发言整理中未检测到三级发言人标题")

    bold_terms = body_bold_terms(markdown)
    ambiguity_exists = has_ambiguity_section(markdown)
    if not ambiguity_exists and bold_terms:
        preview = "、".join(bold_terms[:8])
        errors.append(f"正文存在加粗存疑词但缺少 ## 二、存疑与待确认: {preview}")

    if ambiguity_exists:
        table_lines = ambiguity_table_lines(markdown)
        if not table_lines:
            errors.append("存疑与待确认章节存在时必须包含 Markdown 表格；无存疑内容时应省略整节")
        else:
            headers = [cell.strip() for cell in table_lines[0].strip("|").split("|")]
            ambiguity_headers = headers
            allowed_headers = expected_ambiguity_headers(normalized_source_mode, normalized_timestamp_mode)
            if headers not in allowed_headers:
                errors.append(f"存疑与待确认表格表头必须固定为: {format_allowed_headers(allowed_headers)}")
            if real_ambiguity_row_count(markdown, headers) == 0:
                errors.append("存疑与待确认章节存在时必须包含至少一条真实存疑；无存疑内容时应省略整节")
            manual_confirmation_rows = ambiguity_rows_with_manual_confirmation(markdown, headers)
            if manual_confirmation_rows:
                preview = "；".join(manual_confirmation_rows[:4])
                errors.append(f"人工确认列必须存在于最后且保持空白，供人工填写: {preview}")

    if normalized_source_mode == "document" or normalized_timestamp_mode == "unavailable":
        timestamp_markers = contains_timestamp_markers(markdown)
        if timestamp_markers:
            preview = "；".join(timestamp_markers[:4])
            errors.append(f"文稿或无可靠时间戳模式不应出现存疑时间戳或时间戳列: {preview}")

    invalid_inline_timestamps = invalid_inline_timestamp_markers(markdown)
    invalid_table_timestamps = invalid_table_timestamp_values(markdown, ambiguity_headers)
    if invalid_inline_timestamps or invalid_table_timestamps:
        preview = "；".join((invalid_inline_timestamps + invalid_table_timestamps)[:4])
        errors.append(f"存疑时间戳格式必须为 MM:SS、HH:MM:SS 或同格式范围: {preview}")

    if normalized_timestamp_mode == "reliable":
        missing_body_timestamps = body_bold_terms_missing_timestamps(markdown)
        if missing_body_timestamps:
            preview = "、".join(missing_body_timestamps[:8])
            errors.append(f"可靠音频模式下正文加粗存疑词后必须补充存疑时间戳: {preview}")

    if normalized_source_mode == "audio" or normalized_timestamp_mode == "reliable":
        fallback_timestamps = audio_timestamp_fallbacks(markdown)
        if fallback_timestamps:
            preview = "；".join(fallback_timestamps[:4])
            errors.append(f"音频来源不允许将存疑时间戳写成未提供: {preview}")

    for term in required_terms or []:
        if term not in markdown:
            errors.append(f"缺少样例关键锚点: {term}")
    for term in forbidden_terms or []:
        if term in markdown:
            errors.append(f"包含样例禁止锚点: {term}")

    return {
        "ok": not errors,
        "errors": errors,
        "warnings": warnings,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="校验会议纪要 Markdown 是否符合当前输出契约")
    parser.add_argument("markdown_file", help="待校验 Markdown 文件")
    parser.add_argument("--word", help="可选：同时校验对应导出 Word 文件")
    parser.add_argument("--verification", help="可选：同时校验同 stem 的 verification JSON/JSONL 结构")
    parser.add_argument("--require-verification", action="store_true", help="要求 verification sidecar 存在且至少包含一条记录")
    parser.add_argument("--timestamp-index", help="可选：同时校验 timestamp_index.json 字段和精度")
    parser.add_argument("--require-reliable-timestamp-index", action="store_true", help="要求 timestamp_index 包含可用于存疑的 sentence/phrase anchor")
    parser.add_argument("--require-term", action="append", default=[], help="必须保留的关键原文锚点，可重复")
    parser.add_argument("--forbid-term", action="append", default=[], help="样例或人工复核中不应出现的改写锚点，可重复")
    parser.add_argument(
        "--source-mode",
        choices=["auto", "document", "document-only", "text", "text-only", "audio", "audio-only", "audio-text", "audio+text"],
        default="auto",
        help="来源模式；需结合 --timestamp-mode 判断存疑表是否包含时间戳列",
    )
    parser.add_argument(
        "--timestamp-mode",
        choices=["auto", "reliable", "unavailable"],
        default="auto",
        help="时间戳可靠性；reliable 要求带时间戳表头，unavailable 要求无时间戳表头，auto 接受两者",
    )
    parser.add_argument("--require-audio-timestamps", action="store_true", help="音频来源不允许把存疑时间戳批量写成未提供")
    parser.add_argument("--json", action="store_true", help="输出 JSON")
    args = parser.parse_args()

    path = Path(args.markdown_file).expanduser()
    try:
        markdown = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        payload = {"ok": False, "errors": [f"文件不是有效 UTF-8: {exc}"], "warnings": []}
        if args.json:
            print(json.dumps(payload, ensure_ascii=False, indent=2))
        else:
            print(payload["errors"][0], file=sys.stderr)
        return 1

    result = validate_contract(
        markdown,
        required_terms=args.require_term,
        forbidden_terms=args.forbid_term,
        source_mode=args.source_mode,
        require_audio_timestamps=args.require_audio_timestamps,
        timestamp_mode=args.timestamp_mode,
    )
    if args.word:
        word_result = validate_word_contract(Path(args.word).expanduser(), markdown)
        result["word"] = word_result
        result["errors"].extend(word_result["errors"])
        result["warnings"].extend(word_result["warnings"])
        result["ok"] = result["ok"] and word_result["ok"]
    if args.verification or args.require_verification:
        verification_path = Path(args.verification).expanduser() if args.verification else None
        verification_result = validate_verification_sidecar(
            verification_path,
            require_verification=args.require_verification,
        )
        result["verification"] = verification_result
        result["errors"].extend(verification_result["errors"])
        result["warnings"].extend(verification_result["warnings"])
        result["ok"] = result["ok"] and verification_result["ok"]
    if args.timestamp_index:
        timestamp_index_result = validate_timestamp_index_file(
            Path(args.timestamp_index).expanduser(),
            require_reliable=args.require_reliable_timestamp_index,
        )
        result["timestamp_index"] = timestamp_index_result
        result["errors"].extend(timestamp_index_result["errors"])
        result["warnings"].extend(timestamp_index_result["warnings"])
        result["ok"] = result["ok"] and timestamp_index_result["ok"]
    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        if result["ok"]:
            print(f"{path}: ok")
        for warning in result["warnings"]:
            print(f"warning: {warning}", file=sys.stderr)
        for error in result["errors"]:
            print(f"error: {error}", file=sys.stderr)
    return 0 if result["ok"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
