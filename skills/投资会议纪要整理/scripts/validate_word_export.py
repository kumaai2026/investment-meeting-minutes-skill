#!/usr/bin/env python3
"""Validate exported Word files for the current meeting-minutes contract."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

REQUIRED_TEXT = ["一、逐发言人原文整理", "二、存疑与待确认"]
FORBIDDEN_TEXT = ["AI结构化总结", "标的汇总表", "三、处理说明", "三、标的汇总表", "四、存疑与待确认", "输入来源", "整理说明"]
METADATA_BOLD_LABELS = {
    "会议日期",
    "整理时间",
    "会议标题",
    "会议类型",
    "会议系列",
    "会议标的",
}


def docx_paragraph_text(path: Path) -> tuple[str, list[dict[str, Any]]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，无法校验 Word 文件") from exc

    document = Document(str(path))
    paragraphs: list[str] = []
    runs: list[dict[str, Any]] = []
    for paragraph in document.paragraphs:
        paragraphs.append(paragraph.text)
        for run in paragraph.runs:
            if run.text:
                runs.append({"text": run.text, "bold": bool(run.bold), "underline": bool(run.underline)})
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            runs.append({"text": run.text, "bold": bool(run.bold), "underline": bool(run.underline)})
    return "\n".join(paragraphs), runs


def markdown_bold_terms(path: Path) -> list[str]:
    if not path or not path.exists():
        return []
    markdown = path.read_text(encoding="utf-8")
    terms = []
    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        metadata_match = re.match(r"^\*\*([^*\n：:]{1,20})\*\*[:：]", stripped)
        if metadata_match and metadata_match.group(1).strip() in METADATA_BOLD_LABELS:
            continue
        if re.match(r"^\*\*(Q|q|提问|问)[:：]", stripped):
            continue
        for match in re.finditer(r"\*\*([^*\n]{1,80})\*\*", stripped):
            term = match.group(1).strip()
            if term and term not in METADATA_BOLD_LABELS:
                terms.append(term)
    return terms


def validate_word(docx_path: Path, markdown_path: Path | None = None) -> dict[str, Any]:
    errors: list[str] = []
    warnings: list[str] = []
    if not docx_path.exists():
        return {"ok": False, "errors": [f"Word 文件不存在: {docx_path}"], "warnings": []}

    try:
        text, runs = docx_paragraph_text(docx_path)
    except Exception as exc:  # noqa: BLE001
        return {"ok": False, "errors": [f"Word 文件无法打开: {exc}"], "warnings": []}

    for required in REQUIRED_TEXT:
        if required not in text:
            errors.append(f"Word 缺少必需章节: {required}")
    for forbidden in FORBIDDEN_TEXT:
        if forbidden in text:
            errors.append(f"Word 包含禁止内容: {forbidden}")
    if not re.search(r"[\u4e00-\u9fff]", text):
        errors.append("Word 未检测到中文内容")

    if markdown_path:
        for term in markdown_bold_terms(markdown_path):
            matched = any(term in run["text"] and run["bold"] and run["underline"] for run in runs)
            if not matched:
                errors.append(f"Markdown 加粗存疑词未在 Word 中检测到加粗+下划线: {term}")

    return {
        "ok": not errors,
        "errors": errors,
        "warnings": warnings,
        "paragraph_count": text.count("\n") + 1 if text else 0,
        "run_count": len(runs),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="校验导出的 Word 会议纪要是否符合当前结构契约")
    parser.add_argument("docx_file", help="导出的 .docx 文件")
    parser.add_argument("--markdown", help="对应 Markdown 文件，用于检查存疑词样式")
    parser.add_argument("--json", action="store_true", help="输出 JSON")
    args = parser.parse_args()

    result = validate_word(
        Path(args.docx_file).expanduser(),
        Path(args.markdown).expanduser() if args.markdown else None,
    )
    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        if result["ok"]:
            print(f"{args.docx_file}: ok")
        for warning in result["warnings"]:
            print(f"warning: {warning}", file=sys.stderr)
        for error in result["errors"]:
            print(f"error: {error}", file=sys.stderr)
    return 0 if result["ok"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
